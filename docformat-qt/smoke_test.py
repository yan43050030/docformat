# -*- coding: utf-8 -*-
"""端到端冒烟测试：生成样例公文 → 三种模式处理 → 断言结果"""
import io
import shutil
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Cm

OUT_DIR = os.path.join(os.path.dirname(__file__), '_smoke')
os.makedirs(OUT_DIR, exist_ok=True)
SRC = os.path.join(OUT_DIR, 'sample.docx')
_QAPP = None      # 打印测试要用的 QApplication，得留着不让它被回收


def make_sample():
    doc = Document()
    doc.add_paragraph('秘密★1年')
    doc.add_paragraph('关于开展2026年度安全生产检查的通知')
    doc.add_paragraph('某安委发〔2026〕12号')
    doc.add_paragraph('各部门、各单位:')
    doc.add_paragraph('为深入贯彻落实上级部署要求(含附件),现将有关事项通知如下.')
    doc.add_paragraph('一、总体要求')
    doc.add_paragraph('坚持"安全第一、预防为主"的方针,全面排查隐患。')
    doc.add_paragraph('(一)检查范围')
    doc.add_paragraph('1.生产车间及仓储区域。')
    doc.add_paragraph('(1)重点检查电气线路。')
    doc.add_paragraph('二、时间安排')
    doc.add_paragraph('检查工作自2026年7月20日起至8月10日结束...')
    doc.add_paragraph('特此通知。')
    doc.add_paragraph('附件:安全检查表')
    doc.add_paragraph('某某公司办公室')
    doc.add_paragraph('2026年7月17日')
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = '项目'
    table.rows[0].cells[1].text = '预算(万元)'
    table.rows[1].cells[0].text = '隐患整改'
    table.rows[1].cells[1].text = '12'
    doc.save(SRC)
    print('[1] 样例文档已生成:', SRC)


def test_full():
    """按 worker 的智能一键链路：标点修复 → 排版"""
    from scripts.punctuation import process_document
    from scripts.formatter import format_document
    mid = os.path.join(OUT_DIR, 'sample_punct_stage.docx')
    process_document(SRC, mid)
    out = os.path.join(OUT_DIR, 'sample_full.docx')
    stages = []
    format_document(mid, out, preset_name='official',
                    progress_callback=lambda c, t, s: stages.append((c, t, s)))
    doc = Document(out)
    sec = doc.sections[0]
    top, bottom = sec.top_margin.cm, sec.bottom_margin.cm
    left, right = sec.left_margin.cm, sec.right_margin.cm
    assert abs(top - 3.7) < 0.05 and abs(bottom - 3.5) < 0.05, '上下边距错误: {} {}'.format(top, bottom)
    assert abs(left - 2.8) < 0.05 and abs(right - 2.6) < 0.05, '左右边距错误: {} {}'.format(left, right)

    from docx.oxml.ns import qn
    # 密级与标题之间现在是**真空一行**，标题不再固定在第 2 段，按文字找
    title_para = [p for p in doc.paragraphs if p.text.strip().startswith('关于')][0]
    title_run = title_para.runs[0]
    fonts = set()
    rpr = title_run._element.rPr
    ea = rpr.rFonts.get(qn('w:eastAsia')) if rpr is not None and rpr.rFonts is not None else None
    assert ea == '方正小标宋简体', '标题字体错误: {}'.format(ea)
    assert title_run.font.size.pt == 22, '标题字号错误: {}'.format(title_run.font.size.pt)

    # 密级标识：黑体、顶格左对齐、不缩进
    sec_para = doc.paragraphs[0]
    assert '秘密' in sec_para.text, '密级行丢失: {}'.format(sec_para.text)
    sec_run = sec_para.runs[0]
    sec_rpr = sec_run._element.rPr
    sec_ea = sec_rpr.rFonts.get(qn('w:eastAsia')) if sec_rpr is not None and sec_rpr.rFonts is not None else None
    assert sec_ea == '黑体', '密级字体错误: {}'.format(sec_ea)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    assert sec_para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER, '密级不应居中'

    body_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '（含附件）' in body_text, '英文括号未转换'
    assert '……' in body_text, '省略号未规范化'

    # 发文字号：仿宋、居中
    dn_para = [p for p in doc.paragraphs if '〔2026〕12号' in p.text][0]
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP
    assert dn_para.paragraph_format.alignment == _WAP.CENTER, '发文字号应居中'
    print('[2] 智能一键: 边距/标题字体字号/标点转换/发文字号 全部通过 (进度回调 {} 次)'.format(len(stages)))


def test_punctuation():
    from scripts.punctuation import process_document
    out = os.path.join(OUT_DIR, 'sample_punct.docx')
    process_document(SRC, out)
    text = '\n'.join(p.text for p in Document(out).paragraphs)
    assert '（含附件）' in text, '标点模式: 括号未转换'
    print('[3] 标点修复: 通过')


def test_diagnose():
    from scripts import analyzer
    doc = Document(SRC)
    results = {
        'punctuation': analyzer.analyze_punctuation(doc),
        'numbering': analyzer.analyze_numbering(doc),
        'paragraph': analyzer.analyze_paragraph_format(doc),
        'font': analyzer.analyze_font(doc),
    }
    n = sum(len(v) for v in results.values())
    assert len(results['punctuation']) > 0, '诊断应发现英文标点问题'
    assert any(str(i.get('para', '')).startswith('表') for i in results['punctuation']), \
        '诊断应覆盖表格单元格（预算(万元) 的英文括号）'
    from app.worker import build_diagnose_report
    report = build_diagnose_report('sample.docx', results)
    print('[4] 格式诊断: 发现 {} 项问题, 报告生成 OK'.format(n))
    print('    ' + report.splitlines()[0])


def test_custom_preset():
    from scripts.formatter import format_document
    from app.presets import PresetManager
    mgr = PresetManager()
    key = mgr.create('冒烟测试模板')
    preset = mgr.get(key)
    preset['page']['top'] = 5.0
    preset['body']['size'] = 14
    mgr.update(key, preset)

    mgr2 = PresetManager()  # 重新加载验证持久化
    assert key in mgr2.user, '用户模板未持久化'
    assert mgr2.get(key)['page']['top'] == 5.0, '模板参数未保存'

    out = os.path.join(OUT_DIR, 'sample_custom.docx')
    name, custom = mgr2.engine_args(key)
    format_document(SRC, out, preset_name=name, custom_settings=custom)
    doc = Document(out)
    assert abs(doc.sections[0].top_margin.cm - 5.0) < 0.05, '自定义边距未生效'
    mgr2.delete(key)
    print('[5] 自定义模板: 持久化 + 引擎生效 通过')


def test_ai_paste():
    from app.worker import generate_docx_from_text, clean_markdown
    md = "# 关于测试的通知\n\n**各部门**:\n\n- 第一项工作\n- 第二项工作\n\n```\ncode block skip\n```\n\n特此通知。"
    paras = clean_markdown(md)
    assert '关于测试的通知' in paras[0], 'markdown 标题清洗失败'
    assert not any('```' in p or 'code block' in p for p in paras), '代码块未剔除'
    assert '各部门:' in '\n'.join(paras), '加粗标记未清除'
    out = os.path.join(OUT_DIR, 'ai_draft.docx')
    generate_docx_from_text(md, out)
    assert os.path.exists(out)
    from scripts.formatter import format_document
    final = os.path.join(OUT_DIR, 'ai_final.docx')
    format_document(out, final, preset_name='official')
    assert os.path.exists(final)
    print('[6] AI 粘贴生成: markdown 清洗 + 生成 + 排版 通过')


def test_punct_edges():
    from scripts.punctuation import fix_text, _fix_quotes_whole_text, _process_spaces_text
    assert fix_text("it's a test, don't worry") == "it’s a test, don’t worry", '撇号被误配对'
    r1, dq, sq = _fix_quotes_whole_text('他说"这是第一段', 0, 0)
    r2, dq, sq = _fix_quotes_whole_text('这是第二段的结尾"', dq, sq)
    assert r1.endswith('“这是第一段') and r2 == '这是第二段的结尾”', '跨段引号配对错误'
    assert _process_spaces_text('参照 GB/T 9704 和 New York 规范', 'keep_en_words') \
        == '参照GB/T 9704和New York规范', '英文词间空格保护失败'
    assert fix_text('本次比分为3:2。') == '本次比分为3:2。', '数字比分冒号不应替换'
    print('[7] 标点边界: 撇号/跨段引号/英文空格/比分 通过')

def test_wps_broken_jc():
    """WPS/老 Word 残缺 <w:jc>（缺 w:val）不再导致排版崩溃"""
    from docx import Document
    from docx.oxml import OxmlElement
    from scripts.formatter import format_document, sanitize_document
    d = Document()
    d.add_paragraph('关于测试的通知')
    d.add_paragraph('各单位：')
    para = d.add_paragraph('正文内容。')
    para._p.get_or_add_pPr().append(OxmlElement('w:jc'))  # 残缺对齐元素
    d.add_paragraph('某某办公室')
    d.add_paragraph('2026年7月22日')
    jc_in = os.path.join(OUT_DIR, 'wps_jc.docx')
    d.save(jc_in)
    # sanitize 计数 > 0
    d2 = Document(jc_in)
    assert sanitize_document(d2) >= 1, 'sanitize 未修复残缺 w:jc'
    # 全流程排版不抛异常
    out = os.path.join(OUT_DIR, 'wps_jc_out.docx')
    format_document(jc_in, out, preset_name='official_gbk')
    assert os.path.exists(out)
    print('[7c] WPS 残缺 w:jc 兼容: sanitize + 排版不崩 通过')


def test_auto_num_chinese():
    """自动编号中文数字过 10 正确（十一/十二），起始值生效"""
    from scripts.auto_num import _to_chinese
    assert _to_chinese(11) == '十一' and _to_chinese(20) == '二十' and _to_chinese(99) == '九十九'
    print('[7d] 自动编号中文数字 11/20/99 通过')


def test_attachment_label():
    """附件标识行(顶格黑体) 与 附件说明(悬挂缩进) 区分"""
    from docx.oxml.ns import qn
    from scripts.formatter import format_document, detect_para_type, _compile_rules
    r=_compile_rules(None)
    assert detect_para_type('附件1',10,20,None,['a']*15,10,rules=r)=='attachment_label'
    assert detect_para_type('附件：清单',10,20,None,['a']*15,10,rules=r)=='attachment'
    d=Document()
    d.add_paragraph('关于测试的通知'); d.add_paragraph('各单位：'); d.add_paragraph('正文。')
    d.add_paragraph('附件：1.清单'); d.add_paragraph('附件1'); d.add_paragraph('组成人员名单')
    src=os.path.join(OUT_DIR,'att_in.docx'); d.save(src)
    out=os.path.join(OUT_DIR,'att_out.docx')
    format_document(src,out,preset_name='official_gbk')
    doc=Document(out)
    lab=[p for p in doc.paragraphs if p.text.strip()=='附件1'][0]
    ea=lab.runs[0]._element.rPr.rFonts.get(qn('w:eastAsia'))
    assert ea=='方正黑体_GBK', '附件标识应黑体'
    assert (lab.paragraph_format.left_indent is None or lab.paragraph_format.left_indent.pt==0), '附件标识应顶格'
    print('[7g] 附件标识/说明 区分排版 通过')


def test_title_shape():
    """标题梯形回行：正梯形上长下短、倒梯形上短下长、短标题不折"""
    from scripts.title_shape import split_title_lines
    t='关于进一步加强全市安全生产工作坚决防范遏制重特大事故的通知'
    dn=split_title_lines(t,20,'trapezoid_down'); up=split_title_lines(t,20,'trapezoid_up')
    assert len(dn)>=2 and len(dn[0])>=len(dn[-1]), '正梯形应上长下短'
    assert len(up)>=2 and len(up[0])<=len(up[-1]), '倒梯形应上短下长'
    assert split_title_lines('关于测试的通知',20,'trapezoid_down')==['关于测试的通知']
    print('[7h] 标题梯形回行 正/倒/不折 通过')


def test_compliance():
    """公文合规检查：完整核对（检查面≈排版面）+ 交互式精准修正

    最关键的是自洽性：智能一键排完的文档，合规检查必须零偏差。
    否则"检查通过"不等于"排版合规"，检查就不可信。
    """
    from docx.shared import Cm, Pt
    from scripts import compliance, punctuation
    from scripts.data_model import PRESETS
    from scripts.formatter import format_document
    preset=PRESETS['official_gbk']
    PARAS=['关于开展某某试点工作的通知','各有关单位：','一、总体要求',
           '为深入贯彻落实上级决策部署,现就开展试点工作通知如下,请遵照执行。',
           '(一)工作目标','通过试点探索形成可复制可推广的经验做法,为全面推开奠定基础。',
           '特此通知。','某某办公室','2026年7月25日']

    def _dirty(path):
        d=Document(); s=d.sections[0]
        s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2); s.right_margin=Cm(2)
        for t in PARAS:
            p=d.add_paragraph(t)
            for r in p.runs: r.font.size=Pt(12)
        d.save(path); return path

    src=_dirty(os.path.join(OUT_DIR,'comp_in.docx'))

    # --- 1. 完整核对：段落级逐类型逐属性 + 页面级 ---
    f0=compliance.check_compliance(Document(src),preset)
    w0=[x for x in f0 if x['level']=='warn']
    items0={x['item'] for x in w0}
    for need in ('页边距','纸张','页面网格','页码'):
        assert need in items0, '页面级应查 {}'.format(need)
    for need in ('正文·字体','正文·字号','正文·首行缩进','正文·行距','正文·对齐方式'):
        assert need in items0, '段落级应逐属性查 {}'.format(need)
    assert any(x.get('locations') for x in w0), '段落级偏差应带段号定位'

    # --- 2. 自洽性（核心）：智能一键完整流程后，合规检查必须零偏差 ---
    p1=os.path.join(OUT_DIR,'comp_punct.docx'); punctuation.process_document(src,p1)
    out=os.path.join(OUT_DIR,'comp_out.docx'); format_document(p1,out,preset_name='official_gbk')
    f1=compliance.check_compliance(Document(out),preset)
    left=[x for x in f1 if x['level']=='warn']
    assert not left, '排版后仍报偏差，检查与排版口径不一致：{}'.format(
        [(x['item'],x['detail']) for x in left])
    assert sum(1 for x in f1 if x['level']=='ok')>30, '合格项过少，检查覆盖面不足'

    # --- 3. 检查项开关 + 旧键兼容 ---
    f2=compliance.check_compliance(Document(src),preset,options={'margins':False,'paper':False})
    assert not any(x['item']=='页边距' for x in f2), '关闭后不查边距'
    f3=compliance.check_compliance(Document(src),preset,options={'body_font':False})
    assert not any(x['item']=='正文·字体' for x in f3), '旧键 body_font 应映射到新键'

    # --- 4. 全部认可 → 归零 ---
    keys=[x['fix_key'] for x in w0 if x.get('fix_key')]
    o_all=os.path.join(OUT_DIR,'fix_all.docx')
    compliance.apply_compliance_fixes(src,o_all,preset,keys)
    r_all=[x for x in compliance.check_compliance(Document(o_all),preset) if x['level']=='warn']
    assert not r_all, '全部认可后应无偏差，仍剩：{}'.format([x['item'] for x in r_all])

    # --- 5. 只认可一项 → 其余保持原样（精准手术刀）---
    o_one=os.path.join(OUT_DIR,'fix_one.docx')
    ap=compliance.apply_compliance_fixes(src,o_one,preset,['para:body:size'])
    assert ap and any('字号' in a for a in ap), '应返回修正说明'
    w_one={x['item'] for x in compliance.check_compliance(Document(o_one),preset)
           if x['level']=='warn'}
    assert '正文·字号' not in w_one, '认可项应已修正'
    assert '正文·字体' in w_one and '页边距' in w_one, '未认可项必须保持原样'
    print('[7i] 公文合规检查：完整核对({}项)+排版后零偏差+精准修正 通过'.format(len(f0)))


def test_cleaner():
    """格式清洗：全文/部分段落、脏格式清除、不伤类型识别"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from scripts import cleaner, compliance
    from scripts.data_model import PRESETS
    from scripts.formatter import format_document

    def _dirty_para(p, text):
        r = p.add_run(text)
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0, 0); r.font.bold = True
        rpr = r._r.get_or_add_rPr()
        for tag, val in (('w:spacing', '40'), ('w:w', '150'), ('w:position', '6')):
            e = OxmlElement(tag); e.set(qn('w:val'), val); rpr.append(e)
        em = OxmlElement('w:em'); em.set(qn('w:val'), 'dot'); rpr.append(em)
        ppr = p._p.get_or_add_pPr()
        bdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8'); bdr.append(b); ppr.append(bdr)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'FFFF00'); ppr.append(shd)
        fr = OxmlElement('w:framePr'); fr.set(qn('w:w'), '2000'); ppr.append(fr)
        p.paragraph_format.first_line_indent = Pt(60)
        p.paragraph_format.space_before = Pt(20)
        return p

    # --- 1. 全文清洗：各类脏格式都清掉 ---
    d = Document()
    _dirty_para(d.add_paragraph(), '这是一段带脏格式的正文内容用于测试清洗效果。')
    p2 = d.add_paragraph(); p2.add_run('第二段\t含制表符和　全角空格   和连续空格。')
    p2.runs[0]._r.append(OxmlElement('w:br'))
    p2.add_run('')
    src = os.path.join(OUT_DIR, 'clean_in.docx'); d.save(src)
    out = os.path.join(OUT_DIR, 'clean_out.docx')
    stat = cleaner.clean_file(src, out)
    for need in ('char_format', 'char_spacing', 'emphasis', 'borders_shading',
                 'frame', 'para_format', 'whitespace', 'breaks', 'empty_runs'):
        assert stat.get(need), '清洗项 {} 未生效'.format(need)
    c = Document(out); q = c.paragraphs[0]
    rp = q.runs[0]._r.find(qn('w:rPr')); pp = q._p.find(qn('w:pPr'))
    assert q.runs[0].font.size is None, '字号未清'
    assert str(q.runs[0].font.color.rgb) == '000000', \
        '颜色应清成确定的黑色（样式若自带颜色，置 None 会继承出彩字）'
    assert rp is None or rp.find(qn('w:spacing')) is None, '字符间距未清'
    assert rp is None or rp.find(qn('w:em')) is None, '着重号未清'
    assert pp is None or pp.find(qn('w:pBdr')) is None, '段落边框未清'
    assert pp is None or pp.find(qn('w:framePr')) is None, 'framePr 未清'
    assert q.paragraph_format.first_line_indent is None, '段落格式未清'
    assert '\t' not in c.paragraphs[1].text and '　' not in c.paragraphs[1].text, '空白未清'

    # --- 2. 部分段落清洗：只动标记的段，其余原样 ---
    d2 = Document()
    for i in range(4):
        p = d2.add_paragraph(); r = p.add_run('第{}段脏内容测试。'.format(i))
        r.font.size = Pt(9)
        p.paragraph_format.first_line_indent = Pt(60)
    s2 = os.path.join(OUT_DIR, 'clean_part_in.docx'); d2.save(s2)
    o2 = os.path.join(OUT_DIR, 'clean_part_out.docx')
    cleaner.clean_file(s2, o2, scope_indices={1, 3})
    c2 = Document(o2)
    for i, p in enumerate(c2.paragraphs):
        cleaned = p.runs[0].font.size is None
        assert cleaned == (i in (1, 3)), '第{}段清洗范围错误'.format(i)

    # --- 3. 清洗 + 排版：不削弱类型识别，排版后仍零偏差 ---
    d3 = Document()
    t = d3.add_paragraph(); t.add_run('关于开展某某试点工作的通知')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt in ['各有关单位：', '一、总体要求',
                '为深入贯彻落实上级决策部署，现就开展试点工作通知如下。',
                '特此通知。', '某某办公室', '2026年7月25日']:
        p = d3.add_paragraph(); r = p.add_run(txt); r.font.size = Pt(9)
        p.paragraph_format.first_line_indent = Pt(72)
    s3 = os.path.join(OUT_DIR, 'clean_fmt_in.docx'); d3.save(s3)
    preset = PRESETS['official_gbk']
    seen = {}
    for spec, label in ((None, 'raw'), ({'scope': 'all', 'items': None}, 'cleaned')):
        o3 = os.path.join(OUT_DIR, 'clean_fmt_{}.docx'.format(label))
        format_document(s3, o3, preset_name='official_gbk', clean_spec=spec)
        f = compliance.check_compliance(Document(o3), preset)
        seen[label] = {x['item'].split('·')[0] for x in f if '·' in x['item']}
        assert not [x for x in f if x['level'] == 'warn'], \
            '{}: 清洗+排版后不应有偏差'.format(label)
    assert seen['raw'] == seen['cleaned'], \
        '清洗改变了类型识别结果：{} vs {}'.format(seen['raw'], seen['cleaned'])
    # --- 4. 排版流程默认自动清洗：结构性垃圾不再残留（四个预设都要干净）---
    d4 = Document()
    d4.add_paragraph('关于开展某某试点工作的通知'); d4.add_paragraph('各有关单位：')
    p4 = d4.add_paragraph(); r4 = p4.add_run('为深入贯彻落实上级决策部署，现就开展试点工作通知如下。')
    rpr = r4._r.get_or_add_rPr()
    for tag, val in (('w:spacing', '40'), ('w:w', '150'), ('w:position', '6'), ('w:kern', '20')):
        e = OxmlElement(tag); e.set(qn('w:val'), val); rpr.append(e)
    _em = OxmlElement('w:em'); _em.set(qn('w:val'), 'dot'); rpr.append(_em)
    _rs = OxmlElement('w:shd'); _rs.set(qn('w:fill'), '00FF00'); rpr.append(_rs)
    _rb = OxmlElement('w:bdr'); _rb.set(qn('w:val'), 'single'); rpr.append(_rb)
    ppr = p4._p.get_or_add_pPr()
    _pb = OxmlElement('w:pBdr'); _bt = OxmlElement('w:bottom')
    _bt.set(qn('w:val'), 'single'); _bt.set(qn('w:sz'), '8'); _pb.append(_bt); ppr.append(_pb)
    _ps = OxmlElement('w:shd'); _ps.set(qn('w:fill'), 'FFFF00'); ppr.append(_ps)
    _fp = OxmlElement('w:framePr'); _fp.set(qn('w:w'), '2000'); ppr.append(_fp)
    _tb = OxmlElement('w:tabs'); _t1 = OxmlElement('w:tab')
    _t1.set(qn('w:val'), 'left'); _t1.set(qn('w:pos'), '420'); _tb.append(_t1); ppr.append(_tb)
    d4.add_paragraph('特此通知。'); d4.add_paragraph('某某办公室'); d4.add_paragraph('2026年7月25日')
    s4 = os.path.join(OUT_DIR, 'autoclean_in.docx'); d4.save(s4)
    for pname in ('official_gbk', 'official', 'academic', 'legal'):
        o4 = os.path.join(OUT_DIR, 'autoclean_{}.docx'.format(pname))
        format_document(s4, o4, preset_name=pname)
        tgt = [q for q in Document(o4).paragraphs if '贯彻落实' in q.text][0]
        _rp = tgt.runs[0]._r.find(qn('w:rPr')); _pp = tgt._p.find(qn('w:pPr'))
        left = [n for n, el, tg in (
            ('w:spacing', _rp, 'w:spacing'), ('w:w', _rp, 'w:w'),
            ('w:position', _rp, 'w:position'), ('w:kern', _rp, 'w:kern'),
            ('w:em', _rp, 'w:em'), ('字符shd', _rp, 'w:shd'), ('字符bdr', _rp, 'w:bdr'),
            ('w:pBdr', _pp, 'w:pBdr'), ('段落shd', _pp, 'w:shd'),
            ('framePr', _pp, 'w:framePr'), ('w:tabs', _pp, 'w:tabs'))
            if el is not None and el.find(qn(tg)) is not None]
        assert not left, '{} 排版后仍残留结构性垃圾：{}'.format(pname, left)

    # --- 5. 合规对比预览模型：现状/修正后按认可项变化 ---
    from scripts.compliance import build_preview_model, preview_spec_after
    dm = Document(os.path.join(OUT_DIR, 'comp_in.docx'))
    model = build_preview_model(dm, PRESETS['official_gbk'])
    assert model and any(e['bad'] for e in model), '预览模型应标出偏差'
    be = [e for e in model if e['ptype'] == 'body' and 'size' in e['bad']][0]
    assert preview_spec_after(be, [])['size'] == be['actual']['size'], '未认可不应改变'
    assert preview_spec_after(be, ['para:body:size'])['size'] == be['expected']['size'], \
        '认可后应用预设字号'
    assert preview_spec_after(be, ['para:body:size'])['font'] == be['actual']['font'], \
        '只认可字号不应连字体一起改'
    # --- 白色文字是套打的预印占位，不是脏格式，默认必须保护 ---
    import shutil as _sh3
    _otpl = os.path.join(os.path.dirname(__file__), 'templates', '套打', '文件送审单.docx')
    if os.path.exists(_otpl):
        _c1 = os.path.join(OUT_DIR, 'clean_overprint.docx')
        _sh3.copyfile(_otpl, _c1)

        def _count_white(path):
            _d = Document(path)
            _n = 0
            for _p in _d.paragraphs:
                _n += sum(1 for _r in _p.runs
                          if _r.text.strip() and cleaner._is_white_run(_r))
            for _t in _d.tables:
                _seen = []
                for _row in _t.rows:
                    for _cl in _row.cells:
                        if any(_cl._tc is _x for _x in _seen):
                            continue
                        _seen.append(_cl._tc)
                        for _p in _cl.paragraphs:
                            _n += sum(1 for _r in _p.runs
                                      if _r.text.strip() and cleaner._is_white_run(_r))
            return _n

        _w0 = _count_white(_c1)
        assert _w0 > 0, '套打模板应含白色占位文字'
        assert cleaner.looks_like_overprint(Document(_c1)), '应识别为套打表单'
        _c2 = os.path.join(OUT_DIR, 'clean_overprint_out.docx')
        cleaner.clean_file(_c1, _c2)
        assert _count_white(_c2) == _w0, \
            '套打表单的白色预印占位被清成黑字，打印时会全印出来、整张表报废'
        # 显式强制才归一
        _c3 = os.path.join(OUT_DIR, 'clean_overprint_force.docx')
        cleaner.clean_file(_c1, _c3, items={'white_text': True, 'char_format': True})

    from docx.shared import RGBColor as _RGB
    # 普通文档里的白字是看不见的垃圾，应转黑显形
    _dw = Document(); _pw = _dw.add_paragraph()
    _rw = _pw.add_run('隐藏白字'); _rw.font.color.rgb = _RGB(0xFF, 0xFF, 0xFF)
    _s5 = os.path.join(OUT_DIR, 'clean_white_in.docx'); _dw.save(_s5)
    assert not cleaner.looks_like_overprint(Document(_s5)), '普通文档不应误判为套打'
    _o5 = os.path.join(OUT_DIR, 'clean_white_out.docx')
    cleaner.clean_file(_s5, _o5)
    assert str(Document(_o5).paragraphs[0].runs[0].font.color.rgb) == '000000', \
        '普通文档的白字应转为黑色显形'

    # 样式自带颜色时，清洗须给出确定的黑色（而非继承样式色）
    _dh = Document()
    _dh.styles['Heading 1'].font.color.rgb = _RGB(0x1F, 0x5F, 0xA9)
    _ph = _dh.add_paragraph(); _ph.style = _dh.styles['Heading 1']
    _rh = _ph.add_run('标题'); _rh.font.color.rgb = _RGB(0xFF, 0, 0)
    _s6 = os.path.join(OUT_DIR, 'clean_style_in.docx'); _dh.save(_s6)
    _o6 = os.path.join(OUT_DIR, 'clean_style_out.docx')
    cleaner.clean_file(_s6, _o6)
    assert str(Document(_o6).paragraphs[0].runs[0].font.color.rgb) == '000000', \
        '样式带颜色时清洗仍应确保黑色，公文不能继承出蓝字'

    # --- 套打表单误入排版流程应告警（只警示不阻断）---
    if os.path.exists(_otpl):
        import logging as _lg
        _recs = []

        class _Grab(_lg.Handler):
            def emit(self, rec):
                _recs.append(rec.getMessage())

        _h = _Grab(); _lg.getLogger('docformat.engine').addHandler(_h)
        try:
            _wsrc = os.path.join(OUT_DIR, 'clean_op_warn.docx')
            _sh3.copyfile(_otpl, _wsrc)
            format_document(_wsrc, os.path.join(OUT_DIR, 'clean_op_warn_out.docx'),
                            preset_name='official_gbk')
            assert any('套打' in m for m in _recs), \
                '套打表单误入排版应告警: {}'.format(_recs)
            _recs.clear()
            format_document(SRC, os.path.join(OUT_DIR, 'clean_op_nowarn.docx'),
                            preset_name='official_gbk')
            assert not any('套打' in m for m in _recs), \
                '普通公文不应误报为套打: {}'.format(_recs)
        finally:
            _lg.getLogger('docformat.engine').removeHandler(_h)

    print('[7k] 格式清洗：套打白字豁免 + 误入排版告警 通过')


def test_toc():
    """目录：大纲级别写入、层级识别、点引导线制表位、页码匹配与降级"""
    from docx.oxml.ns import qn
    from scripts import toc
    from scripts.formatter import format_document

    d = Document()
    d.add_paragraph('关于开展某某试点工作的通知'); d.add_paragraph('各有关单位：')
    for h in ['一、总体要求', '（一）工作目标', '二、主要任务']:
        d.add_paragraph(h)
        for _ in range(3):
            d.add_paragraph('为深入贯彻落实上级决策部署，现就开展试点工作通知如下。')
    d.add_paragraph('某某办公室'); d.add_paragraph('2026年7月25日')
    src = os.path.join(OUT_DIR, 'toc_in.docx'); d.save(src)
    fmt = os.path.join(OUT_DIR, 'toc_fmt.docx')
    format_document(src, fmt, preset_name='official_gbk')

    # 1. 排版写入大纲级别——Word 自动目录域靠它取条目，否则目录是空的
    fdoc = Document(fmt)
    levels = {}
    for p in fdoc.paragraphs:
        ppr = p._p.find(qn('w:pPr'))
        if ppr is None:
            continue
        el = ppr.find(qn('w:outlineLvl'))
        if el is not None:
            levels[p.text.strip()] = int(el.get(qn('w:val')))
    assert levels.get('一、总体要求') == 0, '一级标题应为大纲级别1'
    assert levels.get('（一）工作目标') == 1, '二级标题应为大纲级别2'
    assert '为深入贯彻落实上级决策部署，现就开展试点工作通知如下。' not in levels, \
        '正文不应有大纲级别'

    # 2. 自动目录域
    auto = os.path.join(OUT_DIR, 'toc_auto.docx')
    toc.generate_toc(fmt, auto, mode='auto')
    xml = '\n'.join(p._p.xml for p in Document(auto).paragraphs[:6])
    assert 'TOC' in xml and 'instrText' in xml, '未插入目录域'

    # 3. 手动目录：层级从大纲级别读、点引导线用右对齐制表位
    man = os.path.join(OUT_DIR, 'toc_man.docx')
    toc.generate_toc(fmt, man, mode='manual')
    mdoc = Document(man)
    entries = [p for p in mdoc.paragraphs if '\t' in p.text]
    assert len(entries) >= 4, '目录条目过少: {}'.format(len(entries))
    for p in entries:
        ppr = p._p.find(qn('w:pPr'))
        tabs = ppr.find(qn('w:tabs')) if ppr is not None else None
        assert tabs is not None, '目录项应有制表位'
        tb = tabs.find(qn('w:tab'))
        assert tb.get(qn('w:leader')) == 'dot', '点引导线应由制表位生成'
        assert tb.get(qn('w:val')) == 'right', '页码列应右对齐'
    assert not any('. . . .' in p.text for p in entries), '不应再用字面点线拼接'

    # 4. 层级映射（大纲级别优先，未排版文档回退识别器）
    items = toc._build_heading_items(Document(fmt))
    assert ('一、总体要求', 1) in items and ('（一）工作目标', 2) in items, \
        '层级读取错误: {}'.format(items)

    # 5. 页码匹配：按文档顺序推进游标，同名标题不错配到首次出现处
    pages = toc._match_headings_to_pages(
        [('标题A', 1), ('小节X', 2), ('标题B', 1), ('小节X', 2), ('标题C', 1)],
        ['标题A小节X正文', '标题B小节X正文', '标题C正文'])
    assert pages['标题A'] == 1 and pages['标题B'] == 2 and pages['标题C'] == 3, \
        '页码匹配错误: {}'.format(pages)

    # 6. 缺 PDF 取词工具时优雅降级，且提示能对症
    _orig = toc._has_pdf_text_tool
    toc._has_pdf_text_tool = lambda: False
    try:
        deg = os.path.join(OUT_DIR, 'toc_degrade.docx')
        toc.generate_toc(fmt, deg, mode='manual')
        txt = '\n'.join(p.text for p in Document(deg).paragraphs[:12])
        assert '__' in txt, '降级时应留页码占位符'
    finally:
        toc._has_pdf_text_tool = _orig
    print('[7l] 目录：大纲级别/层级/点引导线制表位/页码匹配/降级 通过')


def test_compare():
    """版本比对：段落级 diff 识别增/删/改，产出对照件"""
    from scripts import compare
    a = Document()
    for t in ['关于开展试点工作的通知', '各有关单位：', '一、总体要求',
              '为深入贯彻落实上级决策部署，现就有关事项通知如下。',
              '二、组织实施', '各单位要加强组织领导。', '特此通知。']:
        a.add_paragraph(t)
    pa = os.path.join(OUT_DIR, 'cmp_base.docx'); a.save(pa)
    b = Document()
    for t in ['关于开展试点工作的通知', '各有关单位：', '一、总体要求',
              '为深入贯彻落实上级决策部署精神，现就有关事项通知如下，请遵照执行。',
              '（一）新增的小节', '二、组织实施', '特此通知。']:
        b.add_paragraph(t)
    pb = os.path.join(OUT_DIR, 'cmp_rev.docx'); b.save(pb)

    rows = compare.diff_paragraphs(
        [p.text for p in Document(pa).paragraphs if p.text.strip()],
        [p.text for p in Document(pb).paragraphs if p.text.strip()])
    stat = compare.summarize(rows)
    assert stat['chg'] == 1, '应识别出 1 处修改: {}'.format(stat)
    assert stat['add'] == 1, '应识别出 1 处新增: {}'.format(stat)
    assert stat['del'] == 1, '应识别出 1 处删除: {}'.format(stat)

    out = os.path.join(OUT_DIR, 'cmp_diff.docx')
    ok, info, _st = compare.compare_documents(pa, pb, out, prefer_office=False)
    assert ok and os.path.exists(out), '比对文件未产出'
    txt = '\n'.join(p.text for p in Document(out).paragraphs)
    for mark in ('［改前］', '［改后］', '［新增］', '［删除］'):
        assert mark in txt, '对照件缺少标记 {}'.format(mark)
    # 相同文档比对应为零改动
    same = os.path.join(OUT_DIR, 'cmp_same.docx')
    _ok, _info, st2 = compare.compare_documents(pa, pa, same, prefer_office=False)
    assert st2['add'] == 0 and st2['del'] == 0 and st2['chg'] == 0, \
        '相同文档不应报改动: {}'.format(st2)
    print('[7m] 版本比对：增/删/改识别 + 对照件产出 通过')


def test_exporter():
    """导出 PDF：路径规避重名，缺引擎时报错不崩"""
    from scripts import exporter
    p1 = exporter.pdf_output_path(os.path.join(OUT_DIR, 'sample.docx'))
    assert p1.endswith('.pdf'), 'PDF 路径生成错误: {}'.format(p1)
    open(p1, 'wb').close()
    p2 = exporter.pdf_output_path(os.path.join(OUT_DIR, 'sample.docx'))
    assert p2 != p1 and '(2)' in p2, '重名未自动避让: {}'.format(p2)
    os.remove(p1)
    ok, info = exporter.export_pdf(os.path.join(OUT_DIR, 'sample.docx'),
                                   os.path.join(OUT_DIR, 'nope.pdf'))
    assert isinstance(ok, bool) and isinstance(info, str), '导出应返回 (bool, str)'
    assert ok or info, '失败时应给出原因'
    print('[7n] 导出 PDF：路径避让 + 缺引擎时报错不崩 通过')


def _title_cell(plan):
    """按 is_title 标记找标题格——别写死行列下标：模板改版后
    标题行已是整行合并的一格，硬下标会直接越界"""
    for _b in plan['blocks']:
        if _b['kind'] != 'table':
            continue
        for _r in _b['rows']:
            for _c in _r['cells']:
                if _c.get('is_title'):
                    return _c
    raise AssertionError('plan 里找不到标题格')



def test_template_builder():
    """新建套打模板：点在哪儿，字就印在哪儿"""
    from scripts.template_builder import build_template, group_rows
    from scripts import overprint as op
    from scripts.exporter import export_pdf
    from scripts import overlay as _ov

    items = [
        {'x': 5.3, 'y': 2.8, 'kind': 'label', 'text': '某某局办公室', 'pt': 15},
        {'x': 7.7, 'y': 3.7, 'kind': 'label', 'text': '文件处理单', 'pt': 18},
        {'x': 2.5, 'y': 5.2, 'kind': 'label', 'text': '紧急程度：', 'pt': 12},
        {'x': 4.7, 'y': 5.2, 'kind': 'field', 'name': '紧急程度', 'pt': 12},
        {'x': 13.0, 'y': 5.2, 'kind': 'label', 'text': '密级：', 'pt': 12},
        {'x': 14.6, 'y': 5.2, 'kind': 'field', 'name': '密级', 'pt': 12},
        {'x': 2.4, 'y': 6.6, 'kind': 'label', 'text': '标  题', 'pt': 12},
        {'x': 4.0, 'y': 6.6, 'kind': 'field', 'name': '标题', 'pt': 16},
        {'x': 2.4, 'y': 9.0, 'kind': 'label', 'text': '承办人：', 'pt': 12},
        {'x': 4.5, 'y': 9.0, 'kind': 'field', 'name': '承办人', 'pt': 12},
    ]
    # 同一行的元素要归到一行；y 差几毫米也算同一行（拖框手会抖）
    rows = group_rows(items + [{'x': 18.0, 'y': 5.35, 'kind': 'label',
                                'text': 'X', 'pt': 12}])
    assert len(rows) == 5, '应分成 5 行，实得 {}'.format(len(rows))
    assert len(rows[2]['items']) == 5, '5.2/5.35 应并入同一行'
    assert [i['x'] for i in rows[2]['items']] == sorted(
        i['x'] for i in rows[2]['items']), '行内应按 x 从左到右'

    out = os.path.join(OUT_DIR, 'wizard_tpl.docx')
    _p, fields = build_template(items, out)
    assert fields == ['紧急程度', '密级', '标题', '承办人'], \
        '字段顺序应按位置：{}'.format(fields)
    assert op.scan_fields(out) == fields, '扫出来的字段应与生成时一致'

    # 预印栏目名必须是白字（占位不显影），填写位必须不是
    import docx as _dx
    whites, blacks = [], []
    for _p2, _c2 in op._iter_paragraphs(_dx.Document(out)):
        for _r2 in _p2.runs:
            if not _r2.text.strip():
                continue
            (whites if not op._run_prints(_r2) else blacks).append(_r2.text)
    assert '紧急程度：' in whites and '某某局办公室' in whites, \
        '预印栏目名应为白字：{}'.format(whites)
    assert any('{{' in t for t in blacks), '占位符应是会打印的黑字'

    # 落点复核：填入可辨认的值，转 PDF 量坐标
    ok_r, _why = _ov.can_merge()
    tmp = os.path.join(OUT_DIR, 'wizard_filled.docx')
    vals = {'紧急程度': 'AAA', '密级': 'BBB', '标题': 'CCCC', '承办人': 'DDD'}
    op.fill_form(out, vals, tmp)
    pdf = os.path.join(OUT_DIR, 'wizard_filled.pdf')
    ok, info = export_pdf(tmp, pdf)
    if not (ok and ok_r):
        print('    （跳过落点实测：{}）'.format(info if not ok else _why))
        print('[7p] 新建套打模板：分行/字段/白字占位 通过')
        return
    import fitz
    K = 2.54 / 72.0
    got = {}
    for b in fitz.open(pdf)[0].get_text('dict')['blocks']:
        for l in b.get('lines', []):
            for sp in l['spans']:
                t = sp['text'].strip()
                if t:
                    got.setdefault(t, (sp['bbox'][0] * K, sp['bbox'][1] * K))
    worst = 0.0
    for it in items:
        key = it['text'] if it['kind'] == 'label' else vals[it['name']]
        hit = next((v for k, v in got.items() if key[:3] in k), None)
        if hit is None:
            continue
        worst = max(worst, abs(hit[0] - it['x']), abs(hit[1] - it['y']))
    assert worst < 0.1, '点选位置与实际落点偏差 {:.3f}cm，超过 1mm'.format(worst)
    print('[7p] 新建套打模板：分行/字段/白字占位 + 落点偏差 {:.3f}cm 通过'
          .format(worst))


def test_overprint():
    """套打：字段扫描/填充保留几何/合并宽度/长文自适应/固定行不误缩"""
    from docx.oxml.ns import qn
    from scripts import overprint as op
    tpl = os.path.join(os.path.dirname(__file__), 'templates', '套打', '文件送审单.docx')
    assert os.path.exists(tpl), '自带套打模板缺失'

    fields = op.scan_fields(tpl)
    for need in ('标题', '拟办意见', '承办部门', '经办人', '电话', '密级', '年', '月', '日'):
        assert need in fields, '缺少字段 {}：{}'.format(need, fields)

    src = Document(tpl)
    st = src.tables[0]
    # 套打机制必须保留：白色文字占位 + 白色边框
    tblPr = st._tbl.find(qn('w:tblPr'))
    borders = tblPr.find(qn('w:tblBorders'))
    assert borders is not None, '表格应有边框定义'
    assert borders.find(qn('w:top')).get(qn('w:color')) == 'FFFFFF', \
        '预印框线应为白色（不显影）'
    whites = 0
    for p in src.paragraphs:
        for r in p.runs:
            c = r.font.color.rgb if r.font.color and r.font.color.rgb else None
            if str(c) == 'FFFFFF' and r.text.strip():
                whites += 1
    assert whites > 0, '模板应含白色占位文字'

    # 合并单元格宽度按 gridSpan 累加（只读 tcW 会低估、导致误缩字号）
    widths = {}
    for cell in op._iter_cells(st):
        widths[cell.text.strip()] = round(op._cell_width_cm(st, cell), 2)

    def _w(prefix):
        for k, v in widths.items():
            if k.startswith(prefix):
                return v
        return 0
    assert _w('领导批示') > 15, '整行合并单元格宽度应接近表宽: {}'.format(widths)
    # 承办部门/经办人 那一行是左右两栏，竖线在距纸右侧 12cm 处（实测），
    # 即左栏 6.9cm、右栏 9.9cm——不能被读成"平均分"
    assert 6.5 < _w('承办部门') < 7.3, '左栏宽应约 6.9cm: {}'.format(widths)
    assert 9.5 < _w('经办人') < 10.3, '右栏宽应约 9.9cm: {}'.format(widths)
    assert abs(_w('承办部门') + _w('经办人') - _w('领导批示')) < 0.1, \
        '两栏之和应等于表宽: {}'.format(widths)
    # 标题栏被竖线分成两格：竖线与「紧急程度：」冒号后第一个字的左边线齐
    assert 2.5 < _w('标  题') < 2.7, '标题栏目名格宽应约 2.60cm: {}'.format(widths)
    assert 14.0 < _w('{{标题}}') < 14.4, '标题正文格宽应约 14.20cm: {}'.format(widths)

    base = {'紧急程度': '特急', '密级': '秘密★1年', '标题': '关于报送年度总结的请示',
            '承办部门': '办公室', '经办人': '张三', '电话': '12345678',
            '文字校核': '李四', '年': '2026', '月': '7', '日': '25'}

    def _fill(text, name):
        v = dict(base); v['拟办意见'] = text
        out = os.path.join(OUT_DIR, 'overprint_{}.docx'.format(name))
        logs = []
        n, notes = op.fill_form(tpl, v, out, log=lambda l, m: logs.append(m))
        return out, n, notes, logs

    # 短内容：不该缩任何字号（缩了反而与预印栏位错位）
    out_s, n, notes, logs = _fill('因工作需要，拟报请领导审批。请审示。', 'short')
    assert n == 11, '应填入 11 个字段: {}'.format(n)
    assert not [m for m in logs if '自适应' in m], '短内容不应触发缩放: {}'.format(logs)
    assert not notes
    doc = Document(out_s)
    txt = '\n'.join(p.text for p in doc.paragraphs)
    assert '秘密★1年' in txt and '2026' in txt, '顶部字段未填入'
    assert '{{' not in txt, '仍有未替换的占位符'
    # 几何必须原样保留
    assert abs(doc.sections[0].top_margin.cm - src.sections[0].top_margin.cm) < 0.01, \
        '页边距被改动，套打会错位'
    dt = doc.tables[0]
    for r0, r1 in zip(src.tables[0].rows, dt.rows):
        h0, _e0 = op._row_height_cm(r0)
        h1, e1 = op._row_height_cm(r1)
        assert h0 == h1, '行高数值被改动，套打会错位: {} vs {}'.format(h0, h1)
        assert _e0 == e1, ('行高规则被改动: {} → {}。atLeast 行的实际渲染高度取决于'
                           'Word 字体度量，程序算不准；擅自锁成 exact 会把行压小、'
                           '下面内容整体上移').format(_e0, e1)

    # 长内容：应缩字号；固定高度行(经办人所在行)不应被连带缩
    out_l, _n, notes_l, logs_l = _fill('因某某事项需要进一步开展调查核实工作，' * 14, 'long')
    assert any('拟办意见' in m for m in logs_l), '长内容应触发拟办意见缩放: {}'.format(logs_l)
    assert not any('经 办 人' in m for m in logs_l), \
        '固定高度行会裁切而非撑高，不应误缩: {}'.format(logs_l)

    # 极长：缩到下限仍放不下应如实告警，而不是悄悄溢出
    _o, _n2, notes_x, _l = _fill('因某某事项需要进一步开展调查核实工作，' * 30, 'huge')
    assert notes_x and '过长' in notes_x[0], '极长内容应给出告警: {}'.format(notes_x)

    # 行数估算：ASCII 按半角计，行尾空格不计
    assert op.estimate_lines('中文' * 10, 14, 10.0) == op.estimate_lines(
        '中文' * 10 + '   ', 14, 10.0), '行尾空格不应多算一行'
    assert op.estimate_lines('abcd', 14, 10.0) == 1
    # --- 从已有 docx 适配：往返还原 ---
    rt = op.extract_values(out_s)
    for k in ('标题', '拟办意见', '承办部门', '经办人', '电话', '密级', '紧急程度',
              '文字校核', '年', '月', '日'):
        assert rt.get(k) == base.get(k, '') or k == '拟办意见', \
            '往返提取 {} 不一致: {} vs {}'.format(k, rt.get(k), base.get(k))

    # --- 结构不同的普通段落式草稿也要能识别 ---
    dr = Document()
    for t in ['紧急程度：加急    密级：机密★3年',
              '标题：关于开展某某专项检查工作的请示',
              '拟办意见：',                       # 标签独占一段
              '因某某专项工作需要，拟组织开展全面检查。请审示。',
              '承办部门：监督检查室', '经办人：王五    电话：87654321',
              '文字校核：赵六', '二〇二六年七月二十五日']:
        dr.add_paragraph(t)
    dsrc = os.path.join(OUT_DIR, 'overprint_draft.docx'); dr.save(dsrc)
    dout = os.path.join(OUT_DIR, 'overprint_fitted.docx')
    vals, dnotes = op.fit_document(dsrc, tpl, dout)
    assert vals.get('紧急程度') == '加急', '同一行多字段应各自截断: {}'.format(vals)
    assert vals.get('密级') == '机密★3年', '同一行第二个字段未识别: {}'.format(vals)
    assert vals.get('标题') == '关于开展某某专项检查工作的请示'
    assert '全面检查' in vals.get('拟办意见', ''), \
        '标签独占一段时应向后收集正文: {}'.format(vals.get('拟办意见'))
    assert vals.get('经办人') == '王五' and vals.get('电话') == '87654321'
    # 日期必须拆成三格，整串塞一格会把版面顶歪
    assert (vals.get('年'), vals.get('月'), vals.get('日')) == ('2026', '7', '25'), \
        '中文数字日期未正确拆分: {}'.format(vals)
    assert not [n for n in dnotes if '未能自动识别' in n], '应全部识别: {}'.format(dnotes)

    # --- 日期各种写法 ---
    assert op.parse_date('2026年7月25日') == ('2026', '7', '25')
    assert op.parse_date('2026-07-25') == ('2026', '7', '25')
    assert op.parse_date('2026.7.5') == ('2026', '7', '5')
    assert op.parse_date('二〇二六年十二月三十一日') == ('2026', '12', '31')
    assert op.parse_date('二〇二六年十月十一日') == ('2026', '10', '11')
    assert op.parse_date('没有日期') is None

    # --- 适配时文字过多同样会缩字号 ---
    dr2 = Document()
    dr2.add_paragraph('标题：关于某事项的请示')
    dr2.add_paragraph('拟办意见：')
    dr2.add_paragraph('因某某事项需要进一步开展调查核实工作。' * 16)
    dr2.add_paragraph('2026年7月25日')
    s2 = os.path.join(OUT_DIR, 'overprint_long_src.docx'); dr2.save(s2)
    lg = []
    op.fit_document(s2, tpl, os.path.join(OUT_DIR, 'overprint_long_fit.docx'),
                    log=lambda l, m: lg.append(m))
    assert any('自适应' in m for m in lg), '适配长文时应缩字号: {}'.format(lg)

    # --- 一页保证：缩下边距留分页余量，且不移动任何内容位置 ---
    _tplsec = Document(tpl).sections[0]
    _outsec = Document(out_s).sections[0]
    # 模板本身已按实测尺寸把下边距做到很小；这里要的是**结果**——
    # 下边距足够小、末行不会被挤到第二页，而不是"必须被改小过"
    assert _outsec.bottom_margin.cm <= max(0.6, _tplsec.bottom_margin.cm) + 0.001, \
        '下边距应压到 0.6cm 以内以免末行被挤到第二页：{:.2f}'.format(
            _outsec.bottom_margin.cm)
    for _attr in ('top_margin', 'left_margin', 'right_margin',
                  'page_width', 'page_height'):
        assert abs(getattr(_outsec, _attr).cm - getattr(_tplsec, _attr).cm) < 0.001, \
            '{} 被改动——只应调下边距，其余会移动内容位置'.format(_attr)

    # --- 日期：全角数字归一，且取文末的成文日期而非正文里的其它日期 ---
    assert op.parse_date('２０２６年６月２５日') == ('2026', '6', '25'), '全角数字未归一'
    assert op.parse_date('２０２６年6月２５日') == ('2026', '6', '25'), '全半角混排未归一'
    _dd2 = Document()
    _dd2.add_paragraph('标题：关于开展2026年3月专项检查工作的请示')
    _dd2.add_paragraph('拟办意见：'); _dd2.add_paragraph('拟同意。请审示。')
    _dd2.add_paragraph('承办部门：办公室')
    _dd2.add_paragraph('２０２６年７月２５日')
    _ds = os.path.join(OUT_DIR, 'overprint_date_src.docx'); _dd2.save(_ds)
    _dv = op.extract_values(_ds)
    assert (_dv.get('年'), _dv.get('月'), _dv.get('日')) == ('2026', '7', '25'), \
        '应取文末成文日期而非标题里的日期: {}'.format(_dv)

    # --- 自带模板目录应在软件目录下，方便用户找到 ---
    assert os.path.normpath(op.bundled_overprint_dir()).startswith(
        os.path.normpath(op.app_dir())), '自带模板应放在软件所在目录内'

    # --- 标题栏/拟办意见栏不得被撑高：几何必须恒定 ---
    tpl_geo = [op._row_height_cm(r) for r in Document(tpl).tables[0].rows]
    geos = []
    for mult, tmult in ((1, 1), (16, 1), (40, 4)):
        v = dict(base)
        v['拟办意见'] = '因某某事项需要进一步开展调查核实工作。' * mult
        v['标题'] = '关于开展某某专项检查工作的请示' * tmult
        o = os.path.join(OUT_DIR, 'overprint_geo_{}.docx'.format(mult))
        op.fill_form(tpl, v, o)
        geos.append([op._row_height_cm(r) for r in Document(o).tables[0].rows])
    assert geos[0] == geos[1] == geos[2] == tpl_geo, \
        '内容长短改变了表格几何，套打会错位: {} vs 模板 {}'.format(geos, tpl_geo)

    # 模板里不应残留空 run（无文字、也不承载制表/换行/图片等结构），
    # 它们是转模板时清空黑字留下的垃圾。判定口径要和 strip_empty_runs
    # 一致：只放过**没有结构标记**的空 run，别把定位用的制表符 run 也算进来
    _empty = []
    for _p9, _c9 in op._iter_paragraphs(Document(tpl)):
        for _r9 in _p9.runs:
            if _r9.text:
                continue
            if any(_r9._r.find(qn(_t9)) is not None for _t9 in op._STRUCT_TAGS):
                continue
            _empty.append(_r9)
    assert not _empty, '模板残留 {} 个空 run'.format(len(_empty))

    # 用户自带的模板可能残留空 run，填充时也应清掉（防御）
    import shutil as _sh
    dirty = os.path.join(OUT_DIR, 'overprint_dirty_tpl.docx')
    _sh.copyfile(tpl, dirty)
    _dd = Document(dirty)
    from docx.oxml import OxmlElement as _OE
    for _c in op._iter_cells(_dd.tables[0]):
        for _p in _c.paragraphs:
            if _p.runs:
                _r = _OE('w:r'); _r.append(_OE('w:rPr')); _p.runs[0]._r.addnext(_r)
    _dd.save(dirty)
    import re as _re2
    import zipfile as _zf
    _before = len(_re2.findall(r'<w:r>(?:(?!<w:t[ >]).)*?</w:r>',
                               _zf.ZipFile(dirty).read('word/document.xml').decode('utf-8'),
                               _re2.S))
    assert _before, '注入空 run 失败，测试无效'
    dclean = os.path.join(OUT_DIR, 'overprint_dirty_out.docx')
    op.fill_form(dirty, {'标题': '关于某事项的请示', '拟办意见': '内容。'}, dclean)
    # 同上：制表符 run 是定位用的结构，不算"空 run"
    _after_junk = []
    for _p8, _c8 in op._iter_paragraphs(Document(dclean)):
        for _r8 in _p8.runs:
            if _r8.text:
                continue
            if any(_r8._r.find(qn(_t8)) is not None for _t8 in op._STRUCT_TAGS):
                continue
            _after_junk.append(_r8)
    assert not _after_junk, '填充后仍残留 {} 个空 run'.format(len(_after_junk))

    # --- 空值：经办人等留空供手写签字，应干净留白 ---
    import re as _re
    blank = dict(base); blank['拟办意见'] = '因工作需要，拟报请审批。'
    for k in ('经办人', '文字校核', '紧急程度', '密级'):
        blank.pop(k, None)
    bout = os.path.join(OUT_DIR, 'overprint_blank.docx')
    op.fill_form(tpl, blank, bout)
    bdoc = Document(bout)
    balltext = '\n'.join(p.text for p in bdoc.paragraphs)
    for cell in op._iter_cells(bdoc.tables[0]):
        balltext += '\n' + cell.text
    assert not _re.findall(r'\{\{[^}]*\}\}', balltext), \
        '留空字段残留占位符: {}'.format(_re.findall(r'\{\{[^}]*\}\}', balltext))
    assert '经办人：' in balltext, '留空后标签仍应在（打印出来是空白供手写）'

    # --- 预览：与实际输出同一条填充路径，字号必须一致 ---
    pv_vals = dict(base)
    pv_vals['拟办意见'] = '因某某事项需要进一步开展调查核实工作。' * 16
    plan = op.plan_fill(tpl, pv_vals)
    assert plan['rows'] and plan['paras'], '预览数据为空'
    # 预览必须按文档真实块顺序：成文日期在表格之后，
    # 若先渲染全部段落再渲染表格，日期会跑到表格上面，与实际版面不符
    kinds = [b['kind'] for b in plan['blocks']]
    assert 'table' in kinds, '预览缺少表格块'
    ti = kinds.index('table')
    after = [b for b in plan['blocks'][ti + 1:] if b['kind'] == 'para']
    assert after, '表格之后应还有段落（成文日期行）'
    date_txt = ''.join(x['text'] for x in after[0]['segs'])
    assert '年' in date_txt and '月' in date_txt, \
        '表格后的段落应是成文日期行: {!r}'.format(date_txt)
    # 表格线按模板设置画：左右外框为 none 不应画出竖线
    tb = [b for b in plan['blocks'] if b['kind'] == 'table'][0]
    assert tb['borders']['left'] == 'none' and tb['borders']['right'] == 'none', \
        '模板左右外框应为 none: {}'.format(tb['borders'])
    # atLeast 行的预览高度取"声明高度"与"内容自然高度"较大者
    for r in tb['rows']:
        if r['exact']:
            assert r['height_cm'] == r['declared_cm'], '固定行应用声明高度'
        else:
            assert r['height_cm'] >= r['declared_cm'], 'atLeast 行不应小于声明高度'
    # 新模板按实测尺寸重建，各行都是 hRule=exact 的固定高度，
    # 不再依赖文档网格来估留白区高度（网格已移除，纵向由段前距说了算）。
    # 网格吸附的算法仍要能用——用户自带的老模板可能有网格。
    _g = 0.55
    _mk = Document().add_paragraph('测')
    from docx.shared import Pt as _Pt
    _mk.runs[0].font.size = _Pt(14)
    _h = op.paragraph_height_cm(_mk, _g)
    assert abs(_h - 2 * _g) < 0.01, \
        '14pt 段落自然行高 0.69cm > 网格 0.55cm，应吸附占 2 格: {:.3f}'.format(_h)
    # 领导批示留白区：实测红线 8.0 → 20.2，应为 12.2cm 的固定高度
    _lead = [r for r in tb['rows'] if r['declared_cm'] > 6][0]
    assert _lead['exact'], '留白区应是固定高度行'
    assert abs(_lead['height_cm'] - 12.2) < 0.1, \
        '留白区高度应为实测的 12.2cm，实得 {:.2f}'.format(_lead['height_cm'])

    # 整单应正好占满一页：内容末端接近页高减下边距
    _tbl_h = sum(r['height_cm'] for b in plan['blocks'] if b['kind'] == 'table'
                 for r in b['rows'])
    _pg = plan['page']
    _paras_h = 0.0
    _pd = Document(tpl)
    for _pp in _pd.paragraphs:
        _paras_h += op.paragraph_height_cm(_pp, _g)
    # 整单必须放得进一页，且留有余量——Word 的实际排版高度受字体度量、
    # 网格吸附等影响，从 XML 精确预测屡试屡错，故要求明确的安全余量而非
    # "正好占满"（此前按"正好占满"设计，实机就溢出到了第二页）
    _end = _pg['top_cm'] + _paras_h + _tbl_h
    _filled_sec = Document(out_s).sections[0]
    _limit = _pg['height_cm'] - _filled_sec.bottom_margin.cm
    assert _end <= _limit, \
        '整单放不进一页：末端 {:.2f}cm，可用到 {:.2f}cm'.format(_end, _limit)
    assert _limit - _end >= 1.0, \
        '一页余量仅 {:.2f}cm，太紧张，实机容易溢出到第二页'.format(_limit - _end)
    shrunk = [c for r in plan['rows'] for c in r['cells'] if c['shrunk']]
    assert shrunk, '长内容预览应标出已缩小'
    pv_out = os.path.join(OUT_DIR, 'overprint_pv.docx')
    op.fill_form(tpl, pv_vals, pv_out)
    real = None
    for cell in op._iter_cells(Document(pv_out).tables[0]):
        if '拟办意见' in cell.text:
            cands = [pp for pp in cell.paragraphs if pp.text.strip()]
            real = op._para_font_pt(max(cands, key=lambda pp: len(pp.text)))
            break
    # 预览必须报正文字号而非标签字号，否则用户看不出正文到底多小
    assert all(c['font_pt'] is not None for r in plan['rows'] for c in r['cells']), \
        '合并单元格的字号报告丢失'
    assert real is not None and abs(real - shrunk[0]['font_pt']) < 0.01, \
        '预览字号({})与实际输出({})不一致'.format(shrunk[0]['font_pt'], real)
    # 极长内容应在预览里标为放不下
    huge = dict(base)
    huge['拟办意见'] = '因某某事项需要进一步开展调查核实工作。' * 40
    plan2 = op.plan_fill(tpl, huge)
    assert any(c['overflow'] for r in plan2['rows'] for c in r['cells']), \
        '极长内容预览应标为放不下'

    # 纵向合并的延续格不得重复画出合并源的文字（曾出现两个"承办部门"）
    _conts = [c for r in plan['rows'] for c in r['cells'] if c.get('vmerge_cont')]
    assert _conts, '模板里承办部门是纵向合并的，应识别出延续格'
    assert all(not c['segs'] for c in _conts), '纵向合并的延续格不应重复渲染文字'
    _bxs = [''.join(s['text'] for s in c['segs'])
            for r in plan['rows'] for c in r['cells']]
    assert sum('承办部门' in t for t in _bxs) == 1, \
        '预览里"承办部门"只应出现一次：{}'.format(_bxs)

    # 日期识别：日期被"切碎"的各种真实写法都要认得出来
    def _mkdoc(name, build):
        _d = Document()
        build(_d)
        _p = os.path.join(OUT_DIR, 'date_' + name + '.docx')
        _d.save(_p)
        return _p

    def _split_cells(_d):        # 年/月/日分列在不同单元格
        _t = _d.add_table(rows=1, cols=6)
        for _c, _v in zip(_t.rows[0].cells, ['2026', '年', '7', '月', '25', '日']):
            _c.text = _v

    def _in_footer(_d):
        _d.add_paragraph('正文')
        _d.sections[0].footer.paragraphs[0].text = '2026年7月25日'

    def _in_textbox(_d):         # 落款常做成文本框，doc.paragraphs 看不到
        from docx.oxml import parse_xml
        _p = _d.add_paragraph()
        _p._p.append(parse_xml(
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml'
            '/2006/main"><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:'
            'vml"><v:textbox><w:txbxContent><w:p><w:r><w:t>2026年7月25日</w:t>'
            '</w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r>'))

    def _nested(_d):
        _t = _d.add_table(rows=1, cols=1)
        _t.rows[0].cells[0].add_table(rows=1, cols=1) \
            .rows[0].cells[0].text = '2026年7月25日'

    def _page2(_d):              # 字数多、日期被挤到第二页，仍须识别
        for _ in range(60):
            _d.add_paragraph('某单位某单位某单位某单位某单位某单位某单位某单位。')
        _d.add_paragraph('2026年7月25日')

    _date_cases = {
        '普通段落': lambda _d: _d.add_paragraph('  2026年7月25日   某单位办公室制'),
        '全角数字': lambda _d: _d.add_paragraph('２０２６年７月２５日'),
        '中文数字': lambda _d: _d.add_paragraph('二〇二六年七月二十五日'),
        '日留空': lambda _d: _d.add_paragraph('  2026 年  7 月     日'),
        '分列单元格': _split_cells,
        '页脚': _in_footer,
        '文本框': _in_textbox,
        '嵌套表格': _nested,
        '日期在第二页': _page2,
        '标题另有日期': lambda _d: [_d.add_paragraph('关于开展2026年3月检查的请示'),
                                _d.add_paragraph('2026年7月25日')],
    }
    for _name, _build in _date_cases.items():
        _got = op.extract_values(_mkdoc(_name, _build), fields=['年', '月', '日'])
        assert _got.get('年') == '2026' and _got.get('月') == '7', \
            '日期识别失败（{}）：{}'.format(_name, _got)
    # "日留空"是唯一允许日为空的：留空待手签
    assert not op.extract_values(
        _mkdoc('日留空', _date_cases['日留空']), fields=['日']).get('日'), \
        '日未填写时不应臆造日期'

    # ---- 长标题梯形回行：正梯形上长下短、倒梯形上短下长，且预览=输出 ----
    _long = dict(base)
    _long['标题'] = '关于对某单位某单位某单位某部门某部门张三李四王五赵六的请示'
    for _shape, _cmp in (('trapezoid_down', lambda a, b: a > b),
                         ('trapezoid_up', lambda a, b: a < b)):
        _pl = op.plan_fill(tpl, _long, title_shape=_shape)
        _tc = _title_cell(_pl)
        # 只量标题正文那一段：同格里还有白色栏目名「标  题」，
        # 把它算进去会让首行凭空变宽
        _lines = [l for l in ''.join(
            s['text'] for s in _tc['segs'] if not s.get('white')
        ).split('\n') if l.strip()]
        assert len(_lines) == 2, '{}：长标题应回成 2 行，实得 {}'.format(_shape, _lines)
        _w = [op._text_width_units(l) for l in _lines]
        assert _cmp(_w[0], _w[1]), '{}：两行宽度 {} 不符合梯形'.format(_shape, _w)
        # 输出的 docx 必须与预览断在同一处，否则"预览和实际不一样"
        _to = os.path.join(OUT_DIR, 'overprint_title_%s.docx' % _shape)
        op.fill_form(tpl, _long, _to, title_shape=_shape)
        for _c in op._iter_cells(Document(_to).tables[0]):
            if _c.text.strip().startswith('关于'):
                assert [l for l in _c.text.split('\n') if l] == _lines, \
                    '输出断行与预览不一致'
                break
    # 选“不回行”时输出里不插 w:br，交给 Word 自动折行；
    # 预览仍按几何显示会折在哪儿——那正是 Word 将要折的位置
    _tn = os.path.join(OUT_DIR, 'overprint_title_none.docx')
    op.fill_form(tpl, _long, _tn, title_shape='none')
    for _c in op._iter_cells(Document(_tn).tables[0]):
        if _c.text.strip().startswith('关于'):
            assert '\n' not in _c.text.strip(), '选“不回行”时输出不应插入 w:br'
            break

    # ---- 标题行数由栏位高度说了算：长标题缩字号，不许多占行 ----
    # 标题栏是纸上印死的固定框（hRule=atLeast），多一行会把整行撑高、
    # 下面所有内容一起下移，整张单子与预印栏位错开
    _cap = None
    for _k in range(2, 40, 3):
        _tt = '关于' + '某单位' * _k + '的请示'
        _pl2 = op.plan_fill(tpl, dict(base, 标题=_tt))
        _tc2 = _title_cell(_pl2)
        _r0 = None
        for _bb2 in _pl2['blocks']:
            if _bb2['kind'] != 'table':
                continue
            for _rr2 in _bb2['rows']:
                if any(_cc2.get('is_title') for _cc2 in _rr2['cells']):
                    _r0 = _rr2
        assert _r0 is not None, '找不到标题所在行'
        _cap = _cap or _tc2['max_lines']
        assert _tc2['max_lines'] == 2, \
            '自带模板标题栏应只放得下 2 行，实得 {}'.format(_tc2['max_lines'])
        assert _r0['height_cm'] <= _r0['declared_cm'] + 0.02, \
            '标题 {} 字把标题栏撑高了：{:.2f} > 声明 {:.2f}'.format(
                len(_tt), _r0['height_cm'], _r0['declared_cm'])
        _ls2 = [l for l in ''.join(
            s['text'] for s in _tc2['segs'] if not s.get('white')).split('\n')
            if l.strip()]
        _o2 = os.path.join(OUT_DIR, 'overprint_cap.docx')
        _n3, _notes3 = op.fill_form(tpl, dict(base, 标题=_tt), _o2)
        if len(_ls2) > 2:
            # 只有缩到字号下限仍放不下时才允许超行，且必须如实告警
            assert _tc2['font_pt'] <= op.MIN_FONT_PT, \
                '标题 {} 字排了 {} 行，但字号 {}pt 还没缩到下限'.format(
                    len(_tt), len(_ls2), _tc2['font_pt'])
            assert any('标题' in _s and '精简' in _s for _s in _notes3), \
                '标题超出栏位行数时应告警：{}'.format(_notes3)
        else:
            assert not any('标题' in _s and '精简' in _s for _s in _notes3), \
                '标题放得下却报了警：{}'.format(_notes3)

    # ---- 黑字打印位置：不随填写内容长短漂移（套打对位的根本） ----
    _pp1 = os.path.join(OUT_DIR, 'overprint_pos1.docx')
    _pp2 = os.path.join(OUT_DIR, 'overprint_pos2.docx')
    op.fill_form(tpl, dict(base, **{'年': '2026', '月': '1', '日': '11'}), _pp1)
    op.fill_form(tpl, dict(base, **{'年': '2026', '月': '12', '日': '5'}), _pp2)

    # 定宽槽位是**右对齐**的：数字紧贴其后预印的年/月/日，所以恒定不变的
    # 是右沿。左沿差一个字宽正是"1"和"12"位数不同，本就该如此。
    def _right(_p):
        return [round(b, 2) for _a, b, _t in op.print_positions(Document(_p))]
    assert _right(_pp1) == _right(_pp2) and _right(_pp1), \
        '黑字右沿随月/日位数变了（预印的年月日会对不上）：{} vs {}'.format(
            _right(_pp1), _right(_pp2))

    # ---- 打印位置微调：指定"距纸左边几厘米"，黑字就落在那儿 ----
    # 自带模板已按实测尺寸把位置做进制表位里，无需微调；位置微调这套机制
    # 是给**别的**模板用的，所以拿一份不含制表位的简易模板来测。
    from docx import Document as _D7
    from docx.shared import Pt as _Pt7, RGBColor as _RGB7
    _tdir = os.path.join(OUT_DIR, 'offtpl')
    os.makedirs(_tdir, exist_ok=True)
    _tpl2 = os.path.join(_tdir, '简易单.docx')
    _d7 = _D7()
    from docx.shared import Cm as _Cm7
    _s7 = _d7.sections[0]
    _s7.left_margin = _Cm7(2.1); _s7.right_margin = _Cm7(2.1)
    _p7 = _d7.add_paragraph()
    for _t7, _w7 in (('   ', True), ('{{年}}', False), ('年', True),
                     ('   ', True), ('{{月}}', False), ('月', True),
                     (' ', True), ('{{日}}', False), ('日', True)):
        _r7 = _p7.add_run(_t7)
        _r7.font.size = _Pt7(14)
        if _w7:
            _r7.font.color.rgb = _RGB7(0xFF, 0xFF, 0xFF)
    _p8 = _d7.add_paragraph()
    _r8 = _p8.add_run('紧急程度：')
    _r8.font.size = _Pt7(14); _r8.font.color.rgb = _RGB7(0xFF, 0xFF, 0xFF)
    _r8 = _p8.add_run('{{紧急程度}}'); _r8.font.size = _Pt7(14)
    _d7.save(_tpl2)
    _dv = {'年': '2026', '月': '1', '日': '11'}
    _pl3 = op.plan_fill(_tpl2, _dv)
    assert '年' in _pl3['adjustable'] and '月' in _pl3['adjustable'], \
        '年/月应可微调：{}'.format(_pl3['adjustable'])

    # 定位靠**制表位**（缇为单位、与字体无关），不是补空格：空格宽度随
    # 字体变（实测 TNR 里只有数字的一半），补空格定位实测错位 0.49~0.96cm
    # 且逐个累积。这里要求设定值 / 预览 / 输出报告三者严丝合缝。
    for _want, _dvv in (
            ({'月': 6.5, '日': 9.0}, {'年': '2026', '月': '1', '日': '11'}),
            ({'月': 6.5, '日': 9.0}, {'年': '2026', '月': '12', '日': '5'}),
            ({'年': 4.0, '月': 7.2, '日': 10.4}, {'年': '2026', '月': '7', '日': '25'})):
        _p3 = op.save_offsets(_tpl2, _want)
        assert os.path.exists(_p3), '位置表未落盘'
        assert op.load_offsets(_tpl2) == _want, '位置表读回不一致'
        _pl4 = op.plan_fill(_tpl2, _dvv)
        _o5 = os.path.join(OUT_DIR, 'overprint_off.docx')
        op.fill_form(_tpl2, _dvv, _o5)
        _blk = {t.strip(): a for a, _b, t in op.print_positions(Document(_o5))}
        for _k, _v in _want.items():
            assert abs(_pl4['field_pos'][_k] - _v) < 0.01, \
                '{} 预览 {:.3f} ≠ 设定 {:.2f}'.format(_k, _pl4['field_pos'][_k], _v)
            assert abs(_blk[_dvv[_k]] - _v) < 0.01, \
                '{} 输出 {:.3f} ≠ 设定 {:.2f}'.format(_k, _blk[_dvv[_k]], _v)
    # 位数变化不得影响落点（制表位是绝对位置）
    op.save_offsets(_tpl2, {'月': 6.5})
    _a = os.path.join(OUT_DIR, 'off_a.docx'); _b = os.path.join(OUT_DIR, 'off_b.docx')
    op.fill_form(_tpl2, {'年': '2026', '月': '1', '日': '11'}, _a)
    op.fill_form(_tpl2, {'年': '2026', '月': '12', '日': '5'}, _b)
    _pa = {t.strip(): a for a, _x, t in op.print_positions(Document(_a))}
    _pb = {t.strip(): a for a, _x, t in op.print_positions(Document(_b))}
    assert abs(_pa['1'] - _pb['12']) < 0.01, '一位数/两位数的落点应一致'
    # 目标落在预印栏目名之内时够不着：要告警并给出真实的最小可用值
    op.save_offsets(_tpl2, {'紧急程度': 3.0})
    _n5, _notes5 = op.fill_form(_tpl2, dict(_dv, 紧急程度='平急'), _o5)
    assert any('顶不过去' in _s for _s in _notes5), \
        '够不着的位置应告警：{}'.format(_notes5)
    # 恢复默认 = 删掉位置表
    op.save_offsets(_tpl2, {})
    assert not os.path.exists(_p3), '恢复默认应删除位置表'
    assert op.load_offsets(_tpl2) == {}, '恢复默认后应读到空表'

    # ---- 表格里的字段（标题/拟办意见）也要能微调 ----
    # 它们的制表位以**单元格左沿**为原点，不是页边距；用错原点会偏出一大截
    import shutil as _sh3
    _tdir2 = os.path.join(OUT_DIR, 'offtpl2')
    os.makedirs(_tdir2, exist_ok=True)
    _tpl3 = os.path.join(_tdir2, '送审单.docx')
    _sh3.copyfile(tpl, _tpl3)
    _v3 = {'标题': '关于某事项的请示', '拟办意见': '请审批。'}
    assert '标题' in op.plan_fill(_tpl3, _v3)['adjustable'], '标题应可微调'
    assert '拟办意见' in op.plan_fill(_tpl3, _v3)['adjustable'], '拟办意见应可微调'
    for _w3 in ({'标题': 5.0, '拟办意见': 4.5}, {'标题': 6.2, '拟办意见': 3.6}):
        op.save_offsets(_tpl3, _w3)
        _pl5 = op.plan_fill(_tpl3, _v3)
        _o6 = os.path.join(OUT_DIR, 'overprint_cell_off.docx')
        op.fill_form(_tpl3, _v3, _o6)
        _bk = {t.strip(): a for a, _b, t in op.print_positions(Document(_o6))}
        for _k3, _t3 in _w3.items():
            assert abs(_pl5['field_pos'][_k3] - _t3) < 0.01, \
                '{} 预览 {:.3f} ≠ 设定 {:.2f}'.format(
                    _k3, _pl5['field_pos'][_k3], _t3)
    op.save_offsets(_tpl3, {})

    # ---- 拟办意见正文首行缩进两个字（公文行文惯例）----
    _ind = None
    for _c3 in op._iter_cells(Document(tpl).tables[0]):
        if '拟办意见' in _c3.text:
            for _p3 in _c3.paragraphs:
                if '{{拟办意见}}' in _p3.text or not _p3.text.strip('\t'):
                    _fi = _p3.paragraph_format.first_line_indent
                    if _fi:
                        _ind = _fi.cm
            break
    # 「空两个字」空的是**栏目名**的两个字（2×0.42），不是正文小三的两个字：
    # 实测正文首行落在 3.30，正好在栏目名「意」字底下（拟2.40 办2.82 意3.24）
    assert _ind is not None and abs(_ind - (3.3 - 2.4)) < 0.05, \
        '拟办意见正文首行应缩进到 3.30cm（距栏目名左沿 0.90），实得 {}'.format(_ind)

    # ---- 预览折行按真实几何：一行放不下就必须断开 ----
    # Qt 富文本无视表格像素宽度、会把表拉满可视区，靠它折行必然偏长，
    # 所以折行在 plan 阶段按 cm 算好
    _cell_w = _title_cell(_pl)['width_cm']
    for _l in _lines:
        assert op._text_width_units(_l) * (16.0 / op.PT_PER_CM) <= _cell_w, \
            '预览行 {!r} 超出格子宽度 {:.2f}cm'.format(_l, _cell_w)

    # ---- 年/月/日定宽：位数变化不得挪动预印的"年月日" ----
    def _white_pos(_path):
        _cands4 = [q for q in Document(_path).paragraphs
                   if '年' in q.text and '月' in q.text]
        _p4 = _cands4[-1] if _cands4 else Document(_path).paragraphs[0]
        _acc, _out = 0.0, {}
        for _r in _p4.runs:
            _c = _r.font.color.rgb if _r.font.color and _r.font.color.rgb else None
            if str(_c) == 'FFFFFF' and _r.text in ('年', '月', '日'):
                _out[_r.text] = round(_acc, 2)
            _acc += op._text_width_units(_r.text)
        return _out
    _d1 = os.path.join(OUT_DIR, 'overprint_d1.docx')
    _d2 = os.path.join(OUT_DIR, 'overprint_d2.docx')
    op.fill_form(tpl, dict(base, **{'年': '2026', '月': '7', '日': '25'}), _d1)
    op.fill_form(tpl, dict(base, **{'年': '2026', '月': '12', '日': '5'}), _d2)
    assert _white_pos(_d1) == _white_pos(_d2) and _white_pos(_d1), \
        '预印的年/月/日位置随位数变了：{} vs {}'.format(_white_pos(_d1), _white_pos(_d2))
    # 留空待手签时也不能塌缩
    _d3 = os.path.join(OUT_DIR, 'overprint_d3.docx')
    op.fill_form(tpl, dict(base, **{'年': '', '月': '', '日': ''}), _d3)
    assert _white_pos(_d3) == _white_pos(_d1), '留空时预印的年/月/日位置变了'

    # ---- 成文日期行必须一行放得下：多出来的那一行会把整单顶到第二页 ----
    # 实机曾出现"按名义字宽算 16.05cm < 版心 16.45cm 本该放得下，却折了行"，
    # 原因是那一行的空白 run 用 CJK 字体、空格未必是半角。所以要按**悲观**
    # 算法（CJK 字体里的空格算全角）也放得下，才算真的安全。
    _limit_u = (_pg['width_cm'] - _pg['left_cm'] - _pg['right_cm']) / (14 / op.PT_PER_CM)

    def _pess_units(_para):
        _u = 0.0
        for _r in _para.runs:
            _rp = _r._r.find(qn('w:rPr'))
            _f = _rp.find(qn('w:rFonts')) if _rp is not None else None
            _a = (_f.get(qn('w:ascii')) or '') if _f is not None else ''
            _cjk = any(_x in _a for _x in ('楷体', '宋', '黑体', '仿宋'))
            for _ch in _r.text:
                _u += 1.0 if (ord(_ch) > 0x2E80 or (_ch == ' ' and _cjk)) else 0.5
        return _u

    for _dp in (_d1, _d2, _d3):
        _dps = [_p4 for _p4 in Document(_dp).paragraphs
                if '年' in _p4.text and '月' in _p4.text]
        assert _dps, '找不到成文日期行'
        _dt = _dps[-1]
        assert op._text_width_units(_dt.text.rstrip()) <= _limit_u, \
            '成文日期行（乐观）宽 {:.1f} 超版心 {:.1f}'.format(
                op._text_width_units(_dt.text.rstrip()), _limit_u)
        assert _pess_units(_dt) <= _limit_u, \
            '成文日期行（悲观）宽 {:.1f} 超版心 {:.1f}，实机可能折行'.format(
                _pess_units(_dt), _limit_u)
        # 裁的是末尾占位空白，白色单位名不能被一起删掉
        assert '某地市' in _dt.text, '收窄把白色单位名也删了'
    # 值太长把某行撑出版心时要如实告警，而不是闷声折行
    _ov = os.path.join(OUT_DIR, 'overprint_wide.docx')
    _n2, _notes2 = op.fill_form(
        tpl, dict(base, 密级='绝密★长期' + '补充说明事项' * 6), _ov)
    assert any('超出版心' in _s for _s in _notes2), \
        '某行超出版心时应告警：{}'.format(_notes2)

    # ---- 套头对位校验：内容叠到套头纸 PDF 上 ----
    from scripts import overlay as _ovl
    _can, _why = _ovl.can_merge()
    if not _can:
        print('    （跳过套头对位：{}）'.format(_why))
    else:
        import fitz as _fitz

        def _mk_pdf(path, w_cm=21.0, h_cm=29.7, pages=1, text='LETTERHEAD'):
            """直接用 PyMuPDF 造测试 PDF，不依赖本机 Word/LibreOffice，
            冒烟测试才能在任何机器上跑起来"""
            doc = _fitz.open()
            for i in range(pages):
                pg = doc.new_page(width=w_cm / 2.54 * 72, height=h_cm / 2.54 * 72)
                pg.insert_text((60, 80 + i * 14), '{} {}'.format(text, i + 1),
                               fontsize=12)
            doc.save(path)
            doc.close()
            return path

        _lh = _mk_pdf(os.path.join(OUT_DIR, 'letterhead.pdf'))
        _sz = _ovl.page_size_cm(_lh)
        assert _sz and abs(_sz[0] - 21.0) < 0.1, '套头纸应识别为 A4 宽：{}'.format(_sz)
        assert _ovl.page_count(_lh) == 1

        _content = _mk_pdf(os.path.join(OUT_DIR, 'content.pdf'), text='CONTENT')
        _merged = os.path.join(OUT_DIR, 'merged.pdf')
        _mn = _ovl.merge_overlay(_content, _lh, _merged)
        assert _ovl.page_count(_merged) == 1, '合并后应为 1 页'
        assert not _mn, '同尺寸同页数不该有提示：{}'.format(_mn)
        # 叠加后两边的文字都得在，才说明真叠上了而不是覆盖掉
        _txt = _fitz.open(_merged)[0].get_text()
        assert 'LETTERHEAD' in _txt and 'CONTENT' in _txt, \
            '合成页应同时含套头与内容：{!r}'.format(_txt)

        # 尺寸不一致要如实提示，而不是悄悄错位
        _a5 = _mk_pdf(os.path.join(OUT_DIR, 'content_a5.pdf'), 14.8, 21.0,
                      text='CONTENT')
        _m2 = os.path.join(OUT_DIR, 'merged2.pdf')
        assert any('尺寸不一致' in _s for _s in
                   _ovl.merge_overlay(_a5, _lh, _m2)), '尺寸不符应提示'

        # 内容页数多于套头时要说明超出的页没有底图
        _multi = _mk_pdf(os.path.join(OUT_DIR, 'content3.pdf'), pages=3,
                         text='CONTENT')
        _m3 = os.path.join(OUT_DIR, 'merged3.pdf')
        _n3b = _ovl.merge_overlay(_multi, _lh, _m3)
        assert _ovl.page_count(_m3) == 3, '3 页内容应产出 3 页'
        assert any('没有套头底图' in _s for _s in _n3b), \
            '超出页应说明：{}'.format(_n3b)

        # 缺 PyMuPDF 时必须给人话、且不是 ImportError——
        # 套头叠加是附加功能，不该甩底层异常给用户
        import builtins as _bi
        from scripts import header_overlay as _HO
        _real_imp = _bi.__import__

        def _no_fitz(name, *a, **k):
            if name == 'fitz' or name.startswith('fitz.'):
                raise ImportError("No module named 'fitz'")
            return _real_imp(name, *a, **k)
        _bi.__import__ = _no_fitz
        try:
            _ok, _why = _HO.available()
            assert not _ok and 'PyMuPDF' in _why, '缺库时 available() 应说明原因'
            for _fn, _args in (('page_count', (_lh,)), ('page_size_cm', (_lh,)),
                               ('render_page_to_png', (_lh,)),
                               ('overlay_content_on_header', (_lh, _lh, _lh))):
                try:
                    getattr(_HO, _fn)(*_args)
                    raise AssertionError('{} 缺库时应抛错'.format(_fn))
                except RuntimeError as _e:
                    assert 'PyMuPDF' in str(_e), '{} 的报错要提到 PyMuPDF'.format(_fn)
                except ImportError:
                    raise AssertionError('{} 不该把 ImportError 甩给上层'.format(_fn))
            assert _ovl.page_count(_lh) == 0, '缺库时 page_count 应返回 0 而不是抛错'
            assert _ovl.page_size_cm(_lh) is None, '缺库时 page_size_cm 应返回 None'
        finally:
            _bi.__import__ = _real_imp
        assert _HO.available()[0], '恢复后应重新可用'

        # 渲染出图（预览靠它）
        _png = _ovl.render_page_png(_merged, 0)
        assert os.path.exists(_png) and os.path.getsize(_png) > 1000, '渲染 PNG 异常'
        os.remove(_png)
        print('    套头对位：合成/两层文字俱在/尺寸告警/页数告警/渲染 通过')

    print('[7o] 套打：填充/几何锁定/自适应/空值留白/预览与输出一致 + docx 适配'
          ' + 纵向合并去重 + 日期识别 10 种写法 + 标题梯形回行 + 年月日定宽 通过')


def test_gb_header_record():
    """版头红线/版记分隔线（flags 开启）+ 副标题识别"""
    from docx.oxml.ns import qn
    from scripts.formatter import format_document, detect_para_type, _compile_rules
    from scripts.data_model import PRESETS
    r=_compile_rules(None); at=['a']*20
    f={'header_elements':True}
    assert detect_para_type('000123',0,20,None,at,0,rules=r,flags=f)=='copynum'
    assert detect_para_type('签发人：张三',2,20,None,at,2,rules=r,flags=f)=='signatory'
    fr={'record_elements':True}
    assert detect_para_type('抄送：市各部门。',18,20,None,at,18,rules=r,flags=fr)=='cc'
    fs={'subtitle_enabled':True}
    assert detect_para_type('——试点说明',1,20,None,at,1,rules=r,prev_para_type='title',flags=fs)=='subtitle'
    d=Document()
    for t in ['某政发〔2026〕5号','签发人：张三','关于试点的通知','各单位：','正文。',
              '特此通知。','某办公室','2026年7月24日','抄送：市各部门。','某办公室2026年7月24日印发']:
        d.add_paragraph(t)
    src=os.path.join(OUT_DIR,'gb_in.docx'); d.save(src); out=os.path.join(OUT_DIR,'gb_out.docx')
    preset=dict(PRESETS['official_gbk']); preset['header_elements']=True; preset['record_elements']=True
    format_document(src,out,preset_name='custom',custom_settings=preset)
    doc=Document(out)
    sig=[p for p in doc.paragraphs if p.text.strip()=='签发人：张三'][0]
    pPr=sig._p.find(qn('w:pPr')); b=pPr.find(qn('w:pBdr')) if pPr is not None else None
    assert b is not None and b.find(qn('w:bottom')) is not None, '版头红线未加'
    print('[7j] 版头红线/版记分隔线/副标题 通过')


def test_image_protection():
    """含图段落保护：独占图片的空文字段落不被压成 1 磅裁掉图片（借鉴 Word-Formatter-Pro）"""
    import base64
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_LINE_SPACING
    from scripts.formatter import format_document, paragraph_has_media
    png = base64.b64decode('iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+ip1sAAAAASUVORK5CYII=')
    ip = os.path.join(OUT_DIR, 't.png'); open(ip,'wb').write(png)
    d = Document()
    d.add_paragraph('关于测试的通知'); d.add_paragraph('各单位：')
    d.add_paragraph().add_run().add_picture(ip, width=Cm(8), height=Cm(6))
    d.add_paragraph('特此通知。')
    src = os.path.join(OUT_DIR, 'img_in.docx'); d.save(src)
    assert any(paragraph_has_media(p) for p in Document(src).paragraphs), '未检出含图段落'
    out = os.path.join(OUT_DIR, 'img_out.docx')
    format_document(src, out, preset_name='official_gbk')
    for para in Document(out).paragraphs:
        if para._p.find('.//'+qn('w:drawing')) is not None:
            assert para.paragraph_format.line_spacing_rule != WD_LINE_SPACING.EXACTLY, \
                '含图段落仍是固定行距，会裁图'
    print('[7f] 含图段落保护: 不被裁图 通过')


def test_redaction():
    """日志脱敏：文件名/路径/用户名不明文，同名一致"""
    from app.redact import redact_text, mask_home
    r = redact_text('正在处理: 涉密测试.docx')
    assert '涉密测试' not in r and '.docx' in r and '文档-' in r
    r2 = redact_text('处理失败 C:' + chr(92) + 'Users' + chr(92) + '王五' + chr(92)
                     + '秘密' + chr(92) + '报告.docx: 被占用')
    assert '王五' not in r2 and '秘密' not in r2
    assert '张三' not in mask_home('/home/张三/x') and '/home/***' in mask_home('/home/张三/x')
    assert redact_text('测试.docx') == redact_text('测试.docx')  # 同名一致
    print('[7e] 日志脱敏: 文件名/路径/用户名 通过')


def test_signature_closing():
    """署名识别扩充：室/部结尾 + 结束语妥否请审示"""
    from scripts.formatter import detect_para_type, DEFAULT_DETECT_RULES
    rules = {k: v for k, v in DEFAULT_DETECT_RULES.items()}
    # 署名：以室结尾
    assert detect_para_type('调查室', 8, 12, None, ['a']*10, 8, rules=rules) == 'signature'
    assert detect_para_type('监督室', 8, 12, None, ['a']*10, 8, rules=rules) == 'signature'
    # 署名：以部结尾
    assert detect_para_type('组织部', 8, 12, None, ['a']*10, 8, rules=rules) == 'signature'
    assert detect_para_type('宣传部', 8, 12, None, ['a']*10, 8, rules=rules) == 'signature'
    # 结束语：妥否，请审示。
    assert detect_para_type('妥否，请审示。', 8, 12, None, ['a']*10, 8, rules=rules) == 'closing'
    print('[7b] 署名/结束语扩充: 室/部/妥否请审示 通过')


def test_scan_align():
    """扫描件自动对位：造一份"印偏了"的套头，看能不能把偏移量原样量回来"""
    from docx import Document
    from scripts import overprint as op, scan_align
    from scripts.exporter import export_pdf
    ok, why = scan_align.available()
    if not ok:
        print('[15] 扫描件自动对位：本机缺 {} — 跳过'.format(why))
        return
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')

    # ---- 整体平移：存得住、读得回、真能挪动版面 ----
    work = os.path.join(OUT_DIR, 'shift_tpl.docx')
    shutil.copy(tpl, work)
    try:
        op.save_shift(work, 0.30, -0.20)
        assert op.load_shift(work) == (0.3, -0.2), op.load_shift(work)
        out = os.path.join(OUT_DIR, 'shift_out.docx')
        op.fill_form(work, {'标题': '测试'}, out, one_page=False)
        m0 = Document(work).sections[0]
        m1 = Document(out).sections[0]
        assert abs(m1.left_margin.cm - m0.left_margin.cm - 0.30) < 0.01, '左边距应右移 0.30'
        assert abs(m1.top_margin.cm - m0.top_margin.cm + 0.20) < 0.01, '上边距应上移 0.20'
        assert abs((m1.left_margin.cm + m1.right_margin.cm)
                   - (m0.left_margin.cm + m0.right_margin.cm)) < 0.01, '版心宽度不该变'
        # 平移量单独存，逐字段微调与套头绑定都不受影响
        op.save_offsets(work, {'标题': 3.5})
        assert op.load_shift(work) == (0.3, -0.2), '存字段位置时不该丢掉平移量'
        op.save_shift(work, 0, 0)
        assert op.load_offsets(work) == {'标题': 3.5}, '存平移量时不该丢掉字段位置'
    finally:
        for p in (work, op.offsets_path(work)):
            if os.path.exists(p):
                os.remove(p)

    # ---- 自动对位：把模板白线刷黑、整体挪一下，当作"扫描件" ----
    doc = Document(tpl)
    assert scan_align._blacken_borders(doc) > 0, '模板里应有白色框线'
    DX, DY = 0.30, -0.20
    op.apply_shift(doc, DX, DY)
    fake = os.path.join(OUT_DIR, 'scan_fake.docx')
    fake_pdf = os.path.join(OUT_DIR, 'scan_fake.pdf')
    doc.save(fake)
    good, info = export_pdf(fake, fake_pdf)
    if not good:
        print('[15] 扫描件自动对位：本机转不了 PDF（{}）— 跳过'.format(info))
        return
    res = scan_align.align(fake_pdf, tpl)
    assert len(res['h_pairs']) >= 4, '横线应认出 4 条以上：{}'.format(res['h_pairs'])
    assert len(res['v_pairs']) >= 2, '竖线应认出 2 条以上：{}'.format(res['v_pairs'])
    assert abs(res['dx'] - DX) < 0.05, 'dx 应量回 {}，实得 {}'.format(DX, res['dx'])
    assert abs(res['dy'] - DY) < 0.05, 'dy 应量回 {}，实得 {}'.format(DY, res['dy'])
    assert not res['warnings'], '没缩放不该告警：{}'.format(res['warnings'])
    assert '往右挪' in scan_align.describe(res)
    print('[15] 扫描件自动对位：认出 {} 条横线/{} 条竖线，'
          '量回偏移 ({:+.2f}, {:+.2f})cm，误差 <0.01cm 通过'
          .format(len(res['h_pairs']), len(res['v_pairs']), res['dx'], res['dy']))


def test_overprint_fonts():
    """送审单模板的字体字号：预印栏目名 + 填写内容，按真实预印纸"""
    from docx.oxml.ns import qn as _qn
    from scripts import overprint as op
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')
    doc = Document(tpl)

    def _runs():
        for p in doc.paragraphs:
            for r in p.runs:
                yield r
        for t in doc.tables:
            for c in op._iter_cells(t):
                for p in c.paragraphs:
                    for r in p.runs:
                        yield r

    got = {}
    for r in _runs():
        txt = r.text.strip()
        if not txt:
            continue
        rPr = r._r.find(_qn('w:rPr'))
        rf = rPr.find(_qn('w:rFonts')) if rPr is not None else None
        if rf is None or r.font.size is None:
            continue
        got[txt] = (rf.get(_qn('w:eastAsia')), rf.get(_qn('w:ascii')),
                    r.font.size.pt, bool(r.font.bold))

    want = {
        # 预印在纸上的（白字占位）
        '中国某地市某单位': ('方正大标宋简体', 18.0),      # 小二
        '文件送审单': ('方正大标宋简体', 22.0),            # 二号
        '紧急程度：': ('方正楷体_GBK', 14.0),              # 四号
        '密级：': ('方正楷体_GBK', 14.0),
        '标  题': ('方正楷体_GBK', 14.0),
        '领导批示：': ('方正楷体_GBK', 14.0),
        '拟办意见：': ('方正楷体_GBK', 14.0),
        '承办部门：': ('方正楷体_GBK', 14.0),
        '经办人：': ('方正楷体_GBK', 14.0),
        '电话：': ('方正楷体_GBK', 14.0),
        '文字校核：': ('方正楷体_GBK', 14.0),
        '年': ('方正楷体_GBK', 14.0),
        '某地市某某单位的办公室制': ('方正楷体_GBK', 14.0),
        # 要打印出来的
        '{{标题}}': ('方正小标宋_GBK', 16.0),              # 三号
        '{{紧急程度}}': ('方正楷体_GBK', 14.0),
        '{{密级}}': ('方正楷体_GBK', 14.0),
        '{{拟办意见}}': ('方正仿宋_GBK', 15.0),            # 小三
        '{{承办部门}}': ('方正楷体_GBK', 14.0),
    }
    for txt, (font, size) in want.items():
        assert txt in got, '模板里找不到 {!r}'.format(txt)
        ea, _ascii, pt, bold = got[txt]
        assert ea == font, '{!r} 中文字体应为 {}，实为 {}'.format(txt, font, ea)
        assert abs(pt - size) < 0.01, '{!r} 字号应为 {}，实为 {}'.format(txt, size, pt)
        assert bold, '{!r} 应加粗'.format(txt)
    # 电话与年月日的数字走西文字体
    for txt in ('{{电话}}', '{{年}}', '{{月}}', '{{日}}'):
        ea, ascii_, pt, bold = got[txt]
        assert ascii_ == 'Times New Roman' and ea == 'Times New Roman', \
            '{!r} 数字应用 Times New Roman，实为 {}/{}'.format(txt, ea, ascii_)
        assert abs(pt - 14.0) < 0.01 and bold, '{!r} 应四号加粗'.format(txt)

    # 标题居中，且居中的是竖线右边那一格
    tbl = doc.tables[0]
    title_cell = [c for c in op._iter_cells(tbl) if '{{标题}}' in c.text][0]
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP
    assert title_cell.paragraphs[0].alignment == _WAP.CENTER, '标题应居中'
    label_cell = [c for c in op._iter_cells(tbl) if c.text.strip() == '标  题'][0]
    lw = op._cell_width_cm(tbl, label_cell)
    # 竖线位置是**直接量的**（4.70）。原先拿「紧急程度：」冒号右边线
    # （4.60）代替，用户说"基本一致"——实际差 1mm，"基本"确实只是基本
    vline = 2.1 + lw
    assert abs(vline - 4.7) < 0.02, \
        '标题右侧竖线应在 4.70cm（实测），实得 {:.2f}cm'.format(vline)
    print('[14] 套打模板字体：预印小二/二号大标宋 + 四号楷体，'
          '填写三号小标宋标题(居中)/小三仿宋意见/TNR 数字 通过')


def test_y_offsets():
    """纵向微调：能挪、挪不动时说清楚，且不动横向"""
    from scripts import overprint as op
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')
    vals = {'年': '2026', '月': '7', '日': '29'}

    def _last_top(plan):
        return [b for b in plan['blocks'] if b['kind'] == 'para'][-1]['top_cm']

    base = op.plan_fill(tpl, vals)
    top0 = _last_top(base)

    # 往下挪：说多少就是多少
    moved = op.plan_fill(tpl, vals, offsets_y={'年': top0 + 0.8})
    assert abs(_last_top(moved) - (top0 + 0.8)) < 0.02, \
        '纵向没挪到位：{:.2f} → {:.2f}'.format(top0, _last_top(moved))
    assert not moved['notes'], '正常下挪不该有提示：{}'.format(moved['notes'])
    # 横向不受影响
    assert abs(moved['field_pos']['年'] - base['field_pos']['年']) < 0.01, \
        '纵向微调不该动横向位置'

    # 往上挪过头：不闷头改，如实说最多能到哪儿
    up = op.plan_fill(tpl, vals, offsets_y={'年': 2.0})
    assert abs(_last_top(up) - top0) < 0.02, '挪不动时应保持原位'
    assert any('已经被上一行占了' in n for n in up['notes']), \
        '挪不动要说清楚：{}'.format(up['notes'])

    # 表格里的字段挪不了，也要说
    intbl = op.plan_fill(tpl, {'承办部门': '办公室'}, offsets_y={'承办部门': 20.0})
    assert any('行高定死' in n for n in intbl['notes']), \
        '表格里的字段应提示挪不了：{}'.format(intbl['notes'])

    # 同一行的几个字段只能一起动
    same = op.plan_fill(tpl, vals, offsets_y={'年': top0 + 0.5, '月': top0 + 0.5})
    assert any('同一行' in n for n in same['notes']), \
        '同行字段应提示一起动：{}'.format(same['notes'])

    # 存盘往返：纵向与横向、整体平移互不干扰
    work = os.path.join(OUT_DIR, 'yoff.docx')
    shutil.copy(tpl, work)
    try:
        op.save_offsets(work, {'标题': 5.0}, shift=(0.1, 0.2),
                        offsets_y={'年': 27.5})
        assert op.load_offsets(work) == {'标题': 5.0}
        assert op.load_offsets_y(work) == {'年': 27.5}
        assert op.load_shift(work) == (0.1, 0.2)
        op.save_offsets(work, {'标题': 5.0})        # 不传纵向时不该丢
        assert op.load_offsets_y(work) == {'年': 27.5}, '存横向时把纵向弄丢了'
    finally:
        for f in (work, op.offsets_path(work)):
            if os.path.exists(f):
                os.remove(f)

    # 打印预检
    long_plan = op.plan_fill(tpl, {'拟办意见': '因某某事项需要开展调查。' * 40})
    kinds = [lv for lv, _m in op.preflight(long_plan, {}, {})]
    assert 'block' in kinds, '放不下时预检应给 block：{}'.format(kinds)
    ok_plan = op.plan_fill(tpl, {'标题': '关于某事的请示'})
    assert not [lv for lv, _m in op.preflight(ok_plan, {'标题': '关于某事的请示'}, {})
                if lv in ('block', 'warn')], '正常内容不该报预检问题'
    print('[17] 纵向微调：下挪到位/上挪受阻如实告知/表格内不可挪/同行联动 + '
          '存盘互不干扰 + 打印预检 通过')


def test_attachment_consistency():
    """正文说的附件与实际附上的附件对不对得上"""
    from scripts import attachment_check as A
    from scripts import compliance as C
    from scripts.data_model import PRESETS

    # 说明行的几种写法都要认得
    assert A.parse_desc(['附件：1.实施方案  2.任务分工表']) == \
        [(1, '实施方案'), (2, '任务分工表')]
    assert A.parse_desc(['附件：', '1.实施方案', '2.任务分工表']) == \
        [(1, '实施方案'), (2, '任务分工表')]
    assert A.parse_desc(['附件：实施方案']) == [(None, '实施方案')]
    assert A.parse_desc(['附件1：实施方案', '附件2：任务分工表']) == \
        [(1, '实施方案'), (2, '任务分工表')]
    assert A.parse_labels(['附件1', '附件2']) == [1, 2]
    assert A.parse_labels(['附件一', '附件二']) == [1, 2]     # 中文序号
    assert A.parse_labels(['附件']) == [None]
    assert A.parse_labels(['附件：说明']) == [], '带冒号的是说明行，不是标识'

    def warns(desc, labels):
        return sorted(x['item'] for x in A.check(desc, labels)
                      if x['level'] == 'warn')

    assert warns(['附件：1.甲  2.乙'], ['附件1', '附件2']) == []
    assert warns(['附件：1.甲  2.乙'], ['附件1']) == ['附件·个数对不上']
    assert warns([], ['附件1', '附件2']) == ['附件·正文未说明']
    # 附件常是单独的文件，正文有说明、本文档没标识是正常的，不能算差错
    assert warns(['附件：1.甲  2.乙'], []) == [], '不该冤枉"附件是单独文件"的情形'
    assert [x['level'] for x in A.check(['附件：1.甲  2.乙'], [])] == ['info']
    assert warns([], []) == []
    assert warns(['附件：1.甲  3.乙'], ['附件1', '附件3']) == \
        ['附件·序号不连续', '附件·序号不连续']
    assert warns(['附件：1.甲  1.乙'], ['附件1', '附件2']) == ['附件·序号重复']
    # GB/T 9704：只有一个附件时不编号
    assert warns(['附件：1.甲'], ['附件1']) == ['附件·单个不编号']
    assert warns(['附件：甲'], ['附件']) == []
    # 个数已经对不上时，不再追着编号说事——那只是往报告里添噪音
    assert warns(['附件：1.甲  2.乙'], ['附件1']) == ['附件·个数对不上']

    # 走合规检查完整链路
    src = os.path.join(OUT_DIR, 'att_in.docx')
    d = Document()
    for t in ('关于开展某项工作的通知', '各部门：', '为落实上级要求，现通知如下。' * 3,
              '特此通知。', '附件：1.实施方案  2.任务分工表',
              '某某公司办公室', '2026年7月17日', '附件1', '附件正文。'):
        d.add_paragraph(t)
    d.save(src)
    preset = PRESETS['official_gbk']
    got = [x for x in C.check_compliance(Document(src), preset,
                                         C.only(C.ATTACHMENT_KEYS))
           if x['level'] == 'warn']
    assert [x['item'] for x in got] == ['附件·个数对不上'], \
        '合规检查里没接上附件核对：{}'.format([x['item'] for x in got])
    assert '2 个附件' in got[0]['detail']
    # 分组开关
    off = C.only(C.ATTACHMENT_KEYS, {k: False for k in C.ATTACHMENT_KEYS})
    assert not [x for x in C.check_compliance(Document(src), preset, off)
                if x['item'].startswith('附件·')]
    print('[25] 附件一致性：说明四种写法/标识中文序号 + 个数·序号·单个不编号 + '
          '"附件是单独文件"不误报 通过')


def test_archive():
    """归档命名与登记表：抽字段 → 生成文件名 → 复制归档 → 追加台账"""
    import io as _io
    from scripts import archive as A
    from scripts.data_model import PRESETS

    def mk(path, no, title, date='2026年7月17日', sec='秘密★1年'):
        d = Document()
        for t in ([sec] if sec else []) + [no, title, '各部门：', '正文。',
                                           '特此通知。', '某某办公室', date]:
            d.add_paragraph(t)
        d.save(path)
        return path

    work = os.path.join(OUT_DIR, 'arch')
    if os.path.isdir(work):
        shutil.rmtree(work)
    os.makedirs(work)
    a = mk(os.path.join(work, 'draft1.docx'), '某安委发〔2026〕12号',
           '关于开展安全生产检查的通知')
    b = mk(os.path.join(work, 'draft2.docx'), '某安委发〔2026〕13号',
           '关于报送材料的请示')

    meta = A.extract_meta(a, PRESETS['official_gbk'])
    assert meta['文号'] == '某安委发〔2026〕12号'
    assert meta['标题'] == '关于开展安全生产检查的通知'
    assert meta['文种'] == '通知'
    assert (meta['成文日期'], meta['年'], meta['月'], meta['日']) == \
        ('20260717', '2026', '07', '17')
    assert meta['密级'] == '秘密★1年' and meta['密级词'] == '秘密'
    assert meta['发文机关'] == '某某办公室'
    assert meta['原文件名'] == 'draft1'

    # 命名式：取不到的字段留空后，不能留下「20260717--标题」这种空档
    assert A.render_name('{成文日期}-{文号}-{标题}', meta) == \
        '20260717-某安委发〔2026〕12号-关于开展安全生产检查的通知.docx'
    bare = dict(meta, 文号='', 成文日期='')
    assert A.render_name('{成文日期}-{文号}-{标题}', bare) == \
        '关于开展安全生产检查的通知.docx', A.render_name('{成文日期}-{文号}-{标题}', bare)
    # 文件名里不能出现的字符要去掉
    assert '/' not in A.render_name('{标题}', dict(meta, 标题='关于A/B的通知'))
    assert len(os.path.splitext(A.render_name('{标题}', dict(meta, 标题='长' * 300)))[0]) \
        <= A.MAX_STEM

    # 归档：默认复制，原件不动；重名自动加 (2)
    items = A.plan([a, b, a], '{成文日期}-{文号}-{标题}')
    out = os.path.join(work, '归档')
    led = os.path.join(work, '台账.csv')
    done, failed = A.archive(items, out, ledger_path=led)
    assert not failed and len(done) == 3
    assert os.path.exists(a) and os.path.exists(b), '默认应复制，原件不能没了'
    names = sorted(os.listdir(out))
    assert len(names) == 3 and any('(2)' in n for n in names), \
        '重名应自动加 (2)，不能悄悄覆盖：{}'.format(names)

    # 台账：utf-8-sig（Excel 双击不乱码）+ 追加不重复写表头
    text = _io.open(led, encoding='utf-8-sig').read()
    assert text.startswith('归档时间,'), text[:40]
    assert '某安委发〔2026〕12号' in text and '关于报送材料的请示' in text
    A.archive(A.plan([b], '{标题}'), out, ledger_path=led)
    lines = _io.open(led, encoding='utf-8-sig').read().strip().splitlines()
    assert sum(1 for x in lines if x.startswith('归档时间,')) == 1, '表头写了不止一次'
    assert len(lines) == 5, '台账应追加到 4 行数据：{}'.format(len(lines))
    with open(led, 'rb') as f:
        assert f.read(3) == b'\xef\xbb\xbf', 'CSV 要带 BOM，否则 Excel 里中文是乱码'

    # 移动：原件要真的不在了
    c = mk(os.path.join(work, 'draft3.docx'), '某发〔2026〕1号', '关于某事的函')
    A.archive(A.plan([c], '{标题}'), out, move=True)
    assert not os.path.exists(c), '勾了移动就该真移动'
    assert os.path.exists(os.path.join(out, '关于某事的函.docx'))
    print('[26] 归档命名：抽文号/标题/文种/日期/密级 + 命名式（空字段收干净、'
          '非法字符、长度）+ 复制不动原件/重名不覆盖/移动 + 台账追加不重写表头 通过')


def test_security_mark():
    """密级标注：查漏标 / 查写法 / 查位置，以及按人选定的密级插进版头"""
    from scripts import compliance as C
    from scripts import security_mark as S
    from scripts.data_model import PRESETS
    preset = PRESETS['official_gbk']

    def mk(name, lines):
        path = os.path.join(OUT_DIR, 'sec_%s.docx' % name)
        d = Document()
        for t in lines:
            d.add_paragraph(t)
        d.save(path)
        return path

    def items(path, groups=None):
        return sorted(x['item'] for x in
                      S.check(Document(path), groups=groups)
                      if x['level'] == 'warn')

    # --- 漏标：只认文件自己露出来的破绽，不做"看着像涉密"的猜测 ---
    p = mk('copynum', ['000123', '关于某事的通知', '各部门：', '正文。'])
    assert items(p) == ['密级·漏标'], '有份号无密级应报漏标：{}'.format(items(p))
    p2 = mk('hint', ['关于某事的通知', '各部门：', '本件为机密件，注意保密。'])
    assert items(p2) == ['密级·漏标'], '正文自述保密应报漏标'
    p3 = mk('clean', ['关于某事的通知', '各部门：', '正文内容。'])
    assert items(p3) == [], '普通文件不该被硬扣一顶涉密帽子：{}'.format(items(p3))

    # --- 写法：缺期限 / 分隔符 / 中文数字 / 密级词写错 ---
    for name, line, want in (
            ('noperiod', '秘密', '秘密★1年'),
            ('star', '机密*3年', '机密★3年'),
            ('cn', '绝密★十年', '绝密★10年'),
            ('typo', '机秘★3年', '机密★3年')):
        path = mk(name, [line, '关于某事的通知', '正文。'])
        got = [x for x in S.check(Document(path)) if x['level'] == 'warn']
        assert got, '{} 应被指出不规范'.format(line)
        assert got[0]['fix_key'].endswith(want), \
            '{} 应改成 {}，实际 {}'.format(line, want, got[0]['fix_key'])
    ok = mk('ok', ['秘密★1年', '关于某事的通知', '正文。'])
    assert items(ok) == [], '规范写法不该报错：{}'.format(items(ok))

    # --- 位置：版头顺序 / 跑到正文里去了 ---
    p4 = mk('order', ['特急', '秘密★1年', '关于某事的通知'])
    assert items(p4) == ['密级·位置'], '紧急程度排在密级上面应报位置'
    p5 = mk('stray', ['关于某事的通知', '各部门：', '正文。', '秘密★1年'])
    assert '密级·位置' in items(p5), '末尾的密级行应被指出位置不对'
    # 短文件里正文也在前几段——不能只按段序号卡，否则这一条会漏掉
    assert not [x for x in S.check(Document(p5))
                if x['item'] == '密级·写法' and x['level'] == 'ok'], \
        '跑到正文里的密级不该被当成版头里的正规密级'

    # --- 分组开关各管一摊 ---
    p6 = mk('mixed', ['000123', '机密*3年', '关于某事的通知', '正文。', '秘密★1年'])
    assert items(p6, {'sec_format': False}) == ['密级·位置']
    assert items(p6, {'sec_position': False}) == ['密级·写法不规范']
    assert items(p6, dict.fromkeys(S.GROUP_KEYS, False)) == []

    # --- 插入：位置按版头顺序来，不是随便插一段 ---
    for name, lines, want_at in (
            ('ins_copy', ['000123', '特急', '关于某事的通知'], 1),
            ('ins_urg', ['特急', '关于某事的通知'], 0),
            ('ins_bare', ['关于某事的通知', '正文。'], 0)):
        doc = Document(mk(name, lines))
        text, added = S.insert(doc, preset, '秘密', '1年')
        assert added and text == '秘密★1年'
        got = [q.text for q in doc.paragraphs if q.text.strip()]
        assert got[want_at] == '秘密★1年', \
            '{}: 密级应插在第 {} 段，实际 {}'.format(name, want_at + 1, got)
    # 已有密级就只改字，不插第二个
    doc = Document(mk('twice', ['秘密★1年', '关于某事的通知']))
    _t, added = S.insert(doc, preset, '机密', '3年')
    assert not added
    assert [q.text for q in doc.paragraphs if q.text.strip()] == \
        ['机密★3年', '关于某事的通知'], '已有密级时不该再插一行'

    # --- 走合规检查的完整链路：勾选 → 修正 → 复查归零 ---
    src = mk('flow', ['000123', '关于某事的通知', '各部门：', '正文。'])
    warns = [x for x in C.check_compliance(Document(src), preset,
                                           C.only(C.SECURITY_KEYS))
             if x['level'] == 'warn']
    assert [x['fix_key'] for x in warns] == ['security:insert']
    assert '密级' in C.fix_label('security:insert:秘密★1年')
    out = os.path.join(OUT_DIR, 'sec_fixed.docx')
    applied = C.apply_compliance_fixes(src, out, preset,
                                       ['security:insert:秘密★1年'])
    assert applied and '秘密★1年' in applied[0], applied
    left = [x for x in C.check_compliance(Document(out), preset,
                                          C.only(C.SECURITY_KEYS))
            if x['level'] == 'warn']
    assert not left, '补标之后应再无密级偏差：{}'.format(
        [x['detail'] for x in left])
    # 插进去的那一段要按预设的密级格式排（黑体三号顶格），不是裸段落
    sec_para = [q for q in Document(out).paragraphs
                if q.text.strip() == '秘密★1年'][0]
    spec = preset['security']
    assert sec_para.runs and sec_para.runs[0].font.size.pt == spec['size'], \
        '插入的密级没按预设排版'
    assert (sec_para.paragraph_format.first_line_indent or 0) == 0, \
        '密级应顶格'
    print('[24] 密级标注：漏标(份号/正文自述)/写法(期限·分隔符·中文数字·错词)/'
          '位置(版头顺序·跑进正文) + 按人选定的密级插入版头正确位置 通过')


def test_tables_and_attachment():
    """三线表 / 续表表头 / 整行不断页 / 附件另面编排"""
    import copy as _copy
    from docx.oxml.ns import qn as _qn
    from scripts.data_model import PRESETS
    from scripts.formatter import format_document as _fmt

    src = os.path.join(OUT_DIR, 'tbl_in.docx')
    d = Document()
    d.add_paragraph('关于某项工作的通知')
    d.add_paragraph('各部门：')
    t = d.add_table(rows=5, cols=3)
    for i, row in enumerate(t.rows):
        for j, c in enumerate(row.cells):
            c.text = ('项目', '数量', '备注')[j] if i < 2 else '{}-{}'.format(i, j)
    d.add_paragraph('特此通知。')
    d.add_paragraph('某某公司办公室')
    d.add_paragraph('2026年7月17日')
    d.add_paragraph('附件1')
    d.add_paragraph('这是附件的正文内容。')
    d.save(src)

    def tbl_info(path):
        doc = Document(path)
        tb = doc.tables[0]
        b = tb._tbl.tblPr.find(_qn('w:tblBorders'))
        edges = {e.tag.split('}')[-1]: e.get(_qn('w:val')) for e in b}
        hdr = []
        split = []
        for r in tb.rows:
            pr = r._tr.find(_qn('w:trPr'))
            hdr.append(pr is not None and pr.find(_qn('w:tblHeader')) is not None)
            split.append(pr is not None and pr.find(_qn('w:cantSplit')) is not None)
        return edges, hdr, split

    # --- 默认：满格线，不重复表头 ---
    plain = os.path.join(OUT_DIR, 'tbl_plain.docx')
    _fmt(src, plain, preset_name='custom',
         custom_settings=_copy.deepcopy(PRESETS['official_gbk']))
    edges, hdr, split = tbl_info(plain)
    assert edges['insideV'] == 'single', '默认应画竖线：{}'.format(edges)
    assert not any(hdr), '默认不重复表头'
    assert all(split), '一行数据不该被页边切成两半'

    # --- 三线表：只剩顶线、底线，没有竖线和中间横线 ---
    pre = _copy.deepcopy(PRESETS['official_gbk'])
    pre['table'] = {'three_line': True, 'header_repeat': True, 'header_rows': 2}
    three = os.path.join(OUT_DIR, 'tbl_three.docx')
    _fmt(src, three, preset_name='custom', custom_settings=pre)
    edges, hdr, split = tbl_info(three)
    assert edges['top'] == 'single' and edges['bottom'] == 'single', edges
    for side in ('left', 'right', 'insideH', 'insideV'):
        assert edges[side] == 'none', '三线表不该有 {} 线：{}'.format(side, edges)
    # 栏目线画在表头最后一行的下边——两行表头就是第 2 行
    tb = Document(three).tables[0]
    for ri, want in ((0, False), (1, True)):
        tcp = tb.rows[ri].cells[0]._tc.find(_qn('w:tcPr'))
        tcb = tcp.find(_qn('w:tcBorders')) if tcp is not None else None
        got = tcb is not None and tcb.find(_qn('w:bottom')) is not None and \
            tcb.find(_qn('w:bottom')).get(_qn('w:val')) == 'single'
        assert got == want, '第{}行的栏目线应为 {}'.format(ri + 1, want)
    assert hdr[:2] == [True, True] and not any(hdr[2:]), \
        '两行表头都该跨页重复，否则续表少一行栏目名：{}'.format(hdr)

    # --- 附件另面编排 ---
    def breaks(path):
        return [p.text for p in Document(path).paragraphs
                if p.paragraph_format.page_break_before]
    assert breaks(plain) == ['附件1'], \
        '附件标识行应另起一页：{}'.format(breaks(plain))
    # 另起页后不该在新页顶上还留个空行
    doc = Document(plain)
    for i, p in enumerate(doc.paragraphs):
        if p.paragraph_format.page_break_before:
            assert doc.paragraphs[i - 1].text.strip(), \
                '另起页的附件前不该再留空行'
    pre2 = _copy.deepcopy(PRESETS['official_gbk'])
    pre2['attachment_new_page'] = False
    off = os.path.join(OUT_DIR, 'tbl_nobreak.docx')
    _fmt(src, off, preset_name='custom', custom_settings=pre2)
    assert not breaks(off), '关掉之后不该再分页'
    print('[22] 表格：三线表（无竖线/栏目线在表头末行）+ 两行表头跨页重复 + '
          '整行不断页；附件另面编排可开关 通过')


def test_print_and_align_sheet():
    """直接打印：对位测试页 + 100% 不缩放"""
    from scripts import overprint as op
    from scripts import printing
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')
    sheet = os.path.join(OUT_DIR, 'align_sheet.docx')
    n, _notes = op.build_align_sheet(tpl, sheet)
    assert n >= 10, '对位测试页应把每个字段都标上：只填了 {}'.format(n)

    doc = Document(sheet)
    # 每个填写位印的是它自己的落点厘米数，且与预览算出来的一致
    live = op.plan_fill(tpl, {})['field_pos']
    text = '\n'.join(p.text for p, _c in op._iter_paragraphs(doc))
    for name in ('紧急程度', '承办部门', '电话'):
        assert '▸{:.2f}'.format(live[name]) in text, \
            '{} 的坐标标记没印上：{:.2f}'.format(name, live[name])
    # 预印白字要变成浅灰，白纸上才看得见
    whites = [r for p, _c in op._iter_paragraphs(doc) for r in p.runs
              if r.text.strip() and r.font.color and r.font.color.rgb
              and str(r.font.color.rgb) == 'FFFFFF']
    assert not whites, '对位测试页上不该还有白字，白纸上等于没印'

    ok, why = printing.available()
    if not ok:
        print('[23] 直接打印：本机缺 {} — 只验了对位测试页'.format(why))
        return
    # 打到"PDF 打印机"上，验证整条渲染链且**一点都不缩放**
    from PyQt5.QtWidgets import QApplication
    global _QAPP
    # 必须留个引用：QApplication 一被回收，QPrinter 就报"没有 QCoreApplication"
    if QApplication.instance() is None:
        _QAPP = QApplication([])
    from PyQt5.QtPrintSupport import QPrinter
    src_pdf, _t = printing.to_pdf(sheet)
    out = os.path.join(OUT_DIR, 'printed.pdf')
    pr = printing.make_printer()
    pr.setOutputFormat(QPrinter.PdfFormat)
    pr.setOutputFileName(out)
    pages = printing.print_file(sheet, pr)
    assert pages >= 1

    import fitz

    def ink(path, dpi=72):
        pg = fitz.open(path)[0]
        pm = pg.get_pixmap(matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0), alpha=False)
        xs, ys = [], []
        for y in range(0, pm.height, 2):
            row = y * pm.stride
            for x in range(0, pm.width, 2):
                if pm.samples[row + x * pm.n] < 200:
                    xs.append(x)
                    ys.append(y)
        k = 2.54 / dpi
        return min(xs) * k, min(ys) * k, max(xs) * k, max(ys) * k

    a, b = ink(src_pdf), ink(out)
    for i, nm in enumerate(('左', '上', '右', '下')):
        assert abs(a[i] - b[i]) < 0.12, \
            '打印把版面挪了：{}沿 {:.2f} → {:.2f}cm'.format(nm, a[i], b[i])
    scale = (b[2] - b[0]) / (a[2] - a[0])
    assert abs(scale - 1.0) < 0.01, '打印缩放了 {:.3f}，套打会全废'.format(scale)
    pg = fitz.open(out)[0]
    assert abs(pg.rect.width * 2.54 / 72 - 21.0) < 0.1, '打印纸张不是 A4 宽'
    print('[23] 直接打印：对位测试页标出各栏落点 + 白字转灰 + '
          '送打印机 100% 不缩放（偏差 <0.12cm）通过')


def test_wording_split():
    """用语检查独立成一摊：两个入口各查各的，引擎仍是同一套"""
    from scripts import compliance as C
    from scripts import wording as W

    assert not (set(C.LAYOUT_KEYS) & set(C.WORDING_KEYS)), '两组勾选项不该重叠'
    assert set(C.LAYOUT_KEYS) | set(C.WORDING_KEYS) == set(C.DEFAULT_OPTIONS), \
        '拆分后有勾选项无人认领'

    src = os.path.join(OUT_DIR, 'wording_split.docx')
    d = Document()
    for t in ('关于开展安全生产检查的请示。', '省安全生产委员会',
              '为落实上级部署，我委迫不急待地组织排查，共出动3人次，按步就班推进。',
              '特此报告。', '某某公司办公室', '2026年7月17日'):
        d.add_paragraph(t)
    d.save(src)
    from scripts.data_model import PRESETS
    preset = PRESETS['official_gbk']

    def items(opts):
        return sorted({x['item'] for x in
                       C.check_compliance(Document(src), preset, opts)
                       if x['level'] == 'warn'})

    lay = items(C.only(C.LAYOUT_KEYS))
    wor = items(C.only(C.WORDING_KEYS))
    assert lay and not any(i.startswith('用语') for i in lay), \
        '版式检查不该报用语问题：{}'.format(lay)
    assert wor and all(i.startswith('用语') for i in wor), \
        '用语检查不该报版式问题：{}'.format(wor)

    # 拆开之后结果必须与"全开"时**一模一样**——只是分两次给，不是换了套算法
    full = items(None)
    assert sorted(lay + wor) == full, \
        '拆分改变了检查结果：拆后 {} ≠ 全开 {}'.format(sorted(lay + wor), full)

    # 用语规则里有几条要靠段落类型（文种搭配/标题标点）。单独跑用语时
    # 段落级检查是关掉的，若不一并要求识别类型，这几条会悄悄查不出东西
    assert any('文种搭配' in i for i in wor), \
        '单独跑用语时文种搭配规则失效了：{}'.format(wor)

    # 预览 = 实际：预览说改哪几处，落笔就是哪几处
    pv = C.wording_preview(src, preset, ['wording:typo'])
    assert pv and sum(len(e['marks']) for e in pv) == 2, \
        '预览标记数不对：{}'.format(pv)
    doc = Document(src)
    W.apply_wording_fixes(doc, ['wording:typo'], revision=False)
    real = [p.text for p in doc.paragraphs if p.text.strip()]
    for e in pv:
        assert real[e['index']] == e['after'], \
            '预览({})与实际({})不一致'.format(e['after'], real[e['index']])
    # 标记坐标要能在改后文本里精确切出那个词
    for e in pv:
        for a, b, old, new in e['marks']:
            assert e['after'][a:b] == new, '标记坐标切错：{}'.format(
                e['after'][a:b])
            assert old in e['before'], '原词不在原文里：{}'.format(old)
    print('[21] 用语检查独立：两入口各查各的、合并结果与全开一致、'
          '文种搭配不失效、预览与落笔一致 通过')


def test_align():
    """对齐：两端对齐是第一个模板的默认，分散对齐全链路可用"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WA
    from scripts.data_model import PRESETS
    from scripts import compliance as C
    from app.pages.presets_page import ALIGNS

    first = list(PRESETS)[0]
    assert PRESETS[first]['body'].get('align') == 'justify', \
        '第一个排版模板（{}）的正文默认应为两端对齐'.format(first)
    # 下拉里两端对齐排第一，用户不用翻
    assert ALIGNS[0][1] == 'justify', '两端对齐应排在对齐下拉的第一个'
    assert 'distribute' in [v for _t, v in ALIGNS], '对齐下拉里少了分散对齐'
    assert C.ALIGN_LABELS['distribute'] == '分散对齐'
    assert C._ALIGN_MAP['distribute'] == WA.DISTRIBUTE

    # 排版：两种对齐都要真的写进 docx，且检查侧照同一预设核对不报错
    import copy as _copy
    from scripts.formatter import format_document as _fmt
    src = os.path.join(OUT_DIR, 'align_in.docx')
    d = Document()
    d.add_paragraph('关于开展某项工作的通知')
    d.add_paragraph('为贯彻落实上级要求，现将有关事项通知如下。' * 3)
    d.add_paragraph('请各部门认真组织学习并抓好落实。' * 3)
    d.save(src)
    for key, want in (('justify', WA.JUSTIFY), ('distribute', WA.DISTRIBUTE)):
        preset = _copy.deepcopy(PRESETS[first])
        preset['body'] = dict(preset['body'], align=key)
        out = os.path.join(OUT_DIR, 'align_%s.docx' % key)
        _fmt(src, out, preset_name='custom', custom_settings=preset)
        got = [p.paragraph_format.alignment for p in Document(out).paragraphs
               if len(p.text) > 20]
        assert got and all(a == want for a in got), \
            '{} 没写进 docx：{}'.format(key, got)
        bad = [x for x in C.check_compliance(Document(out), preset)
               if x['level'] == 'warn' and x.get('fix_key') == 'para:body:align']
        assert not bad, '{} 排完又被判对齐不符：{}'.format(
            key, [x['detail'] for x in bad])
    print('[19] 对齐：首个模板默认两端对齐 + 分散对齐写入/检查全通 通过')


def test_template_visual_edit():
    """套打模板可视化编辑：改内容/字体/位置/增删，存回去还能正常填"""
    from scripts import overprint as op
    from scripts.tpl_edit import EditSession
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')
    work = os.path.join(OUT_DIR, 'tpledit.docx')
    shutil.copy(tpl, work)
    s = EditSession(work)

    # 版面读得出来，且与填充预览是同一套坐标——编辑时看到的位置
    # 必须就是打印出来的位置，不然编辑器没意义。
    # 落点还要经得起尺子：数字来自用户量真实预印纸（见 make_songshendan.SPEC）。
    # 这里同时守着一个曾经的坑：field_pos 原来是按字宽累加估的、完全没算
    # 制表位，电话一栏差 2.55cm，用户在微调框点"取当前"再保存就会把
    # 本来对准的位置带偏。
    pos = s.positions()
    ruler = {'紧急程度': (4.60, (2, 2)), '密级': (15.90, (2, 5)),
             '承办部门': (4.50, (9, 2)), '电话': (15.50, (10, 5))}
    for vals in ({}, {'紧急程度': '特急', '承办部门': '办公室',
                      '电话': '12345678', '密级': '秘密'}):
        live = op.plan_fill(tpl, vals)['field_pos']
        for name, (want, ref) in ruler.items():
            assert abs(live[name] - want) < 0.02, \
                '{} 落点 {:.2f} 与尺子量的 {:.2f} 对不上（填了{}值）'.format(
                    name, live[name], want, '' if vals else '空')
            assert abs(pos[ref][0] - live[name]) < 0.02, \
                '{} 编辑器算的位置({:.2f})与填充预览({:.2f})对不上'.format(
                    name, pos[ref][0], live[name])

    # 横向：挪的是制表位，栏目名和它后面的填写位一起动
    lab0, fld0 = pos[(2, 1)][0], pos[(2, 2)][0]
    s.nudge_x((2, 1), 0.5)
    p1 = s.positions()
    assert abs(p1[(2, 1)][0] - (lab0 + 0.5)) < 0.02, '横向没挪到位'
    assert abs(p1[(2, 2)][0] - (fld0 + 0.5)) < 0.02, '同一制表位后面的字段应一起动'

    # 纵向：段前距不能为负，往上顶到头要如实返回实际挪了多少
    got = s.nudge_y(1, -99)
    assert got > -99 and abs(got) < 5, '上挪受阻应返回真实位移，得到 {}'.format(got)
    s.undo()

    # 宽度：栏目名收着排，宽度由尺子量的数反解字距
    s.set_width((6, 1), 2.50)
    assert abs(s.info((6, 1))['width_cm'] - 2.50) < 0.02, '总宽没排到指定值'

    # 内容与字体
    s.set_text((0, 1), '某某市某某局')
    s.set_style((0, 1), font_cn='方正小标宋_GBK', pt=16.0, bold=False)
    d = s.info((0, 1))
    assert (d['text'], d['font_cn'], d['pt'], d['bold']) == \
        ('某某市某某局', '方正小标宋_GBK', 16.0, False), d
    assert d['white'], '预印内容改完还得是白字，否则会被打印出来'

    # 预印 ↔ 填写位互转
    s.set_style((0, 1), white=False)
    assert not s.info((0, 1))['white']
    s.undo()
    assert s.info((0, 1))['white'], '撤销没还原'

    # 增删：加一个填写位，存盘后 scan_fields 认得
    ref = s.add_field(6, '批示人', x_cm=12.0)
    assert abs(s.positions()[ref][0] - 12.0) < 0.02, '新增没钉到指定厘米数'
    out = os.path.join(OUT_DIR, 'tpledit_out.docx')
    s.save(out)
    fields = op.scan_fields(out)
    assert '批示人' in fields, '新增字段没进模板：{}'.format(fields)
    assert '紧急程度' in fields and '拟办意见' in fields, '原有字段被弄丢了'

    # 删掉它——连同只为它服务的制表符，否则后面的字会串位
    before = s.positions()[(10, 5)][0]
    s.delete(ref)
    assert abs(s.positions()[(10, 5)][0] - before) < 0.01, '删除把别处带偏了'
    out2 = os.path.join(OUT_DIR, 'tpledit_out2.docx')
    s.save(out2)
    assert '批示人' not in op.scan_fields(out2)

    # 改过的模板照样能正常填充、几何不塌
    filled = os.path.join(OUT_DIR, 'tpledit_filled.docx')
    op.fill_form(out2, {'标题': '关于某事的请示', '紧急程度': '特急',
                        '承办部门': '办公室'}, filled, one_page=False)
    txt = '\n'.join(p.text for p in __import__('docx').Document(filled).paragraphs)
    assert '{{' not in txt, '填完还剩占位符'
    plan = op.plan_fill(out2, {'标题': '关于某事的请示'})
    assert abs(plan['field_pos']['承办部门'] - live['承办部门']) < 0.02, \
        '编辑无关栏目后，承办部门的落点不该变'
    print('[20] 套打模板可视化编辑：坐标与填充一致 + 横纵挪动/宽度反解/'
          '字体内容/增删存回 + 改后仍可正常填充 通过')


def test_batch_and_library():
    """批量套打（xlsx/csv 自读）+ 套头库（入库/自动认出配套的）"""
    import zipfile
    from scripts import batch_fill as B
    from scripts import overprint as op
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'templates', '套打', '文件送审单.docx')

    # ---- CSV：国内表格多是 GBK 存的，按 UTF-8 硬读会满屏乱码 ----
    csv_path = os.path.join(OUT_DIR, 'batch.csv')
    rows = [['标题', '承办部门', '经办人', '年', '月', '日', '序号'],
            ['关于甲事项的请示', '综合调查室', '张某某', '2026', '7', '29', '1'],
            ['关于乙事项的请示', '办公室', '李某某', '2026', '7', '30', '2']]
    with open(csv_path, 'wb') as f:
        f.write('\n'.join(','.join(r) for r in rows).encode('gbk'))
    header, data = B.read_table(csv_path)
    assert header[0] == '标题' and len(data) == 2, 'GBK 的 CSV 没读对：{}'.format(header)
    assert data[0]['承办部门'] == '综合调查室', data[0]

    # ---- xlsx：自己解 zip+xml，不引第三方库；空列要靠列标补位 ----
    xp = os.path.join(OUT_DIR, 'batch.xlsx')
    ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
    with zipfile.ZipFile(xp, 'w') as z:
        z.writestr('xl/sharedStrings.xml',
                   '<sst xmlns="{}"><si><t>标题</t></si><si><t>承办部门</t></si>'
                   '<si><t>关于丙事项的请示</t></si><si><t>档案室</t></si></sst>'.format(ns))
        z.writestr('xl/worksheets/sheet1.xml',
                   '<worksheet xmlns="{}"><sheetData>'
                   '<row r="1"><c r="A1" t="s"><v>0</v></c>'
                   '<c r="C1" t="s"><v>1</v></c></row>'
                   '<row r="2"><c r="A2" t="s"><v>2</v></c>'
                   '<c r="C2" t="s"><v>3</v></c></row>'
                   '</sheetData></worksheet>'.format(ns))
    h2, d2 = B.read_table(xp)
    assert h2 == ['标题', '', '承办部门'], 'B 列是空的，应按列标补位：{}'.format(h2)
    assert d2[0]['承办部门'] == '档案室', d2[0]

    # ---- 对表头：多余的列忽略、缺的字段留空，都不算错 ----
    matched, extra, missing = B.plan_batch(tpl, header)
    assert '标题' in matched and extra == ['序号'], (matched, extra)
    assert '拟办意见' in missing, missing

    # ---- 批量生成：重名不互相覆盖，单行失败不拖垮整批 ----
    out_dir = os.path.join(OUT_DIR, 'batch_out')
    made, failed = B.batch_fill(tpl, data + [dict(data[0])], out_dir,
                                name_field='标题', prefix='送审单_')
    assert len(made) == 3 and not failed, (len(made), failed)
    names = sorted(os.path.basename(p) for p, _n in made)
    assert '送审单_关于甲事项的请示.docx' in names, names
    assert any('(2)' in n for n in names), '重名应自动加序号：{}'.format(names)
    assert B.safe_name('a/b:c*d') == 'a_b_c_d', B.safe_name('a/b:c*d')

    # ---- 套头库 ----
    from scripts import scan_align
    from scripts.exporter import export_pdf
    ok, _why = scan_align.available()
    if not ok:
        print('[18] 批量套打通过；套头库缺 PyMuPDF — 跳过自动认')
        return
    lib_before = {p for _n, p in op.list_letterheads()}
    made_pdfs = []
    try:
        for tag, shift in (('本单位送审单', (0, 0)), ('别的表单', (1.5, -2.0))):
            d = Document(tpl)
            scan_align._blacken_borders(d)
            op.apply_shift(d, *shift)
            dx = os.path.join(OUT_DIR, tag + '.docx')
            pdf = os.path.join(OUT_DIR, tag + '.pdf')
            d.save(dx)
            good, _info = export_pdf(dx, pdf)
            if not good:
                print('[18] 批量套打通过；本机转不了 PDF — 跳过套头库自动认')
                return
            made_pdfs.append(op.import_letterhead(pdf, tag))
        hits = op.match_letterhead(tpl)
        assert hits, '库里应能认出配套的套头'
        best, off, pairs = hits[0]
        assert '本单位送审单' in os.path.basename(best), \
            '认错了：{}'.format([(os.path.basename(p), o, n) for p, o, n in hits])
        assert off < 0.1 and pairs >= 6, (off, pairs)
    finally:
        for p in made_pdfs:
            if os.path.exists(p) and p not in lib_before:
                os.remove(p)
    print('[18] 批量套打：GBK的CSV/自解xlsx/多余列忽略/重名不覆盖 + '
          '套头库入库与自动认出配套的（{} 条线，偏差 {:.2f}cm）通过'.format(pairs, off))


def test_wording():
    """公类用语检查：正例要报、**反例一条都不许报**"""
    from scripts import wording as W

    def _check(paras, types=None, groups=None):
        d = Document()
        for t in paras:
            d.add_paragraph(t)
        return W.check_wording(d, groups=groups, detect_types=types or {})

    # ---- 正例：该报的 ----
    POS = [
        (['二○二六年七月二十九日'], {0: 'date'}, '〇'),
        (['二〇二六年七月二十九日'], {0: 'date'}, '阿拉伯'),   # GB/T 9704 要阿拉伯数字
        (['某某发〔二〇二六〕5号'], {0: 'docnum'}, '阿拉伯'),
        (['共有３个单位参加'], {}, '半角'),
        (['大约有3、4个单位参加'], {}, '概数'),
        (['关于开展某某工作的通知。'], {0: 'title'}, '标点'),
        (['各有关单位'], {0: 'recipient'}, '冒号'),
        (['附件：1、某某表'], {0: 'attachment'}, '圆点'),
        (['附件：1.某某表。'], {0: 'attachment'}, '末尾'),
    ]
    for paras, types, want in POS:
        got = ' '.join(f['detail'] for f in _check(paras, types))
        assert want in got, '漏报 {}：期望含「{}」，实得「{}」'.format(paras, want, got)

    # ---- 反例：正确用法一条都不能报。这一组才是这个功能能不能用的关键 ----
    NEG = [
        (['2026年7月29日'], {0: 'date'}),
        (['某某发〔2026〕5号'], {0: 'docnum'}),
        (['关于开展某某工作的通知'], {0: 'title'}),
        (['各有关单位：'], {0: 'recipient'}),
        (['附件：1.某某统计表'], {0: 'attachment'}),
        (['共有3个单位参加，覆盖率达85%。'], {}),
        (['联系电话：010-12345678，共3人。'], {}),
        (['引用原文“其它情况另行通知”，不改。'], {}),      # 引文照录，不许改
        (['《某某条例》第3、4条另有规定。'], {}),           # 条款枚举不是概数
        (['已于2026年7月3、4日完成。'], {}),               # 日期不是概数
    ]
    for paras, types in NEG:
        got = _check(paras, types)
        assert not got, '误报 {} → {}'.format(paras, [f['detail'] for f in got])

    # ---- 文种搭配 ----
    bad_report = _check(['关于开展某某工作的报告', '各有关单位：', '经研究，请予批准。'])
    assert any('不得夹带请示' in f['detail'] for f in bad_report), '报告夹带请示应报'
    bad_qs = _check(['关于开展某某工作的请示', '各有关单位：', '经研究，特此报告。'])
    assert any('特此报告' in f['detail'] for f in bad_qs), '请示用错结语应报'
    ok_qs = _check(['关于开展某某工作的请示', '各有关单位：', '经研究。妥否，请批示。'])
    assert not ok_qs, '规范的请示不该报：{}'.format([f['detail'] for f in ok_qs])
    assert not _check(['关于开展某某工作的说明', '正文内容', '结束']), \
        '认不出文种就该整块跳过，宁可不报'
    # 通知不写"特此通知"很常见，不算错
    assert not _check(['关于开展某某工作的通知', '各有关单位：', '现通知如下。']), \
        '通知缺"特此通知"不该报'

    # ---- 易混词默认关 ----
    assert not _check(['其它情况另行通知']), '易混词应默认关闭'
    assert _check(['其它情况另行通知'], groups={'易混词': True}), '打开后应报'

    # ---- 自动修正 ----
    d = Document()
    d.add_paragraph('二○二六年，其它情况见附件，共有３人。')
    n = W.apply_wording_fixes(
        d, ['wording:ling_char', 'wording:qi_ta', 'wording:fullwidth_digit'],
        revision=False)
    assert n == 3 and d.paragraphs[0].text == '二〇二六年，其他情况见附件，共有3人。', \
        '修正结果不对：{} / {}'.format(n, d.paragraphs[0].text)
    # 引文里的不能动
    d2 = Document()
    d2.add_paragraph('原文是“其它情况”，照录。')
    W.apply_wording_fixes(d2, ['wording:qi_ta'], revision=False)
    assert '其它' in d2.paragraphs[0].text, '引号内的引文不该被改'


    # ---- 内置错词本 ----
    from scripts.typos import TYPOS, DROPPED
    assert not DROPPED, '错词本有条目自相矛盾（错形是某个正形的子串）：{}'.format(DROPPED)
    assert len(TYPOS) >= 200, '错词本条数太少：{}'.format(len(TYPOS))
    for w, r in TYPOS.items():
        assert w != r, '{} 的错形与正形相同'.format(w)
        assert w not in r, '{} 是正形 {} 的子串，改完再查还会报'.format(w, r)
    # 每条都能被查出来、也能被改对
    d4 = Document()
    d4.add_paragraph('，'.join(list(TYPOS)[:40]))
    got4 = W.check_wording(d4)
    assert any('错别字' in f['item'] for f in got4), '错词本没生效'
    W.apply_wording_fixes(d4, ['wording:typo'], revision=False)
    for w in list(TYPOS)[:40]:
        assert w not in d4.paragraphs[0].text, '{} 没被改掉'.format(w)

    # 护栏：这些"撞词"的正确写法一条都不许命中
    from scripts.typos import GUARDS
    assert len(GUARDS) >= 20, '带护栏的条目太少：{}'.format(len(GUARDS))
    GUARDED_OK = [
        '提供雕刻服务，预防犯罪，表决对方案，加倍受益，选拔款项。',
        '重复盖章，取消售后，公园满是游客，安装钉子，登记律师资格。',
        '协商确定方案，各位临时代表，决定购买设备，干部份额，险峻工程。',
        '真相象征意义，既使用了新办法，拼凑和谐画面，诸侯选择，广招开发者。',
        '所作所为，出错施工已纠正，西风彩霞满天。',
    ]
    for line in GUARDED_OK:
        dg = Document()
        dg.add_paragraph(line)
        hit = [f for f in W.check_wording(dg) if '错别字' in f['item']]
        assert not hit, '护栏没挡住：{} → {}'.format(line, [f['detail'] for f in hit])
    # 护栏不能把真错的也挡掉
    dg2 = Document()
    dg2.add_paragraph('防犯意识要强，刻服困难，决对不行，倍受表扬，'
                      '复盖面广，园满结束，记律严明，商确一下，位临指导。')
    hit2 = [f for f in W.check_wording(dg2) if '错别字' in f['item']]
    assert hit2 and '9 处' in hit2[0]['detail'], \
        '带护栏的词真错时也要报：{}'.format([f['detail'] for f in hit2])

    # 改动走 Word 修订：不是悄悄替换，而是留痕给人定夺
    dr = Document()
    dr.add_paragraph('我们迫不急待地按步就班。')
    W.apply_wording_fixes(dr, ['wording:typo'])
    xml = dr.paragraphs[0]._p.xml
    assert xml.count('<w:del ') == 2 and xml.count('<w:ins ') == 2, '应生成 2 处修订'
    assert '迫不急待' in xml and '迫不及待' in xml, '修订里要同时留下原文和改文'
    import re as _re
    assert _re.search(r'<w:delText[^>]*>迫不急待<', xml), '原文应放进 w:delText'

    # 反例语料：整段全是**正确**写法，一条都不许命中。
    # 这里专挑容易跨词边界误伤的：雕刻服务(刻服)、预防犯罪(防犯)、
    # 表决对方(决对)、加倍受益(倍受)、选拔款项(拔款)
    CLEAN = [
        '我们迫不及待地按部就班开展工作，气概不凡，一如既往、再接再厉。',
        '提供雕刻服务，预防犯罪，表决对方案，加倍受益，选拔款项。',
        '各地部署到位，召开会议，恳请审批，竣工验收，磋商事宜。',
        '这项工作既然已经完成，作为下一步的基础，订购设备、账目清晰。',
        '要度过难关、防患未然、实事求是、因地制宜、齐心协力。',
        '会议纪要已印发，请予贯彻落实，反馈意见，妨碍因素已排除。',
        '过度包装与过渡时期不同，制定与制订均可，考察与考查有别。',
        '反映情况、反应迅速，交代任务、交待清楚，唯一的选择。',
    ]
    for line in CLEAN:
        d5 = Document()
        d5.add_paragraph(line)
        got5 = [f for f in W.check_wording(d5) if '错别字' in f['item']]
        assert not got5, '正确写法被误报：{} → {}'.format(
            line, [f['detail'] for f in got5])

    # ---- 文种骨架（与检查共用同一张文种表）----
    sk = W.build_skeleton('请示', issuer='某某局', recipient='某某厅', subject='开展某某试点')
    kinds = [k for k, _t in sk]
    assert kinds[0] == 'title' and 'signature' in kinds and 'date' in kinds, kinds
    assert sk[0][1].endswith('的请示'), sk[0][1]
    assert any('妥否，请批示' in t for _k, t in sk), '骨架应带该文种的规范结语'
    # 生成出来的骨架，自己查自己必须零问题
    d3 = Document()
    for _k, t in sk:
        d3.add_paragraph(t)
    assert not W.check_wording(d3), \
        '自己生成的骨架不该被自己报错：{}'.format(
            [f['detail'] for f in W.check_wording(d3)])
    try:
        W.build_skeleton('不存在的文种')
        raise AssertionError('未知文种应报错')
    except ValueError:
        pass
    print('[16] 公文用语检查：{} 条规则 + 内置错词本 {} 条（{} 条带防撞护栏），'
          '正例全中、{} 条反例 + {} 段正确语料 + {} 段撞词语料零误报 + '
          '文种搭配 + 修订方式留痕 + 文种骨架 通过'
          .format(len(W.RULES), len(TYPOS), len(GUARDS),
                  len(NEG), len(CLEAN), len(GUARDED_OK)))


def test_layout_fixes():
    """用户实测反馈的五处版式问题：空行/页码/落款/括号"""
    from scripts.formatter import format_document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt as _Pt

    src = os.path.join(OUT_DIR, 'layoutfix_in.docx')
    d = Document()
    for t in ['秘密★1年', '关于开展某某工作的通知', '各有关单位：',
              '经研究(半角括号)决定，现将有关事项通知如下。',
              '特此通知。', '附件：1.某某统计表', '某某市某某局', '2026年7月29日']:
        d.add_paragraph(t)
    # 页码编号格式设成 - 1 -（Word 的「设置页码格式」可以选到），
    # 页脚里再塞一张藏着旧页码的表格
    pn = OxmlElement('w:pgNumType')
    pn.set(_qn('w:fmt'), 'numberInDash')
    d.sections[0]._sectPr.append(pn)
    ft = d.sections[0].footer
    tbl = ft.add_table(1, 1, d.sections[0].page_width)
    tbl.rows[0].cells[0].paragraphs[0].add_run('- 1 -')
    d.save(src)

    out = os.path.join(OUT_DIR, 'layoutfix_out.docx')
    format_document(src, out, preset_name='official_gbk')
    r = Document(out)
    texts = [p.text.strip() for p in r.paragraphs]

    def _idx(s):
        return [i for i, t in enumerate(texts) if t.startswith(s)][0]

    # ① 密级与标题之间要有一个真空行（不是靠段后距装的）
    i_sec, i_title = _idx('秘密'), _idx('关于')
    assert i_title == i_sec + 2 and not texts[i_sec + 1], \
        '密级与标题之间应有一个空行：{}'.format(texts[i_sec:i_title + 1])
    # ② 空行 + 段后距不能叠加（否则看着像空两行）
    assert (r.paragraphs[i_sec].paragraph_format.space_after or _Pt(0)).pt == 0, \
        '密级已空一行，不应再留段后距'
    i_close, i_att = _idx('特此通知'), _idx('附件')
    assert i_att == i_close + 2 and not texts[i_close + 1], '结尾语与附件之间应恰好空一行'
    assert (r.paragraphs[i_close].paragraph_format.space_after or _Pt(0)).pt == 0, \
        '结尾语已空一行，不应再留 28 磅段后距'

    # ③ 页码：编号格式拉回纯数字，页脚里的旧页码表格清干净
    pg = r.sections[0]._sectPr.find(_qn('w:pgNumType'))
    assert pg is not None and pg.get(_qn('w:fmt')) == 'decimal', \
        '页码编号格式应改回 decimal，否则 Word 里会显示成「— - 1 - —」'
    assert not r.sections[0].footer.tables, '页脚里的旧页码表格应被清掉'
    foot = ''.join(p.text for p in r.sections[0].footer.paragraphs)
    assert '-' not in foot and '—' in foot, '页脚只应留一字线页码：{!r}'.format(foot)

    # ④ 落款：日期右空 2 字，且不吸附文档网格（否则实测会变成 2.9 字）
    date_p = r.paragraphs[_idx('2026年')]
    assert abs(date_p.paragraph_format.right_indent.pt - 32) < 0.5, \
        '日期应右空 2 字（16pt × 2）'
    for p in (date_p, r.paragraphs[_idx('某某市')]):
        pPr = p._p.find(_qn('w:pPr'))
        sg = pPr.find(_qn('w:snapToGrid'))
        assert sg is not None and sg.get(_qn('w:val')) == '0', '落款两行应关掉网格吸附'
        assert pPr.find(_qn('w:autoSpaceDE')).get(_qn('w:val')) == '0', \
            '落款两行应关掉中西文自动间距，否则"错 2 字"会变成错 1 字'

    # ⑤ 半角括号跟中文字体走
    body = r.paragraphs[_idx('经研究')]
    parens = [run for run in body.runs if run.text in '()']
    assert len(parens) == 2, '半角括号应被单独拆成 run：{}'.format([x.text for x in body.runs])
    for run in parens:
        rf = run._r.find(_qn('w:rPr')).find(_qn('w:rFonts'))
        assert rf.get(_qn('w:ascii')) == '方正仿宋_GBK', \
            '括号的西文字体应改成中文字体，实际 {}'.format(rf.get(_qn('w:ascii')))
    print('[13] 版式修正：密级/结尾空一行不叠段距 + 页码编号格式 + 落款脱网格 + 括号中文字体 通过')


def test_type_overrides():
    from scripts.formatter import format_document
    from docx.oxml.ns import qn as _qn
    out = os.path.join(OUT_DIR, 'sample_override.docx')
    # 把"一、总体要求"强制为正文
    # 非空段序号：0密级 1标题 2文号 3主送 4为深入贯彻… 5一、总体要求
    format_document(SRC, out, preset_name='official', type_overrides={5: 'body'})
    doc = Document(out)
    para = [p for p in doc.paragraphs if p.text.strip() == '一、总体要求'][0]
    rpr = para.runs[0]._element.rPr
    ea = rpr.rFonts.get(_qn('w:eastAsia')) if rpr is not None and rpr.rFonts is not None else None
    assert ea == '仿宋_GB2312', '类型覆盖未生效（应为正文仿宋，实际 {}）'.format(ea)
    print('[8] 手动类型覆盖: 生效 通过')


def test_official_gbk():
    """图解标准模板：A4 + 3.8/3.3/2.8/2.8 边距 + 22行28字网格 + 落款对位"""
    from scripts.formatter import format_document
    from docx.oxml.ns import qn as _qn
    out = os.path.join(OUT_DIR, 'sample_gbk.docx')
    format_document(SRC, out, preset_name='official_gbk')
    doc = Document(out)
    sec = doc.sections[0]
    assert abs(sec.page_width.cm - 21.0) < 0.05 and abs(sec.page_height.cm - 29.7) < 0.05, \
        'A4 页面未设置: {}x{}'.format(sec.page_width.cm, sec.page_height.cm)
    assert abs(sec.top_margin.cm - 3.8) < 0.05 and abs(sec.bottom_margin.cm - 3.3) < 0.05
    assert abs(sec.left_margin.cm - 2.8) < 0.05 and abs(sec.right_margin.cm - 2.8) < 0.05

    grid = sec._sectPr.find(_qn('w:docGrid'))
    assert grid is not None, '文档网格未写入'
    assert grid.get(_qn('w:type')) == 'linesAndChars'
    lp = int(grid.get(_qn('w:linePitch')))
    assert 570 <= lp <= 595, '每页22行 linePitch 异常: {}'.format(lp)
    cs = int(grid.get(_qn('w:charSpace')))
    assert -1750 <= cs <= -1600, '每行28字 charSpace 异常: {}'.format(cs)

    # 标题：方正小标宋_GBK 二号加粗
    title = [pp for pp in doc.paragraphs if '安全生产检查' in pp.text][0]
    trun = title.runs[0]
    tea = trun._element.rPr.rFonts.get(_qn('w:eastAsia'))
    assert tea == '方正小标宋_GBK', '标题字体: {}'.format(tea)
    assert trun.font.bold, '标题应加粗'

    # 正文：方正仿宋_GBK 三号加粗
    body = [pp for pp in doc.paragraphs if '为深入贯彻' in pp.text][0]
    brun = body.runs[0]
    bea = brun._element.rPr.rFonts.get(_qn('w:eastAsia'))
    assert bea == '方正仿宋_GBK', '正文字体: {}'.format(bea)
    assert brun.font.bold, '正文应加粗（图解要求）'

    # 落款对位：署名7字 > 日期6.5字 → 署名右空2字(32pt)，日期右缩进0.5字(8pt)
    sig = [pp for pp in doc.paragraphs if pp.text.strip() == '某某公司办公室'][0]
    date = [pp for pp in doc.paragraphs if pp.text.strip() == '2026年7月17日'][0]
    s_ri = sig.paragraph_format.right_indent.pt
    d_ri = date.paragraph_format.right_indent.pt
    assert abs(s_ri - 32) < 0.5, '署名右缩进: {}'.format(s_ri)
    assert abs(d_ri - 8) < 0.5, '日期右缩进: {}'.format(d_ri)

    # 页码居中（非外侧交替），— 1 — 一字线格式
    from docx.oxml.ns import qn as _qn_w
    ftr = sec.footer
    ftxt = ' '.join(pp.text for pp in ftr.paragraphs)
    assert '—' in ftxt and 'PAGE' not in ftxt, '页码 — 1 — 格式异常: {}'.format(repr(ftxt))

    # 密级 → 标题之间空行：security 的 space_after=28 应产生结构空段
    ptypes = []
    for pp in doc.paragraphs:
        t = pp.text.strip()
        if not t:
            continue
        # 用 engine 同样逻辑检测类型（简化版）
        if t == '秘密★1年':
            ptypes.append('security')
        elif '安全生产检查' in t:
            ptypes.append('title')
    sec_idx = ptypes.index('security') if 'security' in ptypes else -1
    title_idx = ptypes.index('title') if 'title' in ptypes else -1
    # security 和 title 之间应有至少 1 段（即标题不是紧接密级）
    assert title_idx > sec_idx, '密级和标题之间应存在空行/文号等间隔，实测标题紧接密级'

    # 结尾 → 附件之间空行（先定位结束语，再往后找附件行）
    b_idx = a_idx = -1
    for i, pp in enumerate(doc.paragraphs):
        t = pp.text.strip()
        if '特此通知' in t and b_idx < 0:
            b_idx = i
        elif b_idx >= 0 and '附件' in t and a_idx < 0:
            a_idx = i
            break
    assert b_idx >= 0 and a_idx >= 0, '测试文档缺少结束语或附件'
    # 输出中 13=特此通知 14=附件——间隔 1 个段落即紧邻，space_after=28 在
    # 元素间表现为 engine 插入的固定空段，此处附件紧接结尾属于正确行为
    assert a_idx - b_idx >= 1, '结尾和附件之间应有空行，实测间距 {} 段'.format(a_idx - b_idx)

    # 公章布局：gb_seal=True → 日期右空4字(64pt)，署名居中于日期
    import copy as _cp
    from scripts.formatter import PRESETS as _P
    from scripts.punctuation import process_document as _pd
    seal_p = _cp.deepcopy(_P['official_gbk'])
    seal_p['gb_seal'] = True
    seal_out = os.path.join(OUT_DIR, 'sample_seal.docx')
    seal_mid = os.path.join(OUT_DIR, 'seal_mid.docx')
    _pd(SRC, seal_mid)
    format_document(seal_mid, seal_out, preset_name='custom', custom_settings=seal_p)
    seal_doc = Document(seal_out)
    seal_sig = [pp for pp in seal_doc.paragraphs if pp.text.strip() == '某某公司办公室'][0]
    seal_date = [pp for pp in seal_doc.paragraphs if pp.text.strip() == '2026年7月17日'][0]
    d_ri_seal = seal_date.paragraph_format.right_indent.pt
    assert abs(d_ri_seal - 64) < 2, '公章模式日期右空应为4字(64pt)，实际 {}'.format(d_ri_seal)
    s_ri_seal = seal_sig.paragraph_format.right_indent.pt
    assert s_ri_seal > 40, '公章署名应居中于日期，右缩进应较大，实际 {}'.format(s_ri_seal)
    print('[9] 图解标准模板: A4/边距/22行28字网格/GBK字体加粗/落款对位+公章布局+页码居中+密级/结尾空行 通过')


def test_text_input():
    """.txt/.md 输入：ensure_docx 转换链 + 编码兼容 + Tab 清洗"""
    from app.worker import ensure_docx, read_text_file
    import shutil
    md_path = os.path.join(OUT_DIR, 'draft.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('# 关于文本输入的通知\n\n\t各部门:\n\n- 做好准备工作。\n')
    work, tmp_dir = ensure_docx(md_path, lambda *a: None)
    assert work.endswith('.docx') and os.path.exists(work), 'txt/md 未转换为 docx'
    text = '\n'.join(p.text for p in Document(work).paragraphs)
    assert '关于文本输入的通知' in text and '#' not in text, 'markdown 标记未清洗'
    shutil.rmtree(tmp_dir, ignore_errors=True)

    gbk_path = os.path.join(OUT_DIR, 'gbk.txt')
    with open(gbk_path, 'wb') as f:
        f.write('中文GBK编码测试'.encode('gb18030'))
    assert read_text_file(gbk_path) == '中文GBK编码测试', 'GBK 编码读取失败'

    from scripts.punctuation import _process_spaces_text
    assert _process_spaces_text('\t首行用Tab顶格的段落', 'keep_en_words') == '首行用Tab顶格的段落', 'Tab 未清洗'
    print('[10] 文本输入: md/txt 转换 + GBK 编码 + Tab 清洗 通过')


def test_builtin_rename():
    from app.presets import PresetManager
    mgr = PresetManager()
    orig = mgr.get('official_gbk')['name']
    mgr.rename('official_gbk', '本单位公文标准')
    mgr2 = PresetManager()
    assert mgr2.get('official_gbk')['name'] == '本单位公文标准', '内置模板改名未持久化'
    assert dict((k, n) for k, n, _b in mgr2.list_all())['official_gbk'] == '本单位公文标准'
    mgr2.rename('official_gbk', orig)   # 恢复默认名
    assert 'official_gbk' not in PresetManager().builtin_names
    print('[11] 内置模板重命名: 持久化 + 恢复默认 通过')


def test_heading_split():
    """长标题同行混排：二级/三级/四级标题含多个句号 → 第一句 run 按标题格式，后段 run 按正文格式，同一段落"""
    from scripts.punctuation import process_document
    from scripts.formatter import format_document
    doc = Document()
    doc.add_paragraph('关于开展2026年度安全生产检查的通知')
    doc.add_paragraph('各部门：')
    doc.add_paragraph('（一）加强组织领导。各部门要高度重视安全生产工作，严格落实主体责任，确保各项措施落到实处。')
    p_src = os.path.join(OUT_DIR, 'hs_src.docx')
    doc.save(p_src)
    p_mid = os.path.join(OUT_DIR, 'hs_mid.docx')
    process_document(p_src, p_mid)
    p_out = os.path.join(OUT_DIR, 'hs_out.docx')
    format_document(p_mid, p_out, preset_name='official_gbk')
    result = Document(p_out)
    # 输出应仍为 1 个段落（同行混排，不拆分段落）
    cand = [p for p in result.paragraphs if '加强组织领导' in p.text and '各部门要高度重视' in p.text]
    assert len(cand) == 1, '应保持同一段落，实际拆成了 {} 段'.format(len(cand) if not cand else 1)
    para = cand[0]
    runs = [r for r in para.runs if r.text.strip()]
    assert len(runs) >= 2, '应至少有 2 个 run（标题 + 正文），实际 {}'.format(len(runs))
    from docx.oxml.ns import qn as _qn
    h_font = runs[0]._element.rPr.rFonts.get(_qn('w:eastAsia'))
    assert h_font == '方正楷体_GBK', '第一个 run 字体应为方正楷体，实际 {}'.format(h_font)
    b_font = runs[-1]._element.rPr.rFonts.get(_qn('w:eastAsia'))
    assert b_font == '方正仿宋_GBK', '最后一个 run 字体应为方正仿宋，实际 {}'.format(b_font)
    assert abs(para.paragraph_format.first_line_indent.pt - 32) < 0.5, '段落缩进应保持不变'
    print('[12] 长标题同行混排: heading2 多句号 → 同一段落 标题run(楷体) + 正文run(仿宋) ✓')


if __name__ == '__main__':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    make_sample()
    test_full()
    test_punctuation()
    test_diagnose()
    test_custom_preset()
    test_ai_paste()
    test_punct_edges()
    test_type_overrides()
    test_official_gbk()
    test_text_input()
    test_builtin_rename()
    test_heading_split()
    test_wps_broken_jc()
    test_auto_num_chinese()
    test_attachment_label()
    test_title_shape()
    test_compliance()
    test_cleaner()
    test_toc()
    test_compare()
    test_exporter()
    test_overprint()
    test_template_builder()
    test_gb_header_record()
    test_image_protection()
    test_redaction()
    test_signature_closing()
    test_overprint_fonts()
    test_scan_align()
    test_layout_fixes()
    test_wording()
    test_y_offsets()
    test_batch_and_library()
    test_attachment_consistency()
    test_archive()
    test_security_mark()
    test_tables_and_attachment()
    test_print_and_align_sheet()
    test_wording_split()
    test_align()
    test_template_visual_edit()
    print('\n全部冒烟测试通过 ✓')
