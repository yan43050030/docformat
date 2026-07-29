# -*- coding: utf-8 -*-
"""套打（套头打印）：把内容精确打到预印红头纸的空白位置上。

套打的实现机制
--------------
预印在纸上的红色内容（单位名、栏目名、表格红线）在 docx 里写成
**白色文字与白色边框**——占住完全相同的位置、保证版式分毫不差，
但白纸上打印时不显影；只有**黑色文字**才真正印到预印纸上。

因此套打模板 = 一份保留全部白色占位与几何的 docx，其中要填的位置
放 {{字段名}} 占位符。填充时只替换占位符文本、不动任何几何，
套准度就由模板本身保证。

内容自适应
----------
预留格子高度固定（trHeight 为 exact 时更严格）。正文过长会把行撑高、
把整张表挤下去，套打就错位了。故填充长文本时按可用高度估算所需行数，
不够就逐档缩小字号与行距，直到放得下（有下限，缩到底仍放不下会如实告警）。
"""
import logging
import os
import re

from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger('docformat.overprint')

PLACEHOLDER_RE = re.compile(r'\{\{\s*([^}]+?)\s*\}\}')

# 自适应下限：再小就影响阅读，宁可告警也不继续缩
MIN_FONT_PT = 9.0
MIN_LINE_SPACING_PT = 12.0
# 每档缩小的幅度
FONT_STEP = 0.5

TWIPS_PER_CM = 566.93
PT_PER_CM = 28.3465

# 裁"末尾不打印的占位空白"时的目标宽度（占名义版心宽的比例）。
#
# 为什么要留这么大余量：真机 Word 里成文日期行按"全角字宽=字号"算只有
# 16.05cm、版心 16.45cm，本该放得下，实际却折了行。原因出在空格——那一行的
# 空白 run 用的是 CJK 字体（方正楷体_GBK），空格未必是半角宽；本机没有这些
# 字体，无从测准。而**表格里**的正文实测是准的（拟办意见一行 Word 排 28 字、
# 预览 27 字），所以不能为了这一行给全局加系数、把格子里的字号无谓缩小。
#
# 于是只在这里留余量：被裁的是末尾纯占位空白，本来就不显影，裁多了零代价，
# 裁少了要赔一整页。0.75 的目标即使按"空格全是全角"的最悲观算法也放得下。
TAIL_TRIM_RATIO = 0.75


def _iter_cells(table):
    """产出表格里每个不重复的单元格。

    不能用 id(cell._tc) 去重——lxml 代理对象被回收后 id 会复用，
    不同单元格会被误判为同一个而漏处理。用保持引用的身份比较。
    """
    seen = []
    for row in table.rows:
        for cell in row.cells:
            if any(cell._tc is x for x in seen):
                continue
            seen.append(cell._tc)
            yield cell


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p, None
    for t in doc.tables:
        for cell in _iter_cells(t):
            for p in cell.paragraphs:
                yield p, cell


def scan_fields(template_path):
    """扫描模板里的 {{字段}}，按出现顺序返回去重后的字段名列表。"""
    from docx import Document
    doc = Document(template_path)
    names = []
    for p, _cell in _iter_paragraphs(doc):
        for m in PLACEHOLDER_RE.finditer(p.text):
            name = m.group(1).strip()
            if name not in names:
                names.append(name)
    return names


# 定宽字段 → 该字段的槽位宽度（全角宽）。
# 套打里"年 月 日"三个字是**预印在纸上的固定位置**。若直接把占位符换成
# "7"，后面的"月"字就跟着左移，与纸上预印的"月"对不上；一位数和两位数
# 的月份还会落在不同位置。把值右对齐补空格到固定槽宽，数字紧贴其后的
# 年/月/日，无论几位数预印字符都纹丝不动，留空待手签时空白也原样保留。
#
# 槽宽按数字本身该占的宽度定（年 4 个半角、月/日各 2 个半角），
# **不是**按占位符 {{月}} 的字面宽度——那是标记的长度，不是纸上空白的
# 宽度，照它补会补出一大段空档（用户反馈"年月日距离远"即出于此）。
_FIXED_WIDTH_FIELDS = {'年': 2.0, '月': 1.0, '日': 1.0}

# 需要梯形回行的字段（公文标题回行要求词意完整、排列对称）
TITLE_FIELDS = ('标题', '题目')


def _pad_to_width(value, slot_units):
    """把值右对齐补空格到 slot_units 个全角宽（数字紧贴其后的年/月/日）。"""
    need = slot_units - _text_width_units(value)
    if need <= 0:
        return value
    return ' ' * int(round(need / 0.5)) + value


def _replace_in_paragraph(para, values, offsets=None, left_cm=0.0):
    """把段落里的占位符替换为取值，保留该 run 的字体/颜色/字号。

    同时记录每个字段最终落在**距纸张左边多少厘米**。offsets 指定目标位置
    时（offsets={'年': 2.5} = "年的数字从距纸左边 2.5cm 处开始印"），
    用**制表位**把值顶过去，而不是补空格。

    为什么必须用制表位而不是空格
    ----------------------------
    空格的宽度随字体变：实测 Times New Roman 里一个空格正好是数字的 50%，
    而 CJK 字体里往往接近一个全角宽。靠补空格定位，等于把位置押在
    "空格有多宽"这个未知数上——本项目实测过错位 0.49/0.69/0.96cm，
    且逐个累积。制表位的位置以缇为单位绝对指定，与字体无关，说 2.5cm
    就是 2.5cm。

    没有指定位置的字段仍走原来的定宽补空格（行为不变），
    这样没配过位置的模板与此前完全一致。

    返回 (是否改动, {字段: 起始位置cm})。
    """
    char_cm = _para_font_pt(para) / PT_PER_CM
    offsets = offsets or {}
    changed = False
    pos_map = {}
    tab_stops = []                  # 本段要加的制表位（cm，相对左边距）
    seen_ws = []                    # 本段里出现过的纯空白 run（填充用）
    acc = 0.0                       # 已排过的宽度（全角单位）
    # 只计"看得见的字"的宽度。判断目标位置够不够得着时用它，不用 acc：
    # 空格的实际宽度随字体变（实测 TNR 里只有数字的一半），拿含空格的
    # 估算去挡，会把明明够得着的位置误判成"顶不过去"而拒绝设置。
    acc_ink = 0.0
    for run in para.runs:
        if '{{' not in run.text:
            acc += _text_width_units(run.text)
            acc_ink += _text_width_units(run.text.replace(' ', ''))
            if run.text and not run.text.strip():
                seen_ws.append(run)
            continue
        out = []
        cur = 0
        for m in PLACEHOLDER_RE.finditer(run.text):
            lead = run.text[cur:m.start()]
            out.append(lead)
            acc += _text_width_units(lead)
            acc_ink += _text_width_units(lead.replace(' ', ''))
            key = m.group(1).strip()
            val = str(values.get(key, ''))
            target = offsets.get(key)
            if target is not None and char_cm > 0:
                # 该字段前面的填充空格一律撤掉——它们本就是"把字顶到位"的
                # 土办法，制表位一来就该让位。必须连**整个空白 run** 一起撤，
                # 只撤紧邻的那点不够：模板里常有几十个空格的长填充，而空格
                # 的实际宽度随字体变（实测 TNR 里只有数字的一半），留着它
                # 就可能把笔位顶过目标、制表位反而跳到下一站，结果差出一截
                # （实测目标 11.50cm 落到了 12.72cm）。
                # 这些空格和它们顶开的白字都不打印，撤掉不影响印出来的东西。
                while out and out[-1] and not out[-1].strip():
                    acc -= _text_width_units(out.pop())
                # 连本段前面所有纯空白 run 一起撤：它们只是把**白色栏目名**
                # 顶到位的填充，白字不显影，挪了不影响印出来的东西；留着却会
                # 让笔位越过目标、制表位跳到下一站。撤掉后位置才真正说了算。
                for _ws in seen_ws:
                    acc -= _text_width_units(_ws.text)
                    _ws.text = ''
                seen_ws = []
                rel = float(target) - left_cm
                if rel > acc_ink * char_cm:
                    # 制表位顶到绝对位置；顶不过去（目标在已排内容左边）
                    # 就不硬塞，由调用方据 pos_map 告警
                    tab_stops.append(rel)
                    out.append('\t')
                    acc = rel / char_cm
                pos_map[key] = left_cm + acc * char_cm
            else:
                if key in _FIXED_WIDTH_FIELDS:
                    val = _pad_to_width(val, _FIXED_WIDTH_FIELDS[key])
                # 记的是"值本身"的起点，前导空格不算——用户量的是数字的左沿
                lead_sp = len(val) - len(val.lstrip(' '))
                pos_map[key] = left_cm + (acc + lead_sp * 0.5) * char_cm
            out.append(val)
            acc += _text_width_units(val)
            acc_ink += _text_width_units(val.replace(' ', ''))
            cur = m.end()
        tail = run.text[cur:]
        out.append(tail)
        acc += _text_width_units(tail)
        acc_ink += _text_width_units(tail.replace(' ', ''))
        new = ''.join(out)
        if new != run.text:
            run.text = new
            changed = True
    if tab_stops:
        _set_tab_stops(para, tab_stops)
    return changed, pos_map


def _set_tab_stops(para, positions_cm):
    """给段落设置左对齐制表位（cm，相对左边距），升序去重。"""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm
    tabs = para.paragraph_format.tab_stops
    have = set()
    try:
        for t in tabs:
            have.add(round(t.position.cm, 2))
    except Exception:
        pass
    for pos in sorted(set(round(p, 2) for p in positions_cm)):
        if pos <= 0 or pos in have:
            continue
        try:
            tabs.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.LEFT)
        except Exception:
            logger.warning('制表位 %.2fcm 添加失败', pos)


# ---------------- 打印位置微调（随模板存盘） ----------------

def offsets_path(template_path):
    """位置微调表的存放路径：与模板同名的 .位置.json，跟着模板走"""
    base = os.path.splitext(template_path)[0]
    return base + '.位置.json'


def load_offsets(template_path):
    """读取该模板的位置微调表 {字段: 距纸左边cm}，没有就返回空表"""
    import json
    p = offsets_path(template_path)
    try:
        with open(p, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (IOError, OSError, ValueError):
        return {}
    out = {}
    for k, v in (data.get('fields') or data).items():
        try:
            out[k] = float(v)
        except (TypeError, ValueError):
            continue
    return out


def load_letterhead(template_path):
    """读取该模板绑定的套头纸 PDF 路径（用于对位校验），没有返回 ''"""
    import json
    try:
        with open(offsets_path(template_path), 'r', encoding='utf-8') as f:
            return str(json.load(f).get('套头PDF') or '')
    except (IOError, OSError, ValueError, AttributeError):
        return ''


def save_offsets(template_path, offsets, letterhead=None, shift=None):
    """写回位置微调表；offsets 与套头路径都空时删除文件（恢复默认）。

    与模板同名的一个 json 里放齐"对位相关的一切"：各字段的目标位置、
    绑定的套头纸 PDF。以后要做套头自动识别，读的也是这个文件。
    """
    import json
    p = offsets_path(template_path)
    if letterhead is None:
        letterhead = load_letterhead(template_path)
    if shift is None:
        shift = load_shift(template_path)
    dx, dy = float(shift[0] or 0.0), float(shift[1] or 0.0)
    if not offsets and not letterhead and not dx and not dy:
        try:
            os.remove(p)
        except OSError:
            pass
        return p
    payload = {
        '说明': '套打打印位置微调。数值 = 该字段第一个字距纸张左边缘的厘米数。'
                '用尺子量真实预印单上空格的左沿填进来即可；留空/删除本文件恢复默认。',
        '单位': 'cm（厘米，从纸张左边缘量起，含页边距）',
        'fields': {k: round(float(v), 2) for k, v in (offsets or {}).items()},
    }
    if letterhead:
        payload['套头PDF'] = letterhead
        payload['套头PDF说明'] = '套头纸（红头文件纸）的 PDF，用于不打印就校验对位'
    if dx or dy:
        payload['整体平移'] = {
            'dx': round(dx, 3), 'dy': round(dy, 3),
            '说明': '整张纸一起挪的厘米数，正数=往右/往下。'
                    '可由「按扫描件自动对位」量出来，也可自己填。',
        }
    with open(p, 'w', encoding='utf-8') as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return p


def save_letterhead(template_path, letterhead):
    """只改套头 PDF 绑定，保留已有的位置微调"""
    return save_offsets(template_path, load_offsets(template_path),
                        letterhead=letterhead or '',
                        shift=load_shift(template_path))


def load_shift(template_path):
    """整体平移量 (dx, dy)，单位 cm；没设过就是 (0, 0)。

    与逐字段的位置微调分开存：那个管"某个字要印在哪儿"，这个管
    "整张纸印偏了"——打印机走纸误差、套头纸裁切误差都是整体性的，
    一个数就能补，不必逐字段去挪。
    """
    import json
    try:
        with open(offsets_path(template_path), 'r', encoding='utf-8') as f:
            v = json.load(f).get('整体平移') or {}
        return float(v.get('dx') or 0.0), float(v.get('dy') or 0.0)
    except (IOError, OSError, ValueError, AttributeError, TypeError):
        return 0.0, 0.0


def save_shift(template_path, dx, dy):
    """只改整体平移，保留已有的位置微调与套头绑定"""
    return save_offsets(template_path, load_offsets(template_path),
                        letterhead=None, shift=(dx, dy))


def apply_shift(doc, dx_cm, dy_cm):
    """整张纸的内容一起挪：加到页边距上，版心大小不变。

    左边距一动，所有制表位（相对左边距）跟着动；上边距一动，整列内容
    跟着下移。比逐个字段改位置省事，也不会把已经对好的相对关系弄乱。
    """
    from docx.shared import Cm
    if not dx_cm and not dy_cm:
        return False
    for sec in doc.sections:
        if dx_cm:
            sec.left_margin = Cm(max(0.0, sec.left_margin.cm + dx_cm))
            sec.right_margin = Cm(max(0.0, sec.right_margin.cm - dx_cm))
        if dy_cm:
            sec.top_margin = Cm(max(0.0, sec.top_margin.cm + dy_cm))
            sec.bottom_margin = Cm(max(0.0, sec.bottom_margin.cm - dy_cm))
    return True


def _row_of_cell(table, cell):
    for row in table.rows:
        for c in row.cells:
            if c._tc is cell._tc:
                return row
    return None


def _row_height_cm(row):
    """返回 (高度cm, 是否为固定高度)；没设高度返回 (None, False)"""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        return None, False
    he = trPr.find(qn('w:trHeight'))
    if he is None:
        return None, False
    try:
        val = int(he.get(qn('w:val')))
    except (TypeError, ValueError):
        return None, False
    rule = he.get(qn('w:hRule')) or 'atLeast'
    return val / TWIPS_PER_CM, rule == 'exact'


def _grid_widths(table):
    grid = table._tbl.find(qn('w:tblGrid'))
    out = []
    if grid is not None:
        for gc in grid.findall(qn('w:gridCol')):
            try:
                out.append(int(gc.get(qn('w:w'))))
            except (TypeError, ValueError):
                out.append(0)
    return out


def _cell_width_cm(table, cell):
    """单元格实际宽度（cm），横向合并时按 gridSpan 累加所跨列。

    只读 tcW 会把合并单元格当成首列宽度，导致宽度被严重低估、
    进而误判"放不下"而缩字号——套打里这等于自己把版式弄歪。
    """
    tcPr = cell._tc.find(qn('w:tcPr'))
    span = 1
    if tcPr is not None:
        gs = tcPr.find(qn('w:gridSpan'))
        if gs is not None:
            try:
                span = max(1, int(gs.get(qn('w:val'))))
            except (TypeError, ValueError):
                span = 1
    widths = _grid_widths(table)
    if span > 1 and widths:
        # 定位该单元格起始列，累加它跨越的列宽
        for row in table.rows:
            col = 0
            for c in row.cells:
                if c._tc is cell._tc:
                    return sum(widths[col:col + span]) / TWIPS_PER_CM
                sp = 1
                p2 = c._tc.find(qn('w:tcPr'))
                if p2 is not None:
                    g2 = p2.find(qn('w:gridSpan'))
                    if g2 is not None:
                        try:
                            sp = max(1, int(g2.get(qn('w:val'))))
                        except (TypeError, ValueError):
                            sp = 1
                col += sp
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None and tcW.get(qn('w:type')) == 'dxa':
            try:
                w = int(tcW.get(qn('w:w')))
                if w > 0:
                    return w / TWIPS_PER_CM
            except (TypeError, ValueError):
                pass
    if widths:
        return sum(widths) / TWIPS_PER_CM
    return 16.0


def _cell_left_cm(table, cell, page_left_cm):
    """单元格左沿距纸张左边多少 cm（含页边距）。

    表格内字段做位置微调时，制表位是以**单元格左沿**为原点的，
    必须把这段偏移算进去，否则整体偏出一大截。
    """
    widths = _grid_widths(table)
    if not widths:
        return page_left_cm
    for row in table.rows:
        col = 0
        for c in row.cells:
            if c._tc is cell._tc:
                return page_left_cm + sum(widths[:col]) / TWIPS_PER_CM
            sp = 1
            p2 = c._tc.find(qn('w:tcPr'))
            if p2 is not None:
                g2 = p2.find(qn('w:gridSpan'))
                if g2 is not None:
                    try:
                        sp = max(1, int(g2.get(qn('w:val'))))
                    except (TypeError, ValueError):
                        sp = 1
            col += sp
    return page_left_cm


def _cell_margins_cm(cell):
    """单元格左右内边距之和（cm），取不到按 0.19cm×2 估。"""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is not None:
        mar = tcPr.find(qn('w:tcMar'))
        if mar is not None:
            tot = 0
            for side in ('w:left', 'w:right'):
                el = mar.find(qn(side))
                if el is not None:
                    try:
                        tot += int(el.get(qn('w:w'))) / TWIPS_PER_CM
                    except (TypeError, ValueError):
                        pass
            if tot:
                return tot
    return 0.38


def _text_width_units(text):
    """文本宽度（以"全角字"为单位）：中日韩全角计 1，ASCII 计 0.5。

    把空格和数字也按全角算会大幅高估宽度，导致本来一行放得下的
    短字段被误判为两行而缩字号。
    """
    units = 0.0
    for ch in text:
        units += 1.0 if ord(ch) > 0x2E80 else 0.5
    return units


def estimate_lines(text, font_pt, usable_width_cm, first_indent_pt=0):
    """估算文本在给定字号/宽度下占多少行（全角字宽≈字号）。"""
    if not text:
        return 0
    char_w_cm = font_pt / PT_PER_CM
    if char_w_cm <= 0:
        return 1
    per_line = max(1.0, usable_width_cm / char_w_cm)
    lines = 0
    for i, seg in enumerate(text.split('\n')):
        # 行尾空格不会触发换行，计入会导致虚假的"多一行"
        units = _text_width_units(seg.rstrip())
        if i == 0 and first_indent_pt:
            units += first_indent_pt / font_pt
        lines += max(1, int(-(-units // per_line)))    # 向上取整
    return lines


def _para_font_pt(para, default=14.0):
    for r in para.runs:
        if r.font.size is not None:
            return r.font.size.pt
    return default


def _para_line_spacing_pt(para, font_pt):
    ls = para.paragraph_format.line_spacing
    if ls is None:
        return font_pt * 1.4
    return ls.pt if hasattr(ls, 'pt') else font_pt * float(ls)


def _set_para_size(para, font_pt, line_pt):
    from docx.enum.text import WD_LINE_SPACING
    for r in para.runs:
        r.font.size = Pt(font_pt)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)


def autofit_cell(table, cell, warn=None):
    """让单元格内容放进预留高度：不够就逐档缩小字号与行距。

    返回 (是否缩过, 最终字号, 是否仍放不下)。
    """
    row = _row_of_cell(table, cell)
    if row is None:
        return False, None, False
    height_cm, is_exact = _row_height_cm(row)
    if not height_cm:
        return False, None, False

    paras = [p for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return False, None, False

    width_cm = _cell_width_cm(table, cell) - _cell_margins_cm(cell)
    # 上下内边距按行高的一小比例估，固定扣 0.2cm 对 1.2cm 的窄行等于扣掉 17%
    avail_cm = height_cm - min(0.15, height_cm * 0.08)

    base_sizes = [_para_font_pt(p) for p in paras]
    orig = list(base_sizes)

    def total_height(scale):
        total = 0.0
        for p, base in zip(paras, base_sizes):
            fs = max(MIN_FONT_PT, base * scale)
            lsp = max(MIN_LINE_SPACING_PT, _para_line_spacing_pt(p, base) * scale)
            n = estimate_lines(p.text, fs, width_cm,
                               p.paragraph_format.first_line_indent.pt
                               if p.paragraph_format.first_line_indent else 0)
            total += n * lsp / PT_PER_CM
        return total

    # 容差按"超出的后果"区分：
    # - atLeast 行：内容超长会把行撑高、整张表往下挤 → 套打直接错位，
    #   必须收紧，略有超出就缩字号；
    # - exact 行：Word 固定行高、超出部分被裁切，几何不会变、
    #   套打对位不受影响，唯一代价是文字被切。字宽是估算值（误差几个
    #   百分点很正常），这里放宽容差，免得把本来印得好好的短字段
    #   无谓缩小、反而和预印栏位对不齐。
    tolerance = 1.35 if is_exact else 1.05
    if total_height(1.0) <= avail_cm * tolerance:
        return False, orig[0] if orig else None, False

    scale = 1.0
    while scale > 0.55:
        scale -= (FONT_STEP / max(orig)) if max(orig) else 0.05
        if total_height(scale) <= avail_cm:
            break
    fits = total_height(scale) <= avail_cm
    for p, base in zip(paras, base_sizes):
        fs = max(MIN_FONT_PT, round(base * scale * 2) / 2.0)
        lsp = max(MIN_LINE_SPACING_PT, _para_line_spacing_pt(p, base) * scale)
        _set_para_size(p, fs, lsp)
    final = max(MIN_FONT_PT, round(orig[0] * scale * 2) / 2.0) if orig else None
    if not fits and warn:
        warn('内容过长，字号已缩到 {}pt 仍可能超出预留高度，建议精简文字'.format(final))
    return True, final, not fits


def max_title_lines(table, cell):
    """这个标题栏**按模板原始字号**设计成能放几行（至少 1 行）。

    套打的标题栏是纸上印死的固定框。行数要按模板原本的字号算，
    不能按缩小后的字号算：标题一长，自适应会把字缩到 9pt，那时
    2.19cm 的框里塞得下 3 行小字——几何上不撑高，可纸上那个框
    是照两行设计的，印出三行小字就是不对。

    所以这个上限必须在缩字号**之前**取好并一路带下去。标题过长的
    正解是缩字号（用户要的就是"文字多了调整字号"），不是加行；
    缩到下限仍放不下时如实告警，让用户精简标题。
    """
    row = _row_of_cell(table, cell)
    if row is None:
        return 3
    height_cm, _exact = _row_height_cm(row)
    if not height_cm:
        return 3
    paras = [p for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return 3
    para = max(paras, key=lambda p: len(p.text))
    # 用**标题正文**的字号算行高：同段里还有更小的白色栏目名，
    # 按它算会以为一格能塞进四五行小字
    vr = [r for r in para.runs if r.text.strip() and _run_prints(r)]
    fs = (vr[-1].font.size.pt if vr and vr[-1].font.size
          else _para_font_pt(para))
    line_cm = _para_line_spacing_pt(para, fs) / PT_PER_CM
    if line_cm <= 0:
        return 3
    # 扣掉段前距：标题不是从格子顶端开始排的，上面那段留白也占高度
    before = 0.0
    try:
        sb = para.paragraph_format.space_before
        before = sb.pt / PT_PER_CM if sb else 0.0
    except Exception:
        before = 0.0
    avail = height_cm - before - min(0.15, height_cm * 0.08)
    return max(1, int(avail / line_cm))


def shape_title_cell(table, cell, shape='trapezoid_down', lines_n=None,
                     max_lines=None):
    """长标题在格子里按梯形回行，返回 (是否改动, 行数)。

    lines_n 指定分成几行；None 表示按长度自动决定。

    公文标题回行要求"词意完整、排列对称、长短适宜"（GB/T 9704），
    靠 Word 自动折行只会齐头齐尾、还可能把词拆断。这里主动算好断点并
    插入 w:br：输出的 docx 与预览看到的行数、断点因此完全一致——
    交给 Word 自动折行时，预览端无从知道它会断在哪里。

    行数由格子宽度和**当前字号**决定，所以要在自适应缩字号之后再调用；
    缩过字号的话一行能多放几个字，断点要按新字号重算。
    """
    from .title_shape import split_title_lines
    from docx.oxml import OxmlElement
    import copy

    paras = [p for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return False, 0, False
    para = max(paras, key=lambda p: len(p.text))
    # 只对**标题正文那个 run** 回行。新模板里"标  题"这个白色栏目名与
    # 标题正文同在一段，整段重排会把栏目名一起搬走、还会丢掉它的白色与
    # 定位制表符——预印的栏目名一移位，整张单子就废了。
    value_runs = [r for r in para.runs if r.text.strip() and _run_prints(r)]
    if not value_runs:
        return False, 0, False
    vrun = value_runs[-1]
    text = vrun.text.replace('\n', '').strip()
    if not text:
        return False, 0, False
    font_pt = vrun.font.size.pt if vrun.font.size else _para_font_pt(para)
    usable = _cell_width_cm(table, cell) - _cell_margins_cm(cell)
    # 减去同一行上排在标题正文**之前**的东西（白色栏目名、定位制表符），
    # 否则会以为整格都归标题用，长标题该回行却没回
    lead_cm = 0.0
    for _x0, _x1, _r in _run_positions(para):
        # 比 _r 的底层元素，别比代理对象：每次访问 para.runs 都会新建代理，
        # 用 is 比代理永远不相等（本项目已在别处栽过同一个坑）
        if _r._r is vrun._r:
            lead_cm = _x0
            break
    usable = max(1.0, usable - lead_cm)
    per_line = usable / (font_pt / PT_PER_CM) if font_pt else 0
    if per_line <= 0:
        return False, 0, False
    # 行数上限：调用方在缩字号前算好传进来，避免用缩小后的字号
    # 重新推导出"还能再放一行"
    cap = max_lines or max_title_lines(table, cell)
    if lines_n:
        lines_n = min(int(lines_n), cap)
    want = lines_n or cap

    def _split(pl):
        ls = split_title_lines(text, pl, shape, lines_n)
        if lines_n is None and len(ls) > cap:
            ls = split_title_lines(text, pl, shape, cap)
        return ls

    lines = _split(per_line)
    # 光把行数切够还不算完：每一行还得真的放得进格子宽度，
    # 否则 Word 会自己再折一次，行数照样超——自适应只管高度，
    # 3 行小字在 2.19cm 里绰绰有余，它不会因此继续缩字号。
    base_pt = font_pt
    base_ls = _para_line_spacing_pt(para, base_pt)
    while (max(_text_width_units(l) for l in lines) > per_line
           and font_pt > MIN_FONT_PT):
        font_pt = max(MIN_FONT_PT, font_pt - FONT_STEP)
        per_line = usable / (font_pt / PT_PER_CM)
        lines = _split(per_line)
    if font_pt < base_pt:
        _set_para_size(para, font_pt,
                       max(MIN_LINE_SPACING_PT, base_ls * font_pt / base_pt))
    # 缩到下限仍有行放不下 → Word 会自己再折，实际行数就会超过栏位设计的
    # 行数（纸上那个框照几行画的，多出来就压线/串行），据实告知
    too_wide = max(_text_width_units(l) for l in lines) > per_line
    if len(lines) <= 1:
        # 一行放得下：若之前插过 br，要还原成单行
        if '\n' in vrun.text:
            _rebuild_run_lines(vrun, [text])
            return True, 1, too_wide
        return False, 1, too_wide
    _rebuild_run_lines(vrun, lines)
    return True, len(lines), too_wide


def _rebuild_run_lines(run, lines):
    """把单个 run 重排成若干行（w:br 分隔），不动同段的其它 run。"""
    from docx.oxml import OxmlElement
    r = run._r
    for child in list(r):
        if child.tag in (qn('w:t'), qn('w:br'), qn('w:cr')):
            r.remove(child)
    for i, line in enumerate(lines):
        t = OxmlElement('w:t')
        t.text = line
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        if i < len(lines) - 1:
            r.append(OxmlElement('w:br'))


def _rebuild_lines(para, lines):
    """把段落重建成若干行（w:br 分隔），沿用首个 run 的字体属性。"""
    from docx.oxml import OxmlElement
    import copy
    runs = para.runs
    if not runs:
        return
    rpr = runs[0]._r.find(qn('w:rPr'))
    for r in list(runs):
        para._p.remove(r._r)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        if rpr is not None:
            old = run._r.find(qn('w:rPr'))
            if old is not None:
                run._r.remove(old)
            run._r.insert(0, copy.deepcopy(rpr))
        if i < len(lines) - 1:
            run._r.append(OxmlElement('w:br'))




def _run_prints(run):
    """这个 run 会不会真的印到纸上：纯空白不印，白字是预印占位也不印。"""
    if not run.text or not run.text.strip():
        return False
    color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    return str(color) != 'FFFFFF'


def trim_nonprinting_tail(doc, log=None):
    """收窄"末尾纯占位空白"，让行放得下，返回收窄的段落数。

    成文日期行长这样：`…11日[9 个空格][白色的 某地市某某单位的办公室制]`。
    末尾这段既不打印（空格没内容、单位名是白字预印占位），又实实在在占宽度；
    行一旦超出版心就折成两行，落款被顶到第二页——为一段根本不显影的文字
    赔上一整页。

    因此只裁**最后一个会打印的 run 之后**的纯空白 run：它们后面没有任何
    黑字，裁掉不会让任何要打印的内容移位，套打对位分毫不受影响。
    黑字之间的空格一个都不动——那些是把数字顶到预印空格里的，动了就打偏。
    """
    sec = doc.sections[0]
    limit = ((sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm)
             * TAIL_TRIM_RATIO)
    n = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        fs = _para_font_pt(para)
        char_cm = fs / PT_PER_CM
        if char_cm <= 0:
            continue
        runs = para.runs
        last_print = -1
        for i, r in enumerate(runs):
            if _run_prints(r):
                last_print = i
        # last_print < 0 表示整行都不打印（日期全部留空待手签就是这种情况），
        # 那整行的空白都可以裁——不裁的话这一行照样会折行占两行
        changed = False
        # 从最后往前裁尾部的纯空白 run
        for r in reversed(runs[last_print + 1:]):
            if _text_width_units(para.text.rstrip()) * char_cm <= limit:
                break
            if not r.text or r.text.strip():
                continue
            while r.text and _text_width_units(para.text.rstrip()) * char_cm > limit:
                r.text = r.text[1:]
                changed = True
        if changed:
            n += 1
            if log:
                log('info', '成文日期行末尾的占位空白已收窄，避免折行把落款顶到第二页'
                            '（该处不打印，收窄不影响套打对位）')
    return n


def _tab_stops_cm(para):
    """段落的制表位（cm，相对左边距），升序"""
    out = []
    try:
        for t in para.paragraph_format.tab_stops:
            out.append(t.position.cm)
    except Exception:
        pass
    return sorted(out)


def _tab_target(stops, used, x):
    """一个 \t 该跳到哪儿（cm，相对左边距）。

    优先按**顺序**认领制表位：填充时每定位一个字段就添一个制表位，
    两者按升序一一对应，照序取最可靠。不能只用"第一个大于当前位置的
    制表位"——当前位置是按字符宽度估的，而空格的实际宽度随字体变
    （实测 TNR 里只有数字的一半），估过头就会认领不到本该属于它的那个，
    结果差出一大截（实测把 13.00cm 报成 16.67cm）。
    """
    if used < len(stops):
        return stops[used], used + 1
    nxt = next((s for s in stops if s > x + 0.01), None)
    return (nxt if nxt is not None else x + 0.74), used   # Word 默认制表间隔


def _run_positions(para):
    """产出 (墨迹左沿cm, 墨迹右沿cm, run)，相对左边距，已考虑制表位。

    量的是**墨迹**的左右沿，run 里的前后空格不算——用户拿尺子量的是
    看得见的字。
    """
    stops = _tab_stops_cm(para)
    cw = _para_font_pt(para) / PT_PER_CM
    x = 0.0
    used = 0
    for r in para.runs:
        txt = r.text or ''
        pieces = txt.split('\t')
        start = x
        for j, piece in enumerate(pieces):
            if j:
                x, used = _tab_target(stops, used, x)
                start = x        # 跳格后，墨迹从制表位处起算
            x += _text_width_units(piece) * cw
        last = pieces[-1]
        lead = _text_width_units(last[:len(last) - len(last.lstrip())]) * cw
        trail = _text_width_units(txt[len(txt.rstrip()):]) * cw
        yield start + lead, x - trail, r


def print_positions(doc):
    """列出表格外每段里**会真正打印**的文字及其距纸张左边的位置（cm）。

    套打对不对，最终只取决于黑字落在哪儿——白字和空格再怎么排都不显影。
    所以把黑字的实际位置报出来，用户拿尺子量一下真实的预印单就能核对；
    对不上时改模板里那几个空格即可，改完位置也不会再随填写内容变动
    （年/月/日是定宽槽位）。
    """
    sec = doc.sections[0]
    left = sec.left_margin.cm
    out = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        for x0, x1, r in _run_positions(para):
            if r.text.strip() and _run_prints(r):
                out.append((left + x0, left + x1, r.text.strip()))
    return out


def fit_one_page(doc, min_bottom_cm=0.6):
    """给套打表单留足分页余量：把下边距压到较小值。

    为什么改下边距是安全的：内容自上而下排布，位置只由上边距和内容本身
    决定；下边距只影响"哪一行放不下要翻页"这个阈值。缩小它**不会移动
    任何已有内容**，套打对位分毫不变，却能把原本被挤到第二页的末行
    （成文日期）留在第一页。

    这比"预测内容总高再调整"可靠得多——Word 的实际排版高度受字体度量、
    网格吸附、单元格内边距等多重影响，从 XML 精确预测屡试屡错；而这里
    根本不需要预测，只是把阈值放宽到肯定够用。

    返回 (是否调整, 原下边距cm, 新下边距cm)。
    """
    from docx.shared import Cm
    sec = doc.sections[0]
    old = sec.bottom_margin.cm if sec.bottom_margin else None
    if old is None or old <= min_bottom_cm:
        return False, old, old
    sec.bottom_margin = Cm(min_bottom_cm)
    return True, round(old, 2), min_bottom_cm


def lock_row_heights(doc):
    """把所有已设高度的行改为固定高度（hRule=exact），返回锁定的行数。

    默认不启用。原因：atLeast 行的实际渲染高度是
    max(声明高度, 内容自然高度)，而"内容自然高度"取决于 Word 的字体度量，
    程序无法可靠预测（实测领导批示行声明 6.40cm、实际渲染约 11.9cm，
    里面有 10 个空段落撑着）。一旦按声明高度锁成 exact，那一行反而被
    **压小**，下面内容整体上移——正是我们要避免的错位。

    行不会被撑高由自适应保证：自适应以"声明高度"为目标缩字号，而声明高度
    ≤ 实际渲染高度，因此内容永远塞得进原有空间，不会把行顶大。
    """
    n = 0
    for table in doc.tables:
        for row in table.rows:
            trPr = row._tr.find(qn('w:trPr'))
            if trPr is None:
                continue
            he = trPr.find(qn('w:trHeight'))
            if he is None:
                continue
            if he.get(qn('w:hRule')) != 'exact':
                he.set(qn('w:hRule'), 'exact')
                n += 1
    return n


# 承载结构（图片/换行/制表/域）的 run 不能当空 run 删掉
_STRUCT_TAGS = ('w:drawing', 'w:pict', 'w:object', 'w:br', 'w:tab',
                'w:fldChar', 'w:instrText', 'w:sym', 'w:cr')


def strip_empty_runs(doc):
    """删除既无文字、也不承载结构的空 run，返回删除个数。

    把模板里的示例文字换成占位符时，多余的 run 会被清空却留在原地。
    这些空壳在 Word 里通常不显示，但在部分版本/WPS 下可能带出多余标记，
    且会让文件越积越乱，填充时统一清掉。
    """
    n = 0
    for para, _cell in _iter_paragraphs(doc):
        for run in list(para.runs):
            if run.text:
                continue
            if any(run._r.find(qn(t)) is not None for t in _STRUCT_TAGS):
                continue
            parent = run._r.getparent()
            if parent is not None:
                parent.remove(run._r)
                n += 1
    return n


def _fill_doc(doc, values, autofit=True, log=None, lock_heights=False,
              one_page=True, title_shape='trapezoid_down',
              title_lines=None, offsets=None):
    """在内存 Document 上完成填充→自适应→锁高，返回 (已填数, 提示, 单元格报告)。

    预览与实际输出共用这一条路径，保证"预览看到的字号"就是"打印出来的字号"，
    两边各写一套迟早会走样。
    """
    notes = []
    reports = []

    filled = set()
    title_tcs = []          # 标题占位符所在的单元格，梯形回行只作用于它
    field_pos = {}          # 字段 → 实际落在距纸左边多少 cm
    adjustable = []         # 可用"位置微调"调的字段（表格外的段落）
    _left = doc.sections[0].left_margin.cm
    # 每个单元格的左沿：表格内字段的制表位以它为原点
    cell_left = []
    for _t0 in doc.tables:
        for _c0 in _iter_cells(_t0):
            cell_left.append((_c0._tc, _cell_left_cm(_t0, _c0, _left)))

    def _origin_of(cell):
        if cell is None:
            return _left
        for _tc0, _x0 in cell_left:
            if _tc0 is cell._tc:
                return _x0
        return _left

    for p, _cell in _iter_paragraphs(doc):
        in_cell = _cell is not None
        for m in PLACEHOLDER_RE.finditer(p.text):
            key = m.group(1).strip()
            filled.add(key)
            if key in TITLE_FIELDS and in_cell:
                title_tcs.append(_cell._tc)
            if key not in adjustable:
                adjustable.append(key)
        # 表格内的字段也能微调，只是原点是所在单元格的左沿而不是页边距
        _ch, _pos = _replace_in_paragraph(
            p, values, offsets, _origin_of(_cell))
        field_pos.update(_pos)

    def _is_title_cell(cell, _tcs=title_tcs):
        return any(cell._tc is x for x in _tcs)

    # 用户在标题里自己敲了回车 = 指定了断点，自动回行让位
    manual_title = any('\n' in str(values.get(k, '')) for k in TITLE_FIELDS)

    for t in doc.tables:
        for cell in _iter_cells(t):
            label = (cell.text.strip().splitlines() or [''])[0][:12]
            row = _row_of_cell(t, cell)
            h, exact = _row_height_cm(row) if row is not None else (None, False)
            # 取"正文最长的那段"而不是第一段——第一段常是标签
            # （如"拟办意见："），报它的字号会让用户看不出正文到底多小
            def _main_para(_c):
                cands = [pp for pp in _c.paragraphs if pp.text.strip()]
                return max(cands, key=lambda pp: len(pp.text)) if cands else None
            mp = _main_para(cell)
            orig = _para_font_pt(mp) if mp is not None else None
            shrunk = overflow = False
            # 标题梯形回行：先按原字号断行，让自适应看到真实行数；
            # 自适应若缩了字号，一行能多放字，再按新字号重断一次。
            # 用户在标题里自己敲了回车就完全照他的断法来——手动优先于自动。
            is_title = _is_title_cell(cell) and not manual_title
            do_shape = is_title and title_shape and title_shape != 'none'
            # 上限在任何缩放之前取好：缩过字号后行距变小，
            # 再算就会得出"还能多放一行"的错误结论
            cap = max_title_lines(t, cell) if _is_title_cell(cell) else None
            if do_shape:
                shape_title_cell(t, cell, title_shape, title_lines, cap)
            if autofit and cell.text.strip():
                def _warn(msg, _l=label):
                    notes.append('【{}】{}'.format(_l, msg))
                shrunk, _size, overflow = autofit_cell(t, cell, warn=_warn)
                if shrunk and log:
                    log('info', '套打自适应：{} 区字号调整为 {}pt{}'.format(
                        label, _size, '（仍偏长）' if overflow else ''))
            if do_shape:
                _ch, _n, _wide = shape_title_cell(t, cell, title_shape,
                                                  title_lines, cap)
                if _n > 1 and log:
                    log('info', '标题按{}回行为 {} 行'.format(
                        '正梯形' if title_shape == 'trapezoid_down' else '倒梯形', _n))
                if _wide:
                    notes.append('【标题】太长了：字号已缩到 {:.0f}pt 下限，仍要多占行，'
                                 '会超出标题栏设计的 {} 行、压到相邻栏位，'
                                 '建议精简标题'.format(MIN_FONT_PT, cap))
                    if log:
                        log('warning', '标题过长，缩到下限仍超出标题栏行数')
            mp2 = _main_para(cell)
            final = _para_font_pt(mp2) if mp2 is not None else None
            reports.append({
                'tc': cell._tc,
                'is_title': _is_title_cell(cell),
                'max_lines': cap,
                'label': label,
                'text': cell.text,
                'width_cm': _cell_width_cm(t, cell),
                'height_cm': h or 0,
                'font_pt': final,
                'orig_font_pt': orig,
                'shrunk': bool(shrunk),
                'overflow': bool(overflow),
            })

    # 表格外的独立段落（紧急程度/密级行、成文日期行）若超出版心会折成两行，
    # 整单就多占一行、可能被顶到第二页。这里只能告警不能自动收窄：
    # 那些空格是把"密级"等预印标签顶到纸上对应位置的，收窄会让黑字
    # 打偏。真正的解法是用户把内容写短些。
    trim_nonprinting_tail(doc, log=log)

    if log:
        for _a, _b, _t in print_positions(doc):
            log('info', '打印位置：「{}」距纸左边 {:.2f}–{:.2f}cm'
                        '（可拿尺子量预印单核对；对不上就改模板里的空格）'
                .format(_t, _a, _b))

    # 告警门槛用**名义**版心宽（另加 2% 容差），比排版计算用的 0.82 宽松得多：
    # 排版留余量是内部的、代价只是字号小半档；告警是给人看的，
    # 用同样紧的门槛会对"秘密★1年"这类完全正常的值天天报警，
    # 反而把真正该看的提示淹掉。这里只在按最乐观的算法都放不下时才提醒。
    sec0 = doc.sections[0]
    content_w = ((sec0.page_width.cm - sec0.left_margin.cm - sec0.right_margin.cm)
                 * 1.02)
    for p in doc.paragraphs:
        txt = p.text.rstrip()
        if not txt.strip():
            continue
        fs = _para_font_pt(p)
        if _text_width_units(txt) * (fs / PT_PER_CM) > content_w:
            head = txt.strip()[:10]
            notes.append('【{}…】这一行超出版心宽度，会折成两行、整单可能多占一行，'
                         '建议精简该行内容'.format(head))
            if log:
                log('warning', '套打：有一行超出版心宽度，可能折行')

    stripped = strip_empty_runs(doc)
    if stripped and log:
        log('info', '清理模板残留空标记 {} 处'.format(stripped))

    if one_page:
        adjusted, old_b, new_b = fit_one_page(doc)
        if adjusted and log:
            log('info', '下边距 {}cm → {}cm，避免末行（成文日期）被挤到第二页；'
                        '此调整不移动任何内容位置，套打对位不受影响'.format(old_b, new_b))

    if lock_heights:
        locked = lock_row_heights(doc)
        if locked and log:
            log('info', '已锁定 {} 行为固定高度，保证与预印栏位对齐'.format(locked))

    # 目标位置在前面内容的左边时顶不过去，据实告知而不是悄悄忽略
    for _k, _want in (offsets or {}).items():
        _got = field_pos.get(_k)
        if _got is not None and abs(_got - float(_want)) > 0.06:
            notes.append('【{}】位置设为 {:.2f}cm，但前面的内容已经排到 {:.2f}cm，'
                         '顶不过去；请把数值调大，或在模板里减少该字段前面的空格'
                         .format(_k, float(_want), _got))

    used = [k for k in values if k in filled and str(values.get(k, '')).strip()]
    return len(used), notes, reports, field_pos, adjustable


def fill_form(template_path, values, output_path, autofit=True, log=None,
              lock_heights=False, one_page=True, title_shape='trapezoid_down',
              title_lines=None, offsets=None, shift=None):
    """按 values 填充套打模板并另存，返回 (已填字段数, 提示列表)。"""
    from docx import Document
    doc = Document(template_path)
    if offsets is None:
        offsets = load_offsets(template_path)
    used, notes, _r, _fp, _adj = _fill_doc(
        doc, values, autofit=autofit, log=log,
        lock_heights=lock_heights, one_page=one_page,
        title_shape=title_shape, title_lines=title_lines, offsets=offsets)
    if shift is None:
        shift = load_shift(template_path)
    if apply_shift(doc, shift[0], shift[1]):
        notes.append('整张按 ({:+.2f}, {:+.2f})cm 平移'.format(shift[0], shift[1]))
    doc.save(output_path)
    return used, notes


def _table_borders(table):
    """读表格边框设置，预览照此画线——模板左右外框为 none，
    若四边都画会凭空多出竖线，和真实版面对不上。"""
    out = {'top': 'single', 'bottom': 'single',
           'left': 'none', 'right': 'none',
           'insideH': 'single', 'insideV': 'single'}
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return out
    b = tblPr.find(qn('w:tblBorders'))
    if b is None:
        return out
    for el in b:
        name = el.tag.split('}')[-1]
        if name in out:
            out[name] = el.get(qn('w:val')) or 'none'
    return out


def _grid_line_cm(doc):
    """文档网格每行高度（cm）；没设网格返回 None。

    未关闭 snapToGrid 的段落（尤其是留白用的空段）按网格行占位，
    预览要按它估高，否则留白区的比例会明显偏小。
    """
    sect = doc.sections[0]._sectPr
    g = sect.find(qn('w:docGrid'))
    if g is None:
        return None
    try:
        return int(g.get(qn('w:linePitch'))) / TWIPS_PER_CM
    except (TypeError, ValueError):
        return None


def paragraph_height_cm(para, grid_cm):
    """单个段落占用的垂直高度（cm）。

    关键是文档网格的吸附规则：未关闭 snapToGrid 的段落，其每一行会吸附到
    网格行上，**行高超过一个网格行时占两格**（本模板网格 0.55cm，14pt 正文
    自然行高 0.69cm → 占 2 格 = 1.10cm）。按一格算会把留白区高度算成一半，
    整单看起来只占大半页——领导批示区 11 段，一格算 6.05cm、吸附算 12.11cm，
    差出 6cm 之多。
    """
    fs = _para_font_pt(para)
    ppr = para._p.find(qn('w:pPr'))
    snap_off = False
    if ppr is not None:
        sg = ppr.find(qn('w:snapToGrid'))
        if sg is not None and sg.get(qn('w:val')) in ('0', 'false'):
            snap_off = True
        rpr = ppr.find(qn('w:rPr'))
        if rpr is not None and not para.runs:
            sz = rpr.find(qn('w:sz'))
            if sz is not None:
                try:
                    fs = int(sz.get(qn('w:val'))) / 2.0
                except (TypeError, ValueError):
                    pass
    natural = _para_line_spacing_pt(para, fs) / PT_PER_CM
    if grid_cm and not snap_off:
        import math
        return max(1, int(math.ceil(natural / grid_cm))) * grid_cm
    return natural


def wrap_segs(segs, usable_cm):
    """按真实几何把片段序列折行，在断点处插入换行片段。

    预览必须自己算折行，不能交给 Qt：Qt 富文本**无视表格的像素宽度**，
    会把表拉满可视区（实测模板表宽应 428px、实际渲染 1190px），
    格子宽出近三倍，于是 29 个字的标题在预览里挤成一行、
    而 Word 里明明是两行——"预览和实际不一样"正出在这里。

    宽度按每个片段自己的字号算（一格里标签和正文字号常不同），
    与 estimate_lines 用的是同一套度量，因此预览的行数就是
    自适应缩字号时依据的行数。
    """
    if usable_cm <= 0:
        return segs
    out = []
    used = 0.0
    for seg in segs:
        if seg.get('text') == '\n':
            out.append(seg)
            used = 0.0
            continue
        pad = seg.get('pad_cm')
        if pad is not None:
            # 制表位摊出来的空白：宽度是算好的，不参与逐字折行
            out.append(seg)
            used += pad
            continue
        char_cm = (seg.get('pt') or 14) / PT_PER_CM
        buf = ''
        for ch in seg.get('text', ''):
            w = char_cm * (1.0 if ord(ch) > 0x2E80 else 0.5)
            # 行尾空格不触发换行，与 estimate_lines 的 rstrip 一致
            if used + w > usable_cm and ch != ' ':
                if buf:
                    out.append(dict(seg, text=buf))
                    buf = ''
                out.append({'text': '\n', 'white': seg.get('white'),
                            'pt': seg.get('pt')})
                used = 0.0
            buf += ch
            used += w
        if buf:
            out.append(dict(seg, text=buf))
    return out


def _cell_content_cm(cell, grid_cm):
    """单元格内容自然高度（cm），空段也算——它们正是留白区的高度来源。"""
    return sum(paragraph_height_cm(p, grid_cm) for p in cell.paragraphs)


def plan_fill(template_path, values, autofit=True,
              title_shape='trapezoid_down', title_lines=None, offsets=None):
    """只算不存：返回预览所需的版面数据，与 fill_form 走同一条填充路径。

    blocks 按文档真实顺序给出（段落与表格交替）——套打单里成文日期在
    表格**之后**，若先渲染全部段落再渲染表格，日期会跑到表格上面去，
    预览与实际版面不符。
    """
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    doc = Document(template_path)
    if offsets is None:
        offsets = load_offsets(template_path)
    _used, notes, reports, field_pos, adjustable = _fill_doc(
        doc, values, autofit=autofit, log=None, title_shape=title_shape,
        title_lines=title_lines, offsets=offsets)

    sec = doc.sections[0]
    page = {
        'width_cm': sec.page_width.cm, 'height_cm': sec.page_height.cm,
        'left_cm': sec.left_margin.cm, 'right_cm': sec.right_margin.cm,
        'top_cm': sec.top_margin.cm, 'bottom_cm': sec.bottom_margin.cm,
    }
    content_w = page['width_cm'] - page['left_cm'] - page['right_cm']
    grid_cm = _grid_line_cm(doc)

    def _is_white(run):
        c = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        return str(c) == 'FFFFFF'

    def _segs_of(para):
        # 制表位（cm，相对左边距）；定位到位的字段靠它顶过去，
        # 预览要照着算，否则会把 \t 当成一个普通字符画出来
        stops = []
        try:
            for t in para.paragraph_format.tab_stops:
                stops.append(t.position.cm)
        except Exception:
            pass
        stops.sort()
        char_cm = _para_font_pt(para) / PT_PER_CM
        out = []
        x = 0.0                      # 当前排到哪儿（cm，相对左边距）
        used = 0                     # 已认领的制表位个数
        for r in para.runs:
            if not r.text:
                continue
            white = _is_white(r)
            pt = r.font.size.pt if r.font.size else 14
            # run 里的 w:br 在 python-docx 里就是 '\n'，必须拆成换行片段，
            # 否则梯形回行的标题在预览里仍会挤成一行
            parts = r.text.split('\n')
            for i, part in enumerate(parts):
                if i:
                    out.append({'text': '\n', 'white': white, 'pt': pt})
                    x = 0.0
                    used = 0
                if not part:
                    continue
                # 把 \t 摊成"到下一个制表位"的空白，预览里位置才对得上
                for j, piece in enumerate(part.split('\t')):
                    if j:
                        nxt, used = _tab_target(stops, used, x)
                        out.append({'text': ' ', 'white': True, 'pt': pt,
                                    'pad_cm': max(0.0, nxt - x)})
                        x = nxt
                    if piece:
                        out.append({'text': piece, 'white': white, 'pt': pt})
                        x += _text_width_units(piece) * char_cm
        return out

    blocks = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            para = Paragraph(child, doc)
            if not para.text.strip():
                continue
            al = para.paragraph_format.alignment
            blocks.append({'kind': 'para',
                           'segs': wrap_segs(_segs_of(para), content_w),
                           'align': {1: 'center', 2: 'right'}.get(
                               int(al) if al is not None else 0, 'left')})
        elif tag == 'tbl':
            table = Table(child, doc)
            borders = _table_borders(table)
            rows = []
            for row in table.rows:
                h, exact = _row_height_cm(row)
                cells = []
                seen = []
                # 本行自己的 tc（含 vMerge continue 的空壳），用来判断
                # 这一格是不是纵向合并的延续格
                own_tcs = row._tr.findall(qn('w:tc'))
                for ci2, cell in enumerate(row.cells):
                    if any(cell._tc is x for x in seen):
                        continue
                    seen.append(cell._tc)
                    # 纵向合并：python-docx 对 continue 格返回**合并源**的
                    # cell，文字会在每一行各画一遍（预览里出现两个"承办部门"）。
                    # 延续格应画成空格子、且不画与上一行之间的横线。
                    cont = False
                    if ci2 < len(own_tcs):
                        _tcpr = own_tcs[ci2].find(qn('w:tcPr'))
                        _vm = _tcpr.find(qn('w:vMerge')) if _tcpr is not None else None
                        if _vm is not None and (_vm.get(qn('w:val')) or 'continue') == 'continue':
                            cont = True
                    rep = None
                    for r in reports:
                        if r['tc'] is cell._tc:
                            rep = r
                            break
                    cw = _cell_width_cm(table, cell)
                    usable = cw - _cell_margins_cm(cell)
                    segs = []
                    for i2, pp in enumerate(cell.paragraphs):
                        if i2:
                            segs.append({'text': '\n', 'white': True, 'pt': 1})
                        segs.extend(wrap_segs(_segs_of(pp), usable))
                    cells.append({
                        'segs': [] if cont else segs,
                        'vmerge_cont': cont,
                        'width_cm': cw,
                        # 延续格的内容属于合并源那一行，不能再拿来抬高本行
                        'content_cm': 0.0 if cont else _cell_content_cm(cell, grid_cm),
                        'font_pt': (rep or {}).get('font_pt'),
                        'orig_font_pt': (rep or {}).get('orig_font_pt'),
                        'shrunk': (rep or {}).get('shrunk', False),
                        'overflow': (rep or {}).get('overflow', False),
                        'is_title': (rep or {}).get('is_title', False),
                        'max_lines': (rep or {}).get('max_lines'),
                    })
                # 行的实际渲染高度：atLeast 取"声明高度"与"内容自然高度"较大者
                natural = max([c['content_cm'] for c in cells] or [0])
                height = (h or 0) if exact else max(h or 0, natural)
                rows.append({'height_cm': height, 'declared_cm': h or 0,
                             'exact': exact, 'cells': cells})
            blocks.append({'kind': 'table', 'rows': rows, 'borders': borders})

    # 兼容旧字段
    paras = [b for b in blocks if b['kind'] == 'para']
    rows_flat = [r for b in blocks if b['kind'] == 'table' for r in b['rows']]
    return {'page': page, 'content_w_cm': content_w, 'grid_cm': grid_cm,
            'blocks': blocks, 'paras': paras, 'rows': rows_flat, 'notes': notes,
            'field_pos': field_pos, 'adjustable': adjustable}


# ---------------- 模板发现 ----------------

def app_dir():
    """软件所在目录（打包后是 exe 所在目录，开发时是项目目录）"""
    import sys as _sys
    if getattr(_sys, 'frozen', False):
        return os.path.dirname(os.path.abspath(_sys.executable))
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _packed_overprint_dir():
    """打包内置的只读模板目录（PyInstaller 解压临时目录）"""
    import sys as _sys
    if getattr(_sys, 'frozen', False):
        base = getattr(_sys, '_MEIPASS', None)
        if base:
            return os.path.join(base, 'templates', '套打')
    return os.path.join(app_dir(), 'templates', '套打')


def bundled_overprint_dir():
    """自带套打模板目录——放在**软件所在目录**下的 templates/套打，方便找。

    打包后 PyInstaller 会把资源解压到临时目录，用户根本找不到；
    故首次运行时把内置模板复制到 exe 同级的 templates/套打 下，
    之后一直用这个看得见、改得动的目录。
    """
    visible = os.path.join(app_dir(), 'templates', '套打')
    packed = _packed_overprint_dir()
    if os.path.normpath(visible) == os.path.normpath(packed):
        return visible
    try:
        os.makedirs(visible, exist_ok=True)
        if os.path.isdir(packed):
            import shutil
            for name in os.listdir(packed):
                if not name.lower().endswith('.docx'):
                    continue
                dst = os.path.join(visible, name)
                if not os.path.exists(dst):
                    shutil.copyfile(os.path.join(packed, name), dst)
    except Exception as exc:      # 目录不可写（如装在 Program Files）时退回打包目录
        logger.warning('无法在软件目录建立模板文件夹（%s），改用内置目录', exc)
        return packed
    return visible


def user_overprint_dir():
    from app.template_common import config_dir
    d = os.path.join(config_dir(), 'overprint')
    return d


def list_templates():
    """返回 [(显示名, 路径, 是否自带)]，自带的排前面。"""
    out = []
    for d, builtin in ((bundled_overprint_dir(), True), (user_overprint_dir(), False)):
        if not os.path.isdir(d):
            continue
        for name in sorted(os.listdir(d)):
            if name.lower().endswith('.docx') and not name.startswith('~$'):
                out.append((os.path.splitext(name)[0], os.path.join(d, name), builtin))
    return out


# ---------------- 从已有 docx 提取内容 ----------------

import datetime

# 字段的标签写法（同一字段可能写成多种样子，如"经办人"/"经 办 人"）
_FIELD_LABELS = {
    '紧急程度': ['紧急程度'],
    '密级': ['密级', '密  级'],
    '标题': ['标题', '标  题', '题目'],
    '拟办意见': ['拟办意见', '拟办意见'],
    '领导批示': ['领导批示'],
    '承办部门': ['承办部门', '承办单位'],
    '经办人': ['经办人', '经 办 人'],
    '电话': ['电话', '联系电话'],
    '文字校核': ['文字校核', '文字核校'],
}

# 允许标签里夹杂空格（"经 办 人：" 这类排版用的空格）
def _label_pattern(label):
    return r'\s*'.join(re.escape(ch) for ch in label if not ch.isspace())


_DATE_PATTERNS = [
    # 日可缺省：送审单常见"2026 年 7 月   日"，日留空待手签，
    # 要求必须有"日"的数字会让整个日期都识别不出来
    re.compile(r'(\d{4})\s*年\s*(\d{1,2})\s*月(?:\s*(\d{1,2})\s*日)?'),
    re.compile(r'(\d{4})\s*[-./]\s*(\d{1,2})\s*[-./]\s*(\d{1,2})'),
]

_CN_DIGITS = {'〇': 0, '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
              '五': 5, '六': 6, '七': 7, '八': 8, '九': 9}
# 公文成文日期常写成中文数字：二〇二六年七月二十五日
_CN_DATE_RE = re.compile(
    r'([〇零一二三四五六七八九]{4})\s*年\s*([一二三四五六七八九十]{1,3})\s*月'
    r'(?:\s*([一二三四五六七八九十]{1,3})\s*日)?')


def _cn_year(text):
    n = 0
    for ch in text:
        if ch not in _CN_DIGITS:
            return None
        n = n * 10 + _CN_DIGITS[ch]
    return n


def _cn_number(text):
    """解析 一~三十一 这类中文数字（含十/十一/二十/二十五）"""
    if not text:
        return None
    if text == '十':
        return 10
    if text.startswith('十'):                      # 十一 ~ 十九
        rest = _CN_DIGITS.get(text[1:], None) if len(text) == 2 else None
        return 10 + rest if rest is not None else None
    if '十' in text:                                # 二十 / 二十五 / 三十一
        head, _sep, tail = text.partition('十')
        h = _CN_DIGITS.get(head)
        if h is None:
            return None
        if not tail:
            return h * 10
        t = _CN_DIGITS.get(tail)
        return h * 10 + t if t is not None else None
    return _CN_DIGITS.get(text)


_FULLWIDTH_DIGITS = {chr(0xFF10 + i): str(i) for i in range(10)}


def _normalize_digits(text):
    """全角数字转半角——公文里日期常写成 ２０２６年６月２５日"""
    return ''.join(_FULLWIDTH_DIGITS.get(ch, ch) for ch in (text or ''))


def parse_date(text):
    """从文本里抽出 (年, 月, 日) 字符串；抽不到返回 None。

    套打模板里年/月/日是三个独立位置，必须拆开分别落位，
    直接把"2026年6月25日"整串塞进"年"格会把后面全顶歪。

    取**最后**一个匹配：成文日期在文末，标题里可能另有日期
    （"关于开展2026年6月专项检查的请示"），取第一个会抓错。
    """
    text = _normalize_digits(text or '')
    for pat in _DATE_PATTERNS:
        m = None
        for m in pat.finditer(text):
            pass
        if m:
            d = m.group(3)
            return m.group(1), str(int(m.group(2))), (str(int(d)) if d else '')
    m = None
    for m in _CN_DATE_RE.finditer(text):
        pass
    if m:
        y = _cn_year(m.group(1))
        mo = _cn_number(m.group(2))
        day = _cn_number(m.group(3)) if m.group(3) else None
        if y and mo:
            return str(y), str(mo), (str(day) if day else '')
    return None


def _blocks_of(doc):
    """产出文档里所有文本块：(文本, 是否来自表格单元格, 单元格对象或None)"""
    for p in doc.paragraphs:
        yield p.text, False, None
    for t in doc.tables:
        for cell in _iter_cells(t):
            yield cell.text, True, cell


def _strip_placeholders(text):
    return PLACEHOLDER_RE.sub('', text or '')


def _date_candidates(doc):
    """产出可能含成文日期的文本，按"从局部到整篇"由细到粗排列。

    日期识别失败的原因几乎都是"日期被切碎了"，而不是格式不认识：

    * 年/月/日被排在同一行的**不同单元格**里（"2026│年│7│月│25│日"），
      逐格看每格都不成日期 → 补一份"整行拼起来"的文本；
    * 日期在**文本框**里（套打单常把落款做成文本框），
      doc.paragraphs 根本看不到 → 直接遍历 body 下所有 w:p；
    * 日期在**页脚**或**嵌套表格**里 → 分别补上；
    * 上面都不成立时，还有"整篇文字拼成一串"兜底。

    与分页无关：这里读的是 XML 里的全部内容，日期落在第几页都能取到。
    """
    from docx.oxml.ns import qn as _qn
    texts = []

    def _p_text(p_el):
        return ''.join(t.text or '' for t in p_el.iter(_qn('w:t')))

    def _harvest(root):
        # 所有段落（含表格内、嵌套表格内、文本框内）
        for p_el in root.iter(_qn('w:p')):
            texts.append(_p_text(p_el))
        # 每一行整行拼接：年/月/日分列在不同格时唯一能拼出日期的形态
        for tr in root.iter(_qn('w:tr')):
            texts.append(''.join(t.text or '' for t in tr.iter(_qn('w:t'))))

    _harvest(doc.element.body)
    for sec in doc.sections:
        for part in (sec.footer, sec.header,
                     sec.even_page_footer, sec.first_page_footer):
            try:
                if part is not None:
                    _harvest(part._element)
            except (AttributeError, ValueError):
                continue
    # 兜底：整篇拼成一串（日期被拆到相邻段落时仍能拼出）
    texts.append(''.join(texts))
    return [_strip_placeholders(t) for t in texts if t and t.strip()]


def _extract_date(doc):
    """从整篇文档里取成文日期，返回 (年, 月, 日) 或 None。

    带"日"的匹配优先于只有年月的匹配——两者都在时前者更完整。
    """
    best = None
    for txt in _date_candidates(doc):
        d = parse_date(txt)
        if not d:
            continue
        if best is None or (d[2] and not best[2]):
            best = d
        elif bool(d[2]) == bool(best[2]):
            best = d          # 同样完整则取更靠后的
    return best


def extract_values(source_path, fields=None):
    """从一份已有 docx 里按标签抽取各字段内容。

    适用于"同类表单的电子版"——不要求与模板结构完全一致，
    按标签文字定位，兼容"承办部门：X"写在段落里或单元格里两种情况。
    返回 {字段: 值}，抽不到的字段不出现在结果里。
    """
    from docx import Document
    from .paragraph import sanitize_document
    doc = Document(source_path)
    sanitize_document(doc)

    wanted = list(fields) if fields else list(_FIELD_LABELS.keys()) + ['年', '月', '日']
    values = {}
    blocks = [(_strip_placeholders(txt), in_cell, cell)
              for txt, in_cell, cell in _blocks_of(doc)]

    # --- 按标签抽取 ---
    for field in wanted:
        if field in ('年', '月', '日'):
            continue
        labels = _FIELD_LABELS.get(field, [field])
        got = None
        for label in labels:
            pat = re.compile(_label_pattern(label) + r'\s*[：:]\s*(.*)', re.S)
            for txt, _in_cell, _cell in blocks:
                if not txt.strip():
                    continue
                m = pat.search(txt)
                if not m:
                    continue
                val = m.group(1).strip()
                # 截到下一个标签处，避免把同一行后面的"密级：X"一起吞掉
                cut = len(val)
                for other_labels in _FIELD_LABELS.values():
                    for ol in other_labels:
                        om = re.search(_label_pattern(ol) + r'\s*[：:]', val)
                        if om and om.start() < cut:
                            cut = om.start()
                val = val[:cut].strip()
                if val:
                    got = val
                    break
            if got:
                break
        if got:
            values[field] = got

    # --- 标题：同一格里"标  题"后面直接跟正文（无冒号）的写法 ---
    # 送审单的标题栏常是"标  题"加几个空格再写内容，没有冒号，
    # 按"标签：值"的模式抽不到
    if '标题' not in values:
        for txt, _in_cell, _cell in blocks:
            m = re.match(r'^\s*(标\s*题|题\s*目)\s+(\S.*)$', txt.strip(), re.S)
            if m:
                cand = re.sub(r'\s*\n\s*', '', m.group(2)).strip()
                if cand and not PLACEHOLDER_RE.search(cand):
                    values['标题'] = cand
                    break

    # --- 标题：标签抽不到时，取"标题"标签相邻单元格 ---
    if '标题' not in values:
        cells = [(txt, cell) for txt, in_cell, cell in blocks if in_cell]
        for i, (txt, _c) in enumerate(cells):
            t = txt.strip()
            if t in ('标题', '标  题', '题目') and i + 1 < len(cells):
                cand = cells[i + 1][0].strip()
                if cand:
                    values['标题'] = re.sub(r'\s*\n\s*', '', cand)
                break

    # --- 长文本字段（拟办意见/领导批示）---
    # 常见两种写法：标签与正文同块的不同段落；或标签独占一段、正文在后续段落。
    # 后者若只在本块里找，会漏掉全部正文。
    _all_labels = [l for ls in _FIELD_LABELS.values() for l in ls]

    def _starts_with_label(text):
        head = text.strip()[:12]
        for lb in _all_labels:
            if re.match(_label_pattern(lb) + r'\s*[：:]', head):
                return True
        return False

    for field in ('拟办意见', '领导批示'):
        if field not in wanted:
            continue
        if len(str(values.get(field, ''))) >= 4:
            continue
        for bi, (txt, _in_cell, _cell) in enumerate(blocks):
            if field not in txt:
                continue
            after = txt.split(field, 1)[1]
            after = re.sub(r'^\s*[：:]\s*', '', after).strip()
            if len(after) >= 4:
                values[field] = after
                break
            # 标签独占一段：往后收集，直到遇到下一个标签或日期行
            collected = []
            for nxt, _ic, _c in blocks[bi + 1:]:
                t = nxt.strip()
                if not t:
                    continue
                if _starts_with_label(t) or parse_date(t):
                    break
                collected.append(t)
            if collected:
                values[field] = '\n'.join(collected)
            break

    # --- 日期：拆成年/月/日三格，避免整串塞进一格把版面顶歪 ---
    got_date = _extract_date(doc)
    if got_date:
        values['年'], values['月'], values['日'] = got_date

    return {k: v for k, v in values.items() if k in wanted and str(v).strip()}


def fit_document(source_path, template_path, output_path,
                 overrides=None, autofit=True, log=None,
                 title_shape='trapezoid_down', title_lines=None):
    """把一份已有 docx 的内容适配到套打模板并输出。

    overrides: 手工修正/补充的字段，优先级高于自动抽取的值。
    返回 (提取到的值, 提示列表)。
    """
    fields = scan_fields(template_path)
    values = extract_values(source_path, fields)
    if overrides:
        values.update({k: v for k, v in overrides.items() if str(v).strip()})
    if log:
        got = '、'.join('{}={}'.format(k, str(v)[:12]) for k, v in sorted(values.items()))
        log('info', '套打适配：识别到 {} 个字段（{}）'.format(len(values), got or '无'))
    _n, notes = fill_form(template_path, values, output_path,
                          autofit=autofit, log=log, title_shape=title_shape,
                          title_lines=title_lines)
    missing = [f for f in fields if not str(values.get(f, '')).strip()]
    if missing:
        notes = list(notes) + ['未能自动识别：{}（可在对话框里手工补填）'.format('、'.join(missing))]
    return values, notes
