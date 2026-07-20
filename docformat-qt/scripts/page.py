# -*- coding: utf-8 -*-
"""
页面设置 — 从 formatter.py 拆分

包含：_apply_page_grid、_set_normal_style_font、add_page_number、
      remove_background、_strip_autospacing_from_styles
"""

import logging
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .font import set_font

logger = logging.getLogger('docformat.page')


def _apply_page_grid(doc, lines_per_page, chars_per_line, base_font_size_pt):
    """设置 Word 文档网格：每页行数 + 每行字符数"""
    for section in doc.sections:
        sect_pr = section._sectPr
        grid = sect_pr.find(qn('w:docGrid'))
        if grid is None:
            grid = OxmlElement('w:docGrid')
            sect_pr.append(grid)

        page_h = section.page_height or Cm(29.7)
        page_w = section.page_width or Cm(21.0)
        text_h_twips = page_h.twips - section.top_margin.twips - section.bottom_margin.twips
        line_pitch = int(round(float(text_h_twips) / lines_per_page))
        grid.set(qn('w:type'), 'linesAndChars' if chars_per_line else 'lines')
        grid.set(qn('w:linePitch'), str(line_pitch))

        if chars_per_line:
            text_w_pt = page_w.pt - section.left_margin.pt - section.right_margin.pt
            pitch_pt = text_w_pt / chars_per_line
            char_space = int(round((pitch_pt - base_font_size_pt) * 4096))
            grid.set(qn('w:charSpace'), str(char_space))


def _set_normal_style_font(doc, font_cn, font_en, size_pt):
    """把 Normal 样式的基准字体/字号设为正文配置"""
    try:
        style = doc.styles['Normal']
        style.font.name = font_en
        style.font.size = Pt(size_pt)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), font_cn)
    except Exception:
        pass


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)


def _strip_autospacing_from_styles(doc):
    """清理整个文档样式表里的 beforeAutospacing/afterAutospacing 属性"""
    try:
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_element = doc.styles.element
        for spacing in styles_element.iter(f'{ns}spacing'):
            for attr in ('beforeAutospacing', 'afterAutospacing'):
                spacing.attrib.pop(f'{ns}{attr}', None)
    except Exception:
        pass


def add_page_number(
    doc,
    font_name="宋体",
    font_size=14,
    style="dash",
    position="outside",
    offset_from_text_mm=7,
    replace_existing=True,
):
    """按自定义样式添加页码"""

    def _footer_state(footer):
        has_content = False
        has_page_field = False
        for para in footer.paragraphs:
            paragraph_text = para.text.strip()
            if paragraph_text:
                has_content = True
                if re.fullmatch(
                    r"[—\-–\s　]*(?:第\s*)?\d+(?:\s*/\s*\d+)?(?:\s*页)?[—\-–\s　]*",
                    paragraph_text,
                ):
                    has_page_field = True
            for run in para.runs:
                xml = run._r.xml or ""
                if 'fldChar' in xml or 'instrText' in xml:
                    has_content = True
                if re.search(r"\bPAGE\b", xml, re.I):
                    has_page_field = True
        return has_content, has_page_field

    for section in doc.sections:
        states = (
            _footer_state(section.footer),
            _footer_state(section.even_page_footer),
            _footer_state(section.first_page_footer),
        )
        has_non_page_content = any(has_content and not has_page for has_content, has_page in states)
        has_existing_page = any(has_page for _has_content, has_page in states)
        if has_non_page_content:
            logger.warning("页脚含有非页码内容，为避免覆盖已跳过页码重设")
            return
        if has_existing_page and not replace_existing:
            return

    use_even_footer = position == "outside"
    try:
        doc.settings.odd_and_even_pages_header_footer = use_even_footer
    except Exception:
        settings_el = doc.settings._element
        even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
        if use_even_footer and even_odd is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))
        elif not use_even_footer and even_odd is not None:
            settings_el.remove(even_odd)

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = use_even_footer
        bottom_margin_cm = section.bottom_margin.cm if section.bottom_margin else 3.5
        footer_distance_cm = max(0.3, bottom_margin_cm - float(offset_from_text_mm) / 10)
        section.footer_distance = Cm(footer_distance_cm)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        first_footer = section.first_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False
        first_footer.is_linked_to_previous = False

        for f in (odd_footer, even_footer, first_footer):
            for para in f.paragraphs:
                para.clear()

        def _add_field(paragraph, instruction):
            begin_run = paragraph.add_run()
            begin = OxmlElement('w:fldChar')
            begin.set(qn('w:fldCharType'), 'begin')
            begin_run._r.append(begin)
            set_font(begin_run, font_name, font_name, font_size, bold=False)

            instruction_run = paragraph.add_run()
            instruction_text = OxmlElement('w:instrText')
            instruction_text.set(qn('xml:space'), 'preserve')
            instruction_text.text = instruction
            instruction_run._r.append(instruction_text)
            set_font(instruction_run, font_name, font_name, font_size, bold=False)

            end_run = paragraph.add_run()
            end = OxmlElement('w:fldChar')
            end.set(qn('w:fldCharType'), 'end')
            end_run._r.append(end)
            set_font(end_run, font_name, font_name, font_size, bold=False)

        def _build_footer_line(footer, align, leading_space=False, trailing_space=False):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align
            para.paragraph_format.left_indent = None
            para.paragraph_format.right_indent = None
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            if leading_space:
                run0 = para.add_run("\u3000")
                set_font(run0, font_name, font_name, font_size, bold=False)

            if style == "dash":
                run = para.add_run("\u2014\u00a0")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " PAGE ")
                run = para.add_run("\u00a0\u2014")
                set_font(run, font_name, font_name, font_size, bold=False)
            elif style == "page_text":
                run = para.add_run("第 ")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " PAGE ")
                run = para.add_run(" 页")
                set_font(run, font_name, font_name, font_size, bold=False)
            elif style == "page_total":
                _add_field(para, " PAGE ")
                run = para.add_run(" / ")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " NUMPAGES ")
            else:
                _add_field(para, " PAGE ")

            if trailing_space:
                run6 = para.add_run("\u3000")
                set_font(run6, font_name, font_name, font_size, bold=False)

        if position == "outside":
            _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, trailing_space=True)
            _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, leading_space=True)
            if section.different_first_page_header_footer:
                _build_footer_line(first_footer, WD_ALIGN_PARAGRAPH.RIGHT, trailing_space=True)
        else:
            align = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(position, WD_ALIGN_PARAGRAPH.CENTER)
            _build_footer_line(odd_footer, align)
            if section.different_first_page_header_footer:
                _build_footer_line(first_footer, align)
