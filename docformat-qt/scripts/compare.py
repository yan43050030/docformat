# -*- coding: utf-8 -*-
"""文档版本比对：修改前 vs 修改后，输出可读的差异对照件。

两条路径：
1) Windows 有 Word/WPS：用原生 CompareDocuments，产出带真正修订痕迹的
   文档，可在 Word 审阅里逐条接受/拒绝——公文会签场景要的就是这个。
2) 没有 Office：用段落级 diff 自行生成对照件，删除的画删除线标红、
   新增的标蓝、改动段并排列出。不是修订痕迹，但足以看清改了什么。
"""
import difflib
import logging
import os
import re
import sys

logger = logging.getLogger('docformat.compare')

_WS = re.compile(r'\s+')


def _norm(text):
    return _WS.sub('', text or '')


def _read_paragraphs(path):
    from docx import Document
    from .paragraph import sanitize_document
    doc = Document(path)
    sanitize_document(doc)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _compare_via_word(base_path, revised_path, output_path, log=None):
    try:
        import pythoncom
        import win32com.client as win32
    except ImportError:
        return False, '本机缺少 pywin32'
    pythoncom.CoInitialize()
    app = None
    for prog in ('Word.Application', 'Kwps.Application', 'wps.Application'):
        try:
            app = win32.Dispatch(prog)
            break
        except Exception:
            continue
    if app is None:
        pythoncom.CoUninitialize()
        return False, '本机未安装 Word/WPS'
    try:
        try:
            app.Visible = False
        except Exception:
            pass
        # 0 = wdCompareTargetNew：结果放进新文档，不动原文件
        result = app.CompareDocuments(
            app.Documents.Open(os.path.abspath(base_path), ReadOnly=True),
            app.Documents.Open(os.path.abspath(revised_path), ReadOnly=True),
            CompareTarget=0)
        result.SaveAs2(os.path.abspath(output_path), FileFormat=16)  # docx
        result.Close(False)
        for d in list(app.Documents):
            try:
                d.Close(False)
            except Exception:
                pass
        return True, 'Word/WPS 原生比对（含修订痕迹，可在审阅中接受/拒绝）'
    except Exception as exc:
        return False, 'Office 比对失败：{}'.format(exc)
    finally:
        try:
            app.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


def diff_paragraphs(base_paras, revised_paras):
    """返回 [(标记, 修改前文本, 修改后文本)]，标记为 same/del/add/chg。"""
    sm = difflib.SequenceMatcher(
        None, [_norm(t) for t in base_paras], [_norm(t) for t in revised_paras])
    out = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for k in range(i2 - i1):
                out.append(('same', base_paras[i1 + k], revised_paras[j1 + k]))
        elif tag == 'delete':
            for k in range(i1, i2):
                out.append(('del', base_paras[k], ''))
        elif tag == 'insert':
            for k in range(j1, j2):
                out.append(('add', '', revised_paras[k]))
        else:  # replace：等长部分视为逐段修改，多余部分算增/删
            n = min(i2 - i1, j2 - j1)
            for k in range(n):
                out.append(('chg', base_paras[i1 + k], revised_paras[j1 + k]))
            for k in range(i1 + n, i2):
                out.append(('del', base_paras[k], ''))
            for k in range(j1 + n, j2):
                out.append(('add', '', revised_paras[k]))
    return out


def summarize(rows):
    return {
        'same': sum(1 for r in rows if r[0] == 'same'),
        'del': sum(1 for r in rows if r[0] == 'del'),
        'add': sum(1 for r in rows if r[0] == 'add'),
        'chg': sum(1 for r in rows if r[0] == 'chg'),
    }


def _build_diff_docx(base_path, revised_path, output_path, rows, stat):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_paragraph('文档版本比对')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True

    meta = doc.add_paragraph('修改前：{}\n修改后：{}\n共 {} 处改动（新增 {}、删除 {}、修改 {}）'.format(
        os.path.basename(base_path), os.path.basename(revised_path),
        stat['add'] + stat['del'] + stat['chg'], stat['add'], stat['del'], stat['chg']))
    for r in meta.runs:
        r.font.size = Pt(10.5)

    legend = doc.add_paragraph('图例：删除内容标红并加删除线，新增内容标蓝；未改动段落以灰色列出。')
    for r in legend.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph('')

    RED = RGBColor(0xC0, 0x39, 0x2B)
    BLUE = RGBColor(0x1F, 0x5F, 0xA9)
    GREY = RGBColor(0x99, 0x99, 0x99)

    for tag, before, after in rows:
        if tag == 'same':
            p = doc.add_paragraph()
            run = p.add_run(before)
            run.font.size = Pt(10.5)
            run.font.color.rgb = GREY
            continue
        if tag == 'del':
            p = doc.add_paragraph()
            lead = p.add_run('［删除］')
            lead.font.size = Pt(10.5); lead.font.bold = True; lead.font.color.rgb = RED
            run = p.add_run(before)
            run.font.size = Pt(10.5); run.font.color.rgb = RED; run.font.strike = True
        elif tag == 'add':
            p = doc.add_paragraph()
            lead = p.add_run('［新增］')
            lead.font.size = Pt(10.5); lead.font.bold = True; lead.font.color.rgb = BLUE
            run = p.add_run(after)
            run.font.size = Pt(10.5); run.font.color.rgb = BLUE
        else:  # chg：前后并排，看得清改了哪句
            p1 = doc.add_paragraph()
            l1 = p1.add_run('［改前］')
            l1.font.size = Pt(10.5); l1.font.bold = True; l1.font.color.rgb = RED
            r1 = p1.add_run(before)
            r1.font.size = Pt(10.5); r1.font.color.rgb = RED; r1.font.strike = True
            p2 = doc.add_paragraph()
            l2 = p2.add_run('［改后］')
            l2.font.size = Pt(10.5); l2.font.bold = True; l2.font.color.rgb = BLUE
            r2 = p2.add_run(after)
            r2.font.size = Pt(10.5); r2.font.color.rgb = BLUE
    doc.save(output_path)


def compare_documents(base_path, revised_path, output_path, log=None,
                      prefer_office=True):
    """比对两个文档，返回 (成功, 说明, 统计dict 或 None)。"""
    if prefer_office and sys.platform == 'win32':
        ok, info = _compare_via_word(base_path, revised_path, output_path, log)
        if ok:
            return True, info, None
        if log:
            log('info', '{}，改用内置段落比对'.format(info))

    base_paras = _read_paragraphs(base_path)
    rev_paras = _read_paragraphs(revised_path)
    rows = diff_paragraphs(base_paras, rev_paras)
    stat = summarize(rows)
    _build_diff_docx(base_path, revised_path, output_path, rows, stat)
    return True, '内置段落比对（改动 {} 处：新增 {}、删除 {}、修改 {}）'.format(
        stat['add'] + stat['del'] + stat['chg'],
        stat['add'], stat['del'], stat['chg']), stat
