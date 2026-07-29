# -*- coding: utf-8 -*-
"""从"在套头底图上点出来的位置"生成套打模板 docx。

输入是一串元素，每个只说三件事：落在纸上哪个位置（cm）、是什么、多大字号。

    {'x': 2.5, 'y': 5.2, 'kind': 'label', 'text': '紧急程度：', 'pt': 12}
    {'x': 4.7, 'y': 5.2, 'kind': 'field', 'name': '紧急程度',  'pt': 12}

kind='label' 是**预印在纸上的栏目名**，写成白字：占住位置、打印时不显影；
kind='field' 是要填的位置，写成 {{字段名}} 占位符，真正会被打印。

定位办法与手工重建送审单时验证过的一致：
* 横向——制表位。位置以缇为单位绝对指定，与字体无关；靠补空格定位会
  随字体漂移（实测 Times New Roman 里一个空格只有数字的一半宽）。
* 纵向——按 y 把元素分行，每行一个段落，用"精确行距 + 段前距"顶到位，
  并关掉文档网格吸附，否则行高会被网格改写。

字面顶端与行盒顶端之间还有个固定差值（随字号变），这里按经验系数补偿；
真机上若仍有零点几厘米的偏差，用「打印位置微调」横向修正，纵向回到
本向导重新拖框即可。
"""
import logging

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger('docformat.template_builder')

PT_PER_CM = 28.3465

# 同一行的 y 容差：拖框难免差几毫米，差这么点就算同一行
ROW_TOL_CM = 0.35
# 页边距：制表位从这里起算，取小值好让最左的元素也能定位
MARGIN_CM = 1.0
# 字面顶端相对行盒顶端的偏移系数（× 行高cm）。
# 实测标定：12/15/18pt 三档下，字面顶端都比目标高出 0.284×行高，
# 三档比值完全一致（0.284 / 0.284 / 0.283），说明就是个常数系数。
# 解出应取 -0.164，代进去后三档的纵向偏差都落到 0.01cm 以内。
TOP_BIAS = -0.164


def _no_grid(doc):
    sectPr = doc.sections[0]._sectPr
    for g in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(g)


def _exact_line(para, line_pt, before_pt):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(max(0.0, before_pt))
    pf.space_after = Pt(0)
    # snapToGrid 必须按 schema 次序插入：直接 append 会排到 w:spacing 之后，
    # Word/LO 可能连带忽略制表位（这个坑踩过）
    pPr = para._p.get_or_add_pPr()
    sg = OxmlElement('w:snapToGrid')
    sg.set(qn('w:val'), '0')
    pPr.insert_element_before(
        sg, 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
        'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
        'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
        'w:rPr', 'w:sectPr', 'w:pPrChange')


def group_rows(items, tol_cm=ROW_TOL_CM):
    """按 y 把元素分行，行内按 x 从左到右排。"""
    rows = []
    for it in sorted(items, key=lambda d: (float(d['y']), float(d['x']))):
        y = float(it['y'])
        if rows and abs(y - rows[-1]['y']) <= tol_cm:
            rows[-1]['items'].append(it)
            # 行的 y 取本行最小值，避免越加越偏
            rows[-1]['y'] = min(rows[-1]['y'], y)
        else:
            rows.append({'y': y, 'items': [it]})
    for r in rows:
        r['items'].sort(key=lambda d: float(d['x']))
    return rows


def build_template(items, out_path, page_w_cm=21.0, page_h_cm=29.7,
                   margin_cm=MARGIN_CM, top_bias=TOP_BIAS):
    """生成套打模板，返回 (输出路径, 字段名列表)。

    items 里 x/y 是该元素**左上角**距纸张左边/上边的厘米数——正是在
    底图上拖框拖出来的那个角，所见即所得。
    """
    if not items:
        raise ValueError('没有任何元素，无法生成模板')
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(page_w_cm), Cm(page_h_cm)
    sec.left_margin = Cm(margin_cm)
    sec.right_margin = Cm(margin_cm)
    sec.top_margin = Cm(margin_cm)
    sec.bottom_margin = Cm(0.6)
    _no_grid(doc)

    rows = group_rows(items)
    fields = []
    prev_bottom = margin_cm          # 上一行的下沿（cm，距纸上边）

    for ri, row in enumerate(rows):
        pts = [float(it.get('pt') or 12) for it in row['items']]
        line_pt = max(pts)
        line_cm = line_pt / PT_PER_CM
        # 目标：本行字面顶端落在 row['y']
        want_top = float(row['y']) - top_bias * line_cm
        before_pt = max(0.0, (want_top - prev_bottom) * PT_PER_CM)

        para = doc.add_paragraph()
        _exact_line(para, line_pt, before_pt)
        ts = para.paragraph_format.tab_stops
        seen_stop = set()
        for it in row['items']:
            pos = round(float(it['x']) - margin_cm, 3)
            if pos > 0 and pos not in seen_stop:
                ts.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.LEFT)
                seen_stop.add(pos)
        for it in row['items']:
            if float(it['x']) - margin_cm > 0:
                para.add_run('\t')
            if it.get('kind') == 'field':
                name = (it.get('name') or '').strip()
                if not name:
                    continue
                if name not in fields:
                    fields.append(name)
                r = para.add_run('{{%s}}' % name)
                r.font.size = Pt(float(it.get('pt') or 12))
            else:
                text = it.get('text') or ''
                if not text:
                    continue
                r = para.add_run(text)
                r.font.size = Pt(float(it.get('pt') or 12))
                # 预印内容：白字占位，打印时不显影
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        prev_bottom = max(prev_bottom, want_top + line_cm)

    doc.save(out_path)
    logger.info('生成套打模板 %s，%d 行 %d 个字段', out_path, len(rows), len(fields))
    return out_path, fields
