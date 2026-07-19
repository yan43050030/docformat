# -*- coding: utf-8 -*-
"""
页眉页脚格式化引擎

参照 WordFormatter header_footer.py + page.py 的设计，
为 DocFormat 新增独立的页眉页脚配置能力。

功能：
  - 页眉/页脚字体、字号、字形、对齐独立配置
  - 奇偶页不同支持（公文标准）
  - 页眉距顶部距离、页脚距底部距离
  - 与现有页码系统集成
  - 自动处理空页眉/页脚（创建默认 run）
  - 断开"链接到上一节"确保设置独立生效
"""

import logging
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from .font import set_font

logger = logging.getLogger('docformat.header_footer')

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

# 默认页眉页脚配置
DEFAULT_HEADER_FOOTER_CONFIG = {
    "header": {
        "font_cn": "宋体",
        "font_en": "Times New Roman",
        "font_size": 10.5,     # 五号
        "bold": False,
        "italic": False,
        "alignment": "center",
        "distance_from_top_mm": 15.0,
        "enabled": False,       # 公文默认无页眉内容
    },
    "footer": {
        "font_cn": "宋体",
        "font_en": "Times New Roman",
        "font_size": 10.5,     # 五号 (页码通常用小五/五号)
        "bold": False,
        "italic": False,
        "alignment": "center",
        "distance_from_bottom_mm": 17.5,
        "enabled": True,        # 页脚用于页码
    },
    "even_odd_different": False,
    "first_page_different": False,
}


def _zero_paragraph_indent(paragraph) -> None:
    """显式写入零值缩进，清除从原文档继承的缩进"""
    pPr = paragraph._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:firstLine="0" w:firstLineChars="0" '
        f'w:hanging="0" w:hangingChars="0" w:left="0" w:right="0"/>')
    pPr.append(ind)


def _ensure_run_properties(para, font_cn, font_en, font_size,
                           bold, italic) -> None:
    """为无 run 的段落设置默认 run 属性，确保用户后续编辑时继承正确字体"""
    if para.runs:
        return

    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        pPr.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)

    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(font_size * 2)))

    if bold:
        b = rPr.find(qn("w:b"))
        if b is None:
            b = parse_xml(f'<w:b {nsdecls("w")}/>')
            rPr.append(b)

    if italic:
        i = rPr.find(qn("w:i"))
        if i is None:
            i = parse_xml(f'<w:i {nsdecls("w")}/>')
            rPr.append(i)


def _format_header_footer_part(part, config: dict) -> None:
    """对单个 header/footer part 应用格式化

    Args:
        part: section.header 或 section.footer
        config: {'font_cn', 'font_en', 'font_size', 'bold', 'italic', 'alignment'}
    """
    if part is None:
        return

    part.is_linked_to_previous = False

    font_cn = config.get("font_cn", "宋体")
    font_en = config.get("font_en", "Times New Roman")
    font_size = config.get("font_size", 10.5)
    bold = config.get("bold", False)
    italic = config.get("italic", False)
    alignment = ALIGNMENT_MAP.get(config.get("alignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

    paragraphs = part.paragraphs

    # 检查是否有实际内容（文本、run、或域代码如页码）
    def _has_content(p):
        if p.text.strip() or p.runs:
            return True
        if p._element.find(qn('w:fldSimple')) is not None:
            return True
        if p._element.find(qn('w:fldChar')) is not None:
            return True
        return False

    if not paragraphs or not any(_has_content(p) for p in paragraphs):
        # 创建空段落保证字体设置生效
        for p in list(paragraphs):
            p._element.getparent().remove(p._element)
        part.add_paragraph().add_run("")

    for para in part.paragraphs:
        _zero_paragraph_indent(para)
        para.alignment = alignment

        for run in para.runs:
            set_font(run, font_cn, font_en, font_size, bold=bold)

        # 更新页码 fldSimple 中的 run
        for fld in para._element.findall(qn("w:fldSimple")):
            for r in fld.findall(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is None:
                    rPr = parse_xml(f'<w:rPr {nsdecls("w")} />')
                    r.insert(0, rPr)

                sz = rPr.find(qn("w:sz"))
                if sz is None:
                    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(font_size * 2)}"/>')
                    rPr.append(sz)
                else:
                    sz.set(qn("w:val"), str(int(font_size * 2)))

                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), font_cn)
                rFonts.set(qn("w:ascii"), font_en)
                rFonts.set(qn("w:hAnsi"), font_en)
                rFonts.set(qn("w:cs"), font_en)

        # 为无 run 的段落设置默认 run 属性
        _ensure_run_properties(para, font_cn, font_en, font_size, bold, italic)


def apply_header_footer(doc, hf_config: dict = None) -> None:
    """对文档所有 section 应用页眉页脚格式化

    Args:
        doc: python-docx Document 实例
        hf_config: 页眉页脚配置字典，格式见 DEFAULT_HEADER_FOOTER_CONFIG
                   为 None 时使用默认配置
    """
    if hf_config is None:
        hf_config = DEFAULT_HEADER_FOOTER_CONFIG

    header_cfg = hf_config.get("header", {})
    footer_cfg = hf_config.get("footer", {})
    even_odd_diff = hf_config.get("even_odd_different", False)
    first_page_diff = hf_config.get("first_page_different", False)

    # 设置文档级奇偶页不同
    try:
        doc.settings.odd_and_even_pages_header_footer = even_odd_diff
    except Exception:
        settings_el = doc.settings._element
        even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
        if even_odd_diff and even_odd is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))
        elif not even_odd_diff and even_odd is not None:
            settings_el.remove(even_odd)

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = even_odd_diff
        section.different_first_page_header_footer = first_page_diff

        # 页眉距顶部距离
        header_distance = header_cfg.get("distance_from_top_mm")
        if header_distance is not None:
            from docx.shared import Mm
            section.header_distance = Mm(header_distance)

        # 页脚距底部距离
        footer_distance = footer_cfg.get("distance_from_bottom_mm")
        if footer_distance is not None:
            from docx.shared import Mm
            section.footer_distance = Mm(footer_distance)

        # 格式化页眉
        if header_cfg.get("enabled", False):
            _format_header_footer_part(section.header, header_cfg)
            if even_odd_diff:
                _format_header_footer_part(section.even_page_header, header_cfg)
            if first_page_diff:
                _format_header_footer_part(section.first_page_header, header_cfg)

        # 格式化页脚
        if footer_cfg.get("enabled", True):
            _format_header_footer_part(section.footer, footer_cfg)
            if even_odd_diff:
                _format_header_footer_part(section.even_page_footer, footer_cfg)
            if first_page_diff:
                _format_header_footer_part(section.first_page_footer, footer_cfg)

    logger.debug("apply_header_footer: header=%s, footer=%s, even_odd=%s, first_page=%s",
                 header_cfg.get("enabled"), footer_cfg.get("enabled"),
                 even_odd_diff, first_page_diff)


def _get_default_hf_config(preset: dict) -> dict:
    """从预设中提取页眉页脚配置，缺失时使用默认值

    公文标准：
    - 页眉通常为空（无内容）
    - 页脚包含页码，居中或外侧
    - 奇偶页可能不同（外侧页码时）
    """
    hf = preset.get("header_footer", {})

    if not hf:
        # 从 preset 的页码设置推导页脚配置
        position = preset.get("page_number_position", "center")
        footer_align = "center"
        if position == "outside":
            footer_align = "right"  # 奇数页居右
        elif position == "left":
            footer_align = "left"

        return {
            "header": {
                "font_cn": preset.get("page_number_font", "宋体"),
                "font_en": "Times New Roman",
                "font_size": preset.get("page_number_size", 14) * 0.75,  # 略小于页码
                "bold": False,
                "italic": False,
                "alignment": "center",
                "distance_from_top_mm": 15.0,
                "enabled": False,
            },
            "footer": {
                "font_cn": preset.get("page_number_font", "宋体"),
                "font_en": "Times New Roman",
                "font_size": preset.get("page_number_size", 14),
                "bold": False,
                "italic": False,
                "alignment": footer_align,
                "distance_from_bottom_mm": preset.get("page_number_offset_mm", 7) + 20,
                "enabled": True,
            },
            "even_odd_different": position == "outside",
            "first_page_different": False,
        }

    # 合并用户配置和默认值
    result = {
        "header": {**DEFAULT_HEADER_FOOTER_CONFIG["header"], **hf.get("header", {})},
        "footer": {**DEFAULT_HEADER_FOOTER_CONFIG["footer"], **hf.get("footer", {})},
        "even_odd_different": hf.get("even_odd_different", False),
        "first_page_different": hf.get("first_page_different", False),
    }
    return result
