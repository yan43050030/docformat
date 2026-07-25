# -*- coding: utf-8 -*-
"""目录生成：Word 域自动目录 / 手动格式化目录页（可算真实页码）"""
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _first_body_element(doc):
    """返回文档正文第一个可插入位置的元素，找不到返回 None"""
    body = doc._body._body
    children = list(body)
    if not children:
        return None
    # 跳过 sectPr
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            return child
    return children[0]


def _insert_paragraph_before(doc, ref_element, text=''):
    """在 ref_element 之前插入新段落，返回 Paragraph 对象"""
    from docx.text.paragraph import Paragraph
    body = doc._body._body
    p = OxmlElement('w:p')
    body.insert(list(body).index(ref_element), p)
    para = Paragraph(p, body)
    if text:
        para.text = text
    return para


def insert_auto_toc(doc, levels=3, title_text='目  录'):
    """在文首插入 Word 自动目录字段"""
    ref = _first_body_element(doc)
    if ref is None:
        return

    # 目录标题
    tp = _insert_paragraph_before(doc, ref, title_text)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.runs[0].font.size = Pt(22)
    tp.runs[0].font.bold = True

    # TOC 字段
    tp2 = _insert_paragraph_before(doc, ref)
    begin_run = tp2.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    begin_run._r.append(begin)
    instr_run = tp2.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-{}" \\h \\z \\u '.format(levels)
    instr_run._r.append(instr)
    end_run = tp2.add_run()
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(end)

    # 提示行
    note = _insert_paragraph_before(doc, ref,
        '（↑ 此目录为 Word 自动目录域，请在 Word/WPS 中右键点击 → 更新域，即可自动生成页码）')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in note.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = None

    # 分隔空行
    _insert_paragraph_before(doc, ref, '')


_OUTLINE_LEVEL_MAP = {0: 1, 1: 2, 2: 3, 3: 4}


def _build_heading_items(doc, preset=None):
    """扫描文档，返回 [(标题文本, 级别 0-4), ...]，保持文档顺序。

    优先读排版时写入的大纲级别 outlineLvl（准确，且与排版结果一致）；
    未排版的文档回退到识别器按规则判断。
    """
    from scripts.detector import detect_para_type, _compile_rules, _build_text_context
    from scripts.formatter import PRESETS
    if preset is None:
        preset = PRESETS.get('official_gbk', PRESETS['official'])

    # 先看有没有大纲级别（排版产物都有）
    outlined = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        ppr = para._p.find(qn('w:pPr'))
        if ppr is None:
            continue
        el = ppr.find(qn('w:outlineLvl'))
        if el is None:
            continue
        try:
            lvl = int(el.get(qn('w:val')))
        except (TypeError, ValueError):
            continue
        mapped = _OUTLINE_LEVEL_MAP.get(lvl)
        if mapped:
            outlined.append((text, mapped))
    if outlined:
        return outlined

    all_texts, idx_map = _build_text_context(doc)
    rules = _compile_rules(preset.get('detect_rules'))
    items = []
    prev_type = None
    total = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        ptype = detect_para_type(
            text, i, total, para.paragraph_format.alignment,
            all_texts, all_texts_index=idx_map.get(i),
            prev_para_type=prev_type, rules=rules,
        )
        prev_type = ptype
        level = None
        if ptype == 'heading1':       level = 1
        elif ptype == 'heading2':     level = 2
        elif ptype == 'heading3':     level = 3
        elif ptype == 'heading4':     level = 4
        elif ptype == 'title':        level = 0
        if level is not None:
            items.append((text, level))
    return items


def _set_dot_leader_tab(para, right_pos_twips):
    """给段落设一个右对齐、带点引导线的制表位。

    这才是 Word 里目录点线的正确做法：标题与页码之间放一个制表符，
    点线由 Word 按实际排版宽度自动填充并精确对齐右边界；
    用字面的「. . . .」拼接永远对不齐。
    """
    ppr = para._p.get_or_add_pPr()
    tabs = ppr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        # tabs 在 pPr 中的位置需符合 OOXML 顺序
        from .east_asian_typography import _insert_paragraph_property
        _insert_paragraph_property(ppr, tabs, 'tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(int(right_pos_twips)))
    tabs.append(tab)


def _content_width_twips(doc):
    sec = doc.sections[0]
    try:
        return (sec.page_width.twips - sec.left_margin.twips - sec.right_margin.twips)
    except (AttributeError, TypeError):
        return 8306      # A4 默认正文宽度约 14.6cm


def build_manual_toc(doc, title_text='目  录', page_numbers=None, fallback_reason=''):
    """扫描文档标题层级，在文首生成带点引导线的手动目录页。

    page_numbers: {标题文本: 页码} —— 有则填真实页码，无则留占位符。
    点引导线用右对齐制表位实现，由 Word 按实际宽度填充，页码列精确对齐。
    """
    items = _build_heading_items(doc)
    if not items:
        return 0

    ref = _first_body_element(doc)
    if ref is None:
        return 0

    indent_map = {0: 0, 1: 0, 2: 32, 3: 64, 4: 96}   # pt 缩进
    size_map = {0: 16, 1: 16, 2: 16, 3: 14, 4: 14}     # pt 字号
    right_tab = _content_width_twips(doc)
    pages = page_numbers or {}

    # 目录标题
    tp = _insert_paragraph_before(doc, ref, title_text)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.runs[0].font.size = Pt(22)
    tp.runs[0].font.bold = True

    # 分隔
    _insert_paragraph_before(doc, ref, '')

    # 逐条插入目录项（顺序插入 = 文档顺序）
    for text, level in items:
        indent_pt = indent_map.get(level, 0)
        font_size = size_map.get(level, 14)
        bold = (level <= 1)
        page = pages.get(text)
        page_str = str(page) if page else '__'

        p = _insert_paragraph_before(doc, ref)
        p.paragraph_format.first_line_indent = Pt(0)
        if indent_pt > 0:
            p.paragraph_format.left_indent = Pt(indent_pt)
        # 制表位右边界随缩进内收，保证页码列在同一竖线上
        _set_dot_leader_tab(p, right_tab)
        run = p.add_run(text + '\t' + page_str)
        run.font.size = Pt(font_size)
        run.font.bold = bold

    # 末尾提示
    _insert_paragraph_before(doc, ref, '')
    if pages:
        msg = '（此目录由程序自动生成，页码为实际排版页码；正文改动后请重新生成）'
    else:
        msg = '（此目录由程序自动生成，页码位为占位符「__」，请手动填入{}）'.format(
            '；' + fallback_reason if fallback_reason else '')
    note = _insert_paragraph_before(doc, ref, msg)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in note.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = None
    return len(items)


def _page_numbers_via_word(path, log=None):
    """Windows：用 Word/WPS COM 计算每个标题所在页码。"""
    try:
        import win32com.client as win32
    except ImportError:
        return None
    pages = {}
    app = None
    for prog in ('Word.Application', 'Kwps.Application', 'wps.Application'):
        try:
            app = win32.Dispatch(prog)
            break
        except Exception:
            continue
    if app is None:
        return None
    try:
        app.Visible = False
        doc = app.Documents.Open(os.path.abspath(path), ReadOnly=True)
        try:
            doc.Repaginate()
            WD_ACTIVE_PAGE = 3     # wdActiveEndPageNumber
            for para in doc.Paragraphs:
                try:
                    lvl = para.OutlineLevel        # 1-9 = 标题级，10 = 正文
                except Exception:
                    continue
                if lvl is None or int(lvl) > 4:
                    continue
                text = (para.Range.Text or '').strip().rstrip('\r\x07')
                if not text:
                    continue
                pages[text] = int(para.Range.Information(WD_ACTIVE_PAGE))
        finally:
            doc.Close(False)
    except Exception as exc:
        if log:
            log('warning', '用 Office 计算目录页码失败：{}'.format(exc))
        return None
    finally:
        try:
            app.Quit()
        except Exception:
            pass
    return pages or None


def _page_numbers_via_libreoffice(path, log=None):
    """Linux/macOS：让 LibreOffice 转出 PDF，再按 PDF 里的标题文字定位页码。"""
    import re
    import shutil
    import subprocess
    import tempfile
    try:
        from app.converter_linux import find_soffice
    except ImportError:
        return None
    soffice = find_soffice()
    if not soffice:
        return None, '本机未找到 LibreOffice'
    if not _has_pdf_text_tool():
        return None, ('本机缺少 PDF 取词工具（pdftotext），无法定位标题页码；'
                      '安装 poppler-utils 后即可自动填真实页码')
    tmp = tempfile.mkdtemp(prefix='docformat_toc_')
    try:
        subprocess.run(
            [soffice, '--headless', '--norestore', '--convert-to', 'pdf',
             '--outdir', tmp, os.path.abspath(path)],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
        pdfs = [f for f in os.listdir(tmp) if f.lower().endswith('.pdf')]
        if not pdfs:
            return None, 'LibreOffice 未能转出 PDF'
        pdf_path = os.path.join(tmp, pdfs[0])
        texts = _pdf_page_texts(pdf_path)
        if not texts:
            return None, 'PDF 取词失败'
        doc = Document(path)
        pages = _match_headings_to_pages(_build_heading_items(doc), texts)
        if not pages:
            return None, '未能在 PDF 中定位到标题文字'
        return pages, ''
    except Exception as exc:
        return None, '用 LibreOffice 计算页码失败：{}'.format(exc)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _match_headings_to_pages(items, page_texts):
    """把标题按文档顺序匹配到 PDF 页文字上，返回 {标题: 页码}。

    页文字已去空白。标题在文档中顺序出现，故页码单调不减——用游标
    从上一个命中页往后找，避免同名小标题（如多处「（一）工作目标」）
    全部错配到第一次出现的位置。
    """
    import re as _re
    pages = {}
    cursor = 0
    for title, _lvl in items:
        key = _re.sub(r'\s+', '', title)
        if not key:
            continue
        for pno in range(cursor, len(page_texts)):
            if key in page_texts[pno]:
                pages[title] = pno + 1
                cursor = pno
                break
    return pages


def _has_pdf_text_tool():
    import shutil
    if shutil.which('pdftotext'):
        return True
    try:
        import pdfminer.high_level  # noqa: F401
        return True
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException:
        # 不只是 ImportError：pdfminer 的 cryptography 依赖损坏时会抛
        # pyo3 的 PanicException，它继承自 BaseException，Exception 接不住。
        # 探测可选依赖绝不能把主程序带崩，故一律视为不可用。
        return False


def _pdf_page_texts(pdf_path):
    """提取 PDF 每页文字（去空白），失败返回 None。优先用 pdftotext。"""
    import re
    import shutil
    import subprocess
    exe = shutil.which('pdftotext')
    if exe:
        try:
            out = subprocess.run([exe, '-layout', pdf_path, '-'],
                                 stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                                 timeout=120)
            raw = out.stdout.decode('utf-8', errors='replace')
            return [re.sub(r'\s+', '', pg) for pg in raw.split('\f')]
        except Exception:
            pass
    try:
        from pdfminer.high_level import extract_text
    except (KeyboardInterrupt, SystemExit):
        raise
    except BaseException:
        return None
    try:
        pages = []
        idx = 0
        while True:
            txt = extract_text(pdf_path, page_numbers=[idx])
            if not txt:
                break
            pages.append(re.sub(r'\s+', '', txt))
            idx += 1
            if idx > 500:
                break
        return pages or None
    except Exception:
        return None


def compute_page_numbers(path, log=None):
    """算出 {标题文本: 页码}，返回 (pages, 失败原因)。

    pages 为 None 时 reason 说明具体缺什么，便于用户对症解决，
    而不是笼统地说"本机没有排版引擎"。
    """
    if sys.platform == 'win32':
        pages = _page_numbers_via_word(path, log)
        if pages:
            return pages, ''
        return _page_numbers_via_libreoffice(path, log)
    return _page_numbers_via_libreoffice(path, log)


def generate_toc(input_path, output_path, mode='auto', levels=3,
                 real_page_numbers=True, log=None):
    """外部调用入口。

    mode='auto'   → 插入 Word 目录域（配合排版写入的大纲级别，
                    在 Word/WPS 中更新域即得真实页码）
    mode='manual' → 生成静态目录页；real_page_numbers=True 时先用
                    Word/WPS 或 LibreOffice 计算真实页码再写入
    """
    pages, reason = None, ''
    if mode != 'auto' and real_page_numbers:
        pages, reason = compute_page_numbers(input_path, log)
        if log:
            log('info', '目录页码：{}'.format(
                '已按实际排版计算 {} 个标题'.format(len(pages)) if pages
                else '使用占位符（{}）'.format(reason or '未知原因')))
    doc = Document(input_path)
    if mode == 'auto':
        insert_auto_toc(doc, levels=levels)
    else:
        build_manual_toc(doc, page_numbers=pages, fallback_reason=reason)
    doc.save(output_path)
