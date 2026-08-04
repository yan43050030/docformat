# -*- coding: utf-8 -*-
"""密级标注检查：该标没标、标得对不对，以及一键补上。

漏标密级不是排版问题，是事故。但机器无法判断一份文件"内容上"该不该定密——
那是人的责任。所以这里只查**文件自己露出来的破绽**，每一条都有明文依据或
自相矛盾之处，不做任何"看着像涉密"的猜测：

  · 有份号却没有密级——GB/T 9704 规定份号只标在涉密公文上，两者必然同进同出；
  · 正文自己说了"注意保密""不得公开""本件为机密件"，版头却没有密级行；
  · 有密级但没有保密期限，或分隔符不是「★」（GB/T 9704 要求「秘密★1年」）；
  · 密级词写错（绝秘/机秘/秘级）；
  · 密级行的位置不对——它应当在版头，且排在紧急程度上面。

插入密级这件事，**永远由人选定密级和期限**，软件不替人拍板：标错密级
比不标更麻烦。软件只负责把选定的内容排到正确的位置、用正确的格式。
"""
import logging
import re

logger = logging.getLogger('docformat.security')

LEVELS = ['绝密', '机密', '秘密']
PERIODS = ['1年', '3年', '5年', '10年', '20年', '30年', '长期']

# 规范写法：绝密★10年
_OK_RE = re.compile(r'^(绝密|机密|秘密)★(\d+个?月|\d+年|长期)$')
# 宽松写法：把各种不规范的也认出来，才谈得上"指出它不规范"
_LOOSE_RE = re.compile(
    r'^(绝密|机密|秘密)\s*[★*✩☆\-—·]?\s*([一二三四五六七八九十0-9]+\s*(?:年|个月|月)|长期)?\s*$')
# 常见写错的密级词
_TYPO_RE = re.compile(r'^(绝秘|机秘|秘级|绝密级|机密级|秘密级)\s*[★*]?\s*'
                      r'([一二三四五六七八九十0-9]+\s*(?:年|个月|月)|长期)?\s*$')
# 份号：版头最前的一串数字，GB/T 9704 用 6 位
_COPYNUM_RE = re.compile(r'^\d{4,8}$')
_URGENCY_RE = re.compile(r'^(特急|加急|平急|急件)$')

# 正文里"这份东西是要保密的"的自述。只认这些明确说法，
# 不做"涉及人事财务所以大概涉密"这种推断——那是定密工作，不是排版软件的事
_HINT_WORDS = [
    '注意保密', '不得公开', '不得外传', '不得复制', '严禁外传',
    '本件为秘密', '本件为机密', '本件为绝密', '本件属秘密',
    '按秘密件管理', '按机密件管理', '按涉密文件管理',
    '密件管理', '涉密文件', '涉密载体',
]
_HINT_RE = re.compile('|'.join(re.escape(w) for w in _HINT_WORDS))

# 版头范围：密级只可能出现在文件最前面这几段里
HEAD_SCAN = 6

FIX_PREFIX = 'security:'


def _nonempty(doc):
    return [p for p in doc.paragraphs if p.text.strip()]


# 版头里还可能有签发人、发文字号——它们不是密级，但也不算"正文开始了"
_SIGNATORY_RE = re.compile(r'^签发人[：:]\s*\S')
_DOCNUM_RE = re.compile(r'^\S{0,20}〔\d{4}〕\d+\s*号$|^\S{0,20}\[\d{4}\]\d+\s*号$')


def _is_head_element(text):
    return bool(_SIGNATORY_RE.match(text) or _DOCNUM_RE.match(text))


def normalize(text):
    """把一行不规范的密级写成规范的「秘密★1年」；认不出返回 None"""
    t = (text or '').strip()
    m = _LOOSE_RE.match(t)
    if not m:
        m2 = _TYPO_RE.match(t)
        if not m2:
            return None
        level = {'绝秘': '绝密', '机秘': '机密', '秘级': '秘密',
                 '绝密级': '绝密', '机密级': '机密', '秘密级': '秘密'}[m2.group(1)]
        period = (m2.group(2) or '').replace(' ', '')
    else:
        level, period = m.group(1), (m.group(2) or '').replace(' ', '')
    if not period:
        return level          # 缺期限，交给调用方决定补什么
    return '{}★{}'.format(level, _cn_to_digits(period))


def _cn_to_digits(period):
    """保密期限里的中文数字转阿拉伯：GB 的示例写的是「绝密★10年」"""
    table = {'一': 1, '二': 2, '两': 2, '三': 3, '四': 4, '五': 5,
             '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
    m = re.match(r'^([一二两三四五六七八九十]+)(年|个月|月)$', period)
    if not m:
        return period
    s, unit = m.group(1), m.group(2)
    if s == '十':
        n = 10
    elif s.startswith('十'):
        n = 10 + table.get(s[1:], 0)
    elif s.endswith('十'):
        n = table.get(s[:-1], 0) * 10
    elif '十' in s:
        a, b = s.split('十', 1)
        n = table.get(a, 0) * 10 + table.get(b, 0)
    else:
        n = table.get(s)
    if not n:
        return period
    return '{}{}'.format(n, '个月' if unit in ('月', '个月') else '年')


def scan(doc):
    """把与密级有关的事实摆出来，判断留给 check()。

    返回 {'sec': (段序, 段落, 文本) 或 None, 'copynum': ..., 'urgency': ...,
          'hints': [(段序, 命中的词)], 'stray': [(段序, 文本)]}
    """
    paras = _nonempty(doc)
    out = {'sec': None, 'copynum': None, 'urgency': None,
           'hints': [], 'stray': [], 'total': len(paras)}
    # 版头到哪儿为止：一旦出现"有实质内容的一段"（标题、主送机关、正文），
    # 版头就结束了。只按段序号卡（比如前 6 段）不行——短文件里正文第 4 段
    # 也还在前 6 段内，末尾误标的密级会被当成版头里的正规密级放过去
    in_head = True
    for i, p in enumerate(paras):
        t = p.text.strip()
        looks_sec = bool(_LOOSE_RE.match(t) or _TYPO_RE.match(t))
        if looks_sec:
            if in_head and i < HEAD_SCAN and out['sec'] is None:
                out['sec'] = (i, p, t)
            else:
                # 密级行跑到正文里去了——位置不对，单独记着
                out['stray'].append((i, t))
            continue
        if in_head and i < HEAD_SCAN:
            if out['copynum'] is None and _COPYNUM_RE.match(t):
                out['copynum'] = (i, p, t)
                continue
            if out['urgency'] is None and _URGENCY_RE.match(t):
                out['urgency'] = (i, p, t)
                continue
            if _is_head_element(t):
                continue
        in_head = False
        m = _HINT_RE.search(t)
        if m:
            out['hints'].append((i, m.group(0)))
    return out


# 勾选项 → 查哪几类。三类各管一摊，谁也不替谁做主
GROUPS = [
    ('sec_missing', '漏标密级（有份号 / 正文自述保密，却没有密级行）'),
    ('sec_format', '密级写法（缺保密期限、分隔符不是「★」、密级词写错）'),
    ('sec_position', '密级位置（版头顺序、跑到正文里去了）'),
]
GROUP_KEYS = [k for k, _ in GROUPS]


def check(doc, default_period='1年', groups=None):
    """返回 findings，字段与 compliance 的一致。

    可自动处理的带 fix_key；fix_key 里带着"补成什么"，因为密级和期限
    必须由人选定，软件只负责排版。
    """
    on = {k: True for k in GROUP_KEYS}
    on.update(groups or {})
    f = scan(doc)
    out = []
    sec = f['sec']

    if sec is None and not on['sec_missing']:
        pass
    elif sec is None:
        why = []
        if f['copynum']:
            why.append('文件有份号（{}）——按 GB/T 9704，份号只标在涉密公文上'
                       .format(f['copynum'][2]))
        for _i, word in f['hints'][:3]:
            why.append('正文出现「{}」'.format(word))
        if why:
            out.append({
                'level': 'warn', 'item': '密级·漏标',
                'detail': '文件看起来是涉密件，却没有密级标识：{}。'
                          '请确认后补标（密级和保密期限须由定密责任人确定）'
                          .format('；'.join(why)),
                'fix_key': FIX_PREFIX + 'insert',
                'locations': [f['copynum'][0] + 1] if f['copynum'] else None,
            })
        else:
            out.append({'level': 'ok', 'item': '密级·漏标',
                        'detail': '无密级标识，也没有发现应当定密的迹象'})
    else:
        idx, _p, text = sec
        norm = normalize(text)
        if not on['sec_format']:
            pass
        elif _OK_RE.match(text):
            out.append({'level': 'ok', 'item': '密级·写法',
                        'detail': '密级标识规范：{}'.format(text)})
        elif norm and '★' not in norm:
            # 只有密级词、没有保密期限
            out.append({
                'level': 'warn', 'item': '密级·缺保密期限',
                'detail': '「{}」没有保密期限。GB/T 9704 要求密级与保密期限'
                          '之间用「★」隔开，如「{}★{}」'
                          .format(text, norm, default_period),
                'fix_key': '{}fix:{}★{}'.format(FIX_PREFIX, norm, default_period),
                'locations': [idx + 1],
            })
        elif norm:
            out.append({
                'level': 'warn', 'item': '密级·写法不规范',
                'detail': '「{}」写法不规范，应为「{}」'.format(text, norm),
                'fix_key': '{}fix:{}'.format(FIX_PREFIX, norm),
                'locations': [idx + 1],
            })

        # 顺序：份号 → 密级 → 紧急程度（GB/T 9704 版头自上而下）
        if on['sec_position'] and f['urgency'] and f['urgency'][0] < idx:
            out.append({'level': 'warn', 'item': '密级·位置',
                        'detail': '紧急程度「{}」排在密级上面了。版头顺序应为'
                                  '份号 → 密级和保密期限 → 紧急程度'
                                  .format(f['urgency'][2]),
                        'locations': [idx + 1]})
        if on['sec_position'] and f['copynum'] and f['copynum'][0] > idx:
            out.append({'level': 'warn', 'item': '密级·位置',
                        'detail': '份号排在密级下面了。版头顺序应为'
                                  '份号 → 密级和保密期限 → 紧急程度',
                        'locations': [idx + 1]})

    for i, text in (f['stray'] if on['sec_position'] else []):
        out.append({'level': 'warn', 'item': '密级·位置',
                    'detail': '第 {} 段出现「{}」，像是密级标识却不在版头。'
                              '密级应标在首页版心右上角（文件最前面）'
                              .format(i + 1, text),
                    'locations': [i + 1]})
    return out


def insert(doc, preset, level, period=''):
    """把密级标识插到该在的位置，按预设的密级格式排版。

    位置：份号之下、紧急程度之上；没有份号就放在最前面。
    这不是随便插一段——版头的顺序是有规定的，插错地方等于换了个错法。
    """
    if level not in LEVELS:
        raise ValueError('密级只能是：{}'.format('、'.join(LEVELS)))
    text = '{}★{}'.format(level, period) if period else level

    f = scan(doc)
    if f['sec'] is not None:
        # 已经有了就只改文字，别插第二个
        para = f['sec'][1]
        _set_text(para, text)
        _apply_fmt(para, preset)
        return text, False

    from docx.text.paragraph import Paragraph
    anchor = None
    after = False
    if f['copynum'] is not None:
        anchor, after = f['copynum'][1], True      # 份号之下
    elif f['urgency'] is not None:
        anchor, after = f['urgency'][1], False     # 紧急程度之上
    else:
        paras = _nonempty(doc)
        if paras:
            anchor, after = paras[0], False        # 全文之首

    if anchor is None:
        para = doc.add_paragraph()
    else:
        new_p = anchor._p.makeelement(anchor._p.tag, {})
        if after:
            anchor._p.addnext(new_p)
        else:
            anchor._p.addprevious(new_p)
        para = Paragraph(new_p, anchor._parent)
    para.add_run(text)
    _apply_fmt(para, preset)
    return text, True


def fix(doc, preset, fix_key):
    """执行 fix_key 指定的修正，返回一句说明。

    fix_key 形如 `security:insert:秘密★1年` 或 `security:fix:秘密★1年`。
    密级和期限就写在 key 里——它是用户在界面上选定的，一路带过来，
    中途谁也不许替他改。
    """
    if not fix_key.startswith(FIX_PREFIX):
        return None
    rest = fix_key[len(FIX_PREFIX):]
    kind, _, arg = rest.partition(':')
    arg = arg.strip()
    if not arg:
        return None
    m = _OK_RE.match(arg) or re.match(r'^(绝密|机密|秘密)$', arg)
    if not m:
        raise ValueError('密级写法不合规范：{}'.format(arg))
    level = m.group(1)
    period = m.group(2) if m.lastindex and m.lastindex >= 2 else ''
    text, added = insert(doc, preset, level, period)
    if kind == 'insert' and added:
        return '插入密级标识「{}」（已按版头顺序排在份号之下）'.format(text)
    return '密级标识改为「{}」'.format(text)


# ---------------------------------------------------------------- 内部

def _set_text(para, text):
    """只换字，保留这一段原有的格式"""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ''


def _apply_fmt(para, preset):
    from .paragraph import format_paragraph
    fmt = (preset or {}).get('security')
    if not fmt:
        return
    try:
        format_paragraph(para, fmt, 'security')
    except Exception as exc:      # 格式化失败不该让"密级已插入"这件事回滚
        logger.warning('密级段落格式化失败：%s', exc)
