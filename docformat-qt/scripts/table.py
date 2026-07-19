# -*- coding: utf-8 -*-
"""
表格处理 — 从 formatter.py 拆分

包含：表格边框、列宽、单元格边距、表格内容格式化、
      表格标题/单位段落识别、智能对齐
"""

import re
import logging
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .font import set_font

logger = logging.getLogger('docformat.table')


def _iter_block_items(doc):
    """迭代文档中的段落和表格（保持顺序）"""
    from docx.oxml.ns import qn as _qn
    body = doc._body._body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            yield Paragraph(child, doc._body)
        elif tag == 'tbl':
            yield Table(child, doc._body)


# 支持的边框样式
_BORDER_STYLES = {
    'single', 'dashed', 'dotted', 'double', 'triple',
    'thinThickSmallGap', 'thickThinSmallGap', 'wave',
    'doubleWave', 'dashDotStroked', 'threeDEmboss',
    'threeDEngrave', 'outset', 'inset', 'nil', 'none',
}


def _set_table_borders(table, size_pt=0.5, color="000000", style="single"):
    size = max(1, int(size_pt * 8))
    if style not in _BORDER_STYLES:
        style = 'single'
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), style)
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if cell_mar is None:
        cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(cell_mar)

    def _set_side(tag, cm_value):
        node = cell_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            cell_mar.append(node)
        node.set(qn('w:type'), 'dxa')
        node.set(qn('w:w'), str(int(Cm(cm_value).twips)))

    _set_side('top', top_cm)
    _set_side('bottom', bottom_cm)
    _set_side('left', left_cm)
    _set_side('right', right_cm)


def _set_table_width_percent(table, percent=100):
    percent = max(1, min(100, int(percent)))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(percent * 50))


def _set_table_indent(table, indent_twips=0):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), str(int(indent_twips)))


def _text_weight(text):
    weight = 0.0
    for ch in text:
        if ord(ch) < 128:
            weight += 0.5
        else:
            weight += 1.0
    return weight


def _normalize_pcts(weights, min_pct, max_pct):
    total = sum(weights) or 1.0
    pcts = [w / total * 100 for w in weights]
    for i, v in enumerate(pcts):
        if v < min_pct:
            pcts[i] = min_pct
    for i, v in enumerate(pcts):
        if v > max_pct:
            pcts[i] = max_pct
    total = sum(pcts) or 1.0
    return [v / total * 100 for v in pcts]


def _set_table_col_widths_by_content(table, min_pct=8, max_pct=45):
    if not table.rows:
        return
    col_count = max(len(row.cells) for row in table.rows)
    if col_count == 0:
        return

    max_weights = [1.0] * col_count
    for row in table.rows:
        for c_idx, cell in enumerate(tuple(row.cells)):
            text = ''.join(p.text for p in cell.paragraphs).strip()
            if text:
                max_weights[c_idx] = max(max_weights[c_idx], _text_weight(text))

    pcts = _normalize_pcts(max_weights, min_pct, max_pct)

    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        tbl.insert(0, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    for pct in pcts:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(pct * 50)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for c_idx, cell in enumerate(tuple(row.cells)):
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'pct')
            tc_w.set(qn('w:w'), str(int(pcts[c_idx] * 50)))


def _set_cell_borders(cell, size_pt=0.5, color="000000", style="single"):
    size = max(1, int(size_pt * 8))
    if style not in _BORDER_STYLES:
        style = 'single'
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), style)
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_header_row_repeat(table, repeat=True):
    """设置表格首行为标题行重复（跨页时重复表头）"""
    if not table.rows:
        return
    first_row = table.rows[0]
    tr_pr = first_row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn('w:tblHeader'))
    if repeat:
        if tbl_header is None:
            tbl_header = OxmlElement('w:tblHeader')
            tr_pr.append(tbl_header)
    else:
        if tbl_header is not None:
            tr_pr.remove(tbl_header)


def _set_cell_shading(cell, color="D9E2F3", fill="auto"):
    """设置单元格背景底纹

    Args:
        color: 填充颜色 (RRGGBB hex)
        fill: 填充模式，通常 'auto' 或 'solid'
    """
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), fill)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)


def _set_cell_vertical_alignment(cell, valign="center"):
    """设置单元格垂直对齐方式

    Args:
        valign: 'top', 'center', 或 'bottom'
    """
    if valign not in ('top', 'center', 'bottom'):
        return
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    v_align = tc_pr.find(qn('w:vAlign'))
    if v_align is None:
        v_align = OxmlElement('w:vAlign')
        tc_pr.append(v_align)
    v_align.set(qn('w:val'), valign)


def _find_nested_tables(table):
    """递归查找表格单元格中嵌套的表格

    Args:
        table: 外层 Table 实例

    Returns:
        list[Table]: 所有嵌套表格（深度优先）
    """
    nested = []
    for row in table.rows:
        for cell in row.cells:
            for child in cell._tc:
                if child.tag == qn('w:tbl'):
                    nested_table = Table(child, cell)
                    nested.append(nested_table)
                    nested.extend(_find_nested_tables(nested_table))
    return nested


def _is_numeric_text(text):
    text = text.replace(',', '').replace('％', '%').strip()
    if not text:
        return False
    return re.match(r'^[-+]?\d+(?:\.\d+)?%?$', text) is not None


def _is_short_text(text, max_len=4):
    return len(text.strip()) <= max_len


def _is_table_title(text):
    """判断是否是表格标题（"表1 xxx" / "附表 xxx" 等）"""
    text = text.strip()
    return bool(re.match(r'^(表\d+|附表)\s', text))


def _is_table_unit(text):
    """判断是否是表格单位行"""
    text = text.strip()
    return bool(re.match(r'^(单位[：:]|注[：:]|数据来源[：:]|备注[：:])', text))


# ===== 段落插入辅助 =====

def _insert_paragraph_after_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_after_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addprevious(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _split_heading_by_punct(paragraph):
    """Split heading like '（三）xxx：正文' into heading paragraph + body paragraph."""
    text = paragraph.text.strip()
    if not text:
        return False

    if not (
        re.match(r'^[一二三四五六七八九十]+、', text) or
        re.match(r'^（[一二三四五六七八九十]+）', text) or
        re.match(r'^\([一二三四五六七八九十]+\)', text) or
        re.match(r'^\d+\.\s*\S', text) or
        re.match(r'^（\d+）', text) or
        re.match(r'^\(\d+\)', text)
    ):
        return False

    punct_positions = []
    for ch in ('：', ':', '。'):
        pos = text.find(ch)
        if pos != -1:
            punct_positions.append(pos)
    if not punct_positions:
        return False
    split_idx = min(punct_positions)
    head = text[:split_idx + 1].strip()
    tail = text[split_idx + 1:].strip()
    if not tail:
        return False

    paragraph.text = head
    new_para = _insert_paragraph_after_paragraph(paragraph, text=tail)
    return new_para is not None
