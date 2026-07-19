# -*- coding: utf-8 -*-
"""
文档水印引擎

支持文字水印（公文密级标注：秘密、机密、绝密、内部资料 等），
通过 VML 形状写入页眉，确保在所有页面显示。

兼容 Word / WPS 等主流文字处理软件。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from typing import Optional

from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

logger = logging.getLogger('docformat.watermark')


# ============================================================
# 配置
# ============================================================

@dataclass
class WatermarkConfig:
    """水印配置"""
    text: str = ""                          # 水印文字（空表示无水印）
    font_cn: str = "宋体"
    font_size: int = 72                     # 磅
    color: str = "#D3D3D3"                 # 浅灰色
    opacity: str = "0.3"                    # 透明度
    rotation: int = -45                     # 旋转角度
    layout: str = "diagonal"               # diagonal / horizontal
    enabled: bool = False

    @classmethod
    def from_dict(cls, d: dict = None) -> WatermarkConfig:
        if not d:
            return cls()
        return cls(
            text=d.get("text", ""),
            font_cn=d.get("font_cn", "宋体"),
            font_size=d.get("font_size", 72),
            color=d.get("color", "#D3D3D3"),
            opacity=d.get("opacity", "0.3"),
            rotation=d.get("rotation", -45),
            layout=d.get("layout", "diagonal"),
            enabled=d.get("enabled", False),
        )


# 预设水印
WATERMARK_PRESETS = {
    "secret": WatermarkConfig(text="秘密", font_size=72, opacity="0.3",
                               rotation=-45, layout="diagonal", enabled=True),
    "confidential": WatermarkConfig(text="机密", font_size=80, opacity="0.35",
                                     rotation=-45, layout="diagonal", enabled=True),
    "top_secret": WatermarkConfig(text="绝密", font_size=86, opacity="0.4",
                                   rotation=-30, layout="diagonal", enabled=True),
    "internal": WatermarkConfig(text="内部资料", font_size=56, opacity="0.25",
                                 rotation=-45, layout="diagonal", enabled=True),
    "draft": WatermarkConfig(text="草稿", font_size=72, opacity="0.25",
                              rotation=-45, layout="diagonal", enabled=True),
    "sample": WatermarkConfig(text="样本", font_size=72, opacity="0.25",
                               rotation=-45, layout="diagonal", enabled=True),
}


DEFAULT_WATERMARK_CONFIG = {
    "text": "",
    "font_cn": "宋体",
    "font_size": 72,
    "color": "#D3D3D3",
    "opacity": "0.3",
    "rotation": -45,
    "layout": "diagonal",
    "enabled": False,
}


# ============================================================
# VML 水印生成
# ============================================================

def _create_vml_watermark(text: str, config: WatermarkConfig) -> OxmlElement:
    """创建 VML 文字水印形状元素

    返回 <w:pict> 元素，包含 VML 矩形 + 文本填充。
    """
    font_size = config.font_size
    color = config.color.lstrip('#')
    # Word uses BGR within VML opacity attribute
    opacity_val = config.opacity

    page_width = 21.0   # cm
    page_height = 29.7  # cm

    if config.layout == "diagonal":
        # 对角水印：大文本居中旋转
        style = (
            f'position:absolute;'
            f'margin-left:0;margin-top:0;'
            f'width:{page_width * 28.35:.0f}pt;'
            f'height:{page_height * 28.35:.0f}pt;'
            f'z-index:-251658240;'
            f'rotation:{config.rotation};'
            f'visibility:visible;'
            f'mso-width-relative:margin;'
            f'mso-height-relative:margin;'
        )
        textpath_style = (
            f'font-family:"{config.font_cn}";'
            f'font-size:{font_size}pt;'
            f'font-weight:normal;'
            f'font-style:normal;'
        )
    else:
        # 水平水印
        style = (
            f'position:absolute;'
            f'margin-left:0;margin-top:200pt;'
            f'width:{page_width * 28.35:.0f}pt;'
            f'height:80pt;'
            f'z-index:-251658240;'
            f'rotation:0;'
            f'visibility:visible;'
        )
        textpath_style = (
            f'font-family:"{config.font_cn}";'
            f'font-size:{font_size}pt;'
            f'font-weight:normal;'
            f'font-style:normal;'
        )

    pict_ns = 'urn:schemas-microsoft-com:vml'
    office_ns = 'urn:schemas-microsoft-com:office:office'

    vml_rect = etree.Element(f'{{{pict_ns}}}rect',
        nsmap={None: pict_ns, 'o': office_ns},
        attrib={
            'id': 'docfmt-watermark',
            f'{{{office_ns}}}spid': '_x0000_s2049',
            'style': style,
            'fillcolor': f'#{color}',
            'strokecolor': f'#{color}',
            f'{{{office_ns}}}allowincell': 'f',
        })

    vml_fill = etree.SubElement(vml_rect, f'{{{pict_ns}}}fill')
    vml_fill.set('opacity', opacity_val)
    vml_fill.set('color2', f'#{color}')

    vml_textpath = etree.SubElement(vml_rect, f'{{{pict_ns}}}textpath')
    vml_textpath.set('on', 't')
    vml_textpath.set('style', textpath_style)
    vml_textpath.set('string', text)

    pict = OxmlElement('w:pict')
    pict.append(vml_rect)

    return pict


def _remove_existing_watermarks(header) -> None:
    """移除页眉中已有的水印形状"""
    if header is None:
        return
    pict_ns = 'urn:schemas-microsoft-com:vml'
    for para in header.paragraphs:
        for pict in para._element.findall(qn('w:pict')):
            rect = pict.find(f'{{{pict_ns}}}rect')
            if rect is not None and rect.get('id') == 'docfmt-watermark':
                para._element.remove(pict)


def apply_watermark(doc, config: WatermarkConfig) -> None:
    """对文档所有 section 应用文字水印

    Args:
        doc: python-docx Document 实例
        config: WatermarkConfig 实例
    """
    if not config.enabled or not config.text.strip():
        logger.debug("Watermark disabled or empty text, skipping")
        return

    text = config.text.strip()
    logger.info("Applying watermark: %s (font=%s, size=%d, layout=%s)",
                text, config.font_cn, config.font_size, config.layout)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        _remove_existing_watermarks(header)

        if not header.paragraphs:
            header.add_paragraph()

        para = header.paragraphs[0]
        vml_pict = _create_vml_watermark(text, config)

        # 水印放在段落开头
        para._element.insert(0, vml_pict)

    logger.debug("Watermark applied to %d section(s)", len(doc.sections))


def remove_watermark(doc) -> None:
    """移除文档中所有水印"""
    for section in doc.sections:
        _remove_existing_watermarks(section.header)
    logger.debug("All watermarks removed")


def get_watermark_preset(name: str) -> Optional[WatermarkConfig]:
    """获取预设水印配置

    Args:
        name: 预设名称 (secret/confidential/top_secret/internal/draft/sample)

    Returns:
        WatermarkConfig 或 None
    """
    return WATERMARK_PRESETS.get(name)
