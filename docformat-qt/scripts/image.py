# -*- coding: utf-8 -*-
"""
图片处理引擎

参照 WordFormatter image.py 设计，为 DocFormat 新增：
  - 图片压缩（Pillow，可选依赖）
  - 图片尺寸控制（固定宽/高，等比缩放，禁止放大）
  - 图片对齐（锚定图片水平/垂直定位）
  - 文字环绕样式转换（inline ↔ anchor 类型）

Pillow 为可选依赖，未安装时静默跳过压缩功能。
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any, Optional

if TYPE_CHECKING:
    from docx import Document

logger = logging.getLogger("docformat.image")


# ============================================================
# 配置数据模型
# ============================================================

@dataclass
class PictureConfig:
    """图片格式化配置"""
    # 尺寸控制
    size_mode: str = "auto"           # auto / width / height
    width: float = 14.0               # cm
    width_unit: str = "cm"
    height: float = 8.0               # cm
    height_unit: str = "cm"
    keep_ratio: bool = True
    no_enlarge: bool = True

    # 对齐
    alignment: str = "center"         # left / center / right

    # 环绕
    wrapping_style: str = "inline"    # inline / square / tight / through /
                                      # topBottom / behindText / inFrontOfText

    # 压缩
    quality: int = 85                 # JPEG quality (1-100)
    max_side_pixels: int = 1600       # 最长边像素上限
    max_file_size: int = 2 * 1024 * 1024  # 2MB
    auto_compress: bool = False       # 是否启用自动压缩

    # 来源：可从 dict 构造
    @classmethod
    def from_dict(cls, d: dict = None) -> PictureConfig:
        if not d:
            return cls()
        return cls(
            size_mode=d.get("size_mode", "auto"),
            width=d.get("width", 14.0),
            width_unit=d.get("width_unit", "cm"),
            height=d.get("height", 8.0),
            height_unit=d.get("height_unit", "cm"),
            keep_ratio=d.get("keep_ratio", True),
            no_enlarge=d.get("no_enlarge", True),
            alignment=d.get("alignment", "center"),
            wrapping_style=d.get("wrapping_style", "inline"),
            quality=d.get("quality", 85),
            max_side_pixels=d.get("max_side_pixels", 1600),
            max_file_size=d.get("max_file_size", 2 * 1024 * 1024),
            auto_compress=d.get("auto_compress", False),
        )


# 默认图片配置（公文场景：印章不缩放、不压缩，仅居中）
DEFAULT_PICTURE_CONFIG = {
    "size_mode": "auto",
    "width": 14.0,
    "width_unit": "cm",
    "height": 8.0,
    "height_unit": "cm",
    "keep_ratio": True,
    "no_enlarge": True,
    "alignment": "center",
    "wrapping_style": "inline",
    "quality": 85,
    "max_side_pixels": 1600,
    "max_file_size": 2 * 1024 * 1024,
    "auto_compress": False,
}


# ============================================================
# 结果封装
# ============================================================

@dataclass
class ImageCompressResult:
    """图片压缩处理结果统计"""
    processed: int = 0
    before_bytes: int = 0
    after_bytes: int = 0
    skipped: int = 0
    errors: list[str] = field(default_factory=list)

    @property
    def saved_bytes(self) -> int:
        return self.before_bytes - self.after_bytes

    @property
    def saved_display(self) -> str:
        saved = self.saved_bytes
        if saved >= 1024 * 1024:
            return f"减少 {saved / (1024 * 1024):.1f}MB"
        elif saved >= 1024:
            return f"减少 {saved / 1024:.0f}KB"
        return f"减少 {saved}B"


# ============================================================
# 图片部件发现
# ============================================================

def _find_image_parts(doc: Document) -> list[Any]:
    """遍历文档包中的所有部件，返回 ImagePart 实例列表。

    通过 doc.part.package.parts 确保页眉、页脚等区域中的图片也被包含。
    """
    from docx.opc.part import Part

    image_parts: list[Any] = []
    for part in doc.part.package.parts:
        if isinstance(part, Part) and hasattr(part, "blob"):
            ct = getattr(part, "content_type", "") or ""
            pn = getattr(part, "partname", "") or ""
            if "image" in ct.lower() or "/media/" in str(pn):
                image_parts.append(part)
    return image_parts


def _replace_image_blob(part: Any, blob: bytes) -> None:
    """安全替换图片部件的二进制数据"""
    part._blob = blob
    if hasattr(part, "_blob_cache"):
        del part._blob_cache


# ============================================================
# 单图压缩
# ============================================================

def _compress_single_image(blob: bytes, max_side: int, quality: int) -> Optional[bytes]:
    """对单张图片执行缩放 + 压缩。

    - 保持原始格式（PNG 保留透明通道）
    - 仅当最长边超过 max_side 时缩放
    - JPEG 使用 quality 参数; PNG 使用 optimize=True
    """
    try:
        from PIL import Image as PILImage
    except ImportError:
        logger.warning("Pillow not installed, cannot compress images")
        return None

    try:
        with PILImage.open(io.BytesIO(blob)) as img:
            orig_fmt = img.format or "JPEG"
            w, h = img.size

            if max(w, h) > max_side:
                ratio = max_side / max(w, h)
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                img = img.resize((new_w, new_h), PILImage.LANCZOS)

            buf = io.BytesIO()
            save_kwargs: dict[str, Any] = {"optimize": True}

            if orig_fmt in ("JPEG", "JPG", "JFIF"):
                if img.mode != "RGB":
                    img = img.convert("RGB")
                save_kwargs["format"] = "JPEG"
                save_kwargs["quality"] = quality
            elif orig_fmt == "PNG":
                save_kwargs["format"] = "PNG"
            else:
                save_kwargs["format"] = orig_fmt

            img.save(buf, **save_kwargs)
            return buf.getvalue()
    except Exception as exc:
        logger.warning("Failed to compress image: %s", exc)
        return None


# ============================================================
# 公开接口 — 压缩
# ============================================================

def compress_images(doc: Document, config: PictureConfig) -> ImageCompressResult:
    """对文档中所有嵌入图片执行压缩处理。

    依据 config.auto_compress 决定是否执行; 若未安装 Pillow 则静默跳过。
    """
    result = ImageCompressResult()

    if not config.auto_compress:
        logger.info("Image compression disabled by config")
        return result

    try:
        from PIL import Image as PILImage  # noqa: F401
    except ImportError:
        logger.warning("Pillow is not installed. Install with: pip install Pillow>=10.0.0")
        return result

    max_side = config.max_side_pixels
    quality = config.quality
    max_file_size = config.max_file_size

    image_parts = _find_image_parts(doc)
    if not image_parts:
        logger.info("No images found in document")
        return result

    logger.info("Compressing %d image(s) (max_side=%dpx, quality=%d, max_file=%dB)",
                len(image_parts), max_side, quality, max_file_size)

    for part in image_parts:
        blob = part.blob
        before_size = len(blob)

        need_compress = False
        try:
            from PIL import Image as PILImage
            with PILImage.open(io.BytesIO(blob)) as img:
                w, h = img.size
                if max(w, h) > max_side:
                    need_compress = True
                if before_size > max_file_size:
                    need_compress = True
        except Exception:
            result.errors.append(f"Failed to inspect image: {part.partname}")
            continue

        if not need_compress:
            result.skipped += 1
            continue

        compressed = _compress_single_image(blob, max_side, quality)
        if compressed is not None:
            _replace_image_blob(part, compressed)
            result.processed += 1
            result.before_bytes += before_size
            result.after_bytes += len(compressed)
            logger.debug("Compressed %s: %dB -> %dB (%.1f%%)",
                        part.partname, before_size, len(compressed),
                        (1 - len(compressed) / before_size) * 100 if before_size else 0)
        else:
            result.errors.append(f"Compression failed: {part.partname}")

    logger.info("Image compression done: processed=%d, skipped=%d, %s",
                result.processed, result.skipped, result.saved_display)
    return result


# ============================================================
# 图片尺寸设置
# ============================================================

def _unit_to_emu(value: float, unit: str) -> int:
    """将用户配置值转换为 EMU（English Metric Unit）"""
    unit_map = {
        "cm": 360000, "厘米": 360000,
        "mm": 36000, "毫米": 36000,
        "pt": 12700, "磅": 12700,
        "inch": 914400, "英寸": 914400,
    }
    multiplier = unit_map.get(unit, 360000)
    return int(round(value * multiplier))


def _calculate_image_size(orig_cx: int, orig_cy: int,
                          config: PictureConfig) -> tuple[Optional[int], Optional[int]]:
    """根据配置计算目标图片尺寸（EMU）。

    Returns:
        (target_cx, target_cy) 或 (None, None) 如果无需修改。
    """
    if config.size_mode == "auto":
        return None, None

    orig_ratio = orig_cx / orig_cy if orig_cy > 0 else 1.0

    if config.size_mode == "width":
        target_cx = _unit_to_emu(config.width, config.width_unit)
        if config.no_enlarge and target_cx > orig_cx:
            target_cx = orig_cx
        target_cy = int(round(target_cx / orig_ratio)) if config.keep_ratio else _unit_to_emu(config.height, config.height_unit)
        if config.no_enlarge and target_cy > orig_cy:
            target_cy = orig_cy
    elif config.size_mode == "height":
        target_cy = _unit_to_emu(config.height, config.height_unit)
        if config.no_enlarge and target_cy > orig_cy:
            target_cy = orig_cy
        target_cx = int(round(target_cy * orig_ratio)) if config.keep_ratio else _unit_to_emu(config.width, config.width_unit)
        if config.no_enlarge and target_cx > orig_cx:
            target_cx = orig_cx
    else:
        return None, None

    return target_cx, target_cy


def _resize_drawing_element(drawing, config: PictureConfig) -> None:
    """修改单个 w:drawing 元素中的图片尺寸。

    同时更新 wp:extent（容器级别）和 pic:spPr/a:xfrm/a:ext（图片变换级别）。
    """
    from docx.oxml.ns import qn

    container = drawing.find(qn("wp:inline"))
    if container is None:
        container = drawing.find(qn("wp:anchor"))
    if container is None:
        return

    extent = container.find(qn("wp:extent"))
    if extent is None:
        return

    try:
        orig_cx = int(extent.get("cx", "0"))
        orig_cy = int(extent.get("cy", "0"))
    except (ValueError, TypeError):
        return

    if orig_cx <= 0 or orig_cy <= 0:
        return

    target_cx, target_cy = _calculate_image_size(orig_cx, orig_cy, config)
    if target_cx is None or target_cx <= 0 or target_cy <= 0:
        return

    target_cx, target_cy = int(target_cx), int(target_cy)

    extent.set("cx", str(target_cx))
    extent.set("cy", str(target_cy))

    for spPr in container.iter(qn("pic:spPr")):
        xfrm = spPr.find(qn("a:xfrm"))
        if xfrm is not None:
            a_ext = xfrm.find(qn("a:ext"))
            if a_ext is not None:
                a_ext.set("cx", str(target_cx))
                a_ext.set("cy", str(target_cy))

    logger.debug("Image resized: %d x %d -> %d x %d EMU", orig_cx, orig_cy, target_cx, target_cy)


def apply_image_size(doc: Document, config: PictureConfig) -> None:
    """遍历文档中所有 w:drawing 元素，设置图片显示尺寸。

    覆盖正文段落、表格单元格、页眉/页脚中的图片。
    """
    if config.size_mode == "auto":
        logger.debug("apply_image_size: size_mode=auto, skip")
        return

    from docx.oxml.ns import qn

    body = doc.element.body
    drawing_count = 0
    for drawing in body.iter(qn("w:drawing")):
        _resize_drawing_element(drawing, config)
        drawing_count += 1

    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.even_page_header, section.even_page_footer,
                     section.first_page_header, section.first_page_footer):
            if part and part.paragraphs:
                for para in part.paragraphs:
                    for drawing in para._element.iter(qn("w:drawing")):
                        _resize_drawing_element(drawing, config)
                        drawing_count += 1

    logger.debug("apply_image_size: %d drawing(s) processed", drawing_count)


# ============================================================
# 图片环绕样式设置
# ============================================================

def apply_image_wrapping(doc: Document, config: PictureConfig) -> None:
    """设置文档中所有图片的文字环绕样式。

    根据 config.wrapping_style 将 <wp:inline> 转换为 <wp:anchor>
    （非嵌入型环绕），或保持 <wp:inline>（嵌入型）。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    wrapping_style = config.wrapping_style
    if wrapping_style == "inline":
        return

    wrap_tags = {
        "square": "wp:wrapSquare",
        "tight": "wp:wrapTight",
        "through": "wp:wrapThrough",
        "topBottom": "wp:wrapTopBottom",
        "behindText": "wp:wrapBehind",
        "inFrontOfText": "wp:wrapInFront",
    }

    wrap_tag = wrap_tags.get(wrapping_style)
    if wrap_tag is None:
        logger.warning("apply_image_wrapping: unknown wrapping_style=%s", wrapping_style)
        return

    def _inline_to_anchor(inline_elem, parent_drawing):
        parent = inline_elem.getparent()
        anchor = etree.Element(qn("wp:anchor"))
        parent.insert(list(parent).index(inline_elem), anchor)
        parent.remove(inline_elem)

        anchor.set("distT", inline_elem.get("distT", "0"))
        anchor.set("distB", inline_elem.get("distB", "0"))
        anchor.set("distL", inline_elem.get("distL", "0"))
        anchor.set("distR", inline_elem.get("distR", "0"))
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "251658240")
        anchor.set("behindDoc", "1" if wrapping_style == "behindText" else "0")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")

        extent = inline_elem.find(qn("wp:extent"))
        if extent is not None:
            anchor.append(extent)

        effect_extent = etree.SubElement(anchor, qn("wp:effectExtent"))
        effect_extent.set("l", "0")
        effect_extent.set("t", "0")
        effect_extent.set("r", "0")
        effect_extent.set("b", "0")

        if wrapping_style == "behindText":
            etree.SubElement(anchor, qn("wp:wrapBehind"))
        elif wrapping_style == "inFrontOfText":
            etree.SubElement(anchor, qn("wp:wrapInFront"))
        elif wrapping_style == "topBottom":
            wrap_elem = etree.SubElement(anchor, qn("wp:wrapTopBottom"))
            wrap_elem.set("distT", "0")
            wrap_elem.set("distB", "0")
            wrap_elem.set("distL", "91440")
            wrap_elem.set("distR", "91440")
        else:
            wrap_elem = etree.SubElement(anchor, qn(wrap_tag))
            wrap_elem.set("wrapText", "both")
            wrap_elem.set("distL", "91440")
            wrap_elem.set("distR", "91440")
            wrap_elem.set("distT", "0")
            wrap_elem.set("distB", "0")

        pos_h = etree.SubElement(anchor, qn("wp:positionH"))
        pos_h.set("relativeFrom", "column")
        pos_h_align = etree.SubElement(pos_h, qn("wp:align"))
        pos_h_align.text = "center"

        pos_v = etree.SubElement(anchor, qn("wp:positionV"))
        pos_v.set("relativeFrom", "paragraph")
        pos_v_offset = etree.SubElement(pos_v, qn("wp:posOffset"))
        pos_v_offset.text = "0"

        graphic = inline_elem.find(qn("a:graphic"))
        if graphic is not None:
            anchor.append(graphic)

        doc_pr = inline_elem.find(qn("wp:docPr"))
        if doc_pr is not None:
            anchor.append(doc_pr)

        cnv_graphic = inline_elem.find(qn("wp:cNvGraphicFramePr"))
        if cnv_graphic is not None:
            anchor.append(cnv_graphic)

        return anchor

    def _process_drawing(drawing):
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            return
        _inline_to_anchor(inline, drawing)

    # 正文
    body = doc.element.body
    for drawing in body.iter(qn("w:drawing")):
        _process_drawing(drawing)

    # 页眉/页脚
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.even_page_header, section.even_page_footer,
                     section.first_page_header, section.first_page_footer):
            if part and part.paragraphs:
                for para in part.paragraphs:
                    for drawing in para._element.iter(qn("w:drawing")):
                        _process_drawing(drawing)


# ============================================================
# 图片对齐设置
# ============================================================

def apply_image_alignment(doc: Document, config: PictureConfig) -> None:
    """设置文档中所有图片的对齐方式。

    对于 <wp:anchor> 图片通过 positionH/align 控制水平对齐，
    对于 <wp:inline> 图片通过段落级 w:jc 控制对齐。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    alignment = config.alignment
    alignment_h_map = {
        "left": "left", "center": "center", "right": "right",
    }
    target_h = alignment_h_map.get(alignment, "center")
    jc_map = {"left": "left", "center": "center", "right": "right"}

    def _set_anchor_alignment(anchor):
        pos_h = anchor.find(qn("wp:positionH"))
        if pos_h is None:
            pos_h = etree.SubElement(anchor, qn("wp:positionH"))
            pos_h.set("relativeFrom", "column")
        else:
            pos_h.set("relativeFrom", "column")
        for child in list(pos_h):
            pos_h.remove(child)
        h_align_elem = etree.SubElement(pos_h, qn("wp:align"))
        h_align_elem.text = target_h

        pos_v = anchor.find(qn("wp:positionV"))
        if pos_v is None:
            pos_v = etree.SubElement(anchor, qn("wp:positionV"))
            pos_v.set("relativeFrom", "paragraph")
        else:
            pos_v.set("relativeFrom", "paragraph")
        for child in list(pos_v):
            pos_v.remove(child)
        v_align_elem = etree.SubElement(pos_v, qn("wp:align"))
        v_align_elem.text = "center"

    def _set_inline_alignment(drawing):
        if target_h not in jc_map:
            return
        parent = drawing.getparent()
        while parent is not None:
            if parent.tag == qn("w:p"):
                pPr = parent.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(parent, qn("w:pPr"))
                    parent.insert(0, pPr)
                old_jc = pPr.find(qn("w:jc"))
                if old_jc is not None:
                    pPr.remove(old_jc)
                jc = etree.SubElement(pPr, qn("w:jc"))
                jc.set(qn("w:val"), jc_map.get(target_h, "left"))
                return
            parent = parent.getparent()

    for drawing in doc.element.body.iter(qn("w:drawing")):
        anchor = drawing.find(qn("wp:anchor"))
        if anchor is not None:
            _set_anchor_alignment(anchor)
        else:
            _set_inline_alignment(drawing)

    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.even_page_header, section.even_page_footer,
                     section.first_page_header, section.first_page_footer):
            if part and part.paragraphs:
                for para in part.paragraphs:
                    for drawing in para._element.iter(qn("w:drawing")):
                        anchor = drawing.find(qn("wp:anchor"))
                        if anchor is not None:
                            _set_anchor_alignment(anchor)
                        else:
                            _set_inline_alignment(drawing)
