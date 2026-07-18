#!/usr/bin/env python3
"""
标点符号修复 v6
- 修复引号处理bug：使用明确的Unicode转义序列
- 正确处理省略号和句号
- v6: 英文撇号保护（it's/don't 不再被配对成中文引号）
- v6: 引号配对状态可跨段落延续（上一段开引号、下一段收引号不再错乱）
- v6: 空格默认策略改为保护英文单词间空格（New York 不再被压成 NewYork）
- v6: print 改为 logging（打包为无控制台 exe 后不再有 stdout 风险）
"""

import logging
import re
import sys
from docx import Document

logger = logging.getLogger('docformat.punctuation')

# 中文标点（使用Unicode转义确保正确）
LEFT_DOUBLE_QUOTE = '“'   # " 左双引号
RIGHT_DOUBLE_QUOTE = '”'  # " 右双引号
LEFT_SINGLE_QUOTE = '‘'   # ' 左单引号
RIGHT_SINGLE_QUOTE = '’'  # ' 右单引号

# 英文撇号（it's / don't）：夹在两个英文字母之间的单引号不参与引号配对，
# 统一规范为弯撇号 U+2019，避免被错误配成中文左单引号
_APOSTROPHE = re.compile("(?<=[A-Za-z])['’](?=[A-Za-z])")

# 基本替换映射
REPLACEMENTS = {
    "(": "（",
    ")": "）",
    ":": "：",
    ";": "；",
    "?": "？",
    "!": "！",
}

# 占位符前缀（使用不可能出现在正常文本中的字符序列）
_PLACEHOLDER_PREFIX = "\x02PROT"


def _protect_special_patterns(text):
    """提取并保护不应被替换的特殊模式，返回 (处理后文本, 保护列表)"""
    protected = []
    counter = [0]

    def _replace_with_placeholder(match):
        placeholder = f"{_PLACEHOLDER_PREFIX}{counter[0]}\x03"
        protected.append((placeholder, match.group()))
        counter[0] += 1
        return placeholder

    result = text

    # 保护 URL（http:// https:// ftp://）
    result = re.sub(r"(?:https?|ftp)://\S+", _replace_with_placeholder, result)

    # 保护邮箱地址
    result = re.sub(r"[\w.+-]+@[\w-]+\.[\w.-]+", _replace_with_placeholder, result)

    # 保护 Windows 文件路径 C:\ D:\
    result = re.sub(r"[A-Za-z]:\\", _replace_with_placeholder, result)

    # 保护标准编号 ISO 9001:2015 等（字母+数字+冒号+数字）
    result = re.sub(r"[A-Za-z]+[\s-]?\d+:\d{2,}", _replace_with_placeholder, result)

    # 保护时间/比分格式 H:MM、HH:MM:SS、3:2 等数字间冒号
    # 注意：不能用 \b，因为 Python3 中中文字符也属于 \w，
    # 导致"上午9:30"中 午 和 9 之间没有 \b 边界
    result = re.sub(r"(?<!\d)(\d{1,2}:\d{1,2}(?::\d{2})?)(?!\d)", _replace_with_placeholder, result)

    return result, protected


def _restore_protected(text, protected):
    """恢复被保护的内容"""
    result = text
    for placeholder, original in protected:
        result = result.replace(placeholder, original)
    return result


def has_chinese(text):
    """检查是否包含中文"""
    return bool(re.search(r"[一-鿿]", text))


def fix_text(text):
    """修复文本中的标点（单段文本入口，引号配对状态不跨调用）"""
    if not text:
        return text

    # 保护特殊模式
    result, protected = _protect_special_patterns(text)

    # 简单替换（省略号/破折号/基本标点/逗号/句号）
    result = _fix_simple_body(result)

    # 引号配对
    result, _dq, _sq = _fix_quotes_whole_text(result)

    # 恢复被保护的内容
    result = _restore_protected(result, protected)
    return result


def _fix_simple_body(result):
    """不涉及配对的简单标点替换（输入应已做过特殊模式保护）"""
    # 省略号（必须在句号之前）
    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)

    # 破折号
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"—(?!—)", "——", result)

    # 基本替换（只在有中文时）
    if has_chinese(result):
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)

    # 逗号
    result = re.sub(r"([一-鿿]),", r"\1，", result)
    result = re.sub(r",([一-鿿])", r"，\1", result)

    # 句号
    result = re.sub(r"([一-鿿])\.(\s|$)", r"\1。\2", result)
    return result


def _fix_simple_punctuation(text):
    """只做不涉及配对的简单标点替换，保留引号不动"""
    if not text:
        return text

    result, protected = _protect_special_patterns(text)
    result = _fix_simple_body(result)
    return _restore_protected(result, protected)


def _fix_quotes_whole_text(text, dq_start=0, sq_start=0):
    """对完整文本做引号配对替换。

    dq_start/sq_start 为此前文本中已出现的双/单引号数量，用于跨段落
    延续配对状态（偶数=下一个是左引号，奇数=下一个是右引号）。
    返回 (结果文本, 双引号累计数, 单引号累计数)。
    """
    result = text

    # 英文撇号先保护，配对结束后统一恢复为弯撇号
    result = _APOSTROPHE.sub("\x04", result)

    # 双引号
    double_quote_chars = ['"', "“", "”", "„", "‟", "「", "」"]
    temp = result
    for q in double_quote_chars:
        temp = temp.replace(q, "\x00")

    dq = dq_start
    if "\x00" in temp:
        chars = list(temp)
        for i, c in enumerate(chars):
            if c == "\x00":
                chars[i] = LEFT_DOUBLE_QUOTE if dq % 2 == 0 else RIGHT_DOUBLE_QUOTE
                dq += 1
        result = "".join(chars)

    # 单引号
    single_quote_chars = ["'", "‘", "’", "‚", "‛"]
    temp = result
    for q in single_quote_chars:
        temp = temp.replace(q, "\x01")

    sq = sq_start
    if "\x01" in temp:
        chars = list(temp)
        for i, c in enumerate(chars):
            if c == "\x01":
                chars[i] = LEFT_SINGLE_QUOTE if sq % 2 == 0 else RIGHT_SINGLE_QUOTE
                sq += 1
        result = "".join(chars)

    result = result.replace("\x04", "’")
    return result, dq, sq


def _redistribute_text_to_runs(runs, new_full_text):
    """将新的完整文本按原 run 的长度边界重新分配回各 run，保留格式"""
    run_lengths = [len(run.text) for run in runs]

    # 如果长度一致（只是字符替换，没有增删），直接按原长度切分
    total_original = sum(run_lengths)
    if len(new_full_text) == total_original:
        pos = 0
        for i, run in enumerate(runs):
            run.text = new_full_text[pos:pos + run_lengths[i]]
            pos += run_lengths[i]
    else:
        # 长度变了（极少情况），回退到旧方案：全塞第一个 run
        runs[0].text = new_full_text
        for run in runs[1:]:
            run.text = ""


def _process_spaces_text(text, mode):
    """根据 mode 处理文本中的空格，返回处理后的文本"""
    if mode == 'keep_all' or not text:
        return text
    if mode == 'remove_all':
        # 删除所有半角空格、全角空格和制表符
        return text.replace('　', '').replace('\t', '').replace(' ', '')
    if mode == 'keep_en_words':
        # 默认策略：删除中文相关的多余空格，但保留英文/数字单词之间的空格
        # （New York、GB/T 9704 类内容不能被压掉）；全角空格与制表符一律
        # 处理掉——首行缩进应由段落格式控制，而不是靠空格/Tab 顶出来。
        text = text.replace('　', '').replace('\t', ' ')
        text = re.sub(r'(?<=[A-Za-z0-9]) +(?=[A-Za-z0-9])', '\x05', text)
        text = text.replace(' ', '')
        return text.replace('\x05', ' ')
    if mode == 'keep_en_boundary':
        CN = r'一-鿿㐀-䶿豈-﫿'
        EN = r'[A-Za-z0-9]'
        CN_cls = f'[{CN}]'

        # 第一步：删除纯中文之间的空格（不涉及英文/数字边界）
        text = re.sub(f'(?<={CN_cls}) +(?={CN_cls})', '', text)

        # 第二步：将中文与英文/数字边界处的空格（0个或多个）统一替换为恰好1个
        # 中文 → 英文/数字
        text = re.sub(f'(?<={CN_cls}) *(?={EN})', ' ', text)
        # 英文/数字 → 中文
        text = re.sub(f'(?<={EN}) *(?={CN_cls})', ' ', text)

        # 第三步：修正因第二步可能在段落首尾产生的多余前导/尾随空格
        # 段落开头的英文/数字不应有前导空格
        text = re.sub(r'^ +', '', text)
        # 段落结尾的英文/数字不应有尾随空格
        text = re.sub(r' +$', '', text)

        return text
    return text


def process_spaces(para, mode='keep_en_words'):
    """处理段落内空格，返回 True 表示有改动"""
    if mode == 'keep_all':
        return False
    full_text = para.text
    if not full_text.strip():
        return False
    new_text = _process_spaces_text(full_text, mode)
    if new_text == full_text:
        return False
    _redistribute_text_to_runs(para.runs, new_text)
    return True


def process_paragraph(para, space_mode='keep_en_words', quote_state=None):
    """处理段落 - 简单替换按 run 做（保留格式），引号配对跨 run 做

    quote_state: 可选 dict {'dq': int, 'sq': int}，跨段落维护引号配对状态。
    传入时会就地更新，使"上一段开引号、下一段收引号"也能正确配对。
    """
    full_text = para.text
    if not full_text.strip():
        return False

    runs = para.runs
    if not runs:
        return False

    changed = False

    # 第一步：逐 run 做简单标点替换（保留所有格式）
    for run in runs:
        original = run.text
        fixed = _fix_simple_punctuation(original)
        if fixed != original:
            run.text = fixed
            changed = True

    # 第二步：引号配对需要整段处理（并可跨段延续状态）
    full_after_simple = para.text
    if quote_state is not None:
        full_after_quotes, dq, sq = _fix_quotes_whole_text(
            full_after_simple, quote_state.get('dq', 0), quote_state.get('sq', 0))
        quote_state['dq'], quote_state['sq'] = dq, sq
    else:
        full_after_quotes, _dq, _sq = _fix_quotes_whole_text(full_after_simple)

    if full_after_quotes != full_after_simple:
        _redistribute_text_to_runs(runs, full_after_quotes)
        changed = True

    # 空格处理
    if process_spaces(para, space_mode):
        changed = True

    return changed


def process_document(input_path, output_path, progress_callback=None):
    """处理文档

    progress_callback: 可选回调，签名为 callback(current, total)
    """
    logger.info('Reading: %s', input_path)
    doc = Document(input_path)

    changes = 0
    paras = doc.paragraphs
    tables = doc.tables
    total_steps = len(paras) + len(tables)
    done = 0

    # 正文段落：引号配对状态跨段延续
    quote_state = {'dq': 0, 'sq': 0}
    for i, para in enumerate(paras):
        if process_paragraph(para, quote_state=quote_state):
            changes += 1
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            logger.info('  Para %d: %s', i + 1, preview)
        done += 1
        if progress_callback and total_steps:
            progress_callback(done, total_steps)

    # 处理表格（每个单元格是独立语境，引号状态不与正文串联）
    table_changes = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        table_changes += 1
        done += 1
        if progress_callback and total_steps:
            progress_callback(done, total_steps)

    if table_changes > 0:
        logger.info('  Tables: %d cells fixed', table_changes)

    logger.info('Total: %d paragraphs + %d table cells fixed', changes, table_changes)
    doc.save(output_path)
    logger.info('Saved: %s', output_path)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format='%(message)s')
    if len(sys.argv) < 3:
        print("Usage: python punctuation.py input.docx output.docx")
        sys.exit(1)

    process_document(sys.argv[1], sys.argv[2])

# 测试用例（验证 fix_text）
# assert fix_text('会议时间:上午9:30至下午14:30,请准时参加.') == '会议时间：上午9:30至下午14:30，请准时参加。'
# assert fix_text('请于14:30前将材料发送至 report@gov.cn,逾期不候.') == '请于14:30前将材料发送至 report@gov.cn，逾期不候。'
# assert fix_text('详情请访问 https://www.example.com:8080/path 了解.') == '详情请访问 https://www.example.com:8080/path 了解。'
# assert fix_text('参照 ISO 9001:2015 执行.') == '参照 ISO 9001:2015 执行。'
# assert fix_text('每日9:00前完成巡检.') == '每日9:00前完成巡检。'
# assert fix_text("it's fine, don't worry") == "it’s fine, don’t worry"   # 撇号不配对
