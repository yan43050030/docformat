# -*- coding: utf-8 -*-
"""格式清洗：把文档里看不见、却会让排版出怪问题的脏格式清干净。

排版"疑难杂症"绝大多数不是引擎的锅，而是原文档里藏着这些东西：
字符缩放/间距、着重号、拼音指南、段落边框底纹、文本框式段落(framePr)、
域代码、书签批注、没接受的修订痕迹、手动换行/分页符、制表符与全角空格、
以及一堆空 run 和残留 rPr。它们在 Word 里肉眼难察，却会挤走缩进、
撑高行距、让字号莫名其妙变化。

支持两种范围：
- 全文清洗（scope_indices=None，含表格单元格）
- 部分行段清洗（scope_indices 传非空段序号集合，与预览界面的段序号一致）

注意：段落对齐(para_align)默认不清。排版引擎依赖对齐作为标题等类型的
识别线索，清掉会削弱自动识别；且排版时本来就会按类型重设对齐。
"""
import logging
import re

from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger('docformat.cleaner')

# 清洗项分组：(组名, [(键, 名称, 说明)])
CLEAN_GROUPS = [
    ('字符层', [
        ('char_format', '字符直接格式', '颜色、高亮、字号、加粗、斜体、下划线、删除线、上下标'),
        ('white_text', '白色文字转为黑色',
         '普通文档里白字是看不见的垃圾，转黑让它显形；'
         '识别为套打模板的文件会自动豁免，不会破坏预印占位'),
        ('char_spacing', '字符间距与缩放', 'w:spacing / w:w / w:position / w:kern，常见的"字挤在一起"元凶'),
        ('emphasis', '着重号与拼音指南', 'w:em 着重号、w:ruby 拼音指南'),
        ('empty_runs', '空 run 与残留属性', '清掉没有文字的 run 和它们携带的格式'),
    ]),
    ('段落层', [
        ('para_format', '段落直接格式', '段前段后距、左右缩进、首行缩进、行距、制表位'),
        ('styles', '段落样式重置', '把自定义/继承样式统一重置为「正文」样式'),
        ('borders_shading', '边框与底纹', 'w:pBdr 段落边框、w:shd 底纹'),
        ('frame', '文本框式段落', 'w:framePr——会让段落脱离正常文流'),
        ('para_align', '段落对齐（默认关）', '排版时会按类型重设；清掉会削弱标题自动识别'),
    ]),
    ('内容层', [
        ('fields', '域代码转文字', '把 PAGE/REF 等域的当前显示结果固化为纯文字'),
        ('bookmarks', '书签', '移除书签标记（不影响文字）'),
        ('comments', '批注引用', '移除正文里的批注锚点'),
        ('revisions', '修订痕迹', '接受插入、落实删除，清掉格式修订记录'),
        ('breaks', '手动换行与分页符', 'w:br——常见的"莫名空行/断页"来源'),
        ('whitespace', '制表符与多余空白', '制表符、全角空格、连续空格、段首尾空白'),
    ]),
]

CLEAN_ITEMS = [(k, n, d) for _g, items in CLEAN_GROUPS for k, n, d in items]

# 默认启用项：覆盖绝大多数疑难杂症，又不动可能有意义的内容
DEFAULT_CLEAN = {k: True for k, _n, _d in CLEAN_ITEMS}
DEFAULT_CLEAN['para_align'] = False      # 会削弱自动识别
DEFAULT_CLEAN['revisions'] = False       # 涉及内容取舍，让用户显式选
DEFAULT_CLEAN['fields'] = False          # 域固化不可逆，让用户显式选
DEFAULT_CLEAN['white_text'] = True       # 普通文档里白字是看不见的垃圾，清成黑字让它显形

CLEAN_LABELS = {k: n for k, n, _d in CLEAN_ITEMS}

# 排版流程里的默认自动清洗集合。
# 只收"排版引擎覆盖不到、且在规范公文里没有正当用途"的结构性垃圾——
# 它们不影响文字内容与语义，清掉零风险，却正是排版怪问题的主要来源。
# 刻意排除：
#   breaks/bookmarks/comments  会改变可见断行或破坏交叉引用、批注；
#   revisions/fields           不可逆，涉及内容取舍；
#   para_align                 是标题等类型的识别线索，清掉削弱自动识别；
#   char_format/styles         排版时本就按类型重设，无需在此重复。
AUTO_CLEAN_ITEMS = {
    'char_spacing': True,      # 字符间距/缩放/位置/字距——"字挤一起"的元凶
    'emphasis': True,          # 着重号、拼音指南
    'empty_runs': True,        # 空 run 及其残留格式
    'borders_shading': True,   # 段落边框与底纹
    'frame': True,             # 文本框式段落，会脱离正常文流
    'para_format': True,       # 制表位、字符数式缩进等（磅值由排版重设）
    'whitespace': True,        # 制表符、全角空格、连续空格
}


def auto_clean_items(preset):
    """排版流程用的清洗项：预设可用 auto_clean_items 覆盖，auto_clean=False 关闭。"""
    if not preset.get('auto_clean', True):
        return None
    items = {k: False for k, _n, _d in CLEAN_ITEMS}
    items.update(AUTO_CLEAN_ITEMS)
    override = preset.get('auto_clean_items')
    if isinstance(override, dict):
        items.update(override)
    items['para_align'] = False      # 永不在排版流程中清对齐
    return items


# ---------------- 工具 ----------------

def _iter_paragraphs(doc, scope_indices=None):
    """产出待清洗段落。

    scope_indices=None → 全文（含表格单元格）；
    否则只取 doc.paragraphs 里第 n 个非空段（n 为预览界面的段序号，从 0 起）。
    """
    if scope_indices is None:
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        return
    want = set(scope_indices)
    ai = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if ai in want:
            yield p
        ai += 1


def _rpr_of(run):
    return run._r.find(qn('w:rPr'))


def _drop(parent, tag):
    """删除 parent 下所有指定标签的子元素，返回删除个数。"""
    n = 0
    for el in parent.findall(qn(tag)):
        parent.remove(el)
        n += 1
    return n


# ---------------- 各清洗项 ----------------

def _is_white_run(run):
    c = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    return str(c) == 'FFFFFF'


def _clean_char_format(para, stat, opts=None):
    for run in para.runs:
        f = run.font
        if f.color is not None and f.color.rgb is not None:
            # 白色文字在套打模板里是"预印内容占位"，不是脏格式：
            # 清成黑字会被真的打印出来，整张套打表作废。普通文档里
            # 白字是看不见的垃圾，清成黑字让它显形才对——故按文档类型区分。
            if _is_white_run(run) and not (opts or {}).get('white_text', True):
                continue
            # 用显式黑而非 rgb=None：样式若自带颜色（如标题样式是蓝色），
            # 移除直接色会让文字继承成蓝色；公文要求黑色，显式设更确定。
            from docx.shared import RGBColor
            f.color.rgb = RGBColor(0, 0, 0)
            stat['char_format'] += 1
        if f.highlight_color is not None:
            f.highlight_color = None
            stat['char_format'] += 1
        for attr in ('size', 'bold', 'italic', 'underline', 'strike',
                     'double_strike', 'subscript', 'superscript',
                     'shadow', 'outline', 'emboss', 'imprint', 'small_caps',
                     'all_caps'):
            try:
                if getattr(f, attr) is not None:
                    setattr(f, attr, None)
                    stat['char_format'] += 1
            except (ValueError, AttributeError):
                continue


def _clean_char_spacing(para, stat, opts=None):
    for run in para.runs:
        rpr = _rpr_of(run)
        if rpr is None:
            continue
        for tag in ('w:spacing', 'w:w', 'w:position', 'w:kern', 'w:fitText'):
            stat['char_spacing'] += _drop(rpr, tag)


def _clean_emphasis(para, stat, opts=None):
    for run in para.runs:
        rpr = _rpr_of(run)
        if rpr is not None:
            stat['emphasis'] += _drop(rpr, 'w:em')
    # 拼音指南：w:ruby → 保留注音基准文字，丢掉注音
    p = para._p
    for ruby in p.findall('.//' + qn('w:ruby')):
        base = ruby.find(qn('w:rubyBase'))
        parent = ruby.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ruby)
        if base is not None:
            for child in list(base):
                parent.insert(idx, child)
                idx += 1
        parent.remove(ruby)
        stat['emphasis'] += 1


def _clean_empty_runs(para, stat, opts=None):
    for run in list(para.runs):
        r = run._r
        if run.text:
            continue
        # 保留承载图片/域/换行等结构的 run
        if len(r.findall(qn('w:drawing'))) or len(r.findall(qn('w:pict'))) \
                or len(r.findall(qn('w:object'))) or len(r.findall(qn('w:br'))) \
                or len(r.findall(qn('w:fldChar'))) or len(r.findall(qn('w:instrText'))) \
                or len(r.findall(qn('w:tab'))):
            continue
        parent = r.getparent()
        if parent is not None:
            parent.remove(r)
            stat['empty_runs'] += 1


def _clean_para_format(para, stat, opts=None):
    pf = para.paragraph_format
    touched = False
    for attr in ('space_before', 'space_after', 'left_indent', 'right_indent',
                 'first_line_indent', 'line_spacing'):
        if getattr(pf, attr) is not None:
            setattr(pf, attr, None)
            touched = True
    if pf.line_spacing_rule is not None:
        pf.line_spacing_rule = None
        touched = True
    for attr in ('keep_together', 'keep_with_next', 'page_break_before', 'widow_control'):
        try:
            if getattr(pf, attr) is not None:
                setattr(pf, attr, None)
                touched = True
        except (ValueError, AttributeError):
            continue
    ppr = para._p.find(qn('w:pPr'))
    if ppr is not None:
        # 制表位、字符数式缩进（firstLineChars 等会覆盖磅值）
        touched = bool(_drop(ppr, 'w:tabs')) or touched
        ind = ppr.find(qn('w:ind'))
        if ind is not None:
            for a in ('w:firstLineChars', 'w:leftChars', 'w:rightChars', 'w:hangingChars'):
                if ind.get(qn(a)) is not None:
                    ind.attrib.pop(qn(a), None)
                    touched = True
    if touched:
        stat['para_format'] += 1


def _clean_para_align(para, stat, opts=None):
    if para.paragraph_format.alignment is not None:
        para.paragraph_format.alignment = None
        stat['para_align'] += 1


def _clean_styles(para, stat, opts=None):
    from .font import _force_normal_style
    try:
        before = para.style.name if para.style is not None else None
    except Exception:
        before = None
    _force_normal_style(para)
    try:
        after = para.style.name if para.style is not None else None
    except Exception:
        after = None
    if before != after:
        stat['styles'] += 1


def _clean_borders_shading(para, stat, opts=None):
    # 字符级边框底纹（rPr 里的 w:bdr/w:shd）——set_font 不会清，
    # 排版后残留会让文字带着黄底/方框，是常见的"排版后还是花的"来源
    for run in para.runs:
        rpr = _rpr_of(run)
        if rpr is not None:
            stat['borders_shading'] += _drop(rpr, 'w:shd') + _drop(rpr, 'w:bdr')
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        return
    stat['borders_shading'] += _drop(ppr, 'w:pBdr') + _drop(ppr, 'w:shd')


def _clean_frame(para, stat, opts=None):
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        return
    stat['frame'] += _drop(ppr, 'w:framePr')


def _clean_fields(para, stat, opts=None):
    """把域代码折叠成它当前显示的文字（丢弃 instrText 与域字符）。"""
    p = para._p
    children = list(p)
    depth = 0
    to_remove = []
    for el in children:
        if el.tag != qn('w:r'):
            continue
        fld = el.find(qn('w:fldChar'))
        instr = el.find(qn('w:instrText'))
        if fld is not None:
            t = fld.get(qn('w:fldCharType'))
            if t == 'begin':
                depth += 1
                to_remove.append(el)
                continue
            if t == 'separate':
                to_remove.append(el)
                continue
            if t == 'end':
                depth = max(0, depth - 1)
                to_remove.append(el)
                continue
        if instr is not None and depth > 0:
            to_remove.append(el)
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            stat['fields'] += 1
    # 简单域（w:fldSimple）：保留其中的 run，去掉域壳
    for fs in p.findall(qn('w:fldSimple')):
        parent = fs.getparent()
        if parent is None:
            continue
        idx = list(parent).index(fs)
        for child in list(fs):
            parent.insert(idx, child)
            idx += 1
        parent.remove(fs)
        stat['fields'] += 1


def _clean_bookmarks(para, stat, opts=None):
    p = para._p
    for tag in ('w:bookmarkStart', 'w:bookmarkEnd'):
        for el in p.findall(qn(tag)):
            p.remove(el)
            stat['bookmarks'] += 1


def _clean_comments(para, stat, opts=None):
    p = para._p
    for tag in ('w:commentRangeStart', 'w:commentRangeEnd'):
        for el in p.findall(qn(tag)):
            p.remove(el)
            stat['comments'] += 1
    for run in list(para.runs):
        if run._r.find(qn('w:commentReference')) is not None:
            parent = run._r.getparent()
            if parent is not None:
                parent.remove(run._r)
                stat['comments'] += 1


def _clean_revisions(para, stat, opts=None):
    """接受修订：保留插入内容、删除被删内容、去掉格式修订记录。"""
    p = para._p
    # 删除标记 w:del：整个 run 连同内容移除
    for d in p.findall('.//' + qn('w:del')):
        parent = d.getparent()
        if parent is not None:
            parent.remove(d)
            stat['revisions'] += 1
    # 插入标记 w:ins：提升其中的 run，去掉标记壳
    for ins in p.findall('.//' + qn('w:ins')):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        stat['revisions'] += 1
    # 格式修订记录
    for tag in ('w:pPrChange', 'w:rPrChange', 'w:sectPrChange', 'w:tblPrChange'):
        for el in p.findall('.//' + qn(tag)):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
                stat['revisions'] += 1


def _clean_breaks(para, stat, opts=None):
    """移除手动换行/分页符（w:br）。分页符所在的空段落由排版流程另行处理。"""
    for run in para.runs:
        for br in run._r.findall(qn('w:br')):
            run._r.remove(br)
            stat['breaks'] += 1


_MULTI_SPACE = re.compile(r'[ 　]{2,}')


def _clean_whitespace(para, stat, opts=None):
    changed = False
    for run in para.runs:
        # 制表符元素
        for tab in run._r.findall(qn('w:tab')):
            run._r.remove(tab)
            changed = True
        t = run.text
        if not t:
            continue
        new = t.replace('\t', ' ')
        new = new.replace('　', ' ')          # 全角空格
        new = _MULTI_SPACE.sub(' ', new)
        if new != t:
            run.text = new
            changed = True
    # 段首尾空白
    runs = para.runs
    if runs:
        first = runs[0].text
        if first != first.lstrip():
            runs[0].text = first.lstrip()
            changed = True
        last = runs[-1].text
        if last != last.rstrip():
            runs[-1].text = last.rstrip()
            changed = True
    if changed:
        stat['whitespace'] += 1


_CLEANERS = [
    ('char_format', _clean_char_format),
    ('char_spacing', _clean_char_spacing),
    ('emphasis', _clean_emphasis),
    ('fields', _clean_fields),
    ('revisions', _clean_revisions),
    ('bookmarks', _clean_bookmarks),
    ('comments', _clean_comments),
    ('breaks', _clean_breaks),
    ('whitespace', _clean_whitespace),
    ('empty_runs', _clean_empty_runs),
    ('borders_shading', _clean_borders_shading),
    ('frame', _clean_frame),
    ('para_format', _clean_para_format),
    ('para_align', _clean_para_align),
    ('styles', _clean_styles),
]


def looks_like_overprint(doc):
    """判断文档是否为套打表单：白色文字 + 白色/无框线的表格。

    套打表靠"白字占位、白线占格"复刻预印纸的版式，两者同时出现基本不会
    是巧合。命中就豁免白字归一，免得把人家的套打模板清成一张会全印出来
    的废纸。
    """
    has_white_text = False
    for p in doc.paragraphs:
        if any(r.text.strip() and _is_white_run(r) for r in p.runs):
            has_white_text = True
            break
    if not has_white_text:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if any(r.text.strip() and _is_white_run(r) for r in p.runs):
                            has_white_text = True
                            break
    if not has_white_text:
        return False
    for t in doc.tables:
        tblPr = t._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            continue
        b = tblPr.find(qn('w:tblBorders'))
        if b is None:
            continue
        for el in b:
            color = el.get(qn('w:color'))
            val = el.get(qn('w:val'))
            if color == 'FFFFFF' or val == 'none':
                return True
    return False


def clean_document(doc, items=None, scope_indices=None, protect_media=True):
    """清洗文档格式，返回各项改动计数 {键: 次数}。

    items: {键: bool}，None 表示用 DEFAULT_CLEAN；
    scope_indices: 非空段序号集合，None 表示全文；
    protect_media: 跳过含图片/嵌入对象的段落（避免破坏图片布局）。
    """
    from .paragraph import paragraph_has_media
    opts = dict(DEFAULT_CLEAN)
    if items:
        opts.update(items)
    active = [(k, fn) for k, fn in _CLEANERS if opts.get(k)]
    if not active:
        return {}

    # 套打表单自动豁免白字归一——它的白字是有效数据，不是垃圾
    if opts.get('white_text', True) and looks_like_overprint(doc):
        opts['white_text'] = False
        logger.info('识别为套打表单，已保护白色预印占位文字')

    stat = {k: 0 for k, _fn in _CLEANERS}
    stat.setdefault('white_text', 0)
    for para in _iter_paragraphs(doc, scope_indices):
        if protect_media and paragraph_has_media(para):
            continue
        for _k, fn in active:
            try:
                fn(para, stat, opts)
            except Exception as exc:      # 单段失败不影响整篇
                logger.warning('清洗段落失败(%s): %s', _k, exc)
    return {k: v for k, v in stat.items() if v}


def format_clean_summary(stat):
    """把清洗计数翻成人话，供日志/提示显示。"""
    if not stat:
        return '未发现需要清洗的格式'
    parts = ['{} {} 处'.format(CLEAN_LABELS.get(k, k), v)
             for k, v in sorted(stat.items(), key=lambda kv: -kv[1])]
    return '、'.join(parts)


def clean_file(input_path, output_path, items=None, scope_indices=None):
    """独立清洗一个文件并另存，返回改动计数。"""
    from docx import Document
    from .paragraph import sanitize_document
    doc = Document(input_path)
    sanitize_document(doc)
    stat = clean_document(doc, items=items, scope_indices=scope_indices)
    doc.save(output_path)
    return stat
