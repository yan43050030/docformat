# -*- coding: utf-8 -*-
"""公文用语检查 + 文种骨架生成。

只查"有明文规定、能判定对错"的东西——数字用法（GB/T 15835）、文种与结语
的搭配（《党政机关公文处理工作条例》）、标题与附件说明的格式。文风一律
不碰："语言要简洁"这种建议正是让工具变讨厌的部分，而且判不了对错。

全部靠正则 + 词表 + 段落类型，不联网、不上模型：软件要在离线的信创机器上
跑，文件还可能涉密。

误报是这类功能唯一的死法——报错了几次人就把它关了，从此再不打开。所以：
  · 先遮蔽后匹配。引号里的引文、书名号、文号、日期、金额、电话一律不查，
    遮蔽时用等长占位符，字符下标才不会错位。
  · 用段落类型限定范围。标题、正文、附件说明的规矩完全不同，一条规则只在
    它该管的段落上跑。
  · 拿不准的宁可不报。规则分 high/medium 两档，medium 的默认不开。
每加一条规则都要在 WORDING_CASES 里配正例**和反例**，反例是"正确用法必须
不报"——回归测试盯的就是这个。
"""
import logging
import re

logger = logging.getLogger('docformat.wording')

# ---------------------------------------------------------------- 文种表
# 一张表两用：查用语搭配靠它，生成文种骨架也靠它。
#   closing   —— 该文种的标准结语（正则），缺了就提示
#   forbidden —— 明确不该出现在该文种里的表述
#   note      —— 提示语里给人看的一句话
DOC_KINDS = {
    '请示': {
        'closing': r'(妥否|当否|可否|如无不妥)[，,]?\s*请?(批示|指示|示复|审批|批复)'
                   r'|请\s*(批示|指示|审批|批准|示复)',
        'forbidden': [(r'特此报告', '请示的结语不能用「特此报告」'),
                      (r'特此通知', '请示的结语不能用「特此通知」')],
        'note': '请示要有「妥否，请批示」一类的结语，且一文一事',
        'skeleton': ['事由', '请示事项', '妥否，请批示。'],
    },
    '报告': {
        'closing': r'特此报告',
        'forbidden': [(r'请(予以?)?(批准|审批|批示|指示)', '报告不得夹带请示事项，'
                                                          '要请示请另行行文'),
                      (r'(妥否|当否)[，,]?\s*请', '报告不得夹带请示事项')],
        'note': '报告只报不请，结语用「特此报告」',
        'skeleton': ['事由', '工作情况', '特此报告。'],
    },
    '通知': {
        'closing': r'特此通知',
        'forbidden': [],
        'note': '通知的结语一般用「特此通知」',
        'skeleton': ['事由', '通知事项', '特此通知。'],
    },
    '函': {
        'closing': r'特此函(达|告|复)|请(予以?)?(支持|协助|函复)|盼复',
        'forbidden': [(r'特此通知', '函是平行文，不能用下行文的「特此通知」')],
        'note': '函用于不相隶属机关之间，结语用「特此函达」「请予支持」一类',
        'skeleton': ['事由', '商洽事项', '特此函达。'],
    },
    '批复': {
        'closing': r'特此批复|此复',
        'forbidden': [],
        'note': '批复要针对来文，结语用「特此批复」',
        'skeleton': ['引述来文', '批复意见', '特此批复。'],
    },
    '意见': {'closing': r'', 'forbidden': [], 'note': '', 'skeleton': ['事由', '意见内容']},
    '纪要': {'closing': r'', 'forbidden': [], 'note': '', 'skeleton': ['会议概况', '议定事项']},
    '决定': {'closing': r'', 'forbidden': [], 'note': '', 'skeleton': ['事由', '决定事项']},
    '通报': {'closing': r'特此通报', 'forbidden': [], 'note': '',
             'skeleton': ['事由', '通报内容', '特此通报。']},
}

# 标题末尾的文种：「关于××的请示」。这个正则很可靠，是文种搭配检查的入口
_KIND_RE = re.compile(r'的\s*(' + '|'.join(DOC_KINDS) + r')\s*$')

# ---------------------------------------------------------------- 遮蔽
# 这些地方一律不查：引文要照录原文、文号日期金额是数据、书名号内是作品名
_MASK_PATTERNS = [
    r'“[^”]{0,200}”',                    # 引号内的引文
    r'《[^》]{0,200}》',                   # 书名号
    r'〔\s*\d{4}\s*〕\s*\d+\s*号',          # 发文字号
    r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日',   # 日期
    r'\d+(\.\d+)?\s*(万|亿)?\s*元',         # 金额
    r'\d[\d\- ]{5,}\d',                   # 电话、长串数字
    r'(?:https?|ftp)://\S+',
    r'[\w.+-]+@[\w-]+\.[\w.-]+',
]
_MASK_RE = re.compile('|'.join(_MASK_PATTERNS))


def mask(text):
    """把不该查的地方换成等长的 \\x00，字符下标保持不变。"""
    out = list(text)
    for m in _MASK_RE.finditer(text):
        for i in range(m.start(), m.end()):
            out[i] = '\x00'
    return ''.join(out)


# ---------------------------------------------------------------- 规则表
# scope 为空表示不限段落类型；level: warn（明确错）/ hint（建议核对）
# fix 给了就能自动改（必须是"改了一定对"的，拿不准的一律不给）
def _fix_ling(m):
    return m.group(0).replace('○', '〇').replace('O', '〇').replace('o', '〇')


RULES = [
    # ---- 数字用法（GB/T 15835）----
    {
        'id': 'ling_char', 'group': '数字用法', 'level': 'warn', 'confidence': 'high',
        'scope': (), 'fix': _fix_ling,
        # 汉字年份里的"〇"是 U+3007，常被打成字母 O 或几何圈 ○，肉眼几乎分不出
        'pattern': r'(?<![0-9A-Za-z])[一二三四五六七八九〇○Oo]*[○Oo][一二三四五六七八九〇○Oo]*(?=年)',
        'message': '汉字年份里的「〇」应为 U+3007，不是字母 O 或几何符号 ○',
    },
    {
        'id': 'fullwidth_digit', 'group': '数字用法', 'level': 'warn', 'confidence': 'high',
        'scope': (), 'pattern': r'[０-９]+',
        'fix': lambda m: m.group(0).translate({ord(c): ord('0') + i
                                               for i, c in enumerate('０１２３４５６７８９')}),
        'message': '数字应用半角，不用全角',
    },
    {
        'id': 'approx_number', 'group': '数字用法', 'level': 'warn', 'confidence': 'high',
        # 「第3、4条」「第1、2项」是条款枚举，不是概数——用 (?<!第) 挡掉，
        # 且不把「条/项」列进后随字：那两个字后面几乎总是枚举
        'scope': (), 'pattern': r'(?<!第)\d\s*[、]\s*\d\s*(?=[个人天次名件处家])',
        'message': '相邻两数字连用表示概数时应用汉字、中间不加顿号（如「三四个」）',
    },
    {
        'id': 'date_arabic', 'group': '数字用法', 'level': 'warn', 'confidence': 'high',
        'scope': ('date',),
        'pattern': r'[一二三四五六七八九十〇]{2,}\s*年',
        'message': '成文日期应用阿拉伯数字（GB/T 9704-2012）',
    },
    {
        'id': 'docnum_arabic', 'group': '数字用法', 'level': 'warn', 'confidence': 'high',
        'scope': ('docnum',),
        'pattern': r'〔\s*[一二三四五六七八九〇]{2,}\s*〕',
        'message': '发文字号的年份应用阿拉伯数字，如〔2026〕5号',
    },
    # ---- 格式化用语 ----
    {
        'id': 'title_punct', 'group': '格式用语', 'level': 'warn', 'confidence': 'high',
        'scope': ('title',), 'pattern': r'[。；;，,、！!？?：:]$',
        'message': '公文标题末尾不加标点（书名号、引号除外）',
    },
    {
        'id': 'recipient_colon', 'group': '格式用语', 'level': 'warn', 'confidence': 'high',
        'scope': ('recipient',), 'pattern': r'[^：:]$',
        'message': '主送机关末尾应加全角冒号「：」',
    },
    {
        'id': 'attach_seq', 'group': '格式用语', 'level': 'warn', 'confidence': 'high',
        'scope': ('attachment',), 'pattern': r'(?<=附件：)\s*\d+\s*[、，,]',
        'message': '附件说明的序号后用圆点，如「附件：1.××」',
    },
    {
        'id': 'attach_tail', 'group': '格式用语', 'level': 'warn', 'confidence': 'high',
        'scope': ('attachment',), 'pattern': r'[。；;]$',
        'message': '附件说明末尾不加标点',
    },
    # ---- 易混词（默认不开：要靠上下文，最容易误报）----
    {
        'id': 'qi_ta', 'group': '易混词', 'level': 'hint', 'confidence': 'medium',
        'scope': (), 'pattern': r'其它', 'fix': lambda m: '其他',
        'message': '公文中统一用「其他」，不用「其它」',
    },
    {
        'id': 'jie_zhi', 'group': '易混词', 'level': 'hint', 'confidence': 'medium',
        'scope': (), 'pattern': r'截止\s*(?=\d|[一二三四五六七八九十]|目前|现在|今)',
        'message': '表示"到某个时间为止"用「截至」，「截止」是停止的意思',
    },
    {
        'id': 'bu_shu', 'group': '易混词', 'level': 'hint', 'confidence': 'medium',
        'scope': (), 'pattern': r'布署', 'fix': lambda m: '部署',
        'message': '应为「部署」',
    },
    {
        'id': 'yi_zhi', 'group': '易混词', 'level': 'hint', 'confidence': 'medium',
        'scope': (), 'pattern': r'以至于', 'fix': lambda m: '以致',
        'message': '表示不好的结果用「以致」',
    },
    {
        'id': 'fa_ding', 'group': '易混词', 'level': 'hint', 'confidence': 'medium',
        'scope': (), 'pattern': r'法订', 'fix': lambda m: '法定',
        'message': '应为「法定」',
    },
]

# ---- 错别字：一张大表，合成一条正则 ----
# 单独成组、默认**开**、可自动改——错的那一形不成词，出现即错，
# 和"要看上下文"的易混词是两回事，不能混在一组里。
# 234 条合成一个正则一次扫完，比 234 条规则各扫一遍快得多；
# 长的排前面，「迫不急待」不会被更短的条目截胡。
from .typos import TYPOS as _TYPOS, pattern_of as _typo_pat   # noqa: E402

# 每条各自带护栏（否定环视），长的排前面免得被短条目截胡
_TYPO_RE = '|'.join(_typo_pat(w) for w in
                    sorted(_TYPOS, key=len, reverse=True))

RULES.append({
    'id': 'typo', 'group': '错别字', 'level': 'warn', 'confidence': 'high',
    'scope': (), 'pattern': _TYPO_RE,
    'fix': lambda m: _TYPOS.get(m.group(0), m.group(0)),
    'message': '错别字',
    'describe': lambda words: '错别字（{} 处）：{}'.format(
        len(words), '、'.join('{}→{}'.format(w, _TYPOS.get(w, ''))
                              for w in sorted(words)[:6])),
})

RULE_BY_ID = {r['id']: r for r in RULES}
GROUPS = []
for _r in RULES:
    if _r['group'] not in GROUPS:
        GROUPS.append(_r['group'])
GROUPS.append('文种搭配')

# 默认开哪些组：medium 的易混词默认关——它最容易误报，让人自己开
DEFAULT_GROUPS = {g: (g != '易混词') for g in GROUPS}


def user_rules_path():
    """用户自己的词表：各单位有自己的用语禁忌，内置表不可能通用"""
    import os
    from app.template_common import config_dir
    return os.path.join(config_dir(), 'wording_rules_user.json')


def load_user_rules():
    """读用户词表，格式与 RULES 同构（fix 只支持"整串替换"的字符串）。"""
    import json
    import os
    p = user_rules_path()
    if not os.path.exists(p):
        return []
    try:
        with open(p, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (IOError, OSError, ValueError) as exc:
        logger.warning('用户词表读取失败：%s', exc.__class__.__name__)
        return []
    out = []
    for i, d in enumerate(data if isinstance(data, list) else []):
        try:
            rule = {
                'id': str(d.get('id') or 'user_%d' % i),
                'group': d.get('group') or '自定义',
                'level': d.get('level') or 'hint',
                'confidence': 'medium',
                'scope': tuple(d.get('scope') or ()),
                'pattern': d['pattern'],
                'message': d.get('message') or '不符合本单位用语规范',
            }
            re.compile(rule['pattern'])
            if d.get('replace'):
                rule['fix'] = (lambda rep: (lambda m: rep))(str(d['replace']))
            out.append(rule)
        except (KeyError, re.error, TypeError) as exc:
            logger.warning('用户词表第 %d 条无效（%s），已跳过', i + 1,
                           exc.__class__.__name__)
    return out


# ---------------------------------------------------------------- 检查
def doc_kind(texts):
    """从标题里认文种；认不出返回 ''。"""
    for t in texts[:6]:
        # 标题末尾可能误加了标点（那本身也是一条要报的错），
        # 认文种时先去掉，别因为一个句号就整块检查都跳过
        m = _KIND_RE.search((t or '').strip().rstrip('。；;，,、！!？?：: '))
        if m:
            return m.group(1)
    return ''


def check_wording(doc, groups=None, detect_types=None, extra_rules=None):
    """返回 findings，结构与 compliance.check_compliance 一致：
    [{'level','item','detail','fix_key','locations'}]

    detect_types: {非空段序号: 段落类型}，没有就只跑不限类型的规则。
    """
    on = dict(DEFAULT_GROUPS)
    if groups:
        on.update({k: bool(v) for k, v in groups.items()})
    rules = [r for r in list(RULES) + list(extra_rules or load_user_rules())
             if on.get(r['group'], True)]

    paras = [p for p in doc.paragraphs]
    texts, idx_map = [], {}
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t:
            idx_map[len(texts)] = i
            texts.append(t)

    hits = {}          # rule_id -> [(非空段序号, 命中的文字)]
    for ai, raw in enumerate(texts):
        ptype = (detect_types or {}).get(ai)
        masked = mask(raw)
        for r in rules:
            if r['scope'] and ptype not in r['scope']:
                continue
            for m in re.finditer(r['pattern'], masked):
                got = raw[m.start():m.end()]
                if not got.strip():
                    continue
                hits.setdefault(r['id'], []).append((ai, got))

    findings = []
    for r in rules:
        got = hits.get(r['id'])
        if not got:
            continue
        where = sorted({ai for ai, _t in got})
        words = {t for _a, t in got}
        if r.get('describe'):
            detail = r['describe'](words)
        else:
            detail = '{}（{} 处，如「{}」）'.format(
                r['message'], len(got), '、'.join(sorted(words)[:3]))
        findings.append({
            'level': r['level'], 'item': '用语·{}'.format(r['group']),
            'detail': detail,
            'locations': where,
            **({'fix_key': 'wording:' + r['id']} if r.get('fix') else {}),
        })

    if on.get('文种搭配', True):
        findings.extend(_check_kind(texts, detect_types))
    return findings


def _check_kind(texts, detect_types):
    """文种与结语的搭配。认不出文种就整段跳过——宁可不报。"""
    kind = doc_kind(texts)
    if not kind:
        return []
    spec = DOC_KINDS.get(kind) or {}
    body = '\n'.join(texts)
    masked = mask(body)
    out = []
    for pat, why in spec.get('forbidden') or []:
        if re.search(pat, masked):
            out.append({'level': 'warn', 'item': '用语·文种搭配',
                        'detail': '这是「{}」：{}'.format(kind, why)})
    closing = spec.get('closing')
    # 「通知/通报」结尾不写"特此通知"极常见，也说不上错，不查缺失；
    # 正文太短（只有标题）也不判——那多半是还没写完
    if kind in ('通知', '通报') or len(texts) < 3:
        closing = ''
    if closing and not re.search(closing, masked):
        out.append({'level': 'hint', 'item': '用语·文种搭配',
                    'detail': '这是「{}」，没找到规范的结语。{}'.format(
                        kind, spec.get('note') or '')})
    if kind == '请示':
        n = len(re.findall(r'(拟请|恳请|请予以?批准|请审批)', masked))
        if n >= 2:
            out.append({'level': 'hint', 'item': '用语·文种搭配',
                        'detail': '请示应一文一事，正文里出现了 {} 处请求事项，'
                                  '请确认是否该拆成多份'.format(n)})
    return out


_REV_ID = [1000]


def _tracked_replace(para, run, pieces):
    """把 run 拆开，改动处以 **Word 修订** 的形式落进去。

    pieces 是 [(原文, 新文或 None)]，None 表示这段不动。改动写成
        <w:del><w:r><w:delText>错的</w:delText></w:r></w:del>
        <w:ins><w:r><w:t>对的</w:t></w:r></w:ins>
    在 Word/WPS 里就是一处修订：谁改的、改成什么，一目了然，
    点"接受/拒绝"就能定夺。这比默默替换要紧——用语这种事，最终得由人拍板。
    """
    from copy import deepcopy
    from datetime import datetime, timezone
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    r = run._r
    parent = r.getparent()
    at = list(parent).index(r)
    now = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    def _mk_run(text, tag='w:t'):
        nr = deepcopy(r)
        for t in nr.findall(_qn('w:t')) + nr.findall(_qn('w:delText')):
            nr.remove(t)
        el = OxmlElement(tag)
        el.set(_qn('xml:space'), 'preserve')
        el.text = text
        nr.append(el)
        return nr

    def _wrap(tag, child):
        w = OxmlElement(tag)
        _REV_ID[0] += 1
        w.set(_qn('w:id'), str(_REV_ID[0]))
        w.set(_qn('w:author'), '公文格式工具')
        w.set(_qn('w:date'), now)
        w.append(child)
        return w

    offset = 0
    for old_text, new_text in pieces:
        if new_text is None:
            parent.insert(at + offset, _mk_run(old_text))
            offset += 1
            continue
        parent.insert(at + offset, _wrap('w:del', _mk_run(old_text, 'w:delText')))
        parent.insert(at + offset + 1, _wrap('w:ins', _mk_run(new_text)))
        offset += 2
    parent.remove(r)


def _fixable_rules(fix_keys):
    ids = {k.split(':', 1)[1] for k in (fix_keys or []) if k.startswith('wording:')}
    return [r for r in list(RULES) + list(load_user_rules())
            if r['id'] in ids and r.get('fix')]


def _spans_in(text, rules, ptype):
    """这段文字里要改哪几处：[(起, 止, 新文)]，按位置升序。

    预览和真正下笔共用这一个函数。两边各写一套匹配，迟早会出现
    "预览说改三处、实际改了两处"——那比不给预览还糟。
    """
    if not text:
        return []
    # 遮蔽是按整段算的，run 里再算一次代价太高；这里只保护
    # run 内可见的引号/书名号，够用
    if re.search(r'[“”《》]', text):
        return []
    spans = []
    for r in rules:
        if r['scope'] and ptype not in r['scope']:
            continue
        for m in re.finditer(r['pattern'], text):
            if any(not (m.end() <= a or m.start() >= b) for a, b, _t in spans):
                continue        # 与已认领的区间重叠，跳过
            spans.append((m.start(), m.end(), r['fix'](m)))
    spans.sort()
    return spans


def preview_wording(doc, fix_keys, detect_types=None, max_paras=400):
    """认可了这些规则之后，哪几段会变成什么样。

    返回 [{'index': 段号, 'before': 原文, 'after': 改后,
           'marks': [(起, 止, 新词)]}]，只给真的会动的段落。
    marks 的坐标是相对**整段**的，界面照它把错词标出来。
    """
    rules = _fixable_rules(fix_keys)
    if not rules:
        return []
    out = []
    ai = -1
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        ai += 1
        if ai >= max_paras:
            break
        ptype = (detect_types or {}).get(ai)
        before, after, marks = [], [], []
        grown = 0       # 改后文本已经写了多少字，标记的坐标按它算
        for run in para.runs:
            text = run.text or ''
            before.append(text)
            cur = 0
            for a, b, t in _spans_in(text, rules, ptype):
                head = text[cur:a]
                after.append(head)
                grown += len(head)
                marks.append((grown, grown + len(t), text[a:b], t))
                after.append(t)
                grown += len(t)
                cur = b
            tail = text[cur:]
            after.append(tail)
            grown += len(tail)
        b_txt, a_txt = ''.join(before), ''.join(after)
        if b_txt != a_txt:
            out.append({'index': ai, 'before': b_txt, 'after': a_txt,
                        'marks': marks})
    return out


def apply_wording_fixes(doc, fix_keys, detect_types=None, revision=True):
    """按 fix_key（wording:规则id）替换文字，返回改动处数。

    revision=True（默认）走 Word 修订：改动不是悄悄替换掉，而是留成一处
    修订痕迹，在 Word/WPS 里能看到"把甲改成了乙"，点接受或拒绝由人定。
    用语对错终究要人拍板，默默改掉是不负责任的。
    只有给了 fix 的规则才动手——那些是"改了一定对"的。
    """
    rules = _fixable_rules(fix_keys)
    if not rules:
        return 0

    ai = -1
    n = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        ai += 1
        ptype = (detect_types or {}).get(ai)
        for run in list(para.runs):
            text = run.text
            spans = _spans_in(text, rules, ptype)
            if not spans:
                continue
            n += len(spans)
            if not revision:
                out, cur = [], 0
                for a, b, t in spans:
                    out.append(text[cur:a]); out.append(t); cur = b
                out.append(text[cur:])
                run.text = ''.join(out)
                continue
            pieces, cur = [], 0
            for a, b, t in spans:
                if a > cur:
                    pieces.append((text[cur:a], None))
                pieces.append((text[a:b], t))
                cur = b
            if cur < len(text):
                pieces.append((text[cur:], None))
            _tracked_replace(para, run, pieces)
    if n:
        logger.info('用语修正 %d 处（%s）', n, '修订方式' if revision else '直接替换')
    return n


# ---------------------------------------------------------------- 骨架生成
def build_skeleton(kind, issuer='', recipient='', subject='', docnum=''):
    """按文种铺一份骨架，返回 [(段落类型, 文字)]。

    与用语检查共用 DOC_KINDS：那边查结语对不对，这边直接把对的结语写进去。
    """
    spec = DOC_KINDS.get(kind)
    if spec is None:
        raise ValueError('不认识的文种：{}（可用：{}）'.format(
            kind, '、'.join(DOC_KINDS)))
    subject = (subject or '××工作').strip()
    title = '关于{}的{}'.format(subject.lstrip('关于').rstrip('的'), kind)
    out = []
    if docnum:
        out.append(('docnum', docnum))
    out.append(('title', title))
    out.append(('recipient', (recipient or '××××').rstrip('：:') + '：'))
    for i, part in enumerate(spec['skeleton']):
        if part.endswith('。'):
            out.append(('closing', part))
        elif i == 0:
            out.append(('body', '为{}，现将有关事项报告如下。'.format(subject)
                        if kind == '报告' else
                        '为{}，现请示如下。'.format(subject) if kind == '请示' else
                        '为{}，现将有关事项通知如下。'.format(subject)))
        else:
            out.append(('body', '一、{}'.format(part)))
            out.append(('body', '（此处填写具体内容）'))
    out.append(('signature', issuer or '××××'))
    out.append(('date', '2026年1月1日'))
    return out
