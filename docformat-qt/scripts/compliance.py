# -*- coding: utf-8 -*-
"""公文合规性检查：对照"当前预设"完整核对文档版式，报告偏差，并可对认可的偏差自动修正。

设计原则
--------
1. 标准来自用户选中的预设，而非死国标——改预设，检查标准跟着变。
2. **检查面必须逼近排版面**：预设里定义了明确规格的项，检查都要覆盖，
   否则"检查通过"不等于"排版合规"，检查就失去意义。因此段落级检查是
   逐段按识别类型 × 逐属性（字体/字号/加粗/对齐/首行缩进/行距/段距）
   对照预设，与排版引擎 format_paragraph 用同一套口径。
3. 交互式修正：可自动修正的偏差带 fix_key，用户勾选认可的项，
   apply_compliance_fixes 只对认可项动手，其余保持原样。

天生"只能做、无法核对"的排版动作（样式清洗、结构性空行、盖章落款布局）
不纳入检查，避免装样子；这些只在"智能一键处理"里执行。
"""
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

# ---------------- 检查项定义 ----------------
# 供 UI 生成勾选面板；分组便于展示
CHECK_GROUPS = [
    ('页面', [
        ('margins', '页边距'),
        ('paper', '纸张大小'),
        ('grid', '页面网格（每页行数 × 每行字数）'),
        ('page_number', '页码（有无 / 字体 / 字号）'),
    ]),
    ('段落（逐段按识别类型对照预设）', [
        ('font', '字体'),
        ('size', '字号'),
        ('bold', '加粗'),
        ('align', '对齐方式'),
        ('indent', '首行缩进'),
        ('line_spacing', '行距'),
        ('spacing', '段前 / 段后间距'),
    ]),
    ('内容', [
        ('structure', '结构完整性（标题 / 主送机关 / 成文日期）'),
        ('numbering', '序号层次是否统一'),
        ('punctuation', '英文 / 不规范标点'),
    ]),
]

CHECK_ITEMS = [item for _g, items in CHECK_GROUPS for item in items]

DEFAULT_OPTIONS = {k: True for k, _ in CHECK_ITEMS}

# 旧版本的检查项键 → 新键（QSettings 里可能存着旧键，做兼容映射）
_LEGACY_OPTION_MAP = {
    'body_font': ('font', 'size'),
    'title_center': ('align',),
}

# 段落级检查的属性 → 中文名
ATTR_LABELS = {
    'font': '字体',
    'size': '字号',
    'bold': '加粗',
    'align': '对齐方式',
    'indent': '首行缩进',
    'line_spacing': '行距',
    'spacing': '段前/段后间距',
}

# 段落类型 → 中文名（与预览界面保持一致）
TYPE_LABELS = {
    'security': '密级', 'docnum': '发文字号', 'title': '标题',
    'subtitle': '副标题', 'recipient': '主送机关',
    'heading1': '一级标题', 'heading2': '二级标题', 'heading3': '三级标题',
    'heading4': '四级标题', 'body': '正文', 'signature': '署名',
    'date': '日期', 'attachment': '附件说明', 'attachment_label': '附件标识',
    'closing': '结尾', 'roster': '组成人员名单',
    'copynum': '份号', 'urgency': '紧急程度', 'signatory': '签发人',
    'cc': '抄送', 'issuer': '印发机关', 'caption': '图表题注',
}

ALIGN_LABELS = {'left': '左对齐', 'center': '居中', 'right': '右对齐',
                'justify': '两端对齐'}

_ALIGN_MAP = {
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# 排版引擎 format_paragraph 的默认行距（fmt 未指定 line_spacing 时用它）
DEFAULT_LINE_SPACING_PT = 28

# roster（组成人员名单）在排版时只设缩进和行距，其余保持原样，
# 检查也只核对这两项，避免误报。
_ROSTER_ATTRS = {'indent', 'line_spacing'}

# 非段落级的修正说明
FIX_LABELS = {
    'margins': '把页边距改为预设值',
    'paper': '把纸张改为 A4',
    'grid': '把页面网格改为预设值',
    'page_number': '按预设重设页码',
    'punctuation': '修复英文/不规范标点',
}


def _cm(v):
    return round(v, 2) if v is not None else None


def normalize_options(options):
    """把可能含旧键的 options 归一到新键集合。"""
    if not options:
        return None
    opts = {}
    for k, v in options.items():
        if k in _LEGACY_OPTION_MAP:
            for nk in _LEGACY_OPTION_MAP[k]:
                # 旧键为真时不强行覆盖新键的显式设置
                opts.setdefault(nk, v)
        else:
            opts[k] = v
    return opts


def _detect_types(doc, preset):
    """返回 [(段序号, paragraph, ptype)]，仅含非空段，供检查与修正共用。"""
    from .detector import detect_para_type, _compile_rules, _build_text_context
    rules = _compile_rules(preset.get('detect_rules'))
    flags = {
        'subtitle_enabled': preset.get('subtitle_enabled', False),
        'header_elements': preset.get('header_elements', False),
        'record_elements': preset.get('record_elements', False),
    }
    all_texts, idx_map = _build_text_context(doc)
    result = []
    prev = None
    total = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        ai = idx_map.get(i)
        ptype = detect_para_type(t, i, total, p.paragraph_format.alignment,
                                 all_texts, all_texts_index=ai, prev_para_type=prev,
                                 rules=rules, flags=flags)
        result.append((i + 1, p, ptype))
        prev = ptype
    return result


# ---------------- 段落级：实际值读取 ----------------

def _first_run(para):
    for r in para.runs:
        if r.text.strip():
            return r
    return None


def _actual_font(para):
    run = _first_run(para)
    if run is None:
        return None
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn('w:eastAsia'))


def _actual_font_en(para):
    """西文字体（数字/英文用），预览里数字要按它渲染才与真实输出一致"""
    run = _first_run(para)
    if run is None:
        return None
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn('w:ascii'))


def _actual_size(para):
    run = _first_run(para)
    if run is None or run.font.size is None:
        return None
    return run.font.size.pt


def _actual_bold(para):
    run = _first_run(para)
    return None if run is None else run.font.bold


def _actual_align(para):
    return para.paragraph_format.alignment


def _actual_indent(para):
    ind = para.paragraph_format.first_line_indent
    return None if ind is None else ind.pt


def _actual_line_spacing(para):
    ls = para.paragraph_format.line_spacing
    if ls is None:
        return None
    return ls.pt if hasattr(ls, 'pt') else None


def _actual_spacing(para):
    pf = para.paragraph_format
    before = pf.space_before.pt if pf.space_before is not None else None
    after = pf.space_after.pt if pf.space_after is not None else None
    return (before, after)


def _expected_line_spacing(fmt):
    return fmt.get('line_spacing', DEFAULT_LINE_SPACING_PT)


def _compare_attr(attr, fmt, para, ptype):
    """比较单个属性，返回 (是否合规, 实际值描述, 期望值描述)。

    返回 None 表示该属性对此类型不适用/无法判定，跳过。
    """
    if attr == 'font':
        exp = fmt.get('font_cn')
        if not exp:
            return None
        got = _actual_font(para)
        if got is None:
            return (False, '未设置', exp)
        return (got == exp, got, exp)

    if attr == 'size':
        exp = fmt.get('size')
        if not exp:
            return None
        got = _actual_size(para)
        if got is None:
            return (False, '未设置', '{}pt'.format(exp))
        return (abs(got - exp) <= 0.3, '{}pt'.format(round(got, 1)), '{}pt'.format(exp))

    if attr == 'bold':
        exp = bool(fmt.get('bold', False))
        got = _actual_bold(para)
        got_b = bool(got)
        return (got_b == exp, '加粗' if got_b else '不加粗', '加粗' if exp else '不加粗')

    if attr == 'align':
        exp_key = fmt.get('align', 'justify')
        exp = _ALIGN_MAP.get(exp_key, WD_ALIGN_PARAGRAPH.JUSTIFY)
        got = _actual_align(para)
        got_label = {v: k for k, v in _ALIGN_MAP.items()}.get(got)
        return (got == exp,
                ALIGN_LABELS.get(got_label, '未设置'),
                ALIGN_LABELS.get(exp_key, exp_key))

    if attr == 'indent':
        # attachment 走悬挂缩进，规则特殊，不做首行缩进核对
        if ptype == 'attachment':
            return None
        exp = fmt.get('indent', 0) or 0
        got = _actual_indent(para) or 0
        return (abs(got - exp) <= 1.0,
                '{}pt'.format(round(got, 1)), '{}pt'.format(exp))

    if attr == 'line_spacing':
        exp = _expected_line_spacing(fmt)
        if not exp:
            return None
        got = _actual_line_spacing(para)
        if got is None:
            return (False, '未设置固定行距', '固定 {}pt'.format(exp))
        return (abs(got - exp) <= 1.0,
                '{}pt'.format(round(got, 1)), '固定 {}pt'.format(exp))

    if attr == 'spacing':
        exp_b = fmt.get('space_before', 0) or 0
        exp_a = fmt.get('space_after', 0) or 0
        got_b, got_a = _actual_spacing(para)
        got_b = got_b or 0
        got_a = got_a or 0
        ok = abs(got_b - exp_b) <= 1.0 and abs(got_a - exp_a) <= 1.0
        return (ok,
                '段前{}pt/段后{}pt'.format(round(got_b, 1), round(got_a, 1)),
                '段前{}pt/段后{}pt'.format(exp_b, exp_a))

    return None


_PARA_ATTRS = ['font', 'size', 'bold', 'align', 'indent', 'line_spacing', 'spacing']


def _check_paragraphs(doc, preset, opts, typed, add):
    """逐段按识别类型 × 逐属性对照预设，按 (类型, 属性) 汇总为一条 finding。"""
    from .paragraph import paragraph_has_media
    active = [a for a in _PARA_ATTRS if opts.get(a)]
    if not active:
        return
    # {(ptype, attr): {'bad': [(段号, 实际)], 'total': n, 'exp': 期望}}
    agg = {}
    for idx, para, ptype in typed:
        fmt = preset.get(ptype)
        if not isinstance(fmt, dict):
            continue
        if paragraph_has_media(para):   # 含图段落排版时被特殊保护，不参与核对
            continue
        for attr in active:
            if ptype == 'roster' and attr not in _ROSTER_ATTRS:
                continue
            res = _compare_attr(attr, fmt, para, ptype)
            if res is None:
                continue
            ok, got, exp = res
            slot = agg.setdefault((ptype, attr), {'bad': [], 'total': 0, 'exp': exp})
            slot['total'] += 1
            if not ok:
                slot['bad'].append((idx, got))

    for (ptype, attr), slot in sorted(agg.items()):
        tname = TYPE_LABELS.get(ptype, ptype)
        aname = ATTR_LABELS.get(attr, attr)
        item = '{}·{}'.format(tname, aname)
        if not slot['bad']:
            add('ok', item, '符合预设（{}，共 {} 段）'.format(slot['exp'], slot['total']))
            continue
        bad = slot['bad']
        locs = [i for i, _g in bad]
        got_vals = []
        for _i, g in bad:
            if g not in got_vals:
                got_vals.append(g)
        preview = '、'.join('第{}段'.format(i) for i in locs[:8])
        more = ' 等 {} 段'.format(len(locs)) if len(locs) > 8 else ''
        add('warn', item,
            '{}/{} 段不符：实际 {}，预设要求 {}（{}{}）'.format(
                len(bad), slot['total'], '/'.join(str(v) for v in got_vals[:3]),
                slot['exp'], preview, more),
            fix_key='para:{}:{}'.format(ptype, attr),
            locations=locs)


# ---------------- 页面级检查 ----------------

def _check_grid(doc, preset, add):
    grid = preset.get('grid') or {}
    lines = grid.get('lines_per_page')
    chars = grid.get('chars_per_line')
    if not lines:
        return
    sec = doc.sections[0]
    g = sec._sectPr.find(qn('w:docGrid'))
    if g is None:
        add('warn', '页面网格', '未设置文档网格，预设要求每页 {} 行 × 每行 {} 字'.format(lines, chars),
            fix_key='grid')
        return
    try:
        line_pitch = int(g.get(qn('w:linePitch')) or 0)
    except (TypeError, ValueError):
        line_pitch = 0
    text_h = sec.page_height.twips - sec.top_margin.twips - sec.bottom_margin.twips
    exp_pitch = int(round(float(text_h) / lines))
    if line_pitch and abs(line_pitch - exp_pitch) <= 2:
        add('ok', '页面网格', '每页 {} 行 × 每行 {} 字'.format(lines, chars))
    else:
        got_lines = int(round(float(text_h) / line_pitch)) if line_pitch else '?'
        add('warn', '页面网格',
            '实际约每页 {} 行，预设要求每页 {} 行 × 每行 {} 字'.format(got_lines, lines, chars),
            fix_key='grid')


def _footer_has_page_field(doc):
    import re as _re
    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            for para in footer.paragraphs:
                for run in para.runs:
                    xml = run._r.xml or ''
                    if _re.search(r'\bPAGE\b', xml, _re.I):
                        return True, para
                txt = para.text.strip()
                if txt and _re.fullmatch(r'[—\-–\s　]*(?:第\s*)?\d+(?:\s*/\s*\d+)?(?:\s*页)?[—\-–\s　]*', txt):
                    return True, para
    return False, None


def _check_page_number(doc, preset, add):
    want = preset.get('page_number', False)
    has, para = _footer_has_page_field(doc)
    if not want:
        if has:
            add('info', '页码', '文档有页码，但预设未要求页码（不做处理）')
        else:
            add('ok', '页码', '预设未要求页码')
        return
    if not has:
        add('warn', '页码', '预设要求页码，但文档页脚未找到页码', fix_key='page_number')
        return
    exp_font = preset.get('page_number_font', '宋体')
    exp_size = preset.get('page_number_size', 14)
    run = _first_run(para) if para is not None else None
    got_font = got_size = None
    if run is not None:
        rpr = run._element.rPr
        if rpr is not None and rpr.rFonts is not None:
            got_font = rpr.rFonts.get(qn('w:eastAsia'))
        if run.font.size is not None:
            got_size = run.font.size.pt
    bad = []
    if got_font and exp_font and got_font != exp_font:
        bad.append('字体实际「{}」应为「{}」'.format(got_font, exp_font))
    if got_size and exp_size and abs(got_size - exp_size) > 0.3:
        bad.append('字号实际 {}pt 应为 {}pt'.format(round(got_size, 1), exp_size))
    if bad:
        add('warn', '页码', '；'.join(bad), fix_key='page_number')
    else:
        add('ok', '页码', '{} {}pt'.format(exp_font, exp_size))


# ---------------- 主检查入口 ----------------

def check_compliance(doc, preset, options=None, detect_types=None):
    """返回 findings 列表：
    [{'level': 'warn'/'info'/'ok', 'item': 名称, 'detail': 说明,
      'fix_key': 可选, 'locations': 可选段号列表}]

    fix_key 存在表示该偏差可被 apply_compliance_fixes 自动修正。
    """
    opts = dict(DEFAULT_OPTIONS)
    norm = normalize_options(options)
    if norm:
        opts.update(norm)
    findings = []

    def add(level, item, detail, fix_key=None, locations=None):
        f = {'level': level, 'item': item, 'detail': detail}
        if fix_key:
            f['fix_key'] = fix_key
        if locations:
            f['locations'] = locations
        findings.append(f)

    sec = doc.sections[0]
    page = preset.get('page', {})

    # --- 页边距 ---
    if opts.get('margins'):
        exp = (page.get('top'), page.get('bottom'), page.get('left'), page.get('right'))
        got = (_cm(sec.top_margin.cm), _cm(sec.bottom_margin.cm),
               _cm(sec.left_margin.cm), _cm(sec.right_margin.cm))
        bad = [n for n, e, g in zip(('上', '下', '左', '右'), exp, got)
               if e is not None and abs((g or 0) - e) > 0.05]
        if bad:
            add('warn', '页边距',
                '实际 上{}/下{}/左{}/右{} cm，预设要求 上{}/下{}/左{}/右{} cm，'
                '不符：{}'.format(got[0], got[1], got[2], got[3],
                                exp[0], exp[1], exp[2], exp[3], '、'.join(bad)),
                fix_key='margins')
        else:
            add('ok', '页边距', '符合预设')

    # --- 纸张 ---
    if opts.get('paper'):
        w, h = _cm(sec.page_width.cm), _cm(sec.page_height.cm)
        want = preset.get('page_size', 'A4')
        is_a4 = abs((w or 0) - 21.0) < 0.2 and abs((h or 0) - 29.7) < 0.2
        if want == 'A4' and not is_a4:
            add('warn', '纸张', '当前 {}×{} cm 非 A4（预设要求 A4）'.format(w, h),
                fix_key='paper')
        else:
            add('ok', '纸张', '{}×{} cm'.format(w, h))

    # --- 页面网格 ---
    if opts.get('grid'):
        _check_grid(doc, preset, add)

    # --- 页码 ---
    if opts.get('page_number'):
        _check_page_number(doc, preset, add)

    # --- 段落级 + 结构完整性（共用一次类型识别）---
    need_typed = any(opts.get(a) for a in _PARA_ATTRS) or opts.get('structure')
    typed = _detect_types(doc, preset) if need_typed else []

    if any(opts.get(a) for a in _PARA_ATTRS):
        _check_paragraphs(doc, preset, opts, typed, add)

    if opts.get('structure'):
        types = {}
        for _i, _p, ptype in typed:
            types[ptype] = types.get(ptype, 0) + 1
        missing = []
        if not types.get('title'):
            missing.append('标题')
        if not types.get('recipient'):
            missing.append('主送机关')
        if not types.get('date'):
            missing.append('成文日期')
        if missing:
            # 结构缺失是内容缺失，无法凭空补出，不给 fix_key
            add('warn', '结构完整性',
                '未识别到：{}（如确有请在预览里核对/指定）'.format('、'.join(missing)))
        else:
            add('ok', '结构完整性', '标题/主送机关/成文日期齐全')

    # --- 序号层次一致性 ---
    if opts.get('numbering'):
        from . import analyzer
        issues = analyzer.analyze_numbering(doc)
        if issues:
            for it in issues:
                add('warn', '序号层次',
                    '{}（{}）；建议用「智能一键处理」统一序号层次'.format(
                        it.get('type', '序号不统一'), it.get('detail', '')))
        else:
            add('ok', '序号层次', '未发现序号风格混用')

    # --- 标点 ---
    if opts.get('punctuation'):
        from . import analyzer
        issues = analyzer.analyze_punctuation(doc)
        if issues:
            add('warn', '标点规范', '发现 {} 处英文/不规范标点'.format(len(issues)),
                fix_key='punctuation')
        else:
            add('ok', '标点规范', '未发现英文标点')

    return findings


# ---------------- 自动修正 ----------------

def _fix_margins(doc, preset, typed=None):
    page = preset.get('page', {})
    sec = doc.sections[0]
    for attr, key in (('top_margin', 'top'), ('bottom_margin', 'bottom'),
                      ('left_margin', 'left'), ('right_margin', 'right')):
        v = page.get(key)
        if v is not None:
            setattr(sec, attr, Cm(v))
    return '页边距→上{}/下{}/左{}/右{} cm'.format(
        page.get('top', '-'), page.get('bottom', '-'),
        page.get('left', '-'), page.get('right', '-'))


def _fix_paper(doc, preset, typed=None):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    return '纸张→A4（21×29.7 cm）'


def _fix_grid(doc, preset, typed=None):
    from .page import _apply_page_grid
    grid = preset.get('grid') or {}
    lines = grid.get('lines_per_page')
    chars = grid.get('chars_per_line')
    if not lines:
        return ''
    base = (preset.get('body') or {}).get('size', 16)
    _apply_page_grid(doc, lines, chars, base)
    return '页面网格→每页 {} 行 × 每行 {} 字'.format(lines, chars)


def _fix_page_number(doc, preset, typed=None):
    from .page import add_page_number
    if not preset.get('page_number'):
        return ''
    add_page_number(
        doc,
        font_name=preset.get('page_number_font', '宋体'),
        font_size=preset.get('page_number_size', 14),
        style=preset.get('page_number_style', 'dash'),
        position=preset.get('page_number_position', 'center'),
        offset_from_text_mm=preset.get('page_number_offset_mm', 7),
        replace_existing=preset.get('replace_existing_page_number', True),
        bold=preset.get('page_number_bold', False),
    )
    return '页码→{} {}pt'.format(preset.get('page_number_font', '宋体'),
                                preset.get('page_number_size', 14))


def _fix_punctuation(doc, preset, typed=None):
    from . import punctuation
    quote_state = {'dq': 0, 'sq': 0}
    n = 0
    for p in doc.paragraphs:
        if punctuation.process_paragraph(p, quote_state=quote_state):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if punctuation.process_paragraph(p, quote_state=quote_state):
                        n += 1
    return '标点修复（{} 段有改动）'.format(n) if n else '标点修复（无需改动）'


def _apply_attr(para, attr, fmt, ptype):
    """把单个属性按预设写入段落（只动这一个属性，其余不碰）。"""
    from .font import set_font
    pf = para.paragraph_format

    if attr in ('font', 'size', 'bold'):
        font_cn = fmt.get('font_cn')
        font_en = fmt.get('font_en', font_cn)
        size = fmt.get('size')
        bold = bool(fmt.get('bold', False))
        for run in para.runs:
            if not run.text.strip():
                continue
            if attr == 'font':
                rpr = run._element.get_or_add_rPr()
                rf = rpr.find(qn('w:rFonts'))
                if rf is None:
                    from docx.oxml import OxmlElement
                    rf = OxmlElement('w:rFonts')
                    rpr.insert(0, rf)
                rf.set(qn('w:eastAsia'), font_cn)
                rf.set(qn('w:ascii'), font_en)
                rf.set(qn('w:hAnsi'), font_en)
            elif attr == 'size' and size:
                run.font.size = Pt(size)
            elif attr == 'bold':
                run.font.bold = bold
        return

    if attr == 'align':
        pf.alignment = _ALIGN_MAP.get(fmt.get('align', 'justify'),
                                      WD_ALIGN_PARAGRAPH.JUSTIFY)
        return

    if attr == 'indent':
        indent = fmt.get('indent', 0) or 0
        pf.first_line_indent = Pt(indent)
        # 同步 firstLineChars，避免 Word 用字符数覆盖磅值
        try:
            from docx.oxml import OxmlElement
            pPr = para._p.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if indent > 0:
                size = fmt.get('size', 16) or 16
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                ind.set(qn('w:firstLineChars'), str(int(round(indent / size * 100))))
            elif ind is not None:
                ind.attrib.pop(qn('w:firstLineChars'), None)
        except Exception:
            pass
        return

    if attr == 'line_spacing':
        ls = _expected_line_spacing(fmt)
        if ls:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(ls)
        return

    if attr == 'spacing':
        from .paragraph import _set_paragraph_spacing_points
        _set_paragraph_spacing_points(para, fmt.get('space_before', 0),
                                      fmt.get('space_after', 0))
        return


def _fix_paragraph_attr(doc, preset, typed, ptype, attr):
    """只修正指定类型段落的指定属性。"""
    from .paragraph import paragraph_has_media
    fmt = preset.get(ptype)
    if not isinstance(fmt, dict):
        return ''
    if ptype == 'roster' and attr not in _ROSTER_ATTRS:
        return ''
    n = 0
    for _idx, para, t in typed:
        if t != ptype or paragraph_has_media(para):
            continue
        if attr == 'indent' and ptype == 'attachment':
            continue   # 悬挂缩进规则特殊，不在此处理
        _apply_attr(para, attr, fmt, ptype)
        n += 1
    if not n:
        return ''
    return '{}·{}→按预设（{} 段）'.format(
        TYPE_LABELS.get(ptype, ptype), ATTR_LABELS.get(attr, attr), n)


_DOC_FIXERS = {
    'margins': _fix_margins,
    'paper': _fix_paper,
    'grid': _fix_grid,
    'page_number': _fix_page_number,
    'punctuation': _fix_punctuation,
}


def fix_label(fix_key):
    """把 fix_key 翻成一句人话，供 UI 显示。"""
    if fix_key in FIX_LABELS:
        return FIX_LABELS[fix_key]
    if fix_key and fix_key.startswith('para:'):
        parts = fix_key.split(':')
        if len(parts) == 3:
            _p, ptype, attr = parts
            return '把{}的{}改为预设值'.format(
                TYPE_LABELS.get(ptype, ptype), ATTR_LABELS.get(attr, attr))
    return '修正'


def apply_compliance_fixes(input_path, output_path, preset, fix_keys):
    """打开文档，仅对 fix_keys 指定的偏差自动修正，另存为 output_path。

    返回已执行修正的说明列表（每项一句话），供 UI 反馈。
    """
    from docx import Document
    from .paragraph import sanitize_document
    doc = Document(input_path)
    sanitize_document(doc)

    para_keys = [k for k in fix_keys if k.startswith('para:')]
    doc_keys = [k for k in fix_keys if k in _DOC_FIXERS]

    # 段落修正依赖类型识别；页边距等页面改动不影响识别结果，
    # 故一次识别即可，且在页面修正之前做（避免 grid 依赖新边距时错序）。
    typed = _detect_types(doc, preset) if para_keys else None

    applied = []

    for key in para_keys:
        parts = key.split(':')
        if len(parts) != 3:
            continue
        _p, ptype, attr = parts
        try:
            desc = _fix_paragraph_attr(doc, preset, typed, ptype, attr)
        except Exception as e:
            applied.append('{}：修正失败（{}）'.format(fix_label(key), e))
            continue
        if desc:
            applied.append(desc)

    # 页面级修正放在后面：网格依赖最终的页边距/纸张
    for key in ('margins', 'paper', 'grid', 'page_number', 'punctuation'):
        if key not in doc_keys:
            continue
        try:
            desc = _DOC_FIXERS[key](doc, preset)
        except Exception as e:
            applied.append('{}：修正失败（{}）'.format(fix_label(key), e))
            continue
        if desc:
            applied.append(desc)

    doc.save(output_path)
    return applied


def build_preview_model(doc, preset, max_paras=400):
    """为"现状 vs 修正后"对比预览提供逐段数据。

    返回 [{'index': 段号, 'text': 文字, 'ptype': 类型,
           'actual': {属性: 值}, 'expected': {属性: 值},
           'bad': {属性: True}}]
    属性值是可直接用于渲染的规格字典（字体/字号/加粗/对齐/缩进/行距/段距）。
    """
    from .paragraph import paragraph_has_media
    typed = _detect_types(doc, preset)
    out = []
    for idx, para, ptype in typed[:max_paras]:
        fmt = preset.get(ptype)
        if not isinstance(fmt, dict):
            fmt = preset.get('body', {})
        has_media = paragraph_has_media(para)
        exp_ls = _expected_line_spacing(fmt)
        got_b, got_a = _actual_spacing(para)
        align_rev = {v: k for k, v in _ALIGN_MAP.items()}
        actual = {
            'font': _actual_font(para),
            'font_en': _actual_font_en(para),
            'size': _actual_size(para),
            'bold': bool(_actual_bold(para)),
            'align': align_rev.get(_actual_align(para)),
            'indent': _actual_indent(para) or 0,
            'line_spacing': _actual_line_spacing(para),
            'space_before': got_b or 0,
            'space_after': got_a or 0,
        }
        expected = {
            'font': fmt.get('font_cn'),
            'font_en': fmt.get('font_en', 'Times New Roman'),
            'size': fmt.get('size'),
            'bold': bool(fmt.get('bold', False)),
            'align': fmt.get('align', 'justify'),
            'indent': fmt.get('indent', 0) or 0,
            'line_spacing': exp_ls,
            'space_before': fmt.get('space_before', 0) or 0,
            'space_after': fmt.get('space_after', 0) or 0,
        }
        bad = {}
        if not has_media:
            for attr in _PARA_ATTRS:
                if ptype == 'roster' and attr not in _ROSTER_ATTRS:
                    continue
                res = _compare_attr(attr, fmt, para, ptype)
                if res is not None and not res[0]:
                    bad[attr] = True
        out.append({
            'index': idx, 'text': para.text.strip(), 'ptype': ptype,
            'actual': actual, 'expected': expected, 'bad': bad,
            'media': has_media,
        })
    return out


# 段落级 fix_key 的属性 → 影响预览里的哪些渲染字段
_ATTR_FIELDS = {
    'font': ('font', 'font_en'),      # 中西文字体一起改，数字才跟着变
    'size': ('size',),
    'bold': ('bold',),
    'align': ('align',),
    'indent': ('indent',),
    'line_spacing': ('line_spacing',),
    'spacing': ('space_before', 'space_after'),
}


def preview_spec_after(entry, fix_keys):
    """按已认可的 fix_keys 算出该段"修正后"的渲染规格。"""
    spec = dict(entry['actual'])
    ptype = entry['ptype']
    for key in fix_keys:
        if not key.startswith('para:'):
            continue
        parts = key.split(':')
        if len(parts) != 3 or parts[1] != ptype:
            continue
        for field in _ATTR_FIELDS.get(parts[2], ()):
            spec[field] = entry['expected'].get(field)
    return spec


def format_compliance_report(filename, findings, preset_name=''):
    warns = [f for f in findings if f['level'] == 'warn']
    lines = ['◆ 公文合规性检查：{}'.format(filename)]
    if preset_name:
        lines.append('  对照预设：{}'.format(preset_name))
    lines.append('  {}'.format('存在 {} 项偏差'.format(len(warns)) if warns else '未发现偏差 ✓'))
    lines.append('')
    for f in findings:
        mark = {'warn': '✗', 'ok': '✓', 'info': '·'}.get(f['level'], '·')
        lines.append('  {} 【{}】{}'.format(mark, f['item'], f['detail']))
    return '\n'.join(lines)
