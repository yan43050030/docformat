# -*- coding: utf-8 -*-
"""
段落格式化 — 从 formatter.py 拆分

包含：format_paragraph、_set_paragraph_spacing_points、
_keep_first_sentence_runs、_append_body_run
_keep_first_sentence_runs、_append_body_run、结构空行处理、deep_clean_document
"""
import re
import logging
from copy import deepcopy
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .font import set_font, _force_normal_style, _add_ppr_change

logger = logging.getLogger('docformat.paragraph')

# python-docx 读取这些"简单类型"元素时，缺少必填属性会抛 InvalidXmlError。
# WPS / 老版 Word 导出的文档里常见残缺的 <w:jc/>（缺 w:val），一旦访问
# paragraph_format.alignment 就崩溃。这里在读取任何对齐前先修掉。
# 键 = 元素标签，值 = 必填属性名。
_REQUIRED_VAL_ELEMENTS = {
    'w:jc': 'w:val',            # 对齐
    'w:textAlignment': 'w:val',  # 垂直文本对齐
}


def sanitize_document(doc):
    """修复 WPS/老 Word 导出文档里缺必填属性的元素，避免 python-docx 读取时崩溃。

    做法：遍历文档 XML（正文、样式表、页眉页脚），对缺 w:val 的
    <w:jc>/<w:textAlignment> 补一个安全默认值（left/auto），保留元素结构。
    返回修复的元素个数。
    """
    fixed = 0

    def _fix_tree(root):
        nonlocal fixed
        if root is None:
            return
        for tag, attr in _REQUIRED_VAL_ELEMENTS.items():
            for el in root.iter(qn(tag)):
                if el.get(qn(attr)) is None:
                    # w:jc 默认按 left（无对齐信息可保留时的自然默认）
                    el.set(qn(attr), 'both' if tag == 'w:textAlignment' else 'left')
                    fixed += 1

    # 正文
    try:
        _fix_tree(doc.element.body)
    except Exception:
        pass
    # 样式表
    try:
        _fix_tree(doc.styles.element)
    except Exception:
        pass
    # 页眉页脚（含奇偶页、首页）
    try:
        for section in doc.sections:
            for part_attr in ('header', 'even_page_header', 'first_page_header',
                              'footer', 'even_page_footer', 'first_page_footer'):
                try:
                    part = getattr(section, part_attr, None)
                    if part is not None and part.paragraphs:
                        _fix_tree(part._element)
                except Exception:
                    continue
    except Exception:
        pass

    if fixed:
        logger.info('sanitize: 修复 %d 个缺属性元素（WPS 兼容）', fixed)
    return fixed


def _set_paragraph_spacing_points(para, before_pt=0, after_pt=0):
    """Set paragraph spacing in points and clear line-based spacing leftovers."""
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)

    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)

    for attr in ('beforeLines', 'afterLines', 'beforeAutospacing', 'afterAutospacing'):
        spacing.attrib.pop(qn(f'w:{attr}'), None)

    spacing.set(qn('w:before'), str(int(round(before_pt * 20))))
    spacing.set(qn('w:after'), str(int(round(after_pt * 20))))


def _compact_empty_paragraph(para):
    """Clear spacing on empty paragraphs."""
    _set_paragraph_spacing_points(para, 0, 0)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(1)


def _mark_structural_blank(para):
    """在段落 pPr 上写自定义标记，标识为结构性空行"""
    pPr = para._p.get_or_add_pPr()
    pPr.set('docfmt-structural-blank', '1')


def _is_structural_blank(para):
    """检查段落是否被标记为结构性空行"""
    pPr = para._p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is None:
        return False
    return pPr.get('docfmt-structural-blank') == '1'


def _format_structural_blank_paragraph(para, line_spacing_pt=28):
    """Format the intentional blank line used between document sections."""
    if not para.runs:
        para.add_run(' ')
    _set_paragraph_spacing_points(para, 0, 0)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    _mark_structural_blank(para)


def _format_empty_paragraphs(doc, structural_blank_ids, line_spacing_pt=28):
    """格式化文档中所有空段落"""
    for para in doc.paragraphs:
        if para.text.strip():
            continue
        if _is_structural_blank(para):
            _format_structural_blank_paragraph(para, line_spacing_pt)
        else:
            _compact_empty_paragraph(para)


def deep_clean_document(doc):
    """深度清洗文档：移除所有段落级用户格式属性"""

    def _clean_paragraph(para):
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = None
        pf.right_indent = None
        pf.first_line_indent = None
        pf.line_spacing = None
        pf.line_spacing_rule = None
        _force_normal_style(para)
        for run in para.runs:
            run.font.color.rgb = None
            run.font.highlight_color = None
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.underline = None
            run.font.strike = None

    for para in doc.paragraphs:
        _clean_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _clean_paragraph(para)


def _keep_first_sentence_runs(para, heading_text):
    """将已格式化段落的文字重置为仅含标题句（保留第一个 run 的格式）"""
    # 收集第一个有效 run 的格式
    first_run = None
    for r in para.runs:
        if r.text.strip():
            first_run = r
            break
    # 清空全部 run 的文字
    for r in para.runs:
        r.text = ''
    # 把标题句写入第一个 run
    if first_run:
        first_run.text = heading_text


def _append_body_run(para, body_text, bfmt, revision_mode=False):
    """在段尾追加一个正文格式的 run"""
    from .font import set_font
    run = para.add_run(body_text)
    set_font(
        run,
        bfmt.get('font_cn', '仿宋_GB2312'),
        bfmt.get('font_en', 'Times New Roman'),
        bfmt.get('size', 16),
        bold=bfmt.get('bold', False),
    )
    if revision_mode:
        from .font import _add_rpr_change, _next_rev_id, _rev_date
        rpr = run._element.get_or_add_rPr()
        rpr.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id',
                str(_next_rev_id()))
        _add_rpr_change(run, para)


def format_paragraph(para, fmt, para_type, line_spacing_pt=28, first_line_bold=False, revision_mode=False, bold_serial=True):
    """格式化段落

    fmt 支持的字段:
        font_cn, font_en, size, bold, align, indent,
        line_spacing  - 行距(磅), None表示使用1.5倍行距
        space_before  - 段前间距(磅), 默认0
        space_after   - 段后间距(磅), 默认0
    """
    # 组成人员名单：保持原格式不动，仅设缩进=0和行距
    if para_type == 'roster':
        pf = para.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = None
        if fmt.get('line_spacing'):
            pf.line_spacing = Pt(fmt['line_spacing'])
        return

    _force_normal_style(para)

    if revision_mode:
        orig_ppr = deepcopy(para._p.pPr)
        orig_ppr_xml = para._p.xml

    pf = para.paragraph_format

    # 对齐方式
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 段落左缩进清零
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)

    # attachment 类型走悬挂缩进
    _attachment_indent_done = False
    if para_type == 'attachment':
        font_size_pt = fmt.get('size', 16) or 16
        pf.left_indent = Pt(font_size_pt * 5)
        if '附件' in para.text:
            pf.first_line_indent = Pt(-font_size_pt * 3)
        else:
            pf.first_line_indent = Pt(0)
        try:
            pPr = para._p.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:firstLineChars'), None)
        except Exception:
            pass
        _attachment_indent_done = True

    # 首行缩进
    if not _attachment_indent_done:
        indent = fmt.get('indent', 0)
        if indent > 0:
            pf.first_line_indent = Pt(indent)
            size = fmt.get('size', 16) or 16
            try:
                chars_100 = int(round(indent / size * 100))
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                ind.set(qn('w:firstLineChars'), str(chars_100))
            except Exception:
                pass
        else:
            pf.first_line_indent = Pt(0)
            try:
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    ind.attrib.pop(qn('w:firstLineChars'), None)
            except Exception:
                pass

    # 行距
    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5

    # 段前段后
    _set_paragraph_spacing_points(
        para,
        fmt.get('space_before', 0),
        fmt.get('space_after', 0)
    )

    # 字体 - 支持首句加粗
    if first_line_bold and para_type == 'body':
        full_text = para.text
        first_sentence_end = full_text.find('。')
        if first_sentence_end != -1:
            split_idx = first_sentence_end + 1
            first_part = full_text[:split_idx]
            rest_part = full_text[split_idx:]

            for run in list(para.runs):
                para._p.remove(run._r)

            run1 = para.add_run(first_part)
            set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True,
                      italic=fmt.get('italic', False), color=fmt.get('color'),
                      revision_mode=revision_mode)

            if rest_part:
                run2 = para.add_run(rest_part)
                set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False),
                          italic=fmt.get('italic', False), color=fmt.get('color'),
                          revision_mode=revision_mode)
        else:
            for run in para.runs:
                set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False),
                          italic=fmt.get('italic', False), color=fmt.get('color'),
                          revision_mode=revision_mode)
    else:
        # 正文里的序列词加粗前缀
        if bold_serial and para_type == 'body':
            _SERIAL_PATTERNS = [
                r'^([一二三四五六七八九十]{1,3}是)([：:、]?)',
                r'^([一二三四五六七八九十]{1,3}要)([：:、]?)',
                r'^(第[一二三四五六七八九十百\d]+[点条项步])([：:、，,]?)',
                r'^([一二三四五六七八九十]{1,3}方面)([：:、]?)',
            ]
            m = None
            for _pat in _SERIAL_PATTERNS:
                m = re.match(_pat, para.text)
                if m:
                    break
            if m:
                lead = m.group(1) + (m.group(2) or '')
                rest = para.text[len(lead):]
                for run in list(para.runs):
                    para._p.remove(run._r)
                run1 = para.add_run(lead)
                set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True,
                      italic=fmt.get('italic', False), color=fmt.get('color'),
                      revision_mode=revision_mode)
                if rest:
                    run2 = para.add_run(rest)
                    set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False),
                          italic=fmt.get('italic', False), color=fmt.get('color'),
                          revision_mode=revision_mode)
                return

        for run in para.runs:
            set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False),
                      italic=fmt.get('italic', False), color=fmt.get('color'),
                      revision_mode=revision_mode)

    # 修订模式
    if revision_mode and para._p.xml != orig_ppr_xml:
        _add_ppr_change(para, orig_ppr)
