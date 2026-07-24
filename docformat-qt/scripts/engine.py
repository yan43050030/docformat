# -*- coding: utf-8 -*-
"""
排版引擎主控 — 从 formatter.py 重构

编排整个格式化流程，各模块职责：
  data_model    - 预设定义、字体适配
  font          - 字体设置、修订标记（支持 italic/color）
  page          - 页面设置、页码、文档网格、多纸张大小
  paragraph     - 段落格式化、空行处理、深度清洗
  detector      - 段落类型识别
  table         - 表格格式化（边框样式、表头重复、底纹、垂直对齐、嵌套表格）
  signature     - 落款对位 (GB/T 9704)
  header_footer - 页眉页脚格式化
  image         - 图片压缩、尺寸、环绕、对齐
  watermark     - 文字水印（密级标注）
"""

import logging
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .data_model import (
    PRESETS, load_custom_preset, _adapt_fonts_for_platform, _merge_preset_settings,
)
from .font import set_font, reset_revision_counter
from .page import (
    _apply_page_grid, _set_normal_style_font, add_page_number,
    remove_background, _strip_autospacing_from_styles,
)
from .paragraph import (
    format_paragraph, _format_structural_blank_paragraph,
    _format_empty_paragraphs, _compact_empty_paragraph,
    _is_structural_blank, deep_clean_document,
    _keep_first_sentence_runs, _append_body_run,
    sanitize_document, paragraph_has_media, protect_media_paragraph,
)
from .detector import (
    detect_para_type, _compile_rules, _build_text_context,
    _standardize_date_text,
)
from .table import (
    _iter_block_items, _set_table_borders, _set_table_cell_margins,
    _set_table_width_percent, _set_table_indent, _set_table_col_widths_by_content,
    _set_cell_borders, _is_numeric_text, _is_short_text, _is_table_title, _is_table_unit,
    _insert_paragraph_after_table, _insert_paragraph_before_table,
    _insert_paragraph_after_paragraph, _insert_paragraph_before_paragraph,
    _split_heading_by_punct,
    _set_header_row_repeat, _set_cell_shading, _set_cell_vertical_alignment,
    _find_nested_tables,
)
from .signature import _apply_gb_signature_layout
from .header_footer import apply_header_footer, _get_default_hf_config
from .image import (
    PictureConfig, DEFAULT_PICTURE_CONFIG,
    compress_images, apply_image_size,
    apply_image_wrapping, apply_image_alignment,
)
from .watermark import WatermarkConfig, apply_watermark, WATERMARK_PRESETS

logger = logging.getLogger('docformat.engine')


def _ensure_structural_blank_lines(doc, line_spacing_pt=28, rules=None, type_overrides=None):
    """确保标题后、落款前的结构性空行"""
    all_texts, all_texts_idx_map = _build_text_context(doc)
    total_paras = len(doc.paragraphs)
    entries = []
    prev_para_type = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        ai = all_texts_idx_map.get(i)
        if type_overrides and ai is not None and ai in type_overrides:
            para_type = type_overrides[ai]
        else:
            para_type = detect_para_type(
                text, i, total_paras,
                para.paragraph_format.alignment,
                all_texts,
                all_texts_index=ai,
                prev_para_type=prev_para_type,
                rules=rules
            )
        entries.append((para, para_type, prev_para_type))
        prev_para_type = para_type

    structural_blank_ids = set()
    p_tag = qn('w:p')
    for para, para_type, prev_para_type in entries:
        needs_blank = (
            (prev_para_type == 'title' and para_type not in ('title', 'docnum')) or
            (prev_para_type == 'docnum' and para_type != 'docnum') or
            (para_type == 'signature' and prev_para_type not in (None, 'signature', 'date')) or
            (para_type == 'attachment' and prev_para_type not in (None, 'attachment'))
        )
        if not needs_blank:
            continue

        prev_el = para._p.getprevious()
        if prev_el is not None and prev_el.tag == p_tag and not Paragraph(prev_el, para._parent).text.strip():
            blank_para = Paragraph(prev_el, para._parent)
        else:
            blank_para = _insert_paragraph_before_paragraph(para)

        _format_structural_blank_paragraph(blank_para, line_spacing_pt)
        structural_blank_ids.add(id(blank_para._p))

    return structural_blank_ids


def format_document(input_path, output_path, preset_name='official', progress_callback=None,
                    revision_mode=False, bold_serial=True, custom_settings=None,
                    type_overrides=None):
    """格式化文档

    Args:
        input_path: 源文件路径
        output_path: 输出文件路径
        preset_name: 预设名称 (official/official_gbk/academic/legal/custom)
        progress_callback: 可选回调 callback(current, total, stage_text)
        revision_mode: 是否启用修订模式
        bold_serial: 是否对序列词加粗
        custom_settings: 自定义预设字典
        type_overrides: {非空段序号: 段落类型}，手动指定的段落类型覆盖
    """
    reset_revision_counter()

    # 处理自定义预设
    if preset_name == 'custom' and custom_settings:
        preset = deepcopy(custom_settings)
        logger.info(f'Preset: {preset.get("name", "自定义格式")}')
    elif preset_name == 'custom':
        preset = load_custom_preset()
        if preset is None:
            logger.warning('Custom preset not found, using official preset')
            preset = PRESETS['official']
        else:
            logger.info(f'Preset: {preset.get("name", "自定义格式")}')
    elif preset_name not in PRESETS:
        raise ValueError(
            "未知预设: {} (可用: {})".format(
                preset_name, ", ".join(PRESETS.keys())))
    else:
        preset = PRESETS[preset_name]
        logger.info(f'Preset: {preset["name"]}')

    if custom_settings and preset_name != 'custom':
        preset = _merge_preset_settings(preset, custom_settings)

    logger.info(f'Input: {input_path}')
    preset = _adapt_fonts_for_platform(preset)

    first_line_bold = preset.get('first_line_bold', False)
    bold_serial = preset.get('bold_serial', bold_serial)

    doc = Document(input_path)

    # v3.0.4: 先修复 WPS/老 Word 残缺元素（缺 w:val 的 <w:jc> 等），
    # 否则后续读取 paragraph_format.alignment 会抛 InvalidXmlError
    sanitize_document(doc)

    # v1.8.0: 强力清洗模式
    if preset.get('deep_clean', False):
        deep_clean_document(doc)

    # 标题+标点拆分
    if preset.get('split_heading_at_punct', False):
        for para in list(doc.paragraphs):
            _split_heading_by_punct(para)

    def _progress(current, total, stage):
        if progress_callback:
            progress_callback(current, total, stage)

    # 1. 移除背景
    logger.info('1. Removing background...')
    _progress(0, 100, '移除背景...')
    remove_background(doc)

    # 2. 设置页面边距和纸张大小
    logger.info('2. Setting page margins...')
    _progress(5, 100, '设置页面边距...')
    page = preset['page']

    # 纸张大小映射 (宽 x 高, cm)
    _PAPER_SIZES = {
        'A3': (29.7, 42.0),
        'A4': (21.0, 29.7),
        'A5': (14.8, 21.0),
        'B4': (25.7, 36.4),
        'B5': (18.2, 25.7),
        '16K': (19.5, 27.0),
        '16k': (19.5, 27.0),
        'Letter': (21.59, 27.94),
        'letter': (21.59, 27.94),
        'Legal': (21.59, 35.56),
        'legal': (21.59, 35.56),
    }
    paper_size = preset.get('page_size', 'A4')
    if paper_size in _PAPER_SIZES:
        pw, ph = _PAPER_SIZES[paper_size]
    else:
        pw, ph = _PAPER_SIZES['A4']

    for section in doc.sections:
        section.page_width = Cm(pw)
        section.page_height = Cm(ph)
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])

    # 2.5 页眉页脚设置（在文档网格之前，确保页面边距已设置）
    logger.info('2.5. Applying header/footer format...')
    _progress(7, 100, '设置页眉页脚...')
    hf_config = _get_default_hf_config(preset)
    apply_header_footer(doc, hf_config)

    # 2.6 水印设置
    wm_cfg = preset.get('watermark', {})
    if wm_cfg.get('enabled') or wm_cfg.get('text'):
        logger.info('2.6. Applying watermark...')
        _progress(8, 100, '设置水印...')
        wm_config = WatermarkConfig.from_dict(wm_cfg)
    elif wm_cfg:
        # 通过预设名称引用 (如 "secret")
        wm_preset_name = wm_cfg if isinstance(wm_cfg, str) else wm_cfg.get('preset', '')
        if wm_preset_name and wm_preset_name in WATERMARK_PRESETS:
            logger.info('2.6. Applying watermark preset: %s', wm_preset_name)
            _progress(8, 100, '设置水印...')
            wm_config = WATERMARK_PRESETS[wm_preset_name]
        else:
            wm_config = WatermarkConfig()
    else:
        wm_config = WatermarkConfig()
    apply_watermark(doc, wm_config)

    body_fmt_cfg = preset.get('body', {})
    body_line_spacing = body_fmt_cfg.get('line_spacing', 28) or 28

    # 文档网格
    grid_cfg = preset.get('grid')
    if grid_cfg and grid_cfg.get('lines_per_page'):
        _set_normal_style_font(
            doc,
            body_fmt_cfg.get('font_cn', '仿宋_GB2312'),
            body_fmt_cfg.get('font_en', 'Times New Roman'),
            body_fmt_cfg.get('size', 16) or 16,
        )
        _apply_page_grid(
            doc,
            grid_cfg.get('lines_per_page'),
            grid_cfg.get('chars_per_line'),
            body_fmt_cfg.get('size', 16) or 16,
        )

    _strip_autospacing_from_styles(doc)
    _active_rules = _compile_rules(preset.get('detect_rules'))
    structural_blank_ids = _ensure_structural_blank_lines(
        doc, body_line_spacing, rules=_active_rules, type_overrides=type_overrides)
    total_paras = len(doc.paragraphs)
    all_texts, all_texts_idx_map = _build_text_context(doc)

    # 3. 格式化段落
    logger.info('3. Formatting paragraphs...')
    _progress(10, 100, '格式化段落...')
    stats = {
        'title': 0, 'recipient': 0, 'heading1': 0, 'heading2': 0,
        'heading3': 0, 'heading4': 0, 'body': 0, 'signature': 0,
        'date': 0, 'attachment': 0, 'closing': 0, 'security': 0, 'docnum': 0,
    }
    errors = []

    prev_para_type = None
    typed_entries = []

    for i, para in enumerate(doc.paragraphs):
        try:
            raw = para.text
            # 清洗前导空白（半角空格、全角空格、Tab），防止首行缩进加倍
            cleaned = raw.lstrip('\x20　\t')
            if cleaned != raw:
                para.text = cleaned
            text = cleaned.strip()
            # 含图片/嵌入对象的段落：保护，绝不按空行压缩或设固定行距（会裁图）
            if paragraph_has_media(para):
                protect_media_paragraph(para)
                continue
            if not text:
                if _is_structural_blank(para):
                    _format_structural_blank_paragraph(para, body_line_spacing)
                else:
                    _compact_empty_paragraph(para)
                continue

            _ai = all_texts_idx_map.get(i)
            if type_overrides and _ai is not None and _ai in type_overrides:
                para_type = type_overrides[_ai]
            else:
                para_type = detect_para_type(
                    text, i, total_paras,
                    para.paragraph_format.alignment,
                    all_texts,
                    all_texts_index=_ai,
                    prev_para_type=prev_para_type,
                    rules=_active_rules
                )

            if para_type == 'date':
                standardized_date = _standardize_date_text(text)
                if standardized_date != text:
                    para.text = standardized_date
                    text = standardized_date

            fmt_key = para_type if para_type in preset else 'body'
            fmt = preset.get(fmt_key, preset['body'])

            # 长标题同行混排：二级/三级/四级标题含多个句号时，
            # 第一句按标题格式、后段内容按正文格式，仍在同一段落中
            heading_part = body_tail = None
            if para_type in ('heading2', 'heading3', 'heading4') and text.count('。') > 1:
                idx = text.index('。')
                heading_part = text[:idx + 1].strip()
                body_tail = text[idx + 1:].strip()
                if not body_tail or len(heading_part) < 4:
                    heading_part = body_tail = None

            if heading_part is not None:
                bfmt = preset.get('body', {})
                # 先按标题格式套整个段落（缩进/行距等）
                format_paragraph(
                    para, fmt, para_type,
                    first_line_bold=first_line_bold,
                    revision_mode=revision_mode,
                    bold_serial=bold_serial
                )
                # 清除所有 run 的文字，仅保留标题句
                _keep_first_sentence_runs(para, heading_part)
                # 在段落末尾追加正文 run
                _append_body_run(para, body_tail, bfmt, revision_mode)
                stats[para_type] = stats.get(para_type, 0) + 1
                typed_entries.append((para, para_type))
            else:
                format_paragraph(
                    para, fmt, para_type,
                    first_line_bold=first_line_bold,
                    revision_mode=revision_mode,
                    bold_serial=bold_serial
                )
                stats[para_type] = stats.get(para_type, 0) + 1
                typed_entries.append((para, para_type))

            prev_para_type = para_type

            preview = text[:35] + '...' if len(text) > 35 else text
            logger.info(f'   [{para_type:10}] {preview}')
        except Exception as exc:
            error_info = {
                'index': i,
                'text': text[:50] if 'text' in dir() else '<empty>',
                'error': str(exc),
            }
            errors.append(error_info)
            logger.warning('   [ERROR    ] paragraph %d: %s', i, exc)

        if total_paras > 0:
            pct = 10 + int(70 * (i + 1) / total_paras)
            _progress(pct, 100, f'格式化段落 ({i + 1}/{total_paras})')

    # 格式化后复查结构空行
    structural_blank_ids.update(_ensure_structural_blank_lines(
        doc, body_line_spacing, rules=_active_rules, type_overrides=type_overrides))
    _format_empty_paragraphs(doc, structural_blank_ids, body_line_spacing)

    # 落款对位
    if preset.get('gb_signature_layout'):
        _apply_gb_signature_layout(typed_entries, preset)

    # 标题梯形回行（可选，默认关闭）：title_shape = trapezoid_down/trapezoid_up
    title_shape = preset.get('title_shape', 'none')
    if title_shape in ('trapezoid_down', 'trapezoid_up'):
        try:
            from .title_shape import apply_title_shape
            title_size = preset.get('title', {}).get('size', 22) or 22
            sec0 = doc.sections[0]
            pw = (sec0.page_width or Cm(21.0)).pt
            usable_pt = pw - sec0.left_margin.pt - sec0.right_margin.pt
            cpl = max(6, int(usable_pt / title_size))   # 每行可容纳全角字数
            for para, ptype in typed_entries:
                if ptype == 'title':
                    apply_title_shape(para, cpl, title_shape)
        except Exception as e:
            logger.warning('标题梯形回行失败: %s', e)

    # 4. 处理表格
    logger.info('4. Formatting tables...')
    _progress(82, 100, '格式化表格...')
    body_fmt = preset.get('body', {})
    table_fmt = preset.get('table', {})
    table_defaults = {
        'optimize': True,
        'border_size_pt': 0.5,
        'border_style': 'single',
        'border_color': '000000',
        'width_percent': 100,
        'auto_col_width': True,
        'col_min_pct': 8,
        'col_max_pct': 45,
        'row_height_cm': 0.7,
        'cell_margin_top_cm': 0.0,
        'cell_margin_bottom_cm': 0.0,
        'cell_margin_left_cm': 0.05,
        'cell_margin_right_cm': 0.05,
        'paragraph_single': True,
        'after_table_blank_line': True,
        'title_align': 'center',
        'unit_align': 'right',
        'unit_space_before_lines': 0.5,
        'short_text_len': 4,
        'smart_align': False,
        'header_repeat': False,
        'header_shading_color': None,
        'cell_valign': None,
    }
    table_cfg = {**table_defaults, **table_fmt}

    tbl_font_cn = table_fmt.get('font_cn', body_fmt.get('font_cn', '仿宋_GB2312'))
    tbl_font_en = table_fmt.get('font_en', body_fmt.get('font_en', 'Times New Roman'))
    tbl_size = table_fmt.get('size', body_fmt.get('size', 16))
    tbl_bold = table_fmt.get('bold', False)
    tbl_line_spacing = table_fmt.get('line_spacing', body_fmt.get('line_spacing', 28))
    tbl_header_bold = table_fmt.get('header_bold', False)
    tbl_first_line_indent = table_fmt.get('first_line_indent', 0)

    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue

        table = block
        if table_cfg.get('optimize', True):
            table.autofit = not table_cfg.get('auto_col_width', True)
            _set_table_width_percent(table, table_cfg.get('width_percent', 100))
            _set_table_indent(table, 0)
            _set_table_borders(table, size_pt=table_cfg.get('border_size_pt', 0.5),
                               color=table_cfg.get('border_color', '000000'),
                               style=table_cfg.get('border_style', 'single'))
            _set_table_cell_margins(
                table,
                top_cm=table_cfg.get('cell_margin_top_cm', 0.0),
                bottom_cm=table_cfg.get('cell_margin_bottom_cm', 0.0),
                left_cm=table_cfg.get('cell_margin_left_cm', 0.05),
                right_cm=table_cfg.get('cell_margin_right_cm', 0.05),
            )
            if table_cfg.get('auto_col_width', True):
                _set_table_col_widths_by_content(
                    table,
                    min_pct=table_cfg.get('col_min_pct', 8),
                    max_pct=table_cfg.get('col_max_pct', 45),
                )

        # 表格前空一行
        prev_block = blocks[idx - 1] if idx - 1 >= 0 else None
        prev_para_is_title = isinstance(prev_block, Paragraph) and _is_table_title(prev_block.text)
        prev_para_is_unit = isinstance(prev_block, Paragraph) and _is_table_unit(prev_block.text)

        if isinstance(prev_block, Paragraph):
            if prev_block.text.strip():
                if prev_para_is_title or prev_para_is_unit:
                    _insert_paragraph_before_paragraph(prev_block, text="")
                else:
                    _insert_paragraph_before_table(table, text="")
        elif isinstance(prev_block, Table):
            _insert_paragraph_after_table(prev_block, text="")
        else:
            if idx == 0:
                _insert_paragraph_before_table(table, text="")

        # 标题段落（表格前一段）
        if prev_para_is_title:
            if table_cfg.get('title_align', 'center') == 'center':
                prev_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_block.paragraph_format.space_before = Pt(0)
            prev_block.paragraph_format.space_after = Pt(0)
            prev_block.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 单位段落（表格后一段）
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        unit_para = None
        if isinstance(next_block, Paragraph) and _is_table_unit(next_block.text):
            unit_para = next_block
            if table_cfg.get('unit_align', 'right') == 'right':
                unit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_space_lines = table_cfg.get('unit_space_before_lines', 0.5)
            unit_para.paragraph_format.space_before = Pt(tbl_size * unit_space_lines)
            unit_para.paragraph_format.space_after = Pt(0)
            unit_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 表格内内容
        serial_col_idx = None
        if table.rows:
            header_cells = tuple(table.rows[0].cells)
            for c_idx, cell in enumerate(header_cells):
                head_text = ''.join(p.text for p in cell.paragraphs).strip()
                if '序号' in head_text or head_text == '序':
                    serial_col_idx = c_idx
                    break

        # 标题行重复
        if table_cfg.get('header_repeat', False) and table.rows:
            _set_header_row_repeat(table, True)

        for row_idx, row in enumerate(table.rows):
            if table_cfg.get('row_height_cm'):
                row.height = Cm(table_cfg.get('row_height_cm'))
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            for col_idx, cell in enumerate(tuple(row.cells)):
                if table_cfg.get('optimize', True):
                    _set_cell_borders(cell, size_pt=table_cfg.get('border_size_pt', 0.5),
                                      color=table_cfg.get('border_color', '000000'),
                                      style=table_cfg.get('border_style', 'single'))

                # 标题行底纹
                if row_idx == 0 and table_cfg.get('header_shading_color'):
                    _set_cell_shading(cell, color=table_cfg['header_shading_color'])

                # 垂直对齐
                if table_cfg.get('cell_valign'):
                    _set_cell_vertical_alignment(cell, valign=table_cfg['cell_valign'])

                cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                for para in cell.paragraphs:
                    if para.text.strip():
                        is_header = (row_idx == 0 and tbl_header_bold)
                        for run in para.runs:
                            set_font(run, tbl_font_cn, tbl_font_en, tbl_size,
                                      bold=(tbl_bold or is_header),
                                      italic=table_fmt.get('italic', False),
                                      color=table_fmt.get('color'))

                    para.paragraph_format.first_line_indent = Pt(tbl_first_line_indent)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_cfg.get('paragraph_single', True):
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    else:
                        if tbl_line_spacing:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            para.paragraph_format.line_spacing = Pt(tbl_line_spacing)
                        else:
                            para.paragraph_format.line_spacing = 1.5

                    smart_align = table_cfg.get('smart_align', False)
                    if smart_align:
                        if row_idx == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif '合计' in cell_text or '总计' in cell_text:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif serial_col_idx is not None and col_idx == serial_col_idx:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif _is_numeric_text(cell_text):
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif _is_short_text(cell_text, table_cfg.get('short_text_len', 4)):
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 嵌套表格处理
        nested_tables = _find_nested_tables(table)
        for nested in nested_tables:
            if table_cfg.get('optimize', True):
                nested.autofit = not table_cfg.get('auto_col_width', True)
                _set_table_borders(nested, size_pt=table_cfg.get('border_size_pt', 0.5),
                                   color=table_cfg.get('border_color', '000000'),
                                   style=table_cfg.get('border_style', 'single'))
                _set_table_cell_margins(
                    nested,
                    top_cm=table_cfg.get('cell_margin_top_cm', 0.0),
                    bottom_cm=table_cfg.get('cell_margin_bottom_cm', 0.0),
                    left_cm=table_cfg.get('cell_margin_left_cm', 0.05),
                    right_cm=table_cfg.get('cell_margin_right_cm', 0.05),
                )
            for n_row_idx, n_row in enumerate(nested.rows):
                for n_cell in n_row.cells:
                    _set_cell_borders(n_cell, size_pt=table_cfg.get('border_size_pt', 0.5),
                                      color=table_cfg.get('border_color', '000000'),
                                      style=table_cfg.get('border_style', 'single'))
                    if n_row_idx == 0 and table_cfg.get('header_shading_color'):
                        _set_cell_shading(n_cell, color=table_cfg['header_shading_color'])
                    if table_cfg.get('cell_valign'):
                        _set_cell_vertical_alignment(n_cell, valign=table_cfg['cell_valign'])
                    for n_para in n_cell.paragraphs:
                        if n_para.text.strip():
                            for n_run in n_para.runs:
                                set_font(n_run, tbl_font_cn, tbl_font_en, tbl_size,
                                          bold=tbl_bold,
                                          italic=table_fmt.get('italic', False),
                                          color=table_fmt.get('color'))

        # 表格后空一行
        if table_cfg.get('after_table_blank_line', True):
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if unit_para is not None:
                after_unit = blocks[idx + 2] if idx + 2 < len(blocks) else None
                if not (isinstance(after_unit, Paragraph) and not after_unit.text.strip()):
                    _insert_paragraph_after_paragraph(unit_para, text="")
            else:
                if not (isinstance(next_block, Paragraph) and not next_block.text.strip()):
                    _insert_paragraph_after_table(table, text="")

    # 4.5 中文禁则处理（标点不溢出、遵循中文换行规则）
    _progress(74, 100, '应用中文排版规则...')
    try:
        from .east_asian_typography import apply_chinese_line_break_rules as _apply_cn_rules
        _apply_cn_rules(doc)
    except ImportError:
        pass

    # 4.6 图片处理
    image_cfg_dict = preset.get('image', {})
    if image_cfg_dict:
        logger.info('4.5. Processing images...')
        _progress(76, 100, '处理图片...')
        image_config = PictureConfig.from_dict(image_cfg_dict)
    else:
        image_config = PictureConfig.from_dict(DEFAULT_PICTURE_CONFIG)

    apply_image_size(doc, image_config)
    apply_image_wrapping(doc, image_config)
    apply_image_alignment(doc, image_config)
    compress_images(doc, image_config)

    # 5. 添加页码
    _progress(78, 100, '添加页码...')
    if preset.get('page_number', True):
        logger.info('5. Adding page numbers...')
        add_page_number(
            doc,
            font_name=preset.get('page_number_font', '宋体'),
            font_size=preset.get('page_number_size', 14),
            style=preset.get('page_number_style', 'dash'),
            position=preset.get('page_number_position', 'outside'),
            offset_from_text_mm=preset.get('page_number_offset_mm', 7),
            replace_existing=preset.get('replace_existing_page_number', True),
            bold=preset.get('page_number_bold', False),
        )
    else:
        logger.info('5. Skipping page numbers...')

    # 保存
    _progress(82, 100, '保存文件...')
    doc.save(output_path)
    _progress(85, 100, '格式化完成')

    logger.info('=' * 50)
    logger.info('Formatting Report:')
    logger.info(f'  Preset: {preset.get("name", preset_name)}')
    logger.info(f'  Paper: {paper_size} ({pw:.1f}×{ph:.1f} cm)')
    logger.info(f'  Font:  {body_fmt_cfg.get("font_cn", "仿宋_GB2312")} {body_fmt_cfg.get("size", 16)}pt')

    # 字数统计
    total_chars = sum(len(p.text) for p in doc.paragraphs if p.text.strip())
    total_words_cn = sum(
        sum(1 for ch in p.text if ord(ch) >= 128)
        for p in doc.paragraphs
    )
    total_words_en = sum(
        sum(1 for ch in p.text if ord(ch) < 128 and ch.isalpha())
        for p in doc.paragraphs
    )
    logger.info(f'  Characters: {total_chars} (CN: {total_words_cn}, EN words: {total_words_en})')

    # 段落统计
    logger.info('  Paragraphs:')
    for k, v in stats.items():
        if v > 0:
            logger.info(f'    {k}: {v}')
    total_parsed = sum(stats.values())
    logger.info(f'    total: {total_parsed} non-empty')

    # 表格统计
    table_count = sum(1 for b in blocks if isinstance(b, Table))
    if table_count:
        logger.info(f'  Tables: {table_count}')

    # 错误报告
    if errors:
        logger.warning(f'  Errors: {len(errors)} paragraph(s) failed')
        for err in errors:
            logger.warning(f'    ¶{err["index"]}: {err["error"][:60]}')
    else:
        logger.info(f'  Errors: 0')

    logger.info(f'Output: {output_path}')
