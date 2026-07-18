#!/usr/bin/env python3
"""
文档格式统一 v5
修复问题：
- 标题检测更全面
- 主送机关顶格
- 落款右对齐
- 清除斜体、下划线、颜色
- 特殊段落处理（附件、特此说明等）

公文标准：
- 页边距：上37mm，下35mm，左28mm，右26mm
- 主标题：居中，二号（22pt），方正小标宋简体
- 主送机关：顶格，三号仿宋
- 正文：三号仿宋GB2312，首行缩进2字符，行距28磅
- 一级标题："一、" 三号黑体，首行缩进2字符
- 二级标题："（一）" 三号楷体GB2312，首行缩进2字符
- 三级标题："1." 三号仿宋GB2312，首行缩进2字符
- 四级标题："（1）" 三号仿宋GB2312，首行缩进2字符
- 落款：右对齐，三号仿宋
- 附件：顶格，三号仿宋
"""

import sys
import os
import json
import re
import logging
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger('docformat.formatter')

# ===== 修复 PyInstaller 打包后 python-docx 找不到模板文件的问题 =====
# python-docx 内部用 __file__ 相对路径去读 templates/default-footer.xml 等文件，
# 在 PyInstaller --onefile 模式下该路径会指向临时解压目录，经常找不到文件。
# 解决方案：直接将模板 XML 嵌入代码，monkey-patch 掉文件读取方法。
def _patch_docx_templates():
    """将 python-docx 的模板文件内容直接嵌入，消除对文件系统路径的依赖"""
    from docx.parts.hdrftr import FooterPart, HeaderPart

    _DEFAULT_FOOTER_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:ftr
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"
    xmlns:mv="urn:schemas-microsoft-com:mac:vml"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc:Ignorable="w14 wp14"
    >
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Footer"/>
    </w:pPr>
  </w:p>
</w:ftr>
"""

    _DEFAULT_HEADER_XML = b"""\
<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"
    xmlns:mv="urn:schemas-microsoft-com:mac:vml"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc:Ignorable="w14 wp14"
    >
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Header"/>
    </w:pPr>
  </w:p>
</w:hdr>
"""

    @classmethod
    def _patched_footer_xml(cls):
        return _DEFAULT_FOOTER_XML

    @classmethod
    def _patched_header_xml(cls):
        return _DEFAULT_HEADER_XML

    FooterPart._default_footer_xml = _patched_footer_xml
    HeaderPart._default_header_xml = _patched_header_xml
    logger.debug("python-docx 模板补丁已应用（内嵌 XML，跳过文件读取）")

_patch_docx_templates()
# ===== 补丁结束 =====

# macOS 字体回退映射：Windows 字体名 → macOS 系统字体名（仅当原字体未安装时使用）
MACOS_FONT_FALLBACK = {
    '仿宋_GB2312': 'STFangsong',
    '仿宋': 'STFangsong',
    '黑体': 'STHeiti',
    '楷体_GB2312': 'STKaiti',
    '楷体': 'STKaiti',
    '宋体': 'STSong',
    '方正小标宋简体': 'STSong',
    '方正仿宋_GBK': 'STFangsong',
    '华文仿宋': 'STFangsong',
    '华文中宋': 'STZhongsong',
}

# macOS 上有些用户手动安装的公文字体，字体族名并不是标准的 GB2312，
# 例如可能显示为“仿宋_GB32312 / 楷体_GB32312”。这些字体比系统华文字体
# 更接近公文要求，因此在回退到 STFangsong/STKaiti 前优先使用它们。
MACOS_FONT_ALIASES = {
    '仿宋_GB2312': ['仿宋_GB2312', '仿宋_GB32312', '仿宋', 'FangSong_GB2312', 'FangSong'],
    '楷体_GB2312': ['楷体_GB2312', '楷体_GB32312', '楷体', 'KaiTi_GB2312', 'KaiTi'],
    '方正仿宋_GBK': ['方正仿宋_GBK', '仿宋_GB2312', '仿宋_GB32312', '仿宋', 'FangSong_GB2312', 'FangSong'],
}

# macOS 已安装字体缓存（启动时检测一次）
_macos_installed_fonts = None
_macos_font_detection_done = False

def _get_macos_installed_fonts():
    """获取 macOS 上已安装的字体族名集合（结果会缓存）
    
    检测策略（按可靠性排序）：
    1. 调用 macOS 系统 Python 的 AppKit.NSFontManager（最准确）
    2. 调用 fc-list（需要安装 fontconfig）
    3. 如果都失败，返回 None 表示无法检测（调用方应保守处理）
    """
    global _macos_installed_fonts, _macos_font_detection_done
    if _macos_font_detection_done:
        return _macos_installed_fonts

    _macos_font_detection_done = True

    if sys.platform != 'darwin':
        _macos_installed_fonts = set()
        return _macos_installed_fonts

    import subprocess

    # 方法1：通过 macOS 系统 Python 调用 AppKit（最可靠）
    # /usr/bin/python3 自带 PyObjC，能获取精确的字体族名
    try:
        result = subprocess.run(
            ['/usr/bin/python3', '-c',
             'from AppKit import NSFontManager;'
             'fm=NSFontManager.sharedFontManager();'
             'print("\\n".join(fm.availableFontFamilies()))'],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0 and result.stdout.strip():
            fonts = {name.strip() for name in result.stdout.strip().split('\n') if name.strip()}
            if len(fonts) > 10:  # 合理性校验
                _macos_installed_fonts = fonts
                logger.info(f"macOS 字体检测成功（AppKit），共 {len(fonts)} 个字体族")
                return _macos_installed_fonts
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as e:
        logger.debug(f"AppKit 字体检测失败: {e}")

    # 方法2：fc-list（需安装 fontconfig，如通过 Homebrew）
    try:
        result = subprocess.run(
            ['fc-list', '--format=%{family}\n'],
            capture_output=True, text=True, timeout=10
        )
        if result.returncode == 0 and result.stdout.strip():
            fonts = set()
            for line in result.stdout.strip().split('\n'):
                for name in line.split(','):
                    name = name.strip()
                    if name:
                        fonts.add(name)
            if len(fonts) > 10:
                _macos_installed_fonts = fonts
                logger.info(f"macOS 字体检测成功（fc-list），共 {len(fonts)} 个字体族")
                return _macos_installed_fonts
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as e:
        logger.debug(f"fc-list 字体检测失败: {e}")

    # 都失败了，返回 None（调用方据此决定：无法检测时不替换）
    logger.warning("macOS 字体检测失败，将保持原字体名称不替换")
    _macos_installed_fonts = None
    return _macos_installed_fonts

def _resolve_font_for_macos(font_name):
    """解析单个字体名：优先保留 Windows 原字体名，仅在确认未安装时回退
    
    策略：
    - 字体不在映射表中 → 原样返回
    - 无法检测已安装字体 → 原样返回（宁可让 Word 自动替换，也不主动改错）
    - 检测到已安装 → 原样返回
    - 检测到未安装 → 使用 macOS 系统字体回退
    """
    if font_name not in MACOS_FONT_FALLBACK:
        return font_name

    installed = _get_macos_installed_fonts()

    # 无法检测时，保守策略：不替换，让 Word/WPS 自行处理
    if installed is None:
        logger.debug(f"字体检测不可用，保留原字体名 '{font_name}'")
        return font_name

    if font_name in installed:
        logger.debug(f"字体 '{font_name}' 已安装，直接使用")
        return font_name

    for alias in MACOS_FONT_ALIASES.get(font_name, []):
        if alias in installed:
            logger.info(f"字体 '{font_name}' 未安装，使用兼容字体 '{alias}'")
            return alias

    fallback = MACOS_FONT_FALLBACK[font_name]
    logger.info(f"字体 '{font_name}' 未安装，回退到 '{fallback}'")
    return fallback

def get_missing_cn_fonts(preset):
    """Linux（麒麟/UOS）：用 fc-list 检查预设中文字体是否安装，返回缺失字体列表。

    检测不可用（无 fc-list / 超时）时返回空列表，不打扰用户。
    Word/WPS 在字体缺失时会静默替换为宋体，导致输出不合规而用户不知情，
    因此在处理前给出明确警告。
    """
    if sys.platform in ('win32', 'darwin'):
        return []
    needed = set()
    for key, fmt in preset.items():
        if isinstance(fmt, dict) and fmt.get('font_cn'):
            needed.add(fmt['font_cn'])
    if preset.get('page_number_font'):
        needed.add(preset['page_number_font'])
    if not needed:
        return []
    import subprocess
    try:
        result = subprocess.run(['fc-list', '--format=%{family}\n'],
                                capture_output=True, text=True, timeout=10)
        if result.returncode != 0 or not result.stdout.strip():
            return []
        installed = set()
        for line in result.stdout.strip().split('\n'):
            for name in line.split(','):
                name = name.strip()
                if name:
                    installed.add(name)
        if len(installed) < 5:
            return []
        return sorted(needed - installed)
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return []


def _adapt_fonts_for_platform(preset):
    """在 macOS 上适配字体：优先使用 Windows 原字体，确认未安装时才回退到系统字体"""
    if sys.platform != 'darwin':
        return preset
    import copy
    preset = copy.deepcopy(preset)
    for key, fmt in preset.items():
        if isinstance(fmt, dict) and 'font_cn' in fmt:
            fmt['font_cn'] = _resolve_font_for_macos(fmt['font_cn'])
    # 页码字体
    if 'page_number_font' in preset:
        preset['page_number_font'] = _resolve_font_for_macos(preset['page_number_font'])
    return preset

# 字号对照：二号=22pt，三号=16pt，小四=12pt
# 2字符缩进 = 2 × 16pt = 32pt（三号字）

# 自定义配置文件路径
def load_custom_preset():
    """加载自定义预设"""
    if sys.platform == 'darwin':
        custom_config_file = Path.home() / 'Library' / 'Application Support' / 'DocFormatter' / "custom_settings.json"
    elif sys.platform == 'win32':
        base = os.environ.get('APPDATA') or str(Path.home() / 'AppData' / 'Roaming')
        custom_config_file = Path(base) / 'DocFormatter' / "custom_settings.json"
    else:
        base = os.environ.get('XDG_CONFIG_HOME') or str(Path.home() / '.config')
        custom_config_file = Path(base) / 'DocFormatter' / "custom_settings.json"

    legacy_config_file = Path(__file__).parent.parent / "custom_settings.json"
    if not custom_config_file.exists() and legacy_config_file.exists():
        custom_config_file = legacy_config_file

    if custom_config_file.exists():
        try:
            with open(custom_config_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if isinstance(data, dict) and data.get('schema_version') == 2:
                presets = data.get('presets', [])
                active_id = data.get('active_preset_id')
                if active_id:
                    for preset in presets:
                        if preset.get('id') == active_id:
                            return preset
                return presets[0] if presets else None
            return data
        except Exception:
            return None
    return None


def _merge_preset_settings(base, overrides):
    """递归合并预设配置，用于测试或临时覆盖少量字段。"""
    merged = deepcopy(base)
    for key, value in (overrides or {}).items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _merge_preset_settings(merged[key], value)
        else:
            merged[key] = deepcopy(value)
    return merged

PRESETS = {
    'official': {
        'name': '公文格式',
        'deep_clean': False,
        'page_number': True,
        'page_number_font': '宋体',
        'page_number_size': 14,
        'page_number_style': 'dash',
        'page_number_position': 'outside',
        'page_number_offset_mm': 7,
        'replace_existing_page_number': True,
        'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
        # 密级标识：三号黑体，顶格版心左上角（如"秘密★1年"）
        'security': {
            'font_cn': '黑体',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'left',
            'indent': 0,  # 顶格
            'space_before': 0,
            'space_after': 0,
        },
        # 发文字号：三号仿宋，居中（如"×政发〔2026〕12号"）
        'docnum': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'center',
            'indent': 0,
            'space_before': 0,
            'space_after': 0,
        },
        # 主标题：二号方正小标宋简体，居中
        'title': {
            'font_cn': '方正小标宋简体',
            'font_en': 'Times New Roman',
            'size': 22,  # 二号
            'bold': False,
            'align': 'center',
            'indent': 0,
            'space_before': 0,
            'space_after': 0,
        },
        # 主送机关：三号仿宋，顶格
        'recipient': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 0,  # 顶格
        },
        # 一级标题：三号黑体，"一、"，首行缩进2字符
        'heading1': {
            'font_cn': '黑体',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'left',
            'indent': 32,  # 2字符缩进
            'space_before': 0,
            'space_after': 0,
        },
        # 二级标题：三号楷体GB2312，"（一）"，首行缩进2字符
        'heading2': {
            'font_cn': '楷体_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
            'space_before': 0,
            'space_after': 0,
        },
        # 三级标题：三号仿宋GB2312，"1."，首行缩进2字符
        'heading3': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': True,
            'align': 'left',
            'indent': 32,
            'space_before': 0,
            'space_after': 0,
        },
        # 四级标题：三号仿宋GB2312，"（1）"，首行缩进2字符
        'heading4': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
            'space_before': 0,
            'space_after': 0,
        },
        # 正文：三号仿宋GB2312，首行缩进2字符（32pt），行距28磅
        'body': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 2字符 = 2×16pt
            'line_spacing': 28,
        },
        # 落款单位：三号仿宋，右对齐
        'signature': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 落款日期：三号仿宋，右对齐
        'date': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 附件行：三号仿宋，顶格左对齐
        'attachment': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 0,
        },
        # 结束语（特此说明/通知等）：三号仿宋，首行缩进
        'closing': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
    },
    'academic': {
        'name': '学术论文格式',
        'deep_clean': False,
        'page': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5},
        'security': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'left', 'indent': 0},
        'docnum': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'center', 'indent': 0},
        'title': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 18, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 15, 'bold': True, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'justify', 'indent': 24, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'justify', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 24},
    },
    'legal': {
        'name': '法律文书格式',
        'deep_clean': False,
        'page': {'top': 3.0, 'bottom': 2.5, 'left': 3.0, 'right': 2.5},
        'security': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 16, 'bold': False, 'align': 'left', 'indent': 0},
        'docnum': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'center', 'indent': 0},
        'title': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 22, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'justify', 'indent': 28, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'justify', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 28},
    },
}


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)
    
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def _iter_block_items(doc):
    """Yield paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield Table(child, doc)


def _set_table_borders(table, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))  # OOXML border size is in 1/8 pt
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if cell_mar is None:
        cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(cell_mar)

    def _set_side(tag, cm_value):
        node = cell_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            cell_mar.append(node)
        node.set(qn('w:type'), 'dxa')
        node.set(qn('w:w'), str(int(Cm(cm_value).twips)))

    _set_side('top', top_cm)
    _set_side('bottom', bottom_cm)
    _set_side('left', left_cm)
    _set_side('right', right_cm)


def _set_table_width_percent(table, percent=100):
    percent = max(1, min(100, int(percent)))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(percent * 50))  # 50ths of a percent


def _set_table_indent(table, indent_twips=0):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), str(int(indent_twips)))


def _text_weight(text):
    weight = 0.0
    for ch in text:
        if ord(ch) < 128:
            weight += 0.5
        else:
            weight += 1.0
    return weight


def _normalize_pcts(weights, min_pct, max_pct):
    total = sum(weights) or 1.0
    pcts = [w / total * 100 for w in weights]

    # Clamp low
    for i, v in enumerate(pcts):
        if v < min_pct:
            pcts[i] = min_pct
    # Clamp high
    for i, v in enumerate(pcts):
        if v > max_pct:
            pcts[i] = max_pct

    # Renormalize to 100
    total = sum(pcts) or 1.0
    return [v / total * 100 for v in pcts]


def _set_table_col_widths_by_content(table, min_pct=8, max_pct=45):
    if not table.rows:
        return
    col_count = max(len(row.cells) for row in table.rows)
    if col_count == 0:
        return

    max_weights = [1.0] * col_count
    for row in table.rows:
        # row.cells 每次访问都会重新解析 XML，取一次复用
        for c_idx, cell in enumerate(tuple(row.cells)):
            text = ''.join(p.text for p in cell.paragraphs).strip()
            if text:
                max_weights[c_idx] = max(max_weights[c_idx], _text_weight(text))

    pcts = _normalize_pcts(max_weights, min_pct, max_pct)

    # Set table grid + cell widths in pct
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        tbl.insert(0, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    for pct in pcts:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(pct * 50)))  # pct in 1/50th %
        tbl_grid.append(grid_col)

    for row in table.rows:
        for c_idx, cell in enumerate(tuple(row.cells)):
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'pct')
            tc_w.set(qn('w:w'), str(int(pcts[c_idx] * 50)))


def _insert_paragraph_after_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_table(table, text=""):
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_after_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_paragraph(paragraph, text=""):
    p = OxmlElement("w:p")
    paragraph._p.addprevious(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _is_numeric_text(text):
    text = text.replace(',', '').replace('％', '%').strip()
    if not text:
        return False
    return re.match(r'^[-+]?\d+(?:\.\d+)?%?$', text) is not None


def _is_short_text(text, max_len=4):
    text = text.strip()
    return 0 < len(text) <= max_len


def _is_table_title(text):
    text = text.strip()
    if not text:
        return False
    if len(text) > 30:
        return False
    return re.match(r'^表\s*(?:\d+|[一二三四五六七八九十]+)(?:[-—._、]\d+)?', text) is not None


def _is_table_unit(text):
    text = text.strip()
    if not text:
        return False
    if len(text) > 20:
        return False
    return re.match(r'^单位\s*[:：]', text) is not None


def _set_cell_borders(cell, size_pt=0.5, color="000000"):
    size = max(1, int(size_pt * 8))
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _normalize_date_text(text):
    """Normalize common date variants before regex matching."""
    trans = str.maketrans({
        '０': '0', '１': '1', '２': '2', '３': '3', '４': '4',
        '５': '5', '６': '6', '７': '7', '８': '8', '９': '9',
        '．': '.', '／': '/', '－': '-', '—': '-',
    })
    return re.sub(r'\s+', '', text.strip().translate(trans))


def _is_date_text(text, date_patterns):
    normalized = _normalize_date_text(text)
    return any(pattern.match(normalized) for pattern in date_patterns)


def _standardize_date_text(text):
    """Convert recognized numeric date text to Chinese year/month/day form."""
    normalized = _normalize_date_text(text)
    match = re.match(r'^(\d{4})年(\d{1,2})月(\d{1,2})日?$', normalized)
    if not match:
        match = re.match(r'^(\d{4})[./-](\d{1,2})[./-](\d{1,2})$', normalized)
    if match:
        year, month, day = match.groups()
        return f'{year}年{int(month)}月{int(day)}日'

    match = re.match(r'^(\d{4})年(\d{1,2})月$', normalized)
    if not match:
        match = re.match(r'^(\d{4})[./-](\d{1,2})$', normalized)
    if match:
        year, month = match.groups()
        return f'{year}年{int(month)}月'

    return text


def _build_text_context(doc):
    """Collect non-empty paragraph texts and map document indexes to text indexes."""
    all_texts = []
    all_texts_idx_map = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_texts_idx_map[i] = len(all_texts)
            all_texts.append(text)
    return all_texts, all_texts_idx_map


# 日期行的默认识别模式（合并为单个可覆盖的规则字符串）
_DATE_PATTERNS = (
    r'^\d{4}年\d{1,2}月\d{1,2}日$',
    r'^\d{4}年\d{1,2}月\d{1,2}$',
    r'^\d{4}年\d{1,2}月$',
    r'^\d{4}\.\d{1,2}\.\d{1,2}$',
    r'^\d{4}\.\d{1,2}$',
    r'^\d{4}/\d{1,2}/\d{1,2}$',
    r'^\d{4}/\d{1,2}$',
    r'^\d{4}-\d{1,2}-\d{1,2}$',
    r'^\d{4}-\d{1,2}$',
    r'^二[○〇零oO0][一二三四五六七八九零〇○oO0]{2}年.{1,3}月.{1,3}日$',
    r'^二[○〇零oO0][一二三四五六七八九零〇○oO0]{2}年.{1,3}月$',
)

_CLOSING_PATTERNS = (
    r'^特此(说明|通知|报告|函复|函告|批复|公告|通报)。?$',
    r'^此致$',
    r'^敬礼[！!]?$',
    r'^以上(报告|意见|方案).{0,10}$',
    r'^妥否.{0,10}$',
    r'^请.{0,15}(批示|审批|审议|指示|核准)。?$',
)

# 段落类型识别规则（可被预设的 detect_rules 覆盖，正则字符串）。
# 除正则本身外，引擎还会结合段落位置（如密级只看文首、署名只看文末）
# 等上下文条件，避免正文中的相似文字被误判。
DEFAULT_DETECT_RULES = {
    'security': r'^(绝密|机密|秘密)\s*[★\*]?\s*([一二三四五六七八九十0-9]+\s*(年|个月|月))?\s*$',
    'docnum': r'^[一-鿿]{2,20}[〔\[](19|20)\d{2}[〕\]]\s*第?\s*\d+\s*号$',
    'heading1': r'^[一二三四五六七八九十]+、',
    'heading2': r'^（[一二三四五六七八九十]+）|^\([一二三四五六七八九十]+\)',
    'heading3': r'^\d+\.\s*[^\d.\s]',
    'heading4': r'^（\d+）|^\(\d+\)',
    'recipient': r'^[一-鿿\d、，,（）()\s]+[：:]$',
    'attachment': r'^附件\d*([：:．.\s]|$)',
    'closing': '|'.join(_CLOSING_PATTERNS),
    'date': '|'.join(_DATE_PATTERNS),
    'signature': r'(公司|局|委|部|厅|院|所|中心|办公室|集团|银行|学校|大学|医院'
                 r'|指挥部|领导小组|委员会|管理处|管委会)$',
}

# 预编译（detect_para_type 每段调用，避免重复构建）
_CLOSING_RES = [re.compile(p) for p in _CLOSING_PATTERNS]
_DATE_RES = [re.compile(p) for p in _DATE_PATTERNS]


def _compile_rules(overrides):
    """合并用户自定义识别规则，非法正则自动回退默认值"""
    rules = {}
    overrides = overrides or {}
    for key, default in DEFAULT_DETECT_RULES.items():
        pattern = overrides.get(key) or default
        try:
            rules[key] = re.compile(pattern)
        except re.error:
            rules[key] = re.compile(default)
    return rules


def detect_para_type(text, index, total, alignment, all_texts, all_texts_index=None, prev_para_type=None, rules=None):
    """
    检测段落类型
    返回: 'title', 'recipient', 'heading1', 'heading2', 'heading3', 'heading4', 
          'body', 'signature', 'date', 'attachment', 'closing'
    
    参数:
        text: 段落文本
        index: 段落索引
        total: 总段落数
        alignment: 原始对齐方式
        all_texts: 所有非空段落的文本列表，用于上下文判断
    """
    text = text.strip()
    if not text:
        return 'empty'

    _rules = rules if isinstance(rules, dict) and all(hasattr(v, 'match') for v in rules.values()) else _compile_rules(rules)

    # ===== 密级标识检测（GB/T 9704：版心左上角，如"秘密★1年""机密★3年"）=====
    # 仅识别文档最前部（前 3 个非空段落）中整行只有密级(+保密期限)的段落
    _early_idx = all_texts_index if all_texts_index is not None else index
    if _early_idx < 3 and _rules['security'].match(text):
        return 'security'

    # ===== 发文字号检测（如"×政发〔2026〕12号"，位于文档前部版头区域）=====
    if _early_idx < 6 and _rules['docnum'].match(text):
        return 'docnum'

    # 结束语/日期/附件/署名规则：用户自定义时替换默认（留空即默认）
    closing_patterns = [_rules['closing']]
    date_patterns = [_rules['date']]

    # ===== 标题续行检测 =====
    if prev_para_type == 'title':
        heading_prefix = re.match(r'^[一二三四五六七八九十（\(\d]', text)
        is_recipient_end = re.search(r'[：:]\s*$', text)
        is_attachment = _rules['attachment'].match(text)
        is_closing = any(pattern.match(text) for pattern in closing_patterns)
        is_date = _is_date_text(text, date_patterns)
        is_sentence_end = re.search(r'[。！？.!?；;]\s*$', text)
        if not heading_prefix and not is_recipient_end and not is_attachment and not is_closing and not is_date and not is_sentence_end and len(text) < 50:
            return 'title'

    # ===== 早期日期检测（v1.7.2 新增）=====
    # 日期格式如 "2026.04.20" 会被三级标题规则 ^\d+\.\s*\S 误匹配。
    # 把日期检测提前，确保它优先于 heading3 触发。
    # 注意：仅在文档后部（最后 1/3 段落）才认定为日期，避免正文中的版本号、
    # 编号被误判（如 "1.2.3 版本说明"）。
    if all_texts_index is not None:
        is_in_tail = all_texts_index >= len(all_texts) * 2 // 3
    else:
        is_in_tail = index >= total * 2 // 3
    if is_in_tail:
        if _is_date_text(text, date_patterns):
            return 'date'

    # ===== 附件块延续识别（v1.8.1 新增）=====
    # 一旦上一段是 attachment，且当前段以 "N." 或 "N、" 开头，
    # 当前段也归入 attachment，避免被 heading3/body 抢走。
    if prev_para_type == 'attachment':
        if re.match(r'^\s*\d{1,2}[.、]\s*\S', text):
            return 'attachment'
    
    # ===== 一级标题："一、" "二、" 等 =====
    if _rules['heading1'].match(text):
        return 'heading1'

    # ===== 二级标题："（一）" "（二）" 等 =====
    if _rules['heading2'].match(text):
        return 'heading2'

    # ===== 三级标题："1." "2." 等 =====
    # 约束 . 后面不能再是数字，避免误匹配 "2026.04.20" 点分日期或 "1.2.3" 版本号
    if _rules['heading3'].match(text) and len(text) < 60:
        return 'heading3'

    # ===== 四级标题："（1）" "（2）" 等 =====
    if _rules['heading4'].match(text) and len(text) < 60:
        return 'heading4'
    
    # ===== 主送机关：XXX： 或 XXX: =====
    # 特征：以冒号结尾的名词性短语，不含动词/虚词
    # 如"各处室、直属单位：" "各市（州）教育局："
    if _rules['recipient'].match(text) and len(text) < 30:
        # 排除含动词/虚词的正文句子，如"现将有关事项通知如下："
        body_indicators = (
            r'(现将|为了|根据|按照|经研究|为贯彻|为落实|为进一步|为深入|'
            r'如下|以下|特此|兹将|报告如下|说明如下|通知如下|汇报如下|'
            r'的意见|的通知|的报告|的决定|的请示|的函)'
        )
        if not re.search(body_indicators, text):
            # 如果下一段是标题，则当前段可能是标题的第一行，不应识别为主送机关
            if all_texts_index is not None:
                next_texts = all_texts[all_texts_index + 1: all_texts_index + 2]
                for nt in next_texts:
                    nt = nt.strip()
                    if re.match(r'^关于.+的(通知|报告|请示|函|意见|决定|公告|通报|批复)', nt):
                        break  # 跳过 recipient 判断，继续向下走
                    if 15 < len(nt) < 80 and not re.search(r'[。！？，、；：]$', nt):
                        break
                else:
                    return 'recipient'
            else:
                return 'recipient'
    
    # ===== 附件行 =====
    if _rules['attachment'].match(text):
        return 'attachment'
    
    # ===== 结束语 =====
    for pattern in closing_patterns:
        if pattern.match(text):
            return 'closing'
    
    # ===== 落款日期 =====
    # 支持多种日期格式
    if _is_date_text(text, date_patterns):
        return 'date'
    
    # ===== 落款单位 =====
    # 判断逻辑：在文档后部，短文本，且下一段是日期或者是文档末尾附近
    # v1.7.2：长度上限从 30 提到 60，覆盖联合发文长机关名
    # （如"中共xx办公室、xx政府办公室、xx委员会"）
    if index >= total - 10 and len(text) < 60:
        allow_signature_check = True
        # 文档开头的机关名可能是多行标题的一部分，不应被新增的“委员会”等
        # 单位后缀规则抢先识别为落款。已经出现主送机关或一级标题后才允许
        # 在前 5 个非空段内按落款判断。
        title_region_idx = all_texts_index if all_texts_index is not None else index
        if title_region_idx < 5:
            previous_texts = all_texts[:title_region_idx] if all_texts_index is not None else all_texts[:index]
            has_document_body_started = any(
                re.search(r'[：:]\s*$', pt.strip()) or re.match(r'^[一二三四五六七八九十]+、', pt.strip())
                for pt in previous_texts
            )
            if not has_document_body_started:
                allow_signature_check = False

        if allow_signature_check:
            # 检查是否像单位名称（后缀规则可在预设中自定义）
            if _rules['signature'].search(text):
                return 'signature'
            # 正文句即使靠近文末日期，也不应因为“下文有日期”被误判为落款。
            if re.search(r'[。！？.!?；;]\s*$', text):
                return 'body'
            # 或者检查下文是否有日期
            if all_texts_index is not None:
                remaining_texts = all_texts[all_texts_index + 1:]
            else:
                remaining_texts = []
            for next_text in remaining_texts[:3]:
                if _is_date_text(next_text, date_patterns):
                    return 'signature'
    
    # ===== 主标题 =====
    # 公文结构：[密级/文号] → [标题] → [主送机关] → [正文/各级标题]
    # 一旦出现主送机关（以：结尾）或一级标题（一、），后续段落不再可能是主标题
    # 原逻辑使用 doc.paragraphs 索引，文档开头有空段时会把标题推到 index >= 5，
    # 导致整段 title 检测被跳过；这里优先使用 all_texts_index 这个非空段索引。
    title_region_idx = all_texts_index if all_texts_index is not None else index
    if title_region_idx < 5:
        _check_idx = all_texts_index if all_texts_index is not None else 0
        _title_region_ended = False
        for pt in all_texts[:_check_idx]:
            pt_s = pt.strip()
            # 主送机关：以：或:结尾的短句
            if re.search(r'[：:]\s*$', pt_s) and len(pt_s) < 50:
                _title_region_ended = True
                break
            # 一级标题序号
            if re.match(r'^[一二三四五六七八九十]+、', pt_s):
                _title_region_ended = True
                break
        
        if not _title_region_ended:
            # 1. 明确的公文标题模式
            title_patterns = [
                r'^关于.+的(通知|报告|请示|函|意见|决定|公告|通报|批复|说明|方案|总结|汇报|复函|答复|建议)$',
                r'^.{2,30}(通知|报告|请示|函|意见|决定|公告|通报|批复|工作方案|工作总结|实施方案|管理办法|暂行规定)$',
                r'^[\u4e00-\u9fff]{2,20}(委员会|办公室|局|厅|院|部|委|中心|公司|集团|学校|大学)$',
            ]
            for pattern in title_patterns:
                if re.match(pattern, text):
                    return 'title'
            
            # 2. 较长的标题（20-80字符），不以标点结尾
            if 15 < len(text) < 80 and not re.search(r'[。！？，、；：.!?,;:]$', text):
                # 排除以序号开头的
                if not re.match(r'^[一二三四五六七八九十\d（(]', text):
                    return 'title'
            
            # 3. 居中的短文本（原本就是居中的）
            if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 60:
                return 'title'
    
    # ===== 其他都是正文 =====
    return 'body'


def _split_heading_by_punct(paragraph):
    """Split heading like '（三）xxx：正文' or '（三）xxx。正文' into heading paragraph + body paragraph."""
    text = paragraph.text.strip()
    if not text:
        return False

    # Heading prefix patterns
    if not (
        re.match(r'^[一二三四五六七八九十]+、', text) or
        re.match(r'^（[一二三四五六七八九十]+）', text) or
        re.match(r'^\([一二三四五六七八九十]+\)', text) or
        re.match(r'^\d+\.\s*\S', text) or
        re.match(r'^（\d+）', text) or
        re.match(r'^\(\d+\)', text)
    ):
        return False

    # Find split punctuation position
    punct_positions = []
    for ch in ('：', ':', '。'):
        pos = text.find(ch)
        if pos != -1:
            punct_positions.append(pos)
    if not punct_positions:
        return False
    split_idx = min(punct_positions)
    head = text[:split_idx + 1].strip()
    tail = text[split_idx + 1:].strip()
    if not tail:
        return False

    paragraph.text = head
    new_para = _insert_paragraph_after_paragraph(paragraph, text=tail)
    return new_para is not None


def _ensure_structural_blank_lines(doc, line_spacing_pt=28, rules=None, type_overrides=None):
    """
    Ensure the standard visible blank lines:
    - after the title block before recipient/body
    - before the signature block after the final body paragraph
    Other empty paragraphs are handled separately.
    """
    all_texts, all_texts_idx_map = _build_text_context(doc)
    total_paras = len(doc.paragraphs)
    entries = []
    prev_para_type = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        ai = all_texts_idx_map.get(i)
        if type_overrides and ai is not None and ai in type_overrides:
            para_type = type_overrides[ai]
        else:
            para_type = detect_para_type(
                text, i, total_paras,
                para.paragraph_format.alignment,
                all_texts,
                all_texts_index=ai,
                prev_para_type=prev_para_type,
                rules=rules
            )
        entries.append((para, para_type, prev_para_type))
        prev_para_type = para_type

    structural_blank_ids = set()
    p_tag = qn('w:p')
    for para, para_type, prev_para_type in entries:
        needs_blank = (
            (prev_para_type == 'title' and para_type not in ('title', 'docnum')) or
            (prev_para_type == 'docnum' and para_type != 'docnum') or
            (para_type == 'signature' and prev_para_type not in (None, 'signature', 'date'))
        )
        if not needs_blank:
            continue

        # 直接检查 XML 前一个兄弟节点，避免每次全量扫描 doc.paragraphs（O(n²)）
        prev_el = para._p.getprevious()
        if prev_el is not None and prev_el.tag == p_tag and not Paragraph(prev_el, para._parent).text.strip():
            blank_para = Paragraph(prev_el, para._parent)
        else:
            blank_para = _insert_paragraph_before_paragraph(para)

        _format_structural_blank_paragraph(blank_para, line_spacing_pt)
        structural_blank_ids.add(id(blank_para._p))

    return structural_blank_ids


def _format_empty_paragraphs(doc, structural_blank_ids, line_spacing_pt=28):
    """v1.7.2: 不再依赖 id()，改用段落 XML 上的持久标记判断。
    structural_blank_ids 参数保留以维持接口兼容，但不再使用。"""
    for para in doc.paragraphs:
        if para.text.strip():
            continue
        if _is_structural_blank(para):
            _format_structural_blank_paragraph(para, line_spacing_pt)
        else:
            _compact_empty_paragraph(para)


# ===== 修订标记辅助 =====
_revision_counter = [0]   # 列表以便嵌套函数修改


def _next_rev_id():
    _revision_counter[0] += 1
    return _revision_counter[0]


def _rev_date():
    return datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')


def _add_ppr_change(para, orig_ppr):
    """将原始段落格式嵌入 <w:pPrChange>，记录改动前状态"""
    pPr = para._p.get_or_add_pPr()
    # 移除已有的 pPrChange，避免重复
    for old in pPr.findall(qn('w:pPrChange')):
        pPr.remove(old)

    change = OxmlElement('w:pPrChange')
    change.set(qn('w:id'),     str(_next_rev_id()))
    change.set(qn('w:author'), '公文格式工具')
    change.set(qn('w:date'),   _rev_date())

    if orig_ppr is not None:
        snapshot = deepcopy(orig_ppr)
        # 清除快照自身的 pPrChange，防止嵌套
        for old in snapshot.findall(qn('w:pPrChange')):
            snapshot.remove(old)
        change.append(snapshot)
    else:
        change.append(OxmlElement('w:pPr'))

    pPr.append(change)


def _add_rpr_change(run, orig_rpr):
    """将原始字符格式嵌入 <w:rPrChange>，记录改动前状态"""
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rPrChange')):
        rPr.remove(old)

    change = OxmlElement('w:rPrChange')
    change.set(qn('w:id'),     str(_next_rev_id()))
    change.set(qn('w:author'), '公文格式工具')
    change.set(qn('w:date'),   _rev_date())

    if orig_rpr is not None:
        snapshot = deepcopy(orig_rpr)
        for old in snapshot.findall(qn('w:rPrChange')):
            snapshot.remove(old)
        change.append(snapshot)
    else:
        change.append(OxmlElement('w:rPr'))

    rPr.append(change)
# ===== 修订标记辅助结束 =====


def _set_paragraph_spacing_points(para, before_pt=0, after_pt=0):
    """Set paragraph spacing in points and clear line-based spacing leftovers."""
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)

    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)

    for attr in ('beforeLines', 'afterLines', 'beforeAutospacing', 'afterAutospacing'):
        spacing.attrib.pop(qn(f'w:{attr}'), None)

    spacing.set(qn('w:before'), str(int(round(before_pt * 20))))
    spacing.set(qn('w:after'), str(int(round(after_pt * 20))))


def _force_normal_style(para):
    """把段落 style 重置为 Normal，避免内置 Heading 样式带来的属性继承。

    v1.7.2: Word 内置 Heading 1~9 / Normal (Web) 等样式自带
    beforeAutospacing/afterAutospacing="1"，会让段前段后无法清零。
    公文格式应走纯 Normal 样式，所有视觉效果通过段落直接属性控制。
    """
    try:
        pPr = para._p.get_or_add_pPr()
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            pStyle = OxmlElement('w:pStyle')
            # pStyle 必须放在 pPr 的最前面（W3C 规范要求）
            pPr.insert(0, pStyle)
        pStyle.set(qn('w:val'), 'Normal')
    except Exception:
        pass


def deep_clean_document(doc):
    """深度清洗文档：移除所有段落级用户格式属性，保留文字和结构。

    v1.8.0: 处理复制粘贴的脏数据时，原文带的颜色、字号、缩进、段前段后
    等用户级格式会干扰 detect_para_type 的启发式判断。本函数把这些
    属性全部清掉，让后续格式化工作在干净的输入上展开。

    注意：本函数不删除文字、不动表格结构。
    """

    def _clean_paragraph(para):
        pf = para.paragraph_format
        # 段落级属性清零
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = None
        pf.right_indent = None
        pf.first_line_indent = None
        # 行距交给后面的格式化处理，这里先清掉
        pf.line_spacing = None
        pf.line_spacing_rule = None

        # 强制 style 为 Normal（避免 Heading 样式继承）
        _force_normal_style(para)

        # Run 级属性清零
        for run in para.runs:
            run.font.color.rgb = None
            run.font.highlight_color = None
            # size / bold / italic 等让后面 set_font 来重置
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.underline = None
            run.font.strike = None

    for para in doc.paragraphs:
        _clean_paragraph(para)

    # 表格内段落同样清洗
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _clean_paragraph(para)


def _strip_autospacing_from_styles(doc):
    """清理整个文档样式表里的 beforeAutospacing/afterAutospacing 属性。

    v1.7.2: Word 默认内置样式带 Autospacing="1"，会让段落直接属性的
    before/after 无法生效。公文格式不需要 Autospacing。
    """
    try:
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        styles_element = doc.styles.element
        for spacing in styles_element.iter(f'{ns}spacing'):
            for attr in ('beforeAutospacing', 'afterAutospacing'):
                spacing.attrib.pop(f'{ns}{attr}', None)
    except Exception:
        pass


def _compact_empty_paragraph(para):
    """Clear spacing on empty paragraphs so inherited template gaps do not remain."""
    _set_paragraph_spacing_points(para, 0, 0)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(1)


def _format_structural_blank_paragraph(para, line_spacing_pt=28):
    """Format the intentional blank line used between document sections.

    v1.7.2: 在段落 pPr 上写一个自定义属性 docfmt:structural-blank=1
    作为持久标识，避免依赖 Python 对象 id 在段落被修改后失效。
    """
    if not para.runs:
        para.add_run(' ')
    _set_paragraph_spacing_points(para, 0, 0)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    # 持久标记：使用自定义命名空间属性，Word 会忽略它但 python-docx 能读
    _mark_structural_blank(para)


def _mark_structural_blank(para):
    """在段落 pPr 上写自定义标记，标识为结构性空行。"""
    pPr = para._p.get_or_add_pPr()
    pPr.set('docfmt-structural-blank', '1')


def _is_structural_blank(para):
    """检查段落是否被标记为结构性空行。"""
    pPr = para._p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is None:
        return False
    return pPr.get('docfmt-structural-blank') == '1'


def set_font(run, font_cn, font_en, size, bold=False, revision_mode=False):
    """
    设置字体，同时清除原有格式（斜体、下划线、颜色）
    """
    # 修订模式：记录改动前的 rPr
    if revision_mode:
        orig_rpr = deepcopy(run._r.rPr)
        orig_xml = run._r.xml

    # 基本字体设置
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    
    # 清除斜体
    run.font.italic = False
    
    # 清除下划线
    run.font.underline = False
    
    # 清除颜色（设置为黑色）
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 清除删除线
    run.font.strike = False
    run.font.double_strike = False
    
    # 清除上下标
    run.font.subscript = False
    run.font.superscript = False
    
    # 设置中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)

    # 修订模式：若有改动则嵌入 rPrChange
    if revision_mode and run._r.xml != orig_xml:
        _add_rpr_change(run, orig_rpr)


def format_paragraph(para, fmt, para_type, line_spacing_pt=28, first_line_bold=False, revision_mode=False, bold_serial=True):
    """格式化段落
    
    fmt 支持的字段:
        font_cn, font_en, size, bold, align, indent,
        line_spacing  - 行距(磅), None表示使用1.5倍行距
        space_before  - 段前间距(磅), 默认0
        space_after   - 段后间距(磅), 默认0
    """
    # v1.7.2: 重置段落 style 为 Normal，避免继承 Heading 样式上的 Autospacing
    _force_normal_style(para)

    # 修订模式：记录段落格式改动前的 pPr XML
    if revision_mode:
        orig_ppr = deepcopy(para._p.pPr)
        orig_ppr_xml = para._p.xml

    pf = para.paragraph_format
    
    # 对齐方式
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # 段落左缩进清零（重要：确保"文本之前缩进"为0）
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    
    # v1.8.1: attachment 类型走悬挂缩进，不走通用缩进逻辑
    if para_type == 'attachment':
        font_size_pt = fmt.get('size', 16) or 16
        # 左缩进 5 字符（首段折行后的对齐位置，后续段的左边界）
        pf.left_indent = Pt(font_size_pt * 5)

        # 是否是"首段"（含"附件"关键字）—— 走悬挂缩进
        if '附件' in para.text:
            pf.first_line_indent = Pt(-font_size_pt * 3)
        else:
            # 后续 "2.xxx" "3.xxx" —— 顶左缩进（first_line_indent = 0）
            pf.first_line_indent = Pt(0)

        try:
            pPr = para._p.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:firstLineChars'), None)
        except Exception:
            pass
        _attachment_indent_done = True
    else:
        _attachment_indent_done = False

    # 首行缩进
    if not _attachment_indent_done:
        indent = fmt.get('indent', 0)
        if indent > 0:
            pf.first_line_indent = Pt(indent)
            # 同时设置 w:firstLineChars 让 Word 显示为"X 字符"而非厘米
            # w:firstLineChars 单位是 1/100 字符，2字符 = 200
            size = fmt.get('size', 16) or 16
            try:
                chars_100 = int(round(indent / size * 100))
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                ind.set(qn('w:firstLineChars'), str(chars_100))
            except Exception:
                pass
        else:
            pf.first_line_indent = Pt(0)
            try:
                pPr = para._p.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    ind.attrib.pop(qn('w:firstLineChars'), None)
            except Exception:
                pass
    
    # 行距：优先读取当前元素自身的 line_spacing，否则用全局默认值
    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5
    
    # 段前段后（支持自定义，默认0）
    _set_paragraph_spacing_points(
        para,
        fmt.get('space_before', 0),
        fmt.get('space_after', 0)
    )
    
    # 字体 - 支持首句加粗
    if first_line_bold and para_type == 'body':
        # 首句以中文句号“。”作为结束
        full_text = para.text
        first_sentence_end = full_text.find('。')
        if first_sentence_end != -1:
            split_idx = first_sentence_end + 1
            first_part = full_text[:split_idx]
            rest_part = full_text[split_idx:]
            
            # 重新构建 runs，确保只加粗首句
            for run in list(para.runs):
                para._p.remove(run._r)
            
            run1 = para.add_run(first_part)
            set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True, revision_mode=revision_mode)
            
            if rest_part:
                run2 = para.add_run(rest_part)
                set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False), revision_mode=revision_mode)
        else:
            # 没找到中文句号，正常处理
            for run in para.runs:
                set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False), revision_mode=revision_mode)
    else:
        # 正文里的序列词加粗前缀
        if bold_serial and para_type == 'body':
            _SERIAL_PATTERNS = [
                r'^([一二三四五六七八九十]{1,3}是)([：:、]?)',       # 一是、二是
                r'^([一二三四五六七八九十]{1,3}要)([：:、]?)',       # 一要、二要
                r'^(第[一二三四五六七八九十百\d]+[点条项步])([：:、，,]?)',  # 第一点、第二条
                r'^([一二三四五六七八九十]{1,3}方面)([：:、]?)',     # 一方面、二方面
            ]
            m = None
            for _pat in _SERIAL_PATTERNS:
                m = re.match(_pat, para.text)
                if m:
                    break
            if m:
                lead = m.group(1) + (m.group(2) or '')
                rest = para.text[len(lead):]
                for run in list(para.runs):
                    para._p.remove(run._r)
                run1 = para.add_run(lead)
                set_font(run1, fmt['font_cn'], fmt['font_en'], fmt['size'], bold=True, revision_mode=revision_mode)
                if rest:
                    run2 = para.add_run(rest)
                    set_font(run2, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False), revision_mode=revision_mode)
                return

        # 正常处理
        for run in para.runs:
            set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False), revision_mode=revision_mode)

    # 修订模式：若段落格式有改动则嵌入 pPrChange
    if revision_mode and para._p.xml != orig_ppr_xml:
        _add_ppr_change(para, orig_ppr)


def add_page_number(
    doc,
    font_name="宋体",
    font_size=14,
    style="dash",
    position="outside",
    offset_from_text_mm=7,
    replace_existing=True,
):
    """按自定义样式添加页码。

    offset_from_text_mm 表示页码位于版心下边缘以下的距离，不是距纸张底边。
    对标准公文下边距 35mm，偏移 7mm 对应 Word 页脚距底边约 28mm。
    """
    def _footer_state(footer):
        """返回 (是否有内容, 是否包含 PAGE 页码域)。"""
        has_content = False
        has_page_field = False
        for para in footer.paragraphs:
            paragraph_text = para.text.strip()
            if paragraph_text:
                has_content = True
                if re.fullmatch(
                    r"[—\-–\s　]*(?:第\s*)?\d+(?:\s*/\s*\d+)?(?:\s*页)?[—\-–\s　]*",
                    paragraph_text,
                ):
                    has_page_field = True
            for run in para.runs:
                xml = run._r.xml or ""
                if 'fldChar' in xml or 'instrText' in xml:
                    has_content = True
                if re.search(r"\bPAGE\b", xml, re.I):
                    has_page_field = True
        return has_content, has_page_field

    for section in doc.sections:
        states = (
            _footer_state(section.footer),
            _footer_state(section.even_page_footer),
            _footer_state(section.first_page_footer),
        )
        has_non_page_content = any(has_content and not has_page for has_content, has_page in states)
        has_existing_page = any(has_page for _has_content, has_page in states)
        if has_non_page_content:
            logger.warning("页脚含有非页码内容，为避免覆盖已跳过页码重设")
            return
        if has_existing_page and not replace_existing:
            return

    # 启用奇偶页页眉页脚（文档级）
    use_even_footer = position == "outside"
    try:
        doc.settings.odd_and_even_pages_header_footer = use_even_footer
    except Exception:
        settings_el = doc.settings._element
        even_odd = settings_el.find(qn('w:evenAndOddHeaders'))
        if use_even_footer and even_odd is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))
        elif not use_even_footer and even_odd is not None:
            settings_el.remove(even_odd)

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = use_even_footer
        bottom_margin_cm = section.bottom_margin.cm if section.bottom_margin else 3.5
        footer_distance_cm = max(0.3, bottom_margin_cm - float(offset_from_text_mm) / 10)
        section.footer_distance = Cm(footer_distance_cm)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        first_footer = section.first_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False
        first_footer.is_linked_to_previous = False

        for f in (odd_footer, even_footer, first_footer):
            for para in f.paragraphs:
                para.clear()

        def _add_field(paragraph, instruction):
            begin_run = paragraph.add_run()
            begin = OxmlElement('w:fldChar')
            begin.set(qn('w:fldCharType'), 'begin')
            begin_run._r.append(begin)
            set_font(begin_run, font_name, font_name, font_size, bold=False)

            instruction_run = paragraph.add_run()
            instruction_text = OxmlElement('w:instrText')
            instruction_text.set(qn('xml:space'), 'preserve')
            instruction_text.text = instruction
            instruction_run._r.append(instruction_text)
            set_font(instruction_run, font_name, font_name, font_size, bold=False)

            end_run = paragraph.add_run()
            end = OxmlElement('w:fldChar')
            end.set(qn('w:fldCharType'), 'end')
            end_run._r.append(end)
            set_font(end_run, font_name, font_name, font_size, bold=False)

        def _build_footer_line(footer, align, leading_space=False, trailing_space=False):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align
            para.paragraph_format.left_indent = None
            para.paragraph_format.right_indent = None
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            if leading_space:
                run0 = para.add_run("　")
                set_font(run0, font_name, font_name, font_size, bold=False)

            if style == "dash":
                run = para.add_run("— ")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " PAGE ")
                run = para.add_run(" —")
                set_font(run, font_name, font_name, font_size, bold=False)
            elif style == "page_text":
                run = para.add_run("第 ")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " PAGE ")
                run = para.add_run(" 页")
                set_font(run, font_name, font_name, font_size, bold=False)
            elif style == "page_total":
                _add_field(para, " PAGE ")
                run = para.add_run(" / ")
                set_font(run, font_name, font_name, font_size, bold=False)
                _add_field(para, " NUMPAGES ")
            else:
                _add_field(para, " PAGE ")

            if trailing_space:
                run6 = para.add_run("　")
                set_font(run6, font_name, font_name, font_size, bold=False)

        if position == "outside":
            # 奇数页居右，空格在右；偶数页居左，空格在左。
            _build_footer_line(
                odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, trailing_space=True
            )
            _build_footer_line(
                even_footer, WD_ALIGN_PARAGRAPH.LEFT, leading_space=True
            )
            if section.different_first_page_header_footer:
                _build_footer_line(
                    first_footer, WD_ALIGN_PARAGRAPH.RIGHT, trailing_space=True
                )
        else:
            align = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(position, WD_ALIGN_PARAGRAPH.CENTER)
            _build_footer_line(odd_footer, align)
            if section.different_first_page_header_footer:
                _build_footer_line(first_footer, align)


def format_document(input_path, output_path, preset_name='official', progress_callback=None,
                    revision_mode=False, bold_serial=True, custom_settings=None,
                    type_overrides=None):
    """格式化文档

    Args:
        progress_callback: 可选回调函数，签名为 callback(current, total, stage_text)
        type_overrides: 可选 {非空段序号: 段落类型}，用户在预览中手动指定的
            段落类型，优先于自动识别（序号按非空段落从 0 计数）
    """
    _revision_counter[0] = 0   # 每篇文档从 1 开始计 ID

    # 处理自定义预设
    if preset_name == 'custom' and custom_settings:
        preset = deepcopy(custom_settings)
        logger.info(f'Preset: {preset.get("name", "自定义格式")}')
    elif preset_name == 'custom':
        preset = load_custom_preset()
        if preset is None:
            logger.warning('Custom preset not found, using official preset')
            preset = PRESETS['official']
        else:
            logger.info(f'Preset: {preset.get("name", "自定义格式")}')
    elif preset_name not in PRESETS:
        logger.error(f'Unknown preset: {preset_name}')
        logger.error(f'Available: {", ".join(PRESETS.keys())}')
        sys.exit(1)
    else:
        preset = PRESETS[preset_name]
        logger.info(f'Preset: {preset["name"]}')

    if custom_settings and preset_name != 'custom':
        preset = _merge_preset_settings(preset, custom_settings)
    
    logger.info(f'Input: {input_path}')
    preset = _adapt_fonts_for_platform(preset)
    
    # 获取首句加粗选项
    first_line_bold = preset.get('first_line_bold', False)
    # bold_serial 优先使用 preset 的值；preset 没设置时用入参兜底。
    # 修复：原代码 `preset.get('bold_serial', True)` 会硬覆盖入参。
    bold_serial = preset.get('bold_serial', bold_serial)
    
    doc = Document(input_path)

    # v1.8.0: 强力清洗模式（如开启，在格式化前先深度清理段落属性）
    if preset.get('deep_clean', False):
        deep_clean_document(doc)

    # 将“标题+标点+正文”拆分为标题段+正文段
    # v1.7.1：默认关闭，避免破坏用户故意写成一行的合法格式（如
    # "1. 第一阶段：完成报废资产清单整理"）。需要这个拆分的用户
    # 可以在 preset 里设置 split_heading_at_punct=True。
    if preset.get('split_heading_at_punct', False):
        for para in list(doc.paragraphs):
            _split_heading_by_punct(para)

    # 进度回调辅助函数
    def _progress(current, total, stage):
        if progress_callback:
            progress_callback(current, total, stage)
    
    # 1. 移除背景
    logger.info('1. Removing background...')
    _progress(0, 100, '移除背景...')
    remove_background(doc)
    
    # 2. 设置页面边距
    logger.info('2. Setting page margins...')
    _progress(5, 100, '设置页面边距...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])

    body_line_spacing = preset.get('body', {}).get('line_spacing', 28) or 28

    # 标准公文版式保留两处可见空行：标题后、落款前。
    # v1.7.2: 清理 styles.xml 里的 Autospacing 属性，避免内置样式覆盖直接属性
    _strip_autospacing_from_styles(doc)
    _active_rules = _compile_rules(preset.get('detect_rules'))
    structural_blank_ids = _ensure_structural_blank_lines(
        doc, body_line_spacing, rules=_active_rules, type_overrides=type_overrides)
    total_paras = len(doc.paragraphs)
    all_texts, all_texts_idx_map = _build_text_context(doc)
    
    # 3. 格式化段落
    logger.info('3. Formatting paragraphs...')
    _progress(10, 100, '格式化段落...')
    stats = {
        'title': 0, 'recipient': 0, 'heading1': 0, 'heading2': 0, 
        'heading3': 0, 'heading4': 0, 'body': 0, 'signature': 0, 
        'date': 0, 'attachment': 0, 'closing': 0
    }
    
    prev_para_type = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            if _is_structural_blank(para):
                _format_structural_blank_paragraph(para, body_line_spacing)
            else:
                _compact_empty_paragraph(para)
            continue
        
        _ai = all_texts_idx_map.get(i)
        if type_overrides and _ai is not None and _ai in type_overrides:
            para_type = type_overrides[_ai]
        else:
            para_type = detect_para_type(
                text, i, total_paras,
                para.paragraph_format.alignment,
                all_texts,
                all_texts_index=_ai,
                prev_para_type=prev_para_type,
                rules=_active_rules
            )

        if para_type == 'date':
            standardized_date = _standardize_date_text(text)
            if standardized_date != text:
                para.text = standardized_date
                text = standardized_date

        # 选择对应的格式
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        
        format_paragraph(
            para, fmt, para_type,
            first_line_bold=first_line_bold,
            revision_mode=revision_mode,
            bold_serial=bold_serial
        )
        stats[para_type] = stats.get(para_type, 0) + 1
        
        # 打印处理信息
        preview = text[:35] + '...' if len(text) > 35 else text
        logger.info(f'   [{para_type:10}] {preview}')
        
        # 进度：10% ~ 80%
        if total_paras > 0:
            pct = 10 + int(70 * (i + 1) / total_paras)
            _progress(pct, 100, f'格式化段落 ({i + 1}/{total_paras})')

        prev_para_type = para_type

    # 格式化后再复查一次。部分文档的标题依赖原始对齐或字体较难在第一次
    # 识别时命中，完成标题格式化后再补齐结构空行更稳。
    structural_blank_ids.update(_ensure_structural_blank_lines(
        doc, body_line_spacing, rules=_active_rules, type_overrides=type_overrides))
    _format_empty_paragraphs(doc, structural_blank_ids, body_line_spacing)
    
    # 4. 处理表格
    logger.info('4. Formatting tables...')
    _progress(82, 100, '格式化表格...')
    body_fmt = preset.get('body', {})
    # 表格配置：优先使用 preset 中的 table 节点，否则用 body 格式
    table_fmt = preset.get('table', {})
    table_defaults = {
        'optimize': True,
        'border_size_pt': 0.5,
        'width_percent': 100,
        'auto_col_width': True,
        'col_min_pct': 8,
        'col_max_pct': 45,
        'row_height_cm': 0.7,
        'cell_margin_top_cm': 0.0,
        'cell_margin_bottom_cm': 0.0,
        'cell_margin_left_cm': 0.05,
        'cell_margin_right_cm': 0.05,
        'paragraph_single': True,
        'after_table_blank_line': True,
        'title_align': 'center',
        'unit_align': 'right',
        'unit_space_before_lines': 0.5,
        'short_text_len': 4,
        'smart_align': False,
    }
    table_cfg = {**table_defaults, **table_fmt}

    tbl_font_cn = table_fmt.get('font_cn', body_fmt.get('font_cn', '仿宋_GB2312'))
    tbl_font_en = table_fmt.get('font_en', body_fmt.get('font_en', 'Times New Roman'))
    tbl_size = table_fmt.get('size', body_fmt.get('size', 16))
    tbl_bold = table_fmt.get('bold', False)
    tbl_line_spacing = table_fmt.get('line_spacing', body_fmt.get('line_spacing', 28))
    tbl_header_bold = table_fmt.get('header_bold', False)
    tbl_first_line_indent = table_fmt.get('first_line_indent', 0)

    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue

        table = block
        if table_cfg.get('optimize', True):
            table.autofit = not table_cfg.get('auto_col_width', True)
            _set_table_width_percent(table, table_cfg.get('width_percent', 100))
            _set_table_indent(table, 0)
            _set_table_borders(table, size_pt=table_cfg.get('border_size_pt', 0.5))
            _set_table_cell_margins(
                table,
                top_cm=table_cfg.get('cell_margin_top_cm', 0.0),
                bottom_cm=table_cfg.get('cell_margin_bottom_cm', 0.0),
                left_cm=table_cfg.get('cell_margin_left_cm', 0.05),
                right_cm=table_cfg.get('cell_margin_right_cm', 0.05),
            )
            if table_cfg.get('auto_col_width', True):
                _set_table_col_widths_by_content(
                    table,
                    min_pct=table_cfg.get('col_min_pct', 8),
                    max_pct=table_cfg.get('col_max_pct', 45),
                )

        # 表格前空一行（如果已有空行则不重复）
        prev_block = blocks[idx - 1] if idx - 1 >= 0 else None
        prev_para_is_title = isinstance(prev_block, Paragraph) and _is_table_title(prev_block.text)
        prev_para_is_unit = isinstance(prev_block, Paragraph) and _is_table_unit(prev_block.text)

        if isinstance(prev_block, Paragraph):
            if prev_block.text.strip():
                if prev_para_is_title or prev_para_is_unit:
                    _insert_paragraph_before_paragraph(prev_block, text="")
                else:
                    _insert_paragraph_before_table(table, text="")
        elif isinstance(prev_block, Table):
            _insert_paragraph_after_table(prev_block, text="")
        else:
            if idx == 0:
                _insert_paragraph_before_table(table, text="")

        # 标题段落（表格前一段）
        if prev_para_is_title:
            if table_cfg.get('title_align', 'center') == 'center':
                prev_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_block.paragraph_format.space_before = Pt(0)
            prev_block.paragraph_format.space_after = Pt(0)
            prev_block.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 单位段落（表格后一段）
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        unit_para = None
        if isinstance(next_block, Paragraph) and _is_table_unit(next_block.text):
            unit_para = next_block
            if table_cfg.get('unit_align', 'right') == 'right':
                unit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_space_lines = table_cfg.get('unit_space_before_lines', 0.5)
            unit_para.paragraph_format.space_before = Pt(tbl_size * unit_space_lines)
            unit_para.paragraph_format.space_after = Pt(0)
            unit_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 表格内内容
        serial_col_idx = None
        if table.rows:
            header_cells = tuple(table.rows[0].cells)
            for c_idx, cell in enumerate(header_cells):
                head_text = ''.join(p.text for p in cell.paragraphs).strip()
                if '序号' in head_text or head_text == '序':
                    serial_col_idx = c_idx
                    break

        for row_idx, row in enumerate(table.rows):
            # 行高
            if table_cfg.get('row_height_cm'):
                row.height = Cm(table_cfg.get('row_height_cm'))
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            for col_idx, cell in enumerate(tuple(row.cells)):
                if table_cfg.get('optimize', True):
                    _set_cell_borders(cell, size_pt=table_cfg.get('border_size_pt', 0.5))

                cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                for para in cell.paragraphs:
                    # 字体设置
                    if para.text.strip():
                        is_header = (row_idx == 0 and tbl_header_bold)
                        for run in para.runs:
                            set_font(run, tbl_font_cn, tbl_font_en, tbl_size, bold=(tbl_bold or is_header))

                    # 段落格式
                    para.paragraph_format.first_line_indent = Pt(tbl_first_line_indent)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    if table_cfg.get('paragraph_single', True):
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    else:
                        if tbl_line_spacing:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            para.paragraph_format.line_spacing = Pt(tbl_line_spacing)
                        else:
                            para.paragraph_format.line_spacing = 1.5

                    # 对齐策略
                    smart_align = table_cfg.get('smart_align', False)
                    if smart_align:
                        # 智能对齐：按内容类型判断
                        if row_idx == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif '合计' in cell_text or '总计' in cell_text:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif serial_col_idx is not None and col_idx == serial_col_idx:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif _is_numeric_text(cell_text):
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif _is_short_text(cell_text, table_cfg.get('short_text_len', 4)):
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # smart_align=False 时不修改对齐，保留原始格式

        # 表格后空一行（若已有空行则不重复）
        if table_cfg.get('after_table_blank_line', True):
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if unit_para is not None:
                # 单位行后再空一行
                after_unit = blocks[idx + 2] if idx + 2 < len(blocks) else None
                if not (isinstance(after_unit, Paragraph) and not after_unit.text.strip()):
                    _insert_paragraph_after_paragraph(unit_para, text="")
            else:
                if not (isinstance(next_block, Paragraph) and not next_block.text.strip()):
                    _insert_paragraph_after_table(table, text="")
    
    # 5. 添加页码
    _progress(78, 100, '添加页码...')
    if preset.get('page_number', True):
        logger.info('5. Adding page numbers...')
        add_page_number(
            doc,
            font_name=preset.get('page_number_font', '宋体'),
            font_size=preset.get('page_number_size', 14),
            style=preset.get('page_number_style', 'dash'),
            position=preset.get('page_number_position', 'outside'),
            offset_from_text_mm=preset.get('page_number_offset_mm', 7),
            replace_existing=preset.get('replace_existing_page_number', True),
        )
    else:
        logger.info('5. Skipping page numbers...')
    
    # 保存
    _progress(82, 100, '保存文件...')
    doc.save(output_path)
    _progress(85, 100, '格式化完成')
    
    logger.info('=' * 50)
    logger.info('Statistics:')
    for k, v in stats.items():
        if v > 0:
            logger.info(f'  {k}: {v}')
    logger.info(f'Output: {output_path}')


if __name__ == '__main__':
    # CLI模式：日志输出到终端
    logging.basicConfig(level=logging.INFO, format='%(message)s')
    
    if len(sys.argv) < 3:
        print('Usage: python formatter.py input.docx output.docx [--preset official|academic|legal]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    preset = 'official'
    if '--preset' in sys.argv:
        idx = sys.argv.index('--preset')
        if idx + 1 < len(sys.argv):
            preset = sys.argv[idx + 1]
    
    format_document(input_file, output_file, preset)
