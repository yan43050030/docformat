# -*- coding: utf-8 -*-
"""套打模板可视化编辑：把模板 docx 当成"可以点、可以拖"的一张纸。

模板以前只有两条路能改：改 tools/make_songshendan.py 里的 SPEC 重跑脚本，
或者拿 Word 手工调。前者只有开发能干，后者一改就把制表位、精确行距、
白字这些讲究的地方弄坏。这里给第三条路——在画布上点中一个字就能改它的
内容、字体、字号、位置，改完存回 docx，讲究的地方由代码替用户守着。

坐标一律是"距纸张左边 / 距纸张上边"的厘米数，与 overprint.layout_doc
算出来的是同一套，所以编辑时看到的位置就是打印出来的位置。

编辑的落点始终是 OOXML 本身：
  · 横向 → 段落制表位（缇为单位，与字体无关）
  · 纵向 → 段落的段前距（模板一律"精确行距 + 段前距"定位）
  · 宽度 → run 的 w:spacing 字距（栏目名要收着/撑开排到实测宽度）
  · 预印/打印 → 字色是不是 FFFFFF（白字占位不显影，就是套打的全部秘密）
"""
import copy
import io
import os

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from scripts.overprint import (PT_PER_CM, _iter_paragraphs, _run_track_cm,
                               _seg_width_cm, iter_seg_positions, layout_doc)

# 字号 → 公文号数，属性面板照这个显示，用户按号数说话、不按磅数
PT_LABELS = [(42.0, '初号'), (36.0, '小初'), (26.0, '一号'), (24.0, '小一'),
             (22.0, '二号'), (18.0, '小二'), (16.0, '三号'), (15.0, '小三'),
             (14.0, '四号'), (12.0, '小四'), (10.5, '五号'), (9.0, '小五')]


def pt_label(pt):
    for v, name in PT_LABELS:
        if abs(v - float(pt or 0)) < 0.01:
            return '%s %g' % (name, v)
    return '%g' % (pt or 0)


class TemplateEditError(Exception):
    pass


class EditSession(object):
    """一次编辑会话：打开模板 → 若干次改动 → 另存。

    改动都落在内存里的 Document 上，每步之前留一份快照，撤销就是把快照
    读回来。不直接改盘上的文件——用户没点"保存"之前，原模板一个字节都
    不动。
    """

    MAX_UNDO = 40

    def __init__(self, path):
        self.path = path
        self.doc = Document(path)
        self._undo = []
        self._redo = []
        self.dirty = False

    # ---------------- 快照 / 撤销 ----------------
    def _snapshot(self):
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

    def _restore(self, blob):
        self.doc = Document(io.BytesIO(blob))

    def _checkpoint(self):
        self._undo.append(self._snapshot())
        del self._undo[:-self.MAX_UNDO]
        self._redo = []
        self.dirty = True

    def can_undo(self):
        return bool(self._undo)

    def can_redo(self):
        return bool(self._redo)

    def undo(self):
        if not self._undo:
            return False
        self._redo.append(self._snapshot())
        self._restore(self._undo.pop())
        self.dirty = True
        return True

    def redo(self):
        if not self._redo:
            return False
        self._undo.append(self._snapshot())
        self._restore(self._redo.pop())
        self.dirty = True
        return True

    # ---------------- 读 ----------------
    def outline(self):
        """当前版面（含 ref 反查地址），画布直接照它画"""
        return layout_doc(self.doc, with_refs=True)

    def _paras(self):
        return [p for p, _c in _iter_paragraphs(self.doc)]

    def para(self, pi):
        ps = self._paras()
        if not 0 <= pi < len(ps):
            raise TemplateEditError('段落编号超出范围：%s' % pi)
        return ps[pi]

    def run(self, ref):
        pi, ri = ref
        p = self.para(pi)
        rs = p.runs
        if not 0 <= ri < len(rs):
            raise TemplateEditError('文字块已不存在，请重新选择')
        return p, rs[ri]

    def info(self, ref):
        """属性面板要显示的东西"""
        p, r = self.run(ref)
        rPr = r._r.find(qn('w:rPr'))
        rf = rPr.find(qn('w:rFonts')) if rPr is not None else None
        col = r.font.color.rgb if r.font.color and r.font.color.rgb else None
        pt = r.font.size.pt if r.font.size else 14.0
        track = _run_track_cm(r)
        return {
            'text': r.text,
            'font_cn': (rf.get(qn('w:eastAsia')) if rf is not None else '') or '',
            'font_en': (rf.get(qn('w:ascii')) if rf is not None else '') or '',
            'pt': pt,
            'bold': bool(r.font.bold),
            'white': str(col) == 'FFFFFF',
            'track_cm': track,
            'width_cm': _seg_width_cm({'text': r.text, 'pt': pt,
                                       'track': track}, pt / PT_PER_CM),
            'is_field': '{{' in r.text,
        }

    # ---------------- 写：内容与字体 ----------------
    def set_text(self, ref, text):
        self._checkpoint()
        _p, r = self.run(ref)
        r.text = text

    def set_style(self, ref, font_cn=None, font_en=None, pt=None, bold=None,
                  white=None):
        self._checkpoint()
        _p, r = self.run(ref)
        if pt is not None:
            r.font.size = Pt(float(pt))
        if bold is not None:
            r.font.bold = bool(bold)
        if white is not None:
            # 白字＝纸上已预印、不打印；黑字＝要打印上去的。套打的分界
            # 就这一条，别的什么都不用改
            r.font.color.rgb = (RGBColor(0xFF, 0xFF, 0xFF) if white
                                else RGBColor(0x00, 0x00, 0x00))
        if font_cn is not None or font_en is not None:
            rPr = r._r.get_or_add_rPr()
            rf = rPr.find(qn('w:rFonts'))
            if rf is None:
                rf = OxmlElement('w:rFonts')
                rPr.insert(0, rf)
            if font_cn:
                rf.set(qn('w:eastAsia'), font_cn)
            if font_en:
                for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rf.set(qn(a), font_en)

    def set_width(self, ref, width_cm):
        """把这段字**整体**排成指定的宽度（靠字距收放）。

        栏目名在预印纸上是收着排的：「领导批示：」五个字只占 2.10cm，
        按四号足宽排要 2.47cm。宽度不对，紧跟其后的黑字就整体偏。
        用户拿尺子量出这一栏有多宽，填进来，字距由它反解。
        """
        self._checkpoint()
        _p, r = self.run(ref)
        n = len(r.text)
        if not n:
            return
        pt = r.font.size.pt if r.font.size else 14.0
        # w:spacing 是"每个字后面都加"，所以 总宽 = n × (字宽 + 字距)
        track = float(width_cm) / n - pt / PT_PER_CM
        _set_track(r, track)

    def set_track(self, ref, track_cm):
        self._checkpoint()
        _p, r = self.run(ref)
        _set_track(r, track_cm)

    # ---------------- 写：位置 ----------------
    def nudge_x(self, ref, delta_cm):
        """横向挪：动的是这一段字所依附的那个制表位。

        制表位是绝对位置、以缇为单位，与字体无关——这也是模板一直用它
        定位的原因。同一个制表位后面若还跟着别的字（栏目名 + 填写位），
        它们会一起动，这正是纸上的实际情形。
        """
        if abs(delta_cm) < 1e-6:
            return None
        self._checkpoint()
        p, _r = self.run(ref)
        ordinal = self._tab_ordinal(p, ref[1])
        stops = _read_stops(p)
        if ordinal is None or ordinal >= len(stops):
            # 这段字前面没有制表符：补一个，位置＝它现在的落点 + 位移
            x = self._run_x_cm(ref)
            if x is None:
                raise TemplateEditError('这段文字没有可调的定位点')
            _insert_tab_before(p, ref[1])
            stops.append((max(0.0, x - _para_origin_cm(self.doc, p) + delta_cm),
                          'left'))
            _write_stops(p, stops)
            return None
        pos, kind = stops[ordinal]
        stops[ordinal] = (max(0.0, round(pos + delta_cm, 3)), kind)
        _write_stops(p, stops)
        return stops[ordinal][0]

    def nudge_y(self, pi, delta_cm):
        """纵向挪：加到段前距上。

        Word 里没有"这段字放在第几厘米"这种设置，只有"离上一段多远"。
        段前距不能为负，往上挪到头了就顶住——顶住时如实返回实际挪了多少，
        免得画布显示的位置和纸上的对不上。
        """
        self._checkpoint()
        p = self.para(pi)
        pf = p.paragraph_format
        cur = pf.space_before.cm if pf.space_before is not None else 0.0
        new = max(0.0, cur + delta_cm)
        pf.space_before = Cm(new)
        return new - cur

    def set_space_before(self, pi, cm):
        self._checkpoint()
        self.para(pi).paragraph_format.space_before = Cm(max(0.0, float(cm)))

    def set_line_cm(self, pi, cm):
        """行高（精确行距）。改字号时行盒得跟着放大，否则字会被裁掉。"""
        self._checkpoint()
        from docx.enum.text import WD_LINE_SPACING
        pf = self.para(pi).paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(float(cm) * PT_PER_CM)

    # ---------------- 写：增删 ----------------
    def add_run(self, pi, text, x_cm=None, style=None, white=True):
        """在段末添一段字；给了 x_cm 就连制表位一起加，钉在那个厘米数上。

        返回新 run 的 ref。
        """
        self._checkpoint()
        p = self.para(pi)
        if x_cm is not None:
            stops = _read_stops(p)
            stops.append((max(0.0, float(x_cm) - _para_origin_cm(self.doc, p)),
                          'left'))
            _write_stops(p, stops)
            p.add_run('\t')
        r = p.add_run(text)
        st = dict(style or {})
        pt = st.get('pt', 14.0)
        r.font.size = Pt(float(pt))
        r.font.bold = bool(st.get('bold', True))
        r.font.color.rgb = (RGBColor(0xFF, 0xFF, 0xFF) if white
                            else RGBColor(0x00, 0x00, 0x00))
        rPr = r._r.get_or_add_rPr()
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'), st.get('font_cn') or '方正楷体_GBK')
        en = st.get('font_en') or 'Times New Roman'
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rf.set(qn(a), en)
        rPr.insert(0, rf)
        return (pi, len(p.runs) - 1)

    def add_field(self, pi, name, x_cm=None, style=None):
        """加一个填写位 {{名称}}——黑字，会打印出来"""
        if not name or '{{' in name:
            raise TemplateEditError('字段名不能为空、也不能自带大括号')
        return self.add_run(pi, '{{%s}}' % name, x_cm=x_cm, style=style,
                            white=False)

    def delete(self, ref):
        """删掉一段字；它前面那个只为它服务的制表符一并删掉，
        否则后面的字会顺次认领错制表位、整行串位。"""
        self._checkpoint()
        p, r = self.run(ref)
        ri = ref[1]
        runs = p.runs
        prev = runs[ri - 1] if ri else None
        ordinal = self._tab_ordinal(p, ri)
        drop_tab = prev is not None and prev.text == '\t'
        if drop_tab and ordinal is not None:
            stops = _read_stops(p)
            if ordinal < len(stops):
                del stops[ordinal]
                _write_stops(p, stops)
            prev._r.getparent().remove(prev._r)
        r._r.getparent().remove(r._r)

    # ---------------- 存 ----------------
    def save(self, path=None):
        out = path or self.path
        d = os.path.dirname(out)
        if d and not os.path.isdir(d):
            os.makedirs(d)
        self.doc.save(out)
        self.dirty = False
        return out

    # ---------------- 内部 ----------------
    @staticmethod
    def _tab_ordinal(para, ri):
        """第 ri 个 run 之前有几个制表符（按行重新数）。

        制表符是**按顺序**认领制表位的（见 overprint._tab_target），
        所以"第 n 个制表符"对应"升序第 n 个制表位"。数错一个，改的就是
        别人的位置。换行会重新起算，与排版侧保持一致。
        """
        n = 0
        for i, r in enumerate(para.runs):
            if i >= ri:
                break
            for ch in r.text:
                if ch == '\n':
                    n = 0
                elif ch == '\t':
                    n += 1
        return n - 1 if n else None

    def _run_x_cm(self, ref):
        """这段字现在画在距纸左边多少厘米"""
        for seg, x in iter_seg_positions(self.outline()):
            if seg.get('ref') == tuple(ref):
                return x
        return None

    def positions(self):
        """{ref: (距纸左边cm, 距纸上边cm)}，画布和属性面板共用"""
        out = {}
        for seg, x, y in iter_seg_positions(self.outline(), with_y=True):
            if seg.get('ref') is not None:
                out[tuple(seg['ref'])] = (x, y)
        return out


# ---------------------------------------------------------------- 工具

def _set_track(run, cm):
    rPr = run._r.get_or_add_rPr()
    for el in rPr.findall(qn('w:spacing')):
        rPr.remove(el)
    if abs(cm) < 1e-5:
        return
    el = OxmlElement('w:spacing')
    el.set(qn('w:val'), str(int(round(cm * PT_PER_CM * 20))))
    rPr.append(el)


def _read_stops(para):
    out = []
    try:
        for t in para.paragraph_format.tab_stops:
            al = t.alignment
            out.append((t.position.cm,
                        'right' if (al is not None and int(al) == 2) else 'left'))
    except Exception:
        pass
    out.sort(key=lambda z: z[0])
    return out


def _write_stops(para, stops):
    ts = para.paragraph_format.tab_stops
    try:
        ts.clear_all()
    except Exception:
        pass
    for pos, kind in sorted(stops, key=lambda z: z[0]):
        ts.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.RIGHT if kind == 'right'
                        else WD_TAB_ALIGNMENT.LEFT)


def _insert_tab_before(para, ri):
    runs = para.runs
    target = runs[ri]
    tab = copy.deepcopy(target._r)
    for child in list(tab):
        if child.tag != qn('w:rPr'):
            tab.remove(child)
    t = OxmlElement('w:t')
    t.text = '\t'
    tab.append(t)
    target._r.addprevious(tab)


def _para_origin_cm(doc, para):
    """制表位是相对**容器左沿**算的：普通段落是左边距，
    表格单元格里则是单元格左沿。用错原点会整体偏出一大截。"""
    from scripts.overprint import _cell_left_cm
    sec = doc.sections[0]
    left = sec.left_margin.cm
    tc = para._p.getparent()
    if tc.tag != qn('w:tc'):
        return left
    from docx.table import Table, _Cell
    tbl = tc.getparent().getparent()
    table = Table(tbl, doc)
    cell = _Cell(tc, table)
    try:
        return _cell_left_cm(table, cell, left)
    except Exception:
        return left
