# -*- coding: utf-8 -*-
"""按文种生成公文骨架：选文种、填三个空，得到一份结构和用语都合规的初稿。

和用语检查共用同一张文种表（scripts.wording.DOC_KINDS）——那边查结语对不对，
这边直接把对的结语写进去。所以生成出来的骨架，拿用语检查去查必然零问题，
这一条在冒烟测试里是硬约束。
"""
import os

from PyQt5.QtWidgets import (QComboBox, QDialog, QDialogButtonBox, QFileDialog,
                             QFormLayout, QLabel, QLineEdit, QMessageBox,
                             QVBoxLayout)

from scripts.wording import DOC_KINDS, build_skeleton


class SkeletonDialog(QDialog):
    """选文种 → 填空 → 生成 docx"""

    def __init__(self, parent=None, preset_name='official_gbk'):
        super(SkeletonDialog, self).__init__(parent)
        self.setWindowTitle("按文种生成骨架")
        self.resize(520, 300)
        self._preset = preset_name
        self.result_path = None

        root = QVBoxLayout(self)
        tip = QLabel(
            "选一个文种，填上机关和事由，直接得到一份**结构和用语都合规**的初稿："
            "标题按「关于××的×××」拼好，主送机关带冒号，结语按文种配对"
            "（请示配「妥否，请批示」、报告配「特此报告」），"
            "落款和成文日期位置也铺好了。生成后按当前预设排版一遍即可。")
        tip.setWordWrap(True)
        tip.setProperty("muted", "true")
        root.addWidget(tip)

        form = QFormLayout()
        self.kind = QComboBox()
        for k in DOC_KINDS:
            self.kind.addItem(k, k)
        self.kind.currentIndexChanged.connect(self._on_kind)
        form.addRow("文种：", self.kind)
        self.subject = QLineEdit()
        self.subject.setPlaceholderText("如：开展某某专项检查工作")
        form.addRow("事由：", self.subject)
        self.recipient = QLineEdit()
        self.recipient.setPlaceholderText("如：省某某厅（可留空）")
        form.addRow("主送机关：", self.recipient)
        self.issuer = QLineEdit()
        self.issuer.setPlaceholderText("如：某某市某某局（可留空）")
        form.addRow("发文机关：", self.issuer)
        self.docnum = QLineEdit()
        self.docnum.setPlaceholderText("如：某某发〔2026〕5号（可留空）")
        form.addRow("发文字号：", self.docnum)
        root.addLayout(form)

        self.note = QLabel("")
        self.note.setWordWrap(True)
        self.note.setProperty("muted", "true")
        root.addWidget(self.note)
        self._on_kind()

        bb = QDialogButtonBox()
        bb.addButton("取消", QDialogButtonBox.RejectRole)
        bb.addButton("生成…", QDialogButtonBox.AcceptRole)
        bb.accepted.connect(self._make)
        bb.rejected.connect(self.reject)
        root.addWidget(bb)

    def _on_kind(self, *_a):
        spec = DOC_KINDS.get(self.kind.currentData()) or {}
        self.note.setText(spec.get('note') or '')

    def paragraphs(self):
        return build_skeleton(self.kind.currentData(),
                              issuer=self.issuer.text().strip(),
                              recipient=self.recipient.text().strip(),
                              subject=self.subject.text().strip(),
                              docnum=self.docnum.text().strip())

    def _make(self):
        try:
            paras = self.paragraphs()
        except ValueError as exc:
            QMessageBox.warning(self, "生成失败", str(exc))
            return
        default = '{}骨架.docx'.format(self.kind.currentData())
        out, _ = QFileDialog.getSaveFileName(self, "保存骨架", default,
                                             "Word 文档 (*.docx)")
        if not out:
            return
        try:
            from docx import Document
            doc = Document()
            for _ptype, text in paras:
                doc.add_paragraph(text)
            doc.save(out)
        except Exception as exc:
            QMessageBox.warning(self, "保存失败", str(exc))
            return
        self.result_path = out
        QMessageBox.information(
            self, "已生成",
            "骨架已保存：\n{}\n\n里面是纯文字，还没排版——"
            "把它拖进「格式处理」跑一遍智能一键，就是一份成品。"
            .format(os.path.basename(out)))
        self.accept()
