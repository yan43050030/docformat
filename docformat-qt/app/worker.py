# -*- coding: utf-8 -*-
"""批处理线程：四种模式分发 + .doc/.wps 预转换 + 进度/日志信号"""
import os
import re
import shutil
import sys
import tempfile

from PyQt5.QtCore import QThread, pyqtSignal

MODE_FULL = 'full'
MODE_DIAGNOSE = 'diagnose'
MODE_PUNCTUATION = 'punctuation'
MODE_AI_PASTE = 'ai_paste'
MODE_TOC_AUTO = 'toc_auto'
MODE_TOC_MANUAL = 'toc_manual'

# Windows/Linux 文件名中的非法字符
_INVALID_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]')


def friendly_error(e):
    """把引擎/系统异常翻译成用户能看懂的提示，返回 (白话消息, 是否已含原文)"""
    s = str(e)
    if isinstance(e, PermissionError) or 'Errno 13' in s or 'Permission denied' in s:
        return '无法写入输出文件，它可能正被 Word/WPS 打开，请关闭后重试', True
    if isinstance(e, FileNotFoundError) or 'No such file' in s:
        return '找不到文件，可能已被移动、重命名或删除', True
    if 'Package not found' in s or 'not a zip file' in s or 'BadZipFile' in type(e).__name__:
        return ('文件不是有效的 Word 文档或已损坏（若是老版 .doc 文件被改名为 .docx，'
                '请改回 .doc 后再处理）'), True
    if 'No space left' in s or 'Errno 28' in s:
        return '磁盘空间不足，请清理后重试', True
    if 'LibreOffice' in s or 'soffice' in s:
        return s, True   # 转换链错误本身已是中文提示
    return '处理出错：{}'.format(s), True


def sanitize_suffix(suffix):
    """清理输出后缀中的非法文件名字符"""
    cleaned = _INVALID_FILENAME_CHARS.sub('', (suffix or '').strip())
    return cleaned or '_processed'


def output_path_for(input_path, suffix):
    """生成输出路径；目标已存在时自动追加 (2) (3)… 序号，避免静默覆盖"""
    d = os.path.dirname(input_path)
    stem = os.path.splitext(os.path.basename(input_path))[0]
    base = os.path.join(d, '{}{}.docx'.format(stem, suffix))
    if not os.path.exists(base):
        return base
    for n in range(2, 1000):
        cand = os.path.join(d, '{}{}({}).docx'.format(stem, suffix, n))
        if not os.path.exists(cand):
            return cand
    return base


TEXT_EXTS = ('.txt', '.md', '.markdown')


def read_text_file(path):
    """读取文本文件，自动尝试常见中文编码"""
    with open(path, 'rb') as f:
        data = f.read()
    for enc in ('utf-8-sig', 'utf-8', 'gb18030'):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='replace')


def ensure_docx(path, log, session=None):
    """非 docx 输入先转换为临时 docx，返回 (工作路径, 待清理的临时目录或 None)"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        return path, None
    if ext in TEXT_EXTS:
        log('info', '检测到文本/Markdown 文件，正在生成 docx...')
        tmp_dir = tempfile.mkdtemp(prefix='docformat_')
        tmp = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + '.docx')
        generate_docx_from_text(read_text_file(path), tmp)
        return tmp, tmp_dir
    log('info', '检测到 {} 格式，正在转换为 docx...'.format(ext))
    if sys.platform == 'win32':
        tmp_dir = tempfile.mkdtemp(prefix='docformat_')
        tmp = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + '.docx')
        if session is not None:
            session.convert(path, tmp)
        else:
            from scripts import converter
            converter.convert_to_docx(path, tmp)
        return tmp, tmp_dir
    else:
        from app import converter_linux
        produced = converter_linux.convert_to_docx(path)
        return produced, os.path.dirname(produced)


def _cleanup_dir(path):
    if path:
        shutil.rmtree(path, ignore_errors=True)


MD_FENCE = re.compile(r'^\s*```')
MD_HEADING = re.compile(r'^\s*#{1,6}\s+')
MD_BULLET = re.compile(r'^\s*[-*+]\s+')
MD_QUOTE = re.compile(r'^\s*>\s?')
MD_HR = re.compile(r'^\s*([-*_]\s*){3,}$')
MD_BOLD = re.compile(r'(\*\*|__)(.+?)\1')
MD_ITALIC = re.compile(r'(?<![*_])([*_])([^*_\n]+?)\1(?![*_])')
MD_LINK = re.compile(r'\[([^\]]+)\]\([^)]*\)')
MD_CODE = re.compile(r'`([^`]*)`')


def clean_markdown(text):
    """把 AI 生成的 Markdown 清洗为纯文本段落列表"""
    lines = []
    in_fence = False
    for raw in text.splitlines():
        if MD_FENCE.match(raw):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        line = raw.rstrip()
        if MD_HR.match(line):
            continue
        line = MD_HEADING.sub('', line)
        line = MD_QUOTE.sub('', line)
        line = MD_BULLET.sub('', line)
        line = MD_BOLD.sub(r'\2', line)
        line = MD_ITALIC.sub(r'\2', line)
        line = MD_LINK.sub(r'\1', line)
        line = MD_CODE.sub(r'\1', line)
        lines.append(line.strip())
    # 折叠连续空行
    paras = []
    prev_blank = True
    for line in lines:
        if line:
            paras.append(line)
            prev_blank = False
        else:
            if not prev_blank:
                paras.append('')
            prev_blank = True
    while paras and not paras[-1]:
        paras.pop()
    return paras


def generate_docx_from_text(text, out_path):
    """纯文本/Markdown → 临时 docx（后续交给排版引擎）"""
    from docx import Document
    doc = Document()
    for para in clean_markdown(text):
        doc.add_paragraph(para)
    doc.save(out_path)
    return out_path


def build_diagnose_report(filename, results):
    lines = ['◆ {}'.format(filename)]
    total = 0

    punct = results.get('punctuation') or []
    if punct:
        lines.append('  【标点符号】{} 处问题：'.format(len(punct)))
        shown = {}
        for it in punct:
            key = '{}「{}」'.format(it.get('type', '?'), it.get('char', ''))
            shown.setdefault(key, []).append(it.get('para', 0))
        for key, paras in shown.items():
            # 去重（表格问题的位置标签会重复出现）
            uniq = list(dict.fromkeys(paras))
            preview = ', '.join('第 {} 段'.format(p) if isinstance(p, int) else str(p)
                                for p in uniq[:8])
            more = ' 等' if len(uniq) > 8 else ''
            lines.append('    - {}：{}{}'.format(key, preview, more))
        total += len(punct)

    for it in results.get('numbering') or []:
        lines.append('  【序号格式】{}'.format(it.get('detail', it.get('type', ''))))
        total += 1

    for it in results.get('paragraph') or []:
        if 'paras' in it:
            paras = it['paras']
            preview = ', '.join(str(p) for p in paras[:10])
            more = ' 等 {} 段'.format(len(paras)) if len(paras) > 10 else ''
            lines.append('  【段落格式】{}：第 {} 段{}'.format(it.get('type', ''), preview, more))
        else:
            lines.append('  【段落格式】{}：{}'.format(it.get('type', ''), it.get('detail', '')))
        total += 1

    for it in results.get('font') or []:
        lines.append('  【字体】{}：{}'.format(it.get('type', ''), it.get('detail', '')))
        total += 1

    if total == 0:
        lines.append('  ✓ 未发现格式问题')
    else:
        lines[0] = '◆ {} — 共 {} 类问题'.format(filename, total)
    return '\n'.join(lines)


class ProcessWorker(QThread):
    progressChanged = pyqtSignal(int)                # 0-100
    logMessage = pyqtSignal(str, str)                # level, message
    fileStarted = pyqtSignal(str)                    # input
    fileFinished = pyqtSignal(str, str)              # input, output ('' = 无输出)
    fileFailed = pyqtSignal(str, str)                # input, error
    diagnoseReady = pyqtSignal(str)                  # 汇总报告
    allFinished = pyqtSignal(int, int)               # ok, fail

    def __init__(self, files, mode, preset_name, custom_settings, suffix,
                 revision_mode=False, type_overrides=None, title_shape=None, parent=None):
        super(ProcessWorker, self).__init__(parent)
        self.files = list(files)
        self.mode = mode
        self.preset_name = preset_name
        self.custom_settings = custom_settings
        self.suffix = sanitize_suffix(suffix)
        self.revision_mode = revision_mode
        # {文件路径: {非空段序号: 段落类型}}，来自预览中的手动调整
        self.type_overrides = {os.path.normpath(k): v
                               for k, v in (type_overrides or {}).items()}
        self.title_shape = title_shape
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def _log(self, level, msg):
        self.logMessage.emit(level, msg)

    def _warn_missing_fonts(self):
        """Linux 上处理前检查预设字体，缺失则提前警告（Word 会静默替换成宋体）"""
        if self.mode not in (MODE_FULL,):
            return
        try:
            from scripts.formatter import PRESETS, get_missing_cn_fonts
            preset = self.custom_settings if self.custom_settings else PRESETS.get(self.preset_name)
            if not preset:
                return
            missing = get_missing_cn_fonts(preset)
            if missing:
                self._log('warning',
                          '本机未安装以下中文字体：{}。输出文档在打开时会被替换为其他字体，'
                          '建议安装后重新处理。'.format('、'.join(missing)))
        except Exception:
            pass

    def run(self):
        session = None
        if sys.platform == 'win32':
            needs_com = any(not f.lower().endswith('.docx') for f in self.files)
            if needs_com:
                try:
                    from scripts.converter import OfficeSession
                    session = OfficeSession().__enter__()
                    self._log('info', '已启动 {}（本批次转换复用同一实例）'.format(session.app_name))
                except Exception as e:
                    self._log('warning', '无法启动 Office/WPS：{}，.doc/.wps 文件将处理失败'.format(e))
                    session = None

        self._warn_missing_fonts()

        ok, fail = 0, 0
        reports = []
        n = len(self.files)
        try:
            for i, path in enumerate(self.files):
                if self._cancelled:
                    self._log('warning', '已取消，剩余文件未处理')
                    break
                base = os.path.basename(path)
                self._log('info', '正在处理: {}'.format(base))
                self.fileStarted.emit(path)
                tmp_dir = None
                try:
                    work, tmp_dir = ensure_docx(path, self._log, session)
                    # 自动转换 Word 自动编号为纯文字
                    _an_dir = None
                    if os.path.splitext(work)[1].lower() == '.docx':
                        try:
                            from scripts import auto_num
                            if auto_num._has_auto_numbering(work):
                                _an_dir = tempfile.mkdtemp(prefix='docformat_an_')
                                tmp2 = os.path.join(_an_dir, 'num.docx')
                                ok, unconverted = auto_num.convert_auto_numbering(
                                    work, tmp2, log=self._log)
                                if ok and os.path.exists(tmp2):
                                    work = tmp2
                                    if unconverted:
                                        self._log('warning',
                                            '{}: {} 段自动编号未能识别'.format(base, len(unconverted)))
                        except Exception:
                            _cleanup_dir(_an_dir)
                            _an_dir = None
                    if self.mode == MODE_DIAGNOSE:
                        reports.append(self._diagnose(work, base))
                        self.fileFinished.emit(path, '')
                    elif self.mode == MODE_PUNCTUATION:
                        out = output_path_for(path, self.suffix)
                        from scripts import punctuation

                        def _punct_cb(cur, total, _i=i, _n=n):
                            if total:
                                self.progressChanged.emit(int((_i + float(cur) / total) * 100 / _n))

                        punctuation.process_document(work, out, progress_callback=_punct_cb)
                        self._log('success', '已完成: {} → {}'.format(base, os.path.basename(out)))
                        self.fileFinished.emit(path, out)
                    elif self.mode in (MODE_TOC_AUTO, MODE_TOC_MANUAL):
                        out = output_path_for(path, self.suffix)
                        from scripts import toc
                        toc.generate_toc(work, out,
                                         mode='auto' if self.mode == MODE_TOC_AUTO else 'manual')
                        self._log('success', '已生成目录: {} → {}'.format(base, os.path.basename(out)))
                        self.fileFinished.emit(path, out)
                    else:  # MODE_FULL
                        out = output_path_for(path, self.suffix)
                        self._run_format(work, out, i, n,
                                         self.type_overrides.get(os.path.normpath(path)))
                        self._log('success', '已完成: {} → {}'.format(base, os.path.basename(out)))
                        self.fileFinished.emit(path, out)
                    ok += 1
                except Exception as e:
                    fail += 1
                    msg, _ = friendly_error(e)
                    self._log('error', '处理失败 {}: {}'.format(base, msg))
                    if msg != str(e) and str(e) not in msg:
                        self._log('info', '（技术细节：{}）'.format(e))
                    self.fileFailed.emit(path, msg)
                finally:
                    _cleanup_dir(tmp_dir)
                    _cleanup_dir(_an_dir)
                self.progressChanged.emit(int((i + 1) * 100 / n) if n else 100)

            if self.mode == MODE_DIAGNOSE and reports:
                self.diagnoseReady.emit('\n\n'.join(reports))
        finally:
            if session is not None:
                try:
                    session.__exit__(None, None, None)
                except Exception:
                    pass
        self.allFinished.emit(ok, fail)

    def _run_format(self, work, out, file_idx, total_files, type_overrides=None):
        """智能一键 = 标点修复(占单文件进度 0-10%) → 排版规范(10-100%)"""
        from scripts import punctuation
        from scripts.formatter import format_document

        def _overall(inner):
            self.progressChanged.emit(int((file_idx + inner) * 100.0 / total_files))

        tmp_root = tempfile.mkdtemp(prefix='docformat_')
        try:
            tmp = os.path.join(tmp_root, 'punct.docx')

            def punct_cb(cur, total):
                if total:
                    _overall(0.1 * cur / total)

            punctuation.process_document(work, tmp, progress_callback=punct_cb)

            def fmt_cb(cur, total, _stage):
                if total:
                    _overall(0.1 + 0.9 * float(cur) / total)

            format_document(
                tmp, out,
                preset_name=self.preset_name,
                progress_callback=fmt_cb,
                custom_settings=self.custom_settings,
                revision_mode=self.revision_mode,
                type_overrides=type_overrides,
                title_shape=self.title_shape,
            )
        finally:
            _cleanup_dir(tmp_root)

    def _diagnose(self, work, display_name):
        from docx import Document
        from scripts import analyzer
        from scripts.formatter import sanitize_document
        doc = Document(work)
        sanitize_document(doc)   # WPS 残缺 <w:jc> 兼容，诊断读 alignment 不崩

        # 自动编号提示：告诉用户本文档含 Word 自动编号，处理时会转成文字
        auto_num_note = ''
        try:
            from scripts import auto_num
            if auto_num._has_auto_numbering(work):
                from docx.oxml.ns import qn as _qn
                cnt = sum(1 for p in doc.paragraphs
                          if p._element.find(_qn('w:pPr')) is not None
                          and p._element.find(_qn('w:pPr')).find(_qn('w:numPr')) is not None)
                auto_num_note = ('本文档含 {} 段 Word 自动编号（如"一、""1."由软件自动生成），'
                                 '处理时会转换为纯文字，请核对转换后的编号是否正确。'.format(cnt))
        except Exception:
            pass
        results = {
            'punctuation': analyzer.analyze_punctuation(doc),
            'numbering': analyzer.analyze_numbering(doc),
            'paragraph': analyzer.analyze_paragraph_format(doc),
            'font': analyzer.analyze_font(doc),
        }
        report = build_diagnose_report(display_name, results)
        if auto_num_note:
            report += '\n  【自动编号】' + auto_num_note
        self._log('info', '诊断完成: {}'.format(display_name))
        return report


class AiPasteWorker(QThread):
    """AI 粘贴生成：文本 → 临时 docx → 排版引擎 → 输出"""
    logMessage = pyqtSignal(str, str)
    finishedWith = pyqtSignal(bool, str)   # success, out_path_or_error

    def __init__(self, text, out_path, preset_name, custom_settings, parent=None):
        super(AiPasteWorker, self).__init__(parent)
        self.text = text
        self.out_path = out_path
        self.preset_name = preset_name
        self.custom_settings = custom_settings

    def run(self):
        tmp_root = tempfile.mkdtemp(prefix='docformat_ai_')
        try:
            tmp = os.path.join(tmp_root, 'draft.docx')
            self.logMessage.emit('info', '正在从文本生成文档...')
            generate_docx_from_text(self.text, tmp)
            self.logMessage.emit('info', '正在修复标点并应用排版规范...')
            from scripts import punctuation
            from scripts.formatter import format_document
            tmp2 = os.path.join(tmp_root, 'punct.docx')
            punctuation.process_document(tmp, tmp2)
            format_document(tmp2, self.out_path,
                            preset_name=self.preset_name,
                            custom_settings=self.custom_settings)
            self.logMessage.emit('success', '已生成: {}'.format(self.out_path))
            self.finishedWith.emit(True, self.out_path)
        except Exception as e:
            msg, _ = friendly_error(e)
            self.logMessage.emit('error', '生成失败: {}'.format(msg))
            self.finishedWith.emit(False, msg)
        finally:
            shutil.rmtree(tmp_root, ignore_errors=True)
