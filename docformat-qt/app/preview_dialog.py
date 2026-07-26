# -*- coding: utf-8 -*-
"""排版前后对比预览：左侧原文，右侧按当前预设模拟排版效果，确认后才真正处理。

右侧每段前的类型标签可以点击，手动指定该段的段落类型（如把误识别的
正文改成一级标题）；手动调整会在实际排版时生效。
"""
import os
import re
import tempfile

from PyQt5.QtCore import Qt, QUrl
from PyQt5.QtGui import QCursor
from PyQt5.QtWidgets import (QCheckBox, QComboBox, QDialog, QFileDialog,
                             QHBoxLayout, QLabel, QMenu, QPushButton,
                             QSplitter, QTextBrowser, QVBoxLayout, QWidget)

from scripts.formatter import (_build_text_context, _compile_rules,
                               detect_para_type)
from scripts.punctuation import fix_text

TYPE_LABELS = {
    'security': '密级', 'docnum': '发文字号', 'title': '标题',
    'recipient': '主送机关',
    'heading1': '一级标题', 'heading2': '二级标题', 'heading3': '三级标题',
    'heading4': '四级标题', 'body': '正文', 'signature': '署名',
    'date': '日期', 'attachment': '附件说明', 'attachment_label': '附件标识', 'closing': '结尾', 'empty': '',
    'roster': '组成人员名单',
}
# 手动指定类型菜单里展示的顺序
TYPE_MENU_ORDER = ['title', 'docnum', 'security', 'recipient',
                   'heading1', 'heading2', 'heading3', 'heading4',
                   'body', 'roster', 'attachment_label', 'attachment', 'closing', 'signature', 'date']
MAX_PARAS = 500

ALIGN_CSS = {'left': 'left', 'center': 'center', 'right': 'right', 'justify': 'justify'}

# 预览字体回退链：本机未安装方正字体时退到同族常见字体，
# 保证标题黑体/楷体/仿宋/加粗等样式差异在预览里肉眼可见
_CSS_FONT_FALLBACK = {
    '方正小标宋_GBK': '"方正小标宋_GBK","方正小标宋简体","华文中宋","SimSun","宋体"',
    '方正小标宋简体': '"方正小标宋简体","华文中宋","SimSun","宋体"',
    '方正黑体_GBK': '"方正黑体_GBK","黑体","SimHei","Microsoft YaHei"',
    '黑体': '"黑体","SimHei","Microsoft YaHei"',
    '方正楷体_GBK': '"方正楷体_GBK","楷体","楷体_GB2312","KaiTi"',
    '楷体_GB2312': '"楷体_GB2312","楷体","KaiTi"',
    '楷体': '"楷体","KaiTi"',
    '方正仿宋_GBK': '"方正仿宋_GBK","仿宋_GB2312","仿宋","FangSong"',
    '仿宋_GB2312': '"仿宋_GB2312","仿宋","FangSong"',
    '仿宋': '"仿宋","FangSong"',
    '宋体': '"宋体","SimSun"',
    '华文中宋': '"华文中宋","宋体","SimSun"',
}


def _css_font(name):
    raw = _CSS_FONT_FALLBACK.get(name, '"{}"'.format(name))
    # 用单引号包裹字体名，避免与 HTML style="..." 的双引号冲突
    return raw.replace('"', "'")


def _css_font_for_en(name):
    """英文字体直接输出（Times New Roman 等不需要回退链）"""
    return "'{}'".format(name) if name else "'Times New Roman'"


# 连续的 ASCII 字母/数字/常见西文符号（用英文字体渲染，与真实 docx 一致：
# 引擎把 ASCII 设为 font_en，东亚字设为 font_cn）
_EN_RUN_RE = re.compile(r"[0-9A-Za-z][0-9A-Za-z\.\-/:%()+，]*[0-9A-Za-z]|[0-9A-Za-z]")


def _segment_font_html(text, font_en, font_cn):
    """把文本按"西文数字段 vs 其它"切分，各自套正确字体。

    Qt 富文本不做浏览器式逐字体回退，中英文混排时数字未必落到英文字体，
    因此手动分段：数字/英文 → font_en(Times New Roman)，其余 → font_cn。
    """
    en_css = _css_font_for_en(font_en)
    cn_css = _css_font(font_cn)
    out = []
    pos = 0
    for m in _EN_RUN_RE.finditer(text):
        if m.start() > pos:
            seg = text[pos:m.start()]
            out.append("<span style=\"font-family:{}\">{}</span>".format(cn_css, _esc(seg)))
        out.append("<span style=\"font-family:{}\">{}</span>".format(en_css, _esc(m.group())))
        pos = m.end()
    if pos < len(text):
        out.append("<span style=\"font-family:{}\">{}</span>".format(cn_css, _esc(text[pos:])))
    return ''.join(out)


def _read_paragraphs(path):
    """返回 (段落列表[(text, alignment)], 表格数, 总段数)。

    支持 .docx 与 .txt/.md（按 AI 粘贴规则清洗后预览）；
    .doc/.wps 返回 None，由对话框先做格式转换再调用。
    """
    lower = path.lower()
    if lower.endswith(('.txt', '.md', '.markdown')):
        from app.worker import clean_markdown, read_text_file
        lines = clean_markdown(read_text_file(path))
        paras = [(t, None, '', 0) for t in lines[:MAX_PARAS]]
        return paras, 0, len(lines), False
    if not lower.endswith('.docx'):
        return None  # .doc/.wps: no auto-num info
    # 预览前自动转换 Word 自动编号为文字（临时文件用完即删）
    _an_tmp = None
    try:
        from scripts import auto_num
        if auto_num._has_auto_numbering(path):
            import tempfile as _tf
            _an_fd, _an_tmp = _tf.mkstemp(suffix='.docx', prefix='docformat_pv_an_')
            os.close(_an_fd)
            ok, _ = auto_num.convert_auto_numbering(path, _an_tmp)
            if ok and os.path.exists(_an_tmp):
                path = _an_tmp
    except Exception:
        pass
    try:
        from docx import Document
        doc = Document(path)
        # 修复 WPS/老 Word 残缺 <w:jc>，否则读 alignment 会崩溃
        try:
            from scripts.formatter import sanitize_document
            sanitize_document(doc)
        except Exception:
            pass
        from docx.oxml.ns import qn as _qn3
        total = len(doc.paragraphs)
        paras = []
        has_auto_num = False
        for p in doc.paragraphs[:MAX_PARAS]:
            align = p.paragraph_format.alignment
            font_cn = ''
            font_size = 0
            for r in p.runs:
                if r.text.strip():
                    font_cn = r.font.name or ''
                    font_size = r.font.size.pt if r.font.size else 0
                    break
            paras.append((p.text, align, font_cn, font_size))
            if not has_auto_num:
                pPr = p._element.find(_qn3('w:pPr'))
                if pPr is not None and pPr.find(_qn3('w:numPr')) is not None:
                    has_auto_num = True
        return paras, len(doc.tables), total, has_auto_num
    finally:
        if _an_tmp:
            try: os.unlink(_an_tmp)
            except Exception: pass


def _html_shell(body, base_size=12):
    return ('<html><head><style>'
            'body {{ font-family: "SimSun", serif; font-size: {}pt; margin: 14px; }}'
            'p {{ margin: 0 0 4px 0; white-space: pre-wrap; }}'
            'a.tag {{ font-size: 8pt; color: #666; background: #F0EDE6; '
            'border: 1px solid #D8D2C4; border-radius: 3px; padding: 0 4px; '
            'margin-right: 4px; text-decoration: none; }}'
            'a.tagx {{ font-size: 8pt; color: #FFFFFF; background: #C0392B; '
            'border: 1px solid #A93226; border-radius: 3px; padding: 0 4px; '
            'margin-right: 4px; text-decoration: none; }}'
            '.tag0 {{ font-size: 8pt; color: #999; background: #F0EDE6; '
            'border-radius: 3px; padding: 0 4px; margin-right: 4px; }}'
            '.broom {{ font-size: 8pt; margin-right: 4px; }}'
            '</style></head><body>{}</body></html>').format(base_size, body)


def render_before_html(paras):
    parts = []
    for item in paras:
        if len(item) == 4:
            text, align, font_cn, font_size = item
        else:
            text, align = item[0], item[1]
            font_cn, font_size = '', 0
        if not text.strip():
            parts.append('<p>&nbsp;</p>')
            continue
        a = 'center' if align is not None and 'CENTER' in str(align) else (
            'right' if align is not None and 'RIGHT' in str(align) else 'left')
        sty = 'text-align:{}'.format(a)
        if font_cn:
            sty += ";font-family:'{}'".format(font_cn)
        if font_size:
            sty += ';font-size:{}pt'.format(font_size)
        parts.append('<p style="{}">{}</p>'.format(sty, _esc(text)))
    return _html_shell(''.join(parts))


def compute_types(paras, preset, overrides=None):
    """返回 [(非空段序号 or None, 段落类型 or None)]，与 paras 一一对应。

    overrides: {非空段序号: 类型}，手动指定优先，且会作为上下文影响后续段落识别。
    """
    overrides = overrides or {}
    rules = _compile_rules(preset.get('detect_rules'))
    texts = [item[0] for item in paras]
    all_texts = [t.strip() for t in texts if t.strip()]
    idx_map = {}
    n = 0
    for i, t in enumerate(texts):
        if t.strip():
            idx_map[i] = n
            n += 1

    result = []
    prev_type = None
    total = len(paras)
    for i, item in enumerate(paras):
        text = item[0]
        align = item[1] if len(item) > 1 else None
        if not text.strip():
            result.append((None, None))
            continue
        ai = idx_map[i]
        if ai in overrides:
            ptype = overrides[ai]
        else:
            ptype = detect_para_type(
                text.strip(), i, total, align, all_texts,
                all_texts_index=ai, prev_para_type=prev_type, rules=rules)
        prev_type = ptype
        result.append((ai, ptype))
    return result


def _preview_title_cpl(preset):
    """预览用：估算标题每行可容纳的全角字数（版心宽 / 标题字号）"""
    page = preset.get('page', {})
    pw_cm = 21.0
    left = page.get('left', 2.8)
    right = page.get('right', 2.6)
    usable_mm = (pw_cm - left - right) * 10
    usable_pt = usable_mm / 0.3528
    tsize = preset.get('title', {}).get('size', 22) or 22
    return max(6, int(usable_pt / tsize))


def render_after_html(paras, preset, overrides=None, title_shape=None, clean_marks=None):
    overrides = overrides or {}
    clean_marks = clean_marks or set()
    types = compute_types(paras, preset, overrides)
    # 标题梯形：None=按模板；否则覆盖
    shape = title_shape if title_shape is not None else preset.get('title_shape', 'none')
    cpl = _preview_title_cpl(preset)

    parts = []
    for item, (ai, ptype) in zip(paras, types):
        text = item[0]
        if ptype is None:
            parts.append('<p>&nbsp;</p>')
            continue
        # 右侧展示标点修复后的效果，与实际输出一致
        display = fix_text(text.strip())
        fmt = preset.get(ptype if ptype in preset else 'body', preset.get('body', {}))

        style = [
            'font-size:{}pt'.format(fmt.get('size', 16)),
            'text-align:{}'.format(ALIGN_CSS.get(fmt.get('align', 'left'), 'left')),
        ]
        indent = fmt.get('indent', 0) or 0
        if indent:
            style.append('text-indent:{}pt'.format(indent))
        ls = fmt.get('line_spacing')
        if ls:
            style.append('line-height:{}pt'.format(ls))
        if fmt.get('bold'):
            style.append('font-weight:bold')

        # 标题梯形回行：预览里按所选形状折行显示
        if ptype == 'title' and shape in ('trapezoid_down', 'trapezoid_up'):
            from scripts.title_shape import split_title_lines
            lines = split_title_lines(display, cpl, shape)
            inner = '<br>'.join(_segment_font_html(
                ln, fmt.get('font_en', 'Times New Roman'),
                fmt.get('font_cn', '仿宋_GB2312')) for ln in lines)
        else:
            # 数字/英文用 font_en，中文用 font_cn，与真实 docx 输出一致
            inner = _segment_font_html(
                display, fmt.get('font_en', 'Times New Roman'), fmt.get('font_cn', '仿宋_GB2312'))
        tag = TYPE_LABELS.get(ptype, ptype)
        cls = 'tagx' if ai in overrides else 'tag'
        broom = '<span class="broom" title="已标记为待清洗">🧹</span>' if ai in clean_marks else ''
        parts.append(
            '<p style="{}"><a class="{}" href="para:{}" title="点击修改此段类型">{}</a>{}{}</p>'.format(
                '; '.join(style), cls, ai, tag, broom, inner))
    return _html_shell(''.join(parts))


def render_overlay_html(paras, preset, header_png_path, overrides=None,
                        title_shape=None, clean_marks=None, page_cfg=None):
    """生成叠加了套头 PDF 背景的 HTML 预览。

    将套头 PNG 作为 A4 尺寸容器的背景图，文字按 preset 的页面边距排布在内。
    使用 px 单位（与 PNG 的 150 DPI 渲染一致），QTextBrowser 可直接显示。
    """
    from scripts.header_overlay import _HEADER_OVERLAY_DPI, page_pixmap_size

    overrides = overrides or {}
    clean_marks = clean_marks or set()
    types = compute_types(paras, preset, overrides)
    shape = title_shape if title_shape is not None else preset.get('title_shape', 'none')
    cpl = _preview_title_cpl(preset)

    page = page_cfg or {}
    pw_px, ph_px = page_pixmap_size(header_png_path) if header_png_path else (1240, 1754)

    px_per_pt = _HEADER_OVERLAY_DPI / 72.0
    px_per_cm = _HEADER_OVERLAY_DPI / 2.54

    left_cm = page.get('left', 2.8)
    right_cm = page.get('right', 2.8)
    top_cm = page.get('top', 3.8)

    left_px = int(left_cm * px_per_cm)
    right_px = int(right_cm * px_per_cm)
    top_px = int(top_cm * px_per_cm)

    # 背景图 URL（QTextBrowser 用 file:/// 前缀）
    bg_url = 'file:///{}'.format(header_png_path.replace('\\', '/'))

    rows = []
    # 天头留白
    if top_px > 0:
        rows.append('<tr><td colspan="3" height="{}"></td></tr>'.format(top_px))

    for item, (ai, ptype) in zip(paras, types):
        text = item[0]
        if ptype is None:
            rows.append('<tr><td colspan="3"><p style="margin:0;">&nbsp;</p></td></tr>')
            continue

        display = fix_text(text.strip())
        fmt = preset.get(ptype if ptype in preset else 'body', preset.get('body', {}))

        font_size_px = int(fmt.get('size', 16) * px_per_pt)
        indent_pt = fmt.get('indent', 0) or 0
        indent_px = int(indent_pt * px_per_pt)
        ls = fmt.get('line_spacing')
        lh_px = 'line-height:{}px;'.format(int(ls * px_per_pt)) if ls else ''
        align = ALIGN_CSS.get(fmt.get('align', 'left'), 'left')
        bold = 'font-weight:bold;' if fmt.get('bold') else ''
        font_cn = _css_font(fmt.get('font_cn', '仿宋_GB2312'))

        if ptype == 'title' and shape in ('trapezoid_down', 'trapezoid_up'):
            from scripts.title_shape import split_title_lines
            lines = split_title_lines(display, cpl, shape)
            inner = '<br>'.join(_segment_font_html(
                ln, fmt.get('font_en', 'Times New Roman'),
                fmt.get('font_cn', '仿宋_GB2312')) for ln in lines)
        else:
            inner = _segment_font_html(
                display, fmt.get('font_en', 'Times New Roman'),
                fmt.get('font_cn', '仿宋_GB2312'))

        tag = TYPE_LABELS.get(ptype, ptype)
        cls = 'tagx' if ai in overrides else 'tag'
        broom = '<span class="broom" title="已标记为待清洗">🧹</span>' if ai in clean_marks else ''

        rows.append(
            '<tr>'
            '<td width="{}" style="width:{}px;"></td>'
            '<td style="vertical-align:top;">'
            '<p style="margin:0 0 4px 0; white-space:pre-wrap; '
            'font-size:{}px; text-align:{}; text-indent:{}px; '
            '{}{}font-family:{}; color:#111;">'
            '<a class="{}" href="para:{}" title="点击修改此段类型">{}</a>{}{}</p>'
            '</td>'
            '<td width="{}" style="width:{}px;"></td>'
            '</tr>'.format(
                left_px, left_px,
                font_size_px, align, indent_px,
                bold, lh_px, font_cn,
                cls, ai, tag, broom, inner,
                right_px, right_px))

    head = ('<html><head><style>'
            'body {{ font-family: "SimSun", serif; margin: 14px; '
            'background: #C9C4B8; }}'
            'p {{ margin: 0 0 2px 0; white-space: pre-wrap; }}'
            'a.tag {{ font-size: 8pt; color: #666; background: #F0EDE6; '
            'border: 1px solid #D8D2C4; border-radius: 3px; padding: 0 4px; '
            'margin-right: 4px; text-decoration: none; }}'
            'a.tagx {{ font-size: 8pt; color: #FFFFFF; background: #C0392B; '
            'border: 1px solid #A93226; border-radius: 3px; padding: 0 4px; '
            'margin-right: 4px; text-decoration: none; }}'
            '.broom {{ font-size: 8pt; margin-right: 4px; }}'
            '</style></head><body>')

    bg_div = (
        '<div style="width:{}px; '
        'background-image:url(\'{}\'); '
        'background-repeat:no-repeat; '
        'background-color:#FFFFFF;">').format(pw_px, bg_url)

    table = '<table width="{}" cellspacing="0" cellpadding="0" ' \
            'style="border-collapse:collapse;">'.format(pw_px)

    footer = '</table></div>' \
             '<div style="color:#8A8578; font-size:11px; margin-top:4px;">' \
             'A4 预览 — 套头 PDF 叠加（{} DPI 等比缩放）</div>'.format(_HEADER_OVERLAY_DPI)

    return head + bg_div + table + ''.join(rows) + footer + '</body></html>'


def _esc(text):
    return (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))


class PreviewDialog(QDialog):
    def __init__(self, files, preset, parent=None):
        super(PreviewDialog, self).__init__(parent)
        self.setWindowTitle("排版效果预览（尚未修改任何文件）")
        self.resize(1080, 720)
        from app.theme import settings as _settings
        _s = _settings()
        geo = _s.value('preview/geometry')
        if geo:
            self.restoreGeometry(geo)
        self.files = files
        self.preset = preset
        # path -> {非空段序号: 类型}
        self._overrides = {}
        # path -> set(非空段序号)，标记为"待清洗"的段落
        self._clean_marks = {}
        # 清洗项（None = 用 cleaner.DEFAULT_CLEAN）
        self._clean_items = None
        self._current_paras = None
        self._converted = {}    # .doc/.wps 预览转换缓存: 原路径 -> 临时 docx
        self._tmp_dirs = []     # 对话框关闭时清理
        # 套头 PDF 叠加
        self._header_pdf_path = None
        self._header_png_temp = None
        self._tmp_header_pngs = []

        root = QVBoxLayout(self)
        root.setContentsMargins(14, 12, 14, 12)
        root.setSpacing(8)

        top = QHBoxLayout()
        top.addWidget(QLabel("预览文件:"))
        self.file_combo = QComboBox()
        for f in files:
            self.file_combo.addItem(os.path.basename(f), f)
        self.file_combo.currentIndexChanged.connect(self._load_current)
        top.addWidget(self.file_combo, 1)
        tip = QLabel("右侧为预设「{}」的模拟效果，实际以 Word 渲染为准".format(preset.get('name', '')))
        tip.setProperty("muted", "true")
        top.addWidget(tip)
        root.addLayout(top)

        self.notice = QLabel("")
        self.notice.setProperty("muted", "true")
        self.notice.setWordWrap(True)
        self.notice.setVisible(False)
        root.addWidget(self.notice)

        self.seal_check = QCheckBox("加盖公章（落款日期右空4字，署名居中于日期编排）")
        root.addWidget(self.seal_check)

        # 套头 PDF 叠加：选一张扫描好的红头纸 PDF，预览时叠在排版文字后方
        overlay_row = QHBoxLayout()
        overlay_row.addWidget(QLabel("套头 PDF："))
        self.header_pdf_label = QLabel("（未选择）")
        self.header_pdf_label.setProperty("muted", "true")
        self.pick_header_btn = QPushButton("选择…")
        self.pick_header_btn.setCursor(Qt.PointingHandCursor)
        self.pick_header_btn.clicked.connect(self._pick_header_pdf)
        self.clear_header_btn = QPushButton("清除")
        self.clear_header_btn.setCursor(Qt.PointingHandCursor)
        self.clear_header_btn.clicked.connect(self._clear_header_pdf)
        self.clear_header_btn.setVisible(False)
        overlay_row.addWidget(self.header_pdf_label, 1)
        overlay_row.addWidget(self.pick_header_btn)
        overlay_row.addWidget(self.clear_header_btn)
        root.addLayout(overlay_row)

        # 标题梯形回行：预览里可选并即时预览、确认后应用（覆盖模板设置）
        ts_row = QHBoxLayout()
        ts_row.addWidget(QLabel("长标题回行："))
        self.title_shape_combo = QComboBox()
        for label, val in [('按模板设置', '__preset__'),
                           ('不处理', 'none'),
                           ('正梯形（上长下短）', 'trapezoid_down'),
                           ('倒梯形（上短下长）', 'trapezoid_up')]:
            self.title_shape_combo.addItem(label, val)
        self.title_shape_combo.setToolTip("在此选择即在预览中应用；确认排版时按此设置生效，"
                                          "不必修改模板")
        self.title_shape_combo.currentIndexChanged.connect(self._render_after)
        ts_row.addWidget(self.title_shape_combo)
        ts_row.addSpacing(18)

        # 格式清洗：排版疑难杂症多半源于原文档里看不见的脏格式
        ts_row.addWidget(QLabel("格式清洗："))
        self.clean_combo = QComboBox()
        for label, val in [('不清洗', 'none'),
                           ('全文清洗', 'all'),
                           ('仅清洗标记的段落', 'selected')]:
            self.clean_combo.addItem(label, val)
        self.clean_combo.setToolTip(
            "排版出怪问题（缩进被挤走、行距莫名撑高、字号变化）时启用。\n"
            "会清掉原文档里看不见的脏格式：字符缩放/间距、着重号、拼音指南、\n"
            "边框底纹、文本框式段落、域代码、书签批注、修订痕迹、\n"
            "手动换行符、制表符与全角空格等。")
        self.clean_combo.currentIndexChanged.connect(self._on_clean_scope_changed)
        ts_row.addWidget(self.clean_combo)
        self.clean_items_btn = QPushButton("清洗项…")
        self.clean_items_btn.setCursor(Qt.PointingHandCursor)
        self.clean_items_btn.clicked.connect(self._pick_clean_items)
        self.clean_items_btn.setEnabled(False)
        ts_row.addWidget(self.clean_items_btn)
        self.clean_count_label = QLabel("")
        self.clean_count_label.setProperty("muted", "true")
        ts_row.addWidget(self.clean_count_label)
        ts_row.addStretch(1)
        root.addLayout(ts_row)

        hint = QLabel("提示：点击右侧段落前的类型标签，可手动指定该段是标题/正文/附件等（红色标签=已手动指定）；"
                      "菜单里还可把该段标记为「待清洗」（🧹）")
        hint.setProperty("muted", "true")
        hint.setWordWrap(True)
        root.addWidget(hint)

        header = QHBoxLayout()
        lab_l = QLabel("排版前（原文）")
        lab_l.setProperty("sectionTitle", "true")
        lab_r = QLabel("排版后（模拟预览，含段落类型标注）")
        lab_r.setProperty("sectionTitle", "true")
        header.addWidget(lab_l, 1)
        header.addWidget(lab_r, 1)
        root.addLayout(header)

        split = QSplitter(Qt.Horizontal)
        self._splitter = split
        self.view_before = QTextBrowser()
        self.view_after = QTextBrowser()
        self.view_after.setOpenLinks(False)
        self.view_after.anchorClicked.connect(self._on_tag_clicked)
        # 左右分栏同步滚动（按比例，两侧内容高度不同也能对齐大致位置）
        self._sync_lock = False
        self.view_before.verticalScrollBar().valueChanged.connect(
            lambda _v: self._sync_scroll(self.view_before, self.view_after))
        self.view_after.verticalScrollBar().valueChanged.connect(
            lambda _v: self._sync_scroll(self.view_after, self.view_before))
        split.addWidget(self.view_before)
        split.addWidget(self.view_after)
        split.setSizes([500, 500])
        _sp = _s.value('preview/splitter')
        if _sp:
            split.restoreState(_sp)
        root.addWidget(split, 1)

        btns = QHBoxLayout()
        self.reset_btn = QPushButton("清除本文件的手动调整")
        self.reset_btn.clicked.connect(self._clear_overrides)
        self.reset_btn.setEnabled(False)
        btns.addWidget(self.reset_btn)
        btns.addStretch(1)
        cancel = QPushButton("关闭")
        cancel.clicked.connect(self.reject)
        ok = QPushButton("确认无误，开始排版")
        ok.setProperty("primary", "true")
        ok.setCursor(Qt.PointingHandCursor)
        ok.clicked.connect(self.accept)
        btns.addWidget(cancel)
        btns.addWidget(ok)
        root.addLayout(btns)

        self._load_current()

    # ---------- 手动类型调整 ----------
    def get_overrides(self):
        """返回 {文件路径: {非空段序号: 类型}}（仅含有调整的文件）"""
        return {p: dict(m) for p, m in self._overrides.items() if m}

    def get_title_shape(self):
        """返回预览里选择的标题梯形（None=按模板，不覆盖）"""
        val = self.title_shape_combo.currentData()
        return None if val == '__preset__' else val

    def get_header_pdf_path(self):
        """返回选中的套头 PDF 路径，或 None"""
        return self._header_pdf_path

    def _pick_header_pdf(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "选择套头 PDF（红头/分隔线在纸上的扫描件）",
            "", "PDF 文件 (*.pdf)")
        if path:
            self._set_header_pdf(path)

    def _set_header_pdf(self, path):
        self._header_pdf_path = path
        self.header_pdf_label.setText(os.path.basename(path))
        self.header_pdf_label.setProperty("muted", "false")
        self.clear_header_btn.setVisible(True)
        self._refresh_header_png()
        self._render_after()

    def _clear_header_pdf(self):
        self._header_pdf_path = None
        self.header_pdf_label.setText("（未选择）")
        self.header_pdf_label.setProperty("muted", "true")
        self.clear_header_btn.setVisible(False)
        self._cleanup_header_png()
        self._render_after()

    def _refresh_header_png(self):
        """渲染套头 PDF 第一页为临时 PNG"""
        self._cleanup_header_png()
        if not self._header_pdf_path:
            return
        try:
            from scripts.header_overlay import render_page_to_png
            self._header_png_temp = render_page_to_png(self._header_pdf_path, page_number=0)
            self._tmp_header_pngs.append(self._header_png_temp)
        except Exception as e:
            self.notice.setText("套头 PDF 预览失败：{}".format(str(e)[:80]))
            self.notice.setVisible(True)

    def _cleanup_header_png(self):
        if self._header_png_temp:
            try:
                os.unlink(self._header_png_temp)
            except Exception:
                pass
            self._header_png_temp = None

    def _current_path(self):
        return self.file_combo.currentData()

    def _on_tag_clicked(self, url):
        s = url.toString() if isinstance(url, QUrl) else str(url)
        if not s.startswith('para:'):
            return
        try:
            ai = int(s.split(':', 1)[1])
        except ValueError:
            return
        path = self._current_path()
        cur = self._overrides.get(path, {}).get(ai)

        menu = QMenu(self)
        menu.addSection("此段落类型")
        for t in TYPE_MENU_ORDER:
            act = menu.addAction(TYPE_LABELS[t] + (' ✓' if cur == t else ''))
            act.setData(t)
        menu.addSeparator()
        auto = menu.addAction("恢复自动识别")
        auto.setData('__auto__')
        menu.addSeparator()
        marked = ai in self._clean_marks.get(path, set())
        clean_act = menu.addAction(
            '取消「待清洗」标记' if marked else '标记此段为「待清洗」🧹')
        clean_act.setData('__clean__')

        chosen = menu.exec_(QCursor.pos())
        if chosen is None:
            return
        val = chosen.data()
        if val == '__clean__':
            marks = self._clean_marks.setdefault(path, set())
            if ai in marks:
                marks.discard(ai)
            else:
                marks.add(ai)
                # 标了段落却没选范围时，自动切到"仅清洗标记的段落"
                if self.clean_combo.currentData() == 'none':
                    i = self.clean_combo.findData('selected')
                    if i >= 0:
                        self.clean_combo.setCurrentIndex(i)
            self._refresh_clean_state()
            self._render_after()
            return
        m = self._overrides.setdefault(path, {})
        if val == '__auto__':
            m.pop(ai, None)
        else:
            m[ai] = val
        self._render_after()

    def _clear_overrides(self):
        path = self._current_path()
        if path in self._overrides:
            self._overrides[path] = {}
        if path in self._clean_marks:
            self._clean_marks[path] = set()
        self._refresh_clean_state()
        self._render_after()

    # ---------- 格式清洗 ----------
    def _on_clean_scope_changed(self, _idx=None):
        self._refresh_clean_state()
        self._render_after()

    def _refresh_clean_state(self):
        scope = self.clean_combo.currentData()
        self.clean_items_btn.setEnabled(scope != 'none')
        path = self._current_path()
        n = len(self._clean_marks.get(path, set()))
        if scope == 'selected':
            self.clean_count_label.setText(
                '已标记 {} 段'.format(n) if n else '尚未标记段落（点类型标签→标记）')
        elif scope == 'all':
            self.clean_count_label.setText('将清洗全文')
        else:
            self.clean_count_label.setText('')

    def _pick_clean_items(self):
        from app.clean_dialog import CleanItemsDialog
        dlg = CleanItemsDialog(self._clean_items, self)
        if dlg.exec_() == CleanItemsDialog.Accepted:
            self._clean_items = dlg.get_items()

    def get_clean_spec(self):
        """返回 {文件路径: clean_spec}，仅含需要清洗的文件；无则返回 {}"""
        scope = self.clean_combo.currentData()
        if scope == 'none':
            return {}
        out = {}
        for path in self.files:
            if scope == 'all':
                out[path] = {'scope': 'all', 'items': self._clean_items}
            else:
                marks = sorted(self._clean_marks.get(path, set()))
                if marks:
                    out[path] = {'scope': 'selected', 'items': self._clean_items,
                                 'paragraphs': marks}
        return out

    # ---------- 渲染 ----------
    def _render_after(self):
        if self._current_paras is None:
            return
        path = self._current_path()
        ovr = self._overrides.get(path, {})
        bar = self.view_after.verticalScrollBar()
        pos = bar.value()
        ts = self.get_title_shape() if hasattr(self, 'title_shape_combo') else None
        marks = self._clean_marks.get(path, set()) \
            if self.clean_combo.currentData() == 'selected' else set()
        if self._header_pdf_path and self._header_png_temp:
            html = render_overlay_html(
                self._current_paras, self.preset, self._header_png_temp,
                overrides=ovr, title_shape=ts, clean_marks=marks,
                page_cfg=self.preset.get('page', {}))
        else:
            html = render_after_html(self._current_paras, self.preset, ovr, ts, marks)
        self.view_after.setHtml(html)
        bar.setValue(pos)
        self.reset_btn.setEnabled(bool(ovr) or bool(self._clean_marks.get(path)))

    def _convert_for_preview(self, path):
        """.doc/.wps → 临时 docx（结果缓存，对话框关闭时清理）"""
        import os as _os
        import sys as _sys
        if path in self._converted:
            return self._converted[path]
        from PyQt5.QtWidgets import QApplication
        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            if _sys.platform == 'win32':
                from scripts import converter
                import tempfile as _tempfile
                tmp_dir = _tempfile.mkdtemp(prefix='docformat_pv_')
                tmp = _os.path.join(
                    tmp_dir, _os.path.splitext(_os.path.basename(path))[0] + '.docx')
                converter.convert_to_docx(path, tmp)
            else:
                from app import converter_linux
                tmp = converter_linux.convert_to_docx(path)
                tmp_dir = _os.path.dirname(tmp)
            self._tmp_dirs.append(tmp_dir)
            self._converted[path] = tmp
            return tmp
        finally:
            QApplication.restoreOverrideCursor()

    def _sync_scroll(self, src_view, dst_view):
        if self._sync_lock:
            return
        self._sync_lock = True
        try:
            sb = src_view.verticalScrollBar()
            db = dst_view.verticalScrollBar()
            if sb.maximum() > 0:
                db.setValue(int(round(db.maximum() * sb.value() / float(sb.maximum()))))
        finally:
            self._sync_lock = False

    def done(self, r):
        import shutil as _shutil
        from app.theme import settings as _settings
        _s = _settings()
        _s.setValue('preview/geometry', self.saveGeometry())
        if getattr(self, '_splitter', None) is not None:
            _s.setValue('preview/splitter', self._splitter.saveState())
        for d in self._tmp_dirs:
            _shutil.rmtree(d, ignore_errors=True)
        self._tmp_dirs = []
        for p in self._tmp_header_pngs:
            try:
                os.unlink(p)
            except Exception:
                pass
        self._tmp_header_pngs = []
        super(PreviewDialog, self).done(r)

    def _load_current(self):
        path = self._current_path()
        if not path:
            return
        self._current_paras = None
        try:
            result = _read_paragraphs(path)
            if result is None:
                # .doc/.wps：先转换为临时 docx 再预览
                result = _read_paragraphs(self._convert_for_preview(path))
        except Exception as e:
            err_msg = str(e)
            # 给出更具体的解决建议
            hint = ''
            if '未检测到' in err_msg or 'COM' in err_msg:
                hint = '<p><b>可能原因：</b>WPS/Word 正在忙于其他任务（如打开了大文件），<br>或 COM 组件初始化失败。关闭所有 Office 窗口后重试。</p>'
            elif 'Permission' in err_msg or '拒绝' in err_msg:
                hint = '<p><b>可能原因：</b>临时目录无写入权限。<br>请以管理员身份运行或检查杀毒软件是否拦截。</p>'
            msg = _html_shell(
                '<p style=\"color:#C0392B;\"><b>预览失败：</b>{}</p>'
                '{}'
                '<p>不影响实际处理，点击「开始排版」仍会正常转换并排版该文件。</p>'.format(
                    _esc(err_msg), hint))
            self.view_before.setHtml(msg)
            self.view_after.setHtml(msg)
            self.notice.setText('预览失败：{}'.format(err_msg[:80]))
            self.notice.setVisible(True)
            return
        paras, table_count, total_paras, has_auto_num = result
        self._current_paras = paras

        notes = []
        if table_count:
            notes.append('文档含 {} 个表格，预览不显示表格（实际处理时会规范表格格式）'.format(table_count))
        if total_paras > MAX_PARAS:
            notes.append('文档共 {} 段，仅预览前 {} 段'.format(total_paras, MAX_PARAS))
        if has_auto_num:
            notes.append('此文档含 Word 自动编号，导入时已自动转换为纯文字')
        self.notice.setText('；'.join(notes))
        self.notice.setVisible(bool(notes))

        self.view_before.setHtml(render_before_html(paras))
        self._render_after()
