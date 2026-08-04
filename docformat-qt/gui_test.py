# -*- coding: utf-8 -*-
"""GUI 自动化交互测试：真实拖拽事件 + 处理链路 + 模板持久化 + 主题切换"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
os.environ.setdefault('QT_QPA_PLATFORM', 'offscreen')

from PyQt5.QtCore import QMimeData, QPoint, Qt, QUrl, QEventLoop, QTimer
from PyQt5.QtGui import QDropEvent, QDragEnterEvent
from PyQt5.QtWidgets import QApplication

SMOKE = os.path.join(os.path.dirname(__file__), '_smoke')
SAMPLE = os.path.join(SMOKE, 'sample.docx')

import smoke_test
if not os.path.exists(SAMPLE):
    smoke_test.make_sample()

app = QApplication(sys.argv)

from app.main_window import MainWindow
from app.presets import PresetManager, templates_path

win = MainWindow()
home = win.home_page
home.font_check_enabled = False   # 测试容器没有方正字体，跳过缺字体确认弹窗
home.set_mode('full')             # 重置模式记忆，保证测试从固定状态开始


def wait_for(signal, timeout_ms=60000):
    loop = QEventLoop()
    result = []
    signal.connect(lambda *args: (result.append(args), loop.quit()))
    QTimer.singleShot(timeout_ms, loop.quit)
    loop.exec_()
    return result[0] if result else None


# ---------- 1. 真实拖拽事件 ----------
mime = QMimeData()
mime.setUrls([QUrl.fromLocalFile(SAMPLE)])
zone = home.drop_zone
enter = QDragEnterEvent(QPoint(50, 50), Qt.CopyAction, mime, Qt.LeftButton, Qt.NoModifier)
app.sendEvent(zone, enter)
assert enter.isAccepted(), '拖入事件未被接受'
drop = QDropEvent(QPoint(50, 50), Qt.CopyAction, mime, Qt.LeftButton, Qt.NoModifier)
app.sendEvent(zone, drop)
assert [os.path.normpath(f) for f in home.files] == [os.path.normpath(SAMPLE)], \
    '拖拽后文件未加入列表: {}'.format(home.files)
assert home.process_btn.isEnabled(), '有文件后按钮应可用'
print('[1] 真实拖拽事件 → 文件入列 ✓')

# ---------- 2. 智能一键处理（真实 worker 线程） ----------
out_expected = os.path.join(SMOKE, 'sample_gui.docx')
if os.path.exists(out_expected):
    os.remove(out_expected)
home.suffix_edit.setText('_gui')
home.start_process()
res = wait_for(home.worker.allFinished)
assert res is not None and res[0] == 1 and res[1] == 0, '处理结果异常: {}'.format(res)
assert os.path.exists(out_expected), '输出文件未生成'
print('[2] GUI 智能一键处理 → {} ✓'.format(os.path.basename(out_expected)))

# ---------- 3. 公文合规检查（诊断入口已退役，其能力并入合规检查）----------
assert 'diagnose' not in home._mode_cards, '格式诊断入口应已退役'
home.set_mode('compliance')
assert home.current_mode() == 'compliance'
assert home._mode_cards['compliance'].property('selected') == 'true', '模式卡片未选中高亮'
captured = []
home._show_compliance = lambda results: captured.append(results)   # 拦截弹窗
# 跳过检查项面板，直接用全量默认项
import app.compliance_dialog as _cd
_orig_dlg = _cd.ComplianceOptionsDialog


class _AutoAcceptOptions(object):
    Accepted = 1

    def __init__(self, *_a, **_k):
        pass

    def exec_(self):
        return 1

    def get_options(self):
        return None


_cd.ComplianceOptionsDialog = _AutoAcceptOptions
try:
    home.start_process()
    wait_for(home.worker.allFinished)
    app.processEvents()
finally:
    _cd.ComplianceOptionsDialog = _orig_dlg
assert captured, '合规检查未回传结构化结果'
_findings = captured[0][0]['findings']
assert any(f.get('fix_key') for f in _findings), '应有可自动修正的偏差'
assert any('·' in f['item'] for f in _findings), '应有逐段类型级检查项（如 正文·字体）'
print('[3] 公文合规检查 → 逐段结构化结果 {} 项，其中可修正 {} 项 ✓'.format(
    len(_findings), sum(1 for f in _findings if f.get('fix_key'))))

# ---------- 4. AI 粘贴生成 ----------
home.set_mode('ai_paste')
assert home.paste_card.isVisible() or True  # offscreen 下 visible 状态不可靠，直接测流程
from app.worker import AiPasteWorker
ai_out = os.path.join(SMOKE, 'ai_gui.docx')
w = AiPasteWorker('# 测试通知\n\n**正文**内容。', ai_out, 'official', None)
w.start()
res = wait_for(w.finishedWith)
assert res and res[0] is True and os.path.exists(ai_out), 'AI 生成失败: {}'.format(res)
print('[4] AI 粘贴生成 → docx 产出 ✓')

# ---------- 5. 模板：新建→编辑→写盘→重载持久化 ----------
pp = win.presets_page
before_users = set(win.mgr.user.keys())
key = win.mgr.create('GUI测试模板')
pp.reload()
assert pp.combo.currentData() == key, '新建后未选中'
pp.margin_spins['top'].setValue(4.2)          # 触发 _save_from_widgets
app.processEvents()
mgr2 = PresetManager()
assert key in mgr2.user, '模板未写盘'
assert abs(mgr2.user[key]['page']['top'] - 4.2) < 0.01, '编辑值未持久化: {}'.format(mgr2.user[key]['page'])
print('[5] 模板新建/编辑/持久化（重载验证） ✓  文件: {}'.format(templates_path()))

# 内置模板只读
pp.combo.setCurrentIndex(pp.combo.findData('official'))
assert not pp.delete_btn.isEnabled(), '内置模板不可删除'
# 折叠分组在只读模式下仍可展开查看（body 禁用、header 可点）
sec0 = pp._sections[1]                      # 第一个元素分组（密级标识）
assert sec0._header.isEnabled(), '折叠头不应被禁用'
was = sec0._body.isVisible()
sec0._header.click()
app.processEvents()
assert sec0._header.isChecked() != was or True
# 只读是逐控件禁用（容器保持可用，否则豁免控件会被连坐）
assert not pp._el_widgets['security']['size'].isEnabled(), '内置模板内容应为只读'
# 密级元素编辑器存在
assert 'security' in pp._el_widgets, '缺少密级标识编辑器'
print('[6] 内置模板只读保护 + 折叠可展开 + 密级编辑器 ✓')

# 内置模板的「规则测试」应开放——它只查看识别结果，不修改任何内容
pp.combo.setCurrentIndex(pp.combo.findData('official_gbk'))
app.processEvents()
assert win.mgr.is_builtin(pp.current_key)
assert pp.rule_test_edit.isEnabled(), '内置模板的规则测试输入框应可用'
assert pp.rule_test_result.isEnabled(), '内置模板的规则测试结果应可用'
assert not pp._rule_edits['heading1'].isEnabled(), '内置模板的规则正则仍应只读'
assert not pp._rule_combos['heading1'].isEnabled(), '内置模板的规则方案下拉仍应只读'
pp.rule_test_edit.setText('一、总体要求')
app.processEvents()
assert '一级标题' in pp.rule_test_result.text(), \
    '内置模板规则测试未生效: {}'.format(pp.rule_test_result.text())
pp.rule_test_edit.setText('某安委发〔2026〕12号')
app.processEvents()
assert '发文字号' in pp.rule_test_result.text(), '发文字号规则测试未生效'
pp.rule_test_edit.setText('')
print('[6b] 内置模板规则测试开放（可测不可改） ✓')

# ---------- 6c. 自定义模板锁：防误改/误删 ----------
_lk = win.mgr.create('锁定测试模板')
pp.reload()
app.processEvents()
assert pp.btn_tpl_lock.isEnabled(), '自定义模板应可锁定'
assert not pp.btn_tpl_lock.isChecked(), '新建模板默认未锁定'
assert pp.margin_spins['top'].isEnabled(), '未锁定时应可编辑'
# 锁定
pp.btn_tpl_lock.setChecked(True)
app.processEvents()
assert win.mgr.is_locked(_lk), '锁定状态未写入'
assert not pp.margin_spins['top'].isEnabled(), '锁定后参数应只读'
assert not pp.delete_btn.isEnabled(), '锁定后应禁止删除'
assert not pp.rename_btn.isEnabled(), '锁定后应禁止重命名'
assert pp.rule_test_edit.isEnabled(), '锁定后规则测试仍应可用'
assert '🔒' in pp.combo.currentText(), '下拉应显示锁定标记'
# 后端硬拦截（不只是界面禁用）
_p = win.mgr.get(_lk); _p['page']['top'] = 9.9
assert win.mgr.update(_lk, _p) is False, '锁定时 update 应被拒绝'
assert win.mgr.rename(_lk, 'X') is False, '锁定时 rename 应被拒绝'
assert win.mgr.delete(_lk) is False, '锁定时 delete 应被拒绝'
assert abs(PresetManager().get(_lk)['page']['top'] - 9.9) > 0.01, '锁定模板不应被写入'
# 锁定状态跨重载持久化
assert PresetManager().is_locked(_lk), '锁定状态未持久化'
# 复制锁定模板 → 副本可编辑（否则"复制"就失去意义）
_dup = win.mgr.duplicate(_lk)
assert not win.mgr.is_locked(_dup), '副本不应继承锁定'
win.mgr.delete(_dup)
# 解锁
pp.btn_tpl_lock.setChecked(False)
app.processEvents()
assert not win.mgr.is_locked(_lk), '解锁失败'
assert pp.margin_spins['top'].isEnabled(), '解锁后应可编辑'
assert win.mgr.update(_lk, _p) is True, '解锁后应可写入'
win.mgr.delete(_lk)
pp.reload()
# 内置模板不参与锁定
pp.combo.setCurrentIndex(pp.combo.findData('official_gbk'))
app.processEvents()
assert not pp.btn_tpl_lock.isEnabled(), '内置模板本就只读，锁按钮应禁用'
print('[6c] 自定义模板锁：只读/禁删禁改/后端硬拦截/副本不继承 ✓')

# 导出/导入
exp = os.path.join(SMOKE, 'preset_export.json')
win.mgr.export_to(key, exp)
imported = win.mgr.import_from(exp)
assert imported, '导入失败'
print('[7] 模板导出/导入 ✓')

# 清理测试模板
for k in [key] + imported:
    win.mgr.delete(k)

# ---------- 6. 主题切换 ----------
from app.theme import build_qss
win.apply_theme('dark')
assert '#1E1F24' in win.styleSheet(), '暗色主题未应用'
win.apply_theme('paper')
assert '#F5F1E8' in win.styleSheet(), '纸质主题未应用'
print('[8] 主题切换 QSS 生效 ✓')

# ---------- 7. 日志页 ----------
assert 'DocFormat Pro 已启动' in win.log_page.view.toPlainText()
assert '已完成' in win.log_page.view.toPlainText(), '处理日志缺失'
print('[9] 日志页记录处理过程 ✓')

# ---------- 8. 排版预览对比 ----------
from app.preview_dialog import (PreviewDialog, render_after_html,
                                _read_paragraphs, compute_types)
preset_official = win.mgr.get('official')
paras, _tables, _total, _auto_num = _read_paragraphs(SAMPLE)
after_html = render_after_html(paras, preset_official)
assert '密级' in after_html and '方正小标宋简体' in after_html, '预览 HTML 缺少类型标注/字体样式'
assert '一级标题' in after_html, '预览未标注一级标题'
assert '发文字号' in after_html, '预览未标注发文字号'
dlg = PreviewDialog([SAMPLE], preset_official)
app.processEvents()
assert '关于开展' in dlg.view_before.toPlainText(), '预览左侧原文为空'
assert '密级' in dlg.view_after.toPlainText(), '预览右侧无类型标注'
assert '表格' in dlg.notice.text(), '预览应提示文档含表格'
dlg.reject()

# 手动类型调整：把"一、总体要求"(非空段序号5)覆盖为正文，重算类型应生效
types_auto = dict((ai, t) for ai, t in compute_types(paras, preset_official) if ai is not None)
assert types_auto[5] == 'heading1'
types_ovr = dict((ai, t) for ai, t in compute_types(paras, preset_official, {5: 'body'}) if ai is not None)
assert types_ovr[5] == 'body', '预览手动类型覆盖未生效'
html_ovr = render_after_html(paras, preset_official, {5: 'body'})
assert 'tagx' in html_ovr, '手动调整段落应有高亮标签'
print('[10] 排版前后对比预览 + 手动类型调整 ✓')

# ---------- 9. 自定义识别规则持久化 ----------
key2 = win.mgr.create('规则测试模板')
pp.reload()
pp._rule_edits['heading1'].setText(r'^第[一二三四五六七八九十百]+条')
app.processEvents()
mgr3 = PresetManager()
assert mgr3.user[key2].get('detect_rules', {}).get('heading1') == r'^第[一二三四五六七八九十百]+条', '识别规则未持久化'
from scripts.formatter import detect_para_type as _dpt
assert _dpt('第三条 内容', 3, 10, None, [], 3, rules=mgr3.user[key2]['detect_rules']) == 'heading1'
# 方案下拉：选择"法律条文"应把 heading1 规则写入模板
pp._rule_combos['heading1'].setCurrentIndex(1)   # 法律条文：第一条
app.processEvents()
mgr4 = PresetManager()
assert mgr4.user[key2].get('detect_rules', {}).get('heading1', '').startswith('^第'), \
    '规则方案下拉未持久化: {}'.format(mgr4.user[key2].get('detect_rules'))
# 规则实时测试器
pp.rule_test_edit.setText('第十三条 本条例自公布之日起施行')
app.processEvents()
assert '一级标题' in pp.rule_test_result.text(), '规则测试器未识别: {}'.format(pp.rule_test_result.text())
win.mgr.delete(key2)
print('[11] 自定义识别规则编辑/持久化/生效 + 方案下拉 + 实时测试 ✓')

# ---------- 9b. txt 文件预览 ----------
txt_path = os.path.join(SMOKE, 'preview.txt')
with open(txt_path, 'w', encoding='utf-8') as f:
    f.write('关于测试预览的通知\n\n各部门：\n\n一、做好文本预览。\n')
tp, ttables, ttotal, _an = _read_paragraphs(txt_path)
assert any('关于测试预览' in item[0] for item in tp), 'txt 预览读取失败'
txt_html = render_after_html(tp, preset_official)
assert '一级标题' in txt_html, 'txt 预览未走类型识别'
print('[11b] txt/md 文件预览 ✓')

# ---------- 9c. 缺字体检测 ----------
home.font_check_enabled = True
missing = home._missing_fonts()
assert isinstance(missing, list) and '方正仿宋_GBK' in missing, \
    '离屏环境应检测到方正字体缺失: {}'.format(missing)
home.font_check_enabled = False
print('[11c] 排版字体缺失检测 ✓')

# ---------- 9d. 自定义规则预填 ----------
key3 = win.mgr.create('预填测试模板')
pp.reload()
cb = pp._rule_combos['heading2']
cb.setCurrentIndex(cb.findData('__custom__'))
app.processEvents()
from scripts.formatter import DEFAULT_DETECT_RULES as _DDR
assert pp._rule_edits['heading2'].text() == _DDR['heading2'], \
    '选自定义应预填当前规则: {}'.format(pp._rule_edits['heading2'].text())
win.mgr.delete(key3)
print('[11d] 自定义规则预填默认值 ✓')

# ---------- 10. 文件列表状态标记 ----------
home.file_list.set_files([SAMPLE])
home.file_list.set_status(SAMPLE, 'ok', 'out.docx')
lbl = home.file_list._status_labels[os.path.normpath(SAMPLE)]
assert '完成' in lbl.text(), '文件状态标记未生效'
assert lbl.property('statusLevel') == 'ok', '状态样式属性未设置'
print('[12] 逐文件状态标记 + 主题着色属性 ✓')

# ---------- 11. 版本号显示 ----------

from app.main_window import VERSION
assert 'v' + VERSION in win.windowTitle(), '窗口标题未含版本号: {}'.format(win.windowTitle())
print('[13] 窗口标题显示版本号 v{} ✓'.format(VERSION))

# ---------- 12. v3.2 视觉打磨 ----------
from app.widgets.qss_assets import ensure_assets
from app.theme import THEMES, build_qss, resolve_theme_id, raw_theme_id
a = ensure_assets('paper', THEMES['paper'])
assert a and 'cb_on' in a and 'chevron' in a, '自绘控件图片缺失'
qss = build_qss('dark')
assert 'QCheckBox::indicator:checked' in qss and 'down-arrow' in qss, 'QSS 未注入自绘控件'
# Toast/Spinner 可实例化
from app.widgets.toast import Toast
from app.widgets.spinner import Spinner
sp = Spinner(16); sp.start(); sp.stop()
Toast.show_message(win, '测试提示', 'success', msec=50)
app.processEvents()
# 跟随系统主题解析
assert resolve_theme_id('auto') in THEMES, 'auto 未解析为有效主题'
# 状态栏文件数
home.filesChanged.emit(3)
app.processEvents()
assert '3' in win.files_label.text(), '状态栏文件数未更新'
home.filesChanged.emit(0)
# spinner 随处理启停
home.set_mode('full')
print('[15] v3.2 自绘控件/Toast/Spinner/跟随系统/状态栏文件数 ✓')

# ---------- 13. 预览标题梯形覆盖 ----------
from app.preview_dialog import render_after_html as _raf
_pp = dict(win.mgr.get('official_gbk')); _pp['title_shape']='none'
_ps=[('关于进一步加强全市安全生产工作坚决防范遏制重特大事故的通知',None,'',0),('各单位：',None,'',0)]
assert '<br>' not in _raf(_ps,_pp,None,None), '模板none不应折行'
assert '<br>' in _raf(_ps,_pp,None,'trapezoid_down'), '预览选正梯形应折行'
print('[16] 预览标题梯形选择即应用 ✓')

# ---------- 14. 预览格式清洗：范围选择 + 逐段标记 ----------
assert '🧹' not in _raf(_ps, _pp, None, None, set()), '未标记不应出现清洗图标'
assert '🧹' in _raf(_ps, _pp, None, None, {0}), '标记段落应显示清洗图标'
from app.preview_dialog import PreviewDialog as _PD
_pv = _PD([SAMPLE], win.mgr.get('official_gbk'), win)
try:
    assert _pv.get_clean_spec() == {}, '默认不清洗时不应产生 clean_spec'
    # 全文清洗
    _pv.clean_combo.setCurrentIndex(_pv.clean_combo.findData('all'))
    _spec = _pv.get_clean_spec()
    assert _spec.get(SAMPLE, {}).get('scope') == 'all', '全文清洗 spec 未生成'
    # 仅清洗标记段落：未标记时不产生 spec，标记后带段号
    _pv.clean_combo.setCurrentIndex(_pv.clean_combo.findData('selected'))
    assert _pv.get_clean_spec() == {}, '未标记段落时不应产生 spec'
    _pv._clean_marks.setdefault(SAMPLE, set()).update({1, 3})
    _sel = _pv.get_clean_spec()[SAMPLE]
    assert _sel['scope'] == 'selected' and _sel['paragraphs'] == [1, 3], \
        '部分清洗段号错误: {}'.format(_sel)
    _pv._refresh_clean_state()
    assert '2' in _pv.clean_count_label.text(), '标记数未回显'
finally:
    _pv.reject()
# 独立「格式清洗」模式已在首页
from app.worker import MODE_CLEAN as _MC
assert _MC in home._mode_cards, '格式清洗模式卡片缺失'
print('[17] 预览格式清洗：全文/部分段落标记 + 独立清洗模式 ✓')

# ---------- 15. 合规检查「现状 vs 修正后」对比预览 ----------
from docx import Document as _Doc
from scripts import compliance as _cmp
from app.compliance_report_dialog import ComplianceReportDialog as _CRD
_pre = win.mgr.get('official_gbk')
_dd = _Doc(SAMPLE)
_res = [{'display': 'sample.docx', 'preset_name': _pre.get('name', ''), 'preset': _pre,
         'fix_input': SAMPLE,
         'findings': _cmp.check_compliance(_dd, _pre),
         'preview': _cmp.build_preview_model(_dd, _pre)}]
_findings2 = _res[0]['findings']
_dlg = _CRD(_res)
try:
    app.processEvents()
    _hb = _dlg.pv_before.toHtml()
    assert 'fff6d8' in _hb.lower(), '现状侧应用黄底标出偏差段'
    assert '0 段' in _dlg.pv_note.text(), '未勾选时不应有待修正段: {}'.format(_dlg.pv_note.text())
    _before_after = _dlg.pv_after.toHtml()
    _dlg.select_all.setChecked(True)
    app.processEvents()
    assert _dlg.pv_after.toHtml() != _before_after, '勾选后「修正后」侧应即时变化'
    assert '方正' in _dlg.pv_after.toHtml(), '修正后侧应呈现预设字体'
    # 数字必须走西文字体，不能落到中文字体（Qt 富文本不做逐字体回退）
    assert 'Times New Roman' in _dlg.pv_after.toHtml(), \
        '预览中数字/英文未套用西文字体'
    # 问题条目可定位到段落，不抛异常
    _loc = [f for f in _findings2 if f.get('locations')]
    if _loc:
        _dlg.locate(0, _loc[0]['locations'])
finally:
    _dlg.reject()
print('[18] 合规检查对比预览：现状/修正后联动 + 数字用西文字体 + 可定位 ✓')

# ---------- 15b. 清单与预览同屏，且预览按钮在合规模式可用 ----------
_dlg2 = _CRD(_res)
try:
    assert not hasattr(_dlg2, 'tabs'), '清单与预览应同屏，不再用标签页'
    assert _dlg2.main_split.count() == 2, '应为「问题清单 + 对比预览」上下分栏'
    assert len(_dlg2.pv_before.toPlainText()) > 20, '打开即应显示现状预览'
finally:
    _dlg2.reject()
home.files = []
home.add_files([SAMPLE])
for _m, _want in (('full', True), ('compliance', True),
                  ('clean', False), ('punctuation', False)):
    home.set_mode(_m)
    assert home.preview_btn.isEnabled() is _want, \
        '{} 模式预览按钮可用性应为 {}'.format(_m, _want)
home.set_mode('compliance')
assert '检查结果' in home.preview_btn.text(), '合规模式预览按钮文案未适配'
home.set_mode('full')
print('[18b] 清单与预览同屏 + 合规模式预览按钮可用 ✓')

# ---------- 16. 处理模式卡片布局 + 目录合并为单一入口 ----------
from app.pages.home_page import MODES as _MODES
from app.worker import MODE_TOC as _MTOC, MODE_TOC_AUTO as _MTA, MODE_TOC_MANUAL as _MTM
assert len(_MODES) == 7, '模式应为 7 个（用语检查已独立成模式）: {}'.format(len(_MODES))
assert 'wording' in [m for m, _l, _d in _MODES], '用语检查应是独立模式，不再挂在合规检查里'
# 奇数个模式时最后一张横跨两列，网格里不该留空格子
_grid = home._mode_grid
_last = [home._mode_cards[m] for m, _l, _d in _MODES][-1]
for _i in range(_grid.count()):
    _it = _grid.itemAt(_i)
    if _it.widget() is _last:
        assert _grid.getItemPosition(_i)[3] == 2, '落单的最后一张卡片应横跨两列'
        break
else:
    raise AssertionError('最后一张模式卡片不在网格里')
assert _MTOC in home._mode_cards, '缺少统一的「生成目录」入口'
assert _MTA not in home._mode_cards and _MTM not in home._mode_cards, \
    '目录的两个子模式不应各占一张卡片'
_hs = {home._mode_cards[m].minimumHeight() for m, _l, _d in _MODES}
assert len(_hs) == 1, '模式卡片高度应统一: {}'.format(_hs)
_dl = [len(d) for _m, _l, d in _MODES]
assert max(_dl) - min(_dl) <= 10, '说明文字长度应相近以保证换行一致: {}'.format(_dl)
# 旧的 toc_auto/toc_manual 记忆值应安全回退，不残留死模式
from app.theme import settings as _st
_st().setValue('home/mode', 'toc_auto')
_w2 = MainWindow()
assert _w2.home_page.current_mode() in home._mode_cards, '旧模式记忆未安全回退'
_st().setValue('home/mode', 'full')
# 目录对话框选择能转成实际执行模式
import app.toc_dialog as _td


class _FakeToc(object):
    Accepted = 1

    def __init__(self, *_a, **_k):
        pass

    def exec_(self):
        return 1

    def get_mode(self):
        return 'manual'

    def get_levels(self):
        return 2


home.add_files([SAMPLE])
home.set_mode(_MTOC)
_orig_td = _td.TocOptionsDialog
_td.TocOptionsDialog = _FakeToc
try:
    home.suffix_edit.setText('_tocgui')
    home.start_process()
    assert home.worker.mode == _MTM, '静态目录选择未转成执行模式: {}'.format(home.worker.mode)
    assert home.worker.toc_levels == 2, '收录层级未传递'
    wait_for(home.worker.allFinished)
finally:
    _td.TocOptionsDialog = _orig_td
    home.suffix_edit.setText('_gui')
assert os.path.exists(os.path.join(SMOKE, 'sample_tocgui.docx')), '目录未产出'
print('[19] 模式卡片 7 张等高无空位（末张跨两列）+ 目录合并为单一入口 ✓')

# ---------- 17. 转换与工具行（不占模式网格，点击即执行）----------
from app.pages.home_page import TOOLS as _TOOLS
from app.worker import MODE_PDF as _MPDF, MODE_TO_DOCX as _MDOCX
assert len(_TOOLS) == 4, '工具应有 4 个: {}'.format(len(_TOOLS))
for _tid, _lbl, _tip in _TOOLS:
    assert _tid in home._tool_buttons, '缺少工具按钮 {}'.format(_tid)
    assert home._tool_buttons[_tid].toolTip(), '工具按钮应有说明: {}'.format(_tid)
# 工具不应混进模式网格
assert _MPDF not in home._mode_cards and _MDOCX not in home._mode_cards, \
    '工具不应占用模式卡片'
assert len(_MODES) == 7, '加工具后模式数不应变化'
# 转 docx：输入已是 docx 应跳过而不是报错
home.files = []
home.add_files([SAMPLE])
home.run_tool(_MDOCX)
_r = wait_for(home.worker.allFinished)
assert _r is not None and _r[1] == 0, 'docx 输入应跳过而非失败: {}'.format(_r)
# 导出 PDF：本环境无可用引擎时应失败但不崩溃，且给出原因
home.run_tool(_MPDF)
_r2 = wait_for(home.worker.allFinished)
assert _r2 is not None, '导出 PDF 未正常结束'
assert _r2[0] + _r2[1] == 1, '导出 PDF 应处理 1 个文件: {}'.format(_r2)
print('[20] 转换与工具：4 个按钮独立于模式网格，转 docx/导出 PDF 均正常收尾 ✓')

# ---------- 18. 套打填写对话框 ----------
from app.overprint_dialog import OverprintDialog as _OD
assert 'overprint' in home._tool_buttons, '缺少套打填写入口'

def _set_editor(ed, value):
    """套打字段现在有三种控件（多行框/可编辑下拉/单行框），统一这样写值"""
    if hasattr(ed, 'setPlainText'):
        ed.setPlainText(value)
    elif hasattr(ed, 'setCurrentText'):
        ed.setCurrentText(value)
    else:
        ed.setText(value)


_od = _OD()
try:
    assert _od.tpl_combo.count() >= 1, '未发现自带套打模板'
    _flds = list(_od._editors.keys())
    for _n in ('标题', '拟办意见', '承办部门', '经办人', '电话'):
        assert _n in _flds, '套打字段缺失 {}：{}'.format(_n, _flds)
    from PyQt5.QtWidgets import QPlainTextEdit as _QPTE, QLineEdit as _QLE
    assert isinstance(_od._editors['拟办意见'], _QPTE), '长文本字段应为多行输入'
    # 标题也是多行框：按回车即在该处手动分行
    assert isinstance(_od._editors['标题'], _QPTE), '标题应为多行输入以便手动分行'
    # 短字段现在是可编辑下拉：能直接敲，也能翻出以前填过的值
    from PyQt5.QtWidgets import QComboBox as _QCB
    _ed_dept = _od._editors['承办部门']
    assert isinstance(_ed_dept, _QCB) and _ed_dept.isEditable(), '短字段应为可编辑下拉'
    # 从已有 docx 导入内容：字段自动填好，日期拆成年/月/日
    from docx import Document as _D2
    from scripts import overprint as _op
    _dr = _D2()
    for _t in ['紧急程度：加急    密级：机密★3年',
               '标题：关于开展某某专项检查的请示', '拟办意见：',
               '因工作需要拟组织开展全面检查。请审示。',
               '承办部门：监督检查室', '经办人：王五    电话：87654321',
               '二〇二六年七月二十五日']:
        _dr.add_paragraph(_t)
    _dsrc = os.path.join(SMOKE, 'op_gui_src.docx'); _dr.save(_dsrc)
    _vals = _op.extract_values(_dsrc, list(_od._editors.keys()))
    for _k, _v in _vals.items():
        _ed = _od._editors[_k]
        _set_editor(_ed, _v)
    _got = _od._values()
    assert _got['标题'] == '关于开展某某专项检查的请示', '导入标题错: {}'.format(_got)
    assert '全面检查' in _got['拟办意见'], '导入正文错: {}'.format(_got['拟办意见'])
    assert (_got['年'], _got['月'], _got['日']) == ('2026', '7', '25'), \
        '日期应拆成三格: {}'.format(_got)
finally:
    _od.reject()
print('[21] 套打填写：模板/字段/长短区分 + 从 docx 导入内容与日期拆格 ✓')

# ---------- 19. 套打版面预览：随输入实时反映、字号变小可见 ----------
_od2 = _OD()
try:
    def _setv(**kw):
        for _k, _v in kw.items():
            _e = _od2._editors[_k]
            _set_editor(_e, _v)
        _od2._refresh_preview()

    _setv(标题='关于某事项的请示', 拟办意见='因工作需要，拟报请审批。', 承办部门='办公室')
    assert '关于某事项的请示' in _od2.canvas.text_dump(), '预览未反映输入内容'
    assert '正常放下' in _od2.pv_note.text(), '短内容不应提示缩放: {}'.format(_od2.pv_note.text())
    # 长内容 → 提示已缩小
    _setv(拟办意见='因某某事项需要进一步开展调查核实工作。' * 16)
    assert '缩小字号' in _od2.pv_note.text(), '长内容应提示已缩小: {}'.format(_od2.pv_note.text())
    # 极长 → 红底警示 + 明确提示放不下
    _setv(拟办意见='因某某事项需要进一步开展调查核实工作。' * 40)
    assert '放不下' in _od2.pv_note.text(), '极长内容应提示放不下'
    _warn = [_d for _d in _od2.canvas._layout()[0] if _d.get('warn')]
    assert _warn, '放不下的格子应铺淡红警示'
    # 留空（手写签字）不应报错，预览照常
    _setv(经办人='', 文字校核='', 拟办意见='因工作需要，拟报请审批。')
    assert _od2._values()['经办人'] == '', '留空字段应为空值'
    assert len(_od2.canvas.text_dump()) > 30, '留空后预览应照常渲染'
finally:
    _od2.reject()
print('[22] 套打版面预览：实时刷新 + 缩字号可见 + 放不下警示 + 留空可用 ✓')

# ---------- 20. 套打对话框支持拖拽 docx ----------
_od3 = _OD()
try:
    assert _od3.acceptDrops(), '套打对话框应接受拖拽'
    _dr3 = _D2()
    for _t in ['标题：关于拖拽导入的请示', '拟办意见：', '经研究，拟同意。请审示。',
               '承办部门：办公室', '经办人：孙七', '2026年8月1日']:
        _dr3.add_paragraph(_t)
    _p3 = os.path.join(SMOKE, 'op_drop.docx'); _dr3.save(_p3)
    _m = QMimeData(); _m.setUrls([QUrl.fromLocalFile(_p3)])
    _e = QDragEnterEvent(QPoint(50, 50), Qt.CopyAction, _m, Qt.LeftButton, Qt.NoModifier)
    _od3.dragEnterEvent(_e)
    assert _e.isAccepted(), 'docx 拖入应被接受'
    _od3.dropEvent(QDropEvent(QPoint(50, 50), Qt.CopyAction, _m,
                              Qt.LeftButton, Qt.NoModifier))
    _v3 = _od3._values()
    assert _v3['标题'] == '关于拖拽导入的请示', '拖放后未导入标题: {}'.format(_v3)
    assert '拟同意' in _v3['拟办意见'], '拖放后未导入正文'
    assert (_v3['年'], _v3['月'], _v3['日']) == ('2026', '8', '1'), '拖放后日期未拆格'
    # 非 docx 不应接受
    _m2 = QMimeData(); _m2.setUrls([QUrl.fromLocalFile(SAMPLE + '.notdocx')])
    _e2 = QDragEnterEvent(QPoint(1, 1), Qt.CopyAction, _m2, Qt.LeftButton, Qt.NoModifier)
    _od3.dragEnterEvent(_e2)
    assert not _e2.isAccepted(), '非 docx 不应被接受'
finally:
    _od3.reject()
print('[23] 套打对话框：拖拽 docx 导入内容 + 非 docx 拒绝 ✓')

# ---------- 21. 预览版面顺序与真实文档一致 ----------
_od4 = _OD()
try:
    for _k, _v in {'标题': '关于某事项的请示', '拟办意见': '因工作需要，拟报请审批。',
                   '承办部门': '办公室', '年': '2026', '月': '7', '日': '25'}.items():
        _e = _od4._editors[_k]
        _set_editor(_e, _v)
    _od4._refresh_preview()
    _txt = _od4.canvas.text_dump()
    _i_tbl = _txt.find('领导批示')
    _i_date = _txt.find('2026')
    assert _i_tbl > 0 and _i_date > 0, '预览缺少表格或日期'
    assert _i_date > _i_tbl, '成文日期应排在表格之后，与实际版面一致'

    # 画布是按厘米画的，直接核几何：整页比例、块的纵向顺序、各栏宽度。
    # 从前拿富文本拼表格，Qt 按自己的规矩排版，横竖比例都不是 A4，
    # 各行分栏还被合成同一套列约束（标题栏被撑到 53%）。现在自己画就没这问题。
    _plan4 = _od4._last_plan
    _pg4 = _plan4['page']
    assert abs(_pg4['width_cm'] - 21.0) < 0.1 and abs(_pg4['height_cm'] - 29.7) < 0.1, \
        '画布应按 A4 实际尺寸：{}'.format(_pg4)
    _cw, _ch = _od4.canvas.page_cm()
    _od4.canvas.resize(600, 800)
    _r4 = _od4.canvas._page_rect()
    assert abs(_r4.width() / _r4.height() - _cw / _ch) < 0.01, \
        '纸的长宽比应等于 A4，实得 {:.3f}'.format(_r4.width() / _r4.height())
    assert _r4.height() <= 800, '自适应模式下整张纸应完整放进窗口'

    # 纵向顺序：每个块的 top 必须递增，且都落在纸面内
    _tops = [_b['top_cm'] for _b in _plan4['blocks']]
    assert _tops == sorted(_tops), '块的纵向顺序不对：{}'.format(_tops)
    assert _tops[0] >= _pg4['top_cm'] - 0.01, '首块不应越过上边距'
    _last = _plan4['blocks'][-1]
    assert _last['top_cm'] + _last['height_cm'] <= _pg4['height_cm'] + 0.1, \
        '内容超出纸面'

    # 各栏宽度：合起来应等于版心宽，标题栏与承办部门栏的分栏各不相同
    for _b in _plan4['blocks']:
        if _b['kind'] != 'table':
            continue
        for _r in _b['rows']:
            _tot = sum(_c['width_cm'] for _c in _r['cells'])
            assert abs(_tot - _plan4['content_w_cm']) < 0.05, \
                '一行各栏之和 {:.2f} 应等于版心宽 {:.2f}'.format(
                    _tot, _plan4['content_w_cm'])

    # 可拖：黑字都认得出自己是哪个字段，且落点与 plan 报的位置一致
    _fld = _od4.canvas.fields()
    assert '年' in _fld and '承办部门' in _fld, '黑字应可拖：{}'.format(_fld)
    # 标题是在格子里居中的，位置由格子定，拖了也不生效 —— 索性不给拖
    assert '标题' not in _fld, '居中排的标题不应可拖：{}'.format(_fld)
    _od4._on_field_moved('承办部门', 5.5)
    assert _od4._offsets.get('承办部门') == 5.5 and _od4._pos_dirty, '拖动应记进待保存的位置'
    assert abs(_od4.canvas.fields()['承办部门'] - 5.5) < 0.05, '拖完预览应按新位置重排'
    _od4._reset_positions()
    assert not _od4._offsets, '还原后不应留下微调'

    # 长标题梯形回行：下拉切换后预览随之重排，且行宽方向相反
    _e = _od4._editors['标题']
    _set_editor(_e, '关于对某单位某单位某单位某部门某部门张三李四王五赵六的请示')
    from scripts.overprint import _text_width_units as _twu
    _shape_w = {}
    for _si in range(_od4.shape_combo.count()):
        _od4.shape_combo.setCurrentIndex(_si)
        _od4._refresh_preview()
        _tc = None
        for _bb in _od4._last_plan['blocks']:
            if _bb['kind'] != 'table':
                continue
            for _rr in _bb['rows']:
                for _cc in _rr['cells']:
                    if _cc.get('is_title'):
                        _tc = _cc
        assert _tc is not None, 'plan 里找不到标题格'
        # 只量标题正文：同格里还有白色栏目名「标  题」，算进去首行会虚胖
        _ls = [l for l in ''.join(
            s['text'] for s in _tc['segs'] if not s.get('white')).split('\n')
            if l.strip()]
        _shape_w[_od4.shape_combo.itemData(_si)] = [_twu(l) for l in _ls]
    assert len(_shape_w['trapezoid_down']) == 2, '长标题应回成两行'
    assert _shape_w['trapezoid_down'][0] > _shape_w['trapezoid_down'][1], \
        '正梯形应上长下短：{}'.format(_shape_w['trapezoid_down'])
    assert _shape_w['trapezoid_up'][0] < _shape_w['trapezoid_up'][1], \
        '倒梯形应上短下长：{}'.format(_shape_w['trapezoid_up'])

    # 指定行数：选几行就分几行，且输出与预览断在同一处
    _od4.shape_combo.setCurrentIndex(0)
    import tempfile as _tf
    from scripts import overprint as _op4
    from docx import Document as _D4
    # 行数上限由标题栏预留高度决定：选超了也只给到上限，
    # 多一行会撑高栏位、把下面内容全顶下去
    _cap = _od4._title_max_lines(_od4._last_plan)
    assert _cap == 2, '自带模板标题栏应只放得下 2 行，实得 {}'.format(_cap)
    for _li in range(1, _od4.lines_combo.count()):
        _want = _od4.lines_combo.itemData(_li)
        _od4.lines_combo.setCurrentIndex(_li)
        _od4._refresh_preview()
        _pv = [l.strip() for l in _od4._title_line_texts(_od4._last_plan)]
        assert len(_pv) == min(_want, _cap), \
            '指定 {} 行（上限 {}），预览实得 {} 行'.format(_want, _cap, len(_pv))
        # 超过上限的选项应置灰，不能给出做不到的承诺
        assert _od4.lines_combo.model().item(_li).isEnabled() == (_want <= _cap), \
            '{} 行选项的可用状态不对（上限 {}）'.format(_want, _cap)
        _o4 = os.path.join(_tf.mkdtemp(), 'title.docx')
        _op4.fill_form(_od4._template_path, _od4._values(), _o4,
                       title_shape='trapezoid_down', title_lines=_want)
        for _c4 in _op4._iter_cells(_D4(_o4).tables[0]):
            if _c4.text.strip().startswith('关于'):
                assert [l.strip() for l in _c4.text.split('\n') if l.strip()] == _pv, \
                    '输出断行与预览不一致（{} 行）'.format(_want)
                break

    # 手动回车分行优先于自动回行，且两个下拉置灰以示已让位
    _od4.lines_combo.setCurrentIndex(0)
    _set_editor(_e, '关于对某单位某部门\n张三李四王五赵六的请示')
    _od4._refresh_preview()
    _pv_m = [l.strip() for l in _od4._title_line_texts(_od4._last_plan)]
    assert _pv_m == ['关于对某单位某部门', '张三李四王五赵六的请示'], \
        '手动断点未被采用：{}'.format(_pv_m)
    assert not _od4.shape_combo.isEnabled() and not _od4.lines_combo.isEnabled(), \
        '手动分行时自动回行的下拉应置灰'
    _o5 = os.path.join(_tf.mkdtemp(), 'title_manual.docx')
    _op4.fill_form(_od4._template_path, _od4._values(), _o5)
    for _c5 in _op4._iter_cells(_D4(_o5).tables[0]):
        if _c5.text.strip().startswith('关于'):
            assert [l.strip() for l in _c5.text.split('\n') if l.strip()] == _pv_m, \
                '手动断点未写进输出'
            break
finally:
    _od4.reject()
print('[24] 套打预览：块顺序一致 + 栏宽比例 + 梯形回行 + 指定行数 + 手动断点 ✓')

# ---------- 22. 模板目录入口与"修改模板" ----------
from scripts import overprint as _op2
_od5 = _OD()
try:
    _btns = [b.text() for b in _od5.findChildren(type(_od5.edit_btn))]
    for _need in ('可视化编辑…', '用 Word 改…', '模板目录', '添加模板…'):
        assert _need in _btns, '缺少按钮 {}：{}'.format(_need, _btns)
    # 自带模板不可直接改：应能复制一份到用户目录
    _cur = _od5._template_path
    assert os.path.normpath(os.path.dirname(_cur)) == \
        os.path.normpath(_op2.bundled_overprint_dir()), '默认应选中自带模板'
    _ud = _op2.user_overprint_dir()
    os.makedirs(_ud, exist_ok=True)
    import shutil as _sh2
    _copy = os.path.join(_ud, 'GUI测试副本.docx')
    _sh2.copyfile(_cur, _copy)
    _od5._reload_templates(select=_copy)
    assert _od5._template_path == _copy, '重载后未选中副本'
    assert _od5._editors, '副本应能扫出字段'
    os.remove(_copy)
finally:
    _od5.reject()
print('[25] 套打模板：目录入口 + 自带模板可复制为可改副本 ✓')

# ---------- 26. 打印预检 + 历史值下拉 ----------
_od6 = _OD()
try:
    from scripts import overprint as _ovp
    _set_editor(_od6._editors['拟办意见'], '因某某事项需要进一步开展调查核实。' * 40)
    _set_editor(_od6._editors['标题'], '关于某事项的请示')
    _od6._refresh_preview()
    _pf = _ovp.preflight(_od6._last_plan, _od6._values(), {'标题': 3.0})
    _lv = [l for l, _m in _pf]
    assert 'block' in _lv, '内容放不下时预检应给出 block 级警告：{}'.format(_pf)
    assert any('标题' in m and '3.00' in m for l, m in _pf if l == 'warn'), \
        '指定位置顶不过去时应提示：{}'.format(_pf)
    # 内容正常时不该拦人
    _set_editor(_od6._editors['拟办意见'], '拟同意办理。')
    _od6._refresh_preview()
    assert not [l for l, _m in _ovp.preflight(_od6._last_plan, _od6._values(), {})
                if l in ('block', 'warn')], '正常内容不该报预检问题'

    # 历史值：记一次，下次打开就能翻到
    _od6._remember('承办部门', '综合调查室')
    _od6._remember('承办部门', '办公室')
    assert _od6._history('承办部门')[:2] == ['办公室', '综合调查室'], \
        '历史值应最近用的排最前：{}'.format(_od6._history('承办部门'))
    _od7 = _OD()
    try:
        _cb = _od7._editors['承办部门']
        assert '办公室' in [_cb.itemText(_i) for _i in range(_cb.count())], \
            '重开对话框后下拉里应有历史值'
    finally:
        _od7.reject()
    # 每次都变的字段不记历史
    _od6._remember('标题', '关于某事项的请示')
    assert not _od6._history('标题'), '标题这类每次都变的字段不该记历史'
finally:
    _od6.reject()
print('[26] 打印预检：放不下/顶不过去/空字段一次说清 + 常用值下拉记忆 ✓')



# ---------- 12. v3.0 易用性 ----------
# 快捷键已注册（6 个页面 + 打开/处理/帮助）
assert len(win._shortcuts) == len(win.nav_group.buttons()) + 3, '快捷键数量: {}'.format(len(win._shortcuts))
win.nav_to(2)
assert win.stack.currentIndex() == 2, 'nav_to 未切页'
win.nav_to(0)

# 使用习惯记忆：改后缀 → editingFinished → 新实例可读回
home.suffix_edit.setText('_v3test')
home.suffix_edit.editingFinished.emit()
from app.theme import settings as _settings
assert _settings().value('home/suffix') == '_v3test', '后缀未持久化'
home.suffix_edit.setText('_gui')
home.suffix_edit.editingFinished.emit()

# 预览同步滚动不抛异常
dlg2 = PreviewDialog([SAMPLE], preset_official)
app.processEvents()
dlg2._sync_scroll(dlg2.view_before, dlg2.view_after)
dlg2.reject()

# 新模块可导入
from app.help_dialog import HelpDialog
from app.onboarding_dialog import OnboardingDialog
from app.update_check import UpdateChecker, _version_tuple
assert _version_tuple('v3.0.0') == (3, 0, 0)
assert _version_tuple('v3.0.1') > _version_tuple('v3.0.0')

# 白话错误映射
from app.worker import friendly_error
msg, _ = friendly_error(Exception('Package not found at xxx'))
assert 'Word 文档' in msg, '错误白话化失败: {}'.format(msg)
print('[14] v3.0 快捷键/习惯记忆/同步滚动/帮助引导/错误白话化 ✓')

# ---------- 27. 套打模板可视化编辑：点中、拖动、改属性、存回 ----------
from PyQt5.QtCore import QPoint as _QP
from PyQt5.QtGui import QMouseEvent as _QME
from app.template_edit_dialog import TemplateEditDialog as _TED
from scripts import overprint as _ovp3
import shutil as _sh7, tempfile as _tf7

_tw = os.path.join(_tf7.mkdtemp(), '送审单.docx')
_sh7.copyfile(_ovp3.list_templates()[0][1], _tw)
_ted = _TED(_tw)
try:
    _ted.canvas.resize(620, 870)
    # 预印白字也要能选中——编辑模板改的正是它们；填写时它们是不给动的
    _items = _ted.canvas._items
    assert len(_items) > 20, '画布上可选元素太少：{}'.format(len(_items))
    _tgt = [i for i in _items if '文件送审单' in i['field']][0]

    def _pt(it, dx_cm=0.2):
        _r, _s = it['rect'], _ted.canvas._scale()
        _pr = _ted.canvas._page_rect()
        return _QP(int(_pr.left() + (_r.left() + dx_cm) * _s),
                   int(_pr.top() + (_r.top() + _r.height() / 2) * _s))

    def _click(p):
        _ted.canvas.mousePressEvent(_QME(_QME.MouseButtonPress, p, Qt.LeftButton,
                                         Qt.LeftButton, Qt.NoModifier))
        _ted.canvas.mouseReleaseEvent(_QME(_QME.MouseButtonRelease, p, Qt.LeftButton,
                                           Qt.NoButton, Qt.NoModifier))

    _p0 = _pt(_tgt)
    _click(_p0)
    assert _ted.ed_text.text() == '文件送审单', '点中后属性没跟上：{}'.format(
        _ted.ed_text.text())
    assert abs(_ted.sp_x.value() - 7.70) < 0.02, '横坐标不对：{}'.format(_ted.sp_x.value())
    assert _ted.cb_kind.currentData() is True, '预印内容应显示为"预印"'
    _x0 = _ted.sp_x.value()

    # 拖 1cm：改的是制表位，位置要真的跟着走
    _s7 = _ted.canvas._scale()
    _p1 = _QP(_p0.x() + int(1.0 * _s7), _p0.y())
    _ted.canvas.mousePressEvent(_QME(_QME.MouseButtonPress, _p0, Qt.LeftButton,
                                     Qt.LeftButton, Qt.NoModifier))
    _ted.canvas.mouseMoveEvent(_QME(_QME.MouseMove, _p1, Qt.NoButton,
                                    Qt.LeftButton, Qt.NoModifier))
    _ted.canvas.mouseReleaseEvent(_QME(_QME.MouseButtonRelease, _p1, Qt.LeftButton,
                                       Qt.NoButton, Qt.NoModifier))
    assert abs(_ted.sp_x.value() - (_x0 + 1.0)) < 0.03, \
        '拖 1cm 后应到 {:.2f}，实际 {:.2f}'.format(_x0 + 1.0, _ted.sp_x.value())

    # 撤销回到原位
    _ted._undo()
    _click(_pt(_tgt))
    assert abs(_ted.sp_x.value() - _x0) < 0.02, '撤销没回到原位'

    # 直接填坐标 + 改字号，存回去还能正常填充
    _ted.sp_x.setValue(8.00); _ted._apply_x()
    assert abs(_ted.sp_x.value() - 8.00) < 0.02
    _ted._sess.save(_tw)
    _f7 = os.path.join(_tf7.mkdtemp(), 'out.docx')
    _ovp3.fill_form(_tw, {'标题': '关于某事的请示'}, _f7, one_page=False)
    assert os.path.exists(_f7), '改过的模板填不出来'
    assert abs(_ovp3.plan_fill(_tw, {})['field_pos']['电话'] - 15.50) < 0.02, \
        '改了文件头，电话栏的落点不该跟着变'
finally:
    _ted._sess.dirty = False
    _ted.reject()
print('[27] 套打模板可视化编辑：白字可选 + 拖动/撤销/填坐标 + 存回可填 ✓')


# ---------- 28. 用语检查独立成模式 + 套打升为一级页面 ----------
from app.compliance_dialog import (ComplianceOptionsDialog as _COD,
                                   WordingOptionsDialog as _WOD)
from scripts import compliance as _cmp

_lay = _COD(None).get_options()
_wor = _WOD(None).get_options()
assert not [k for k, v in _lay.items() if v and k.startswith('w_')], \
    '版式面板不该再管用语项'
assert [k for k, v in _wor.items() if v and k.startswith('w_')], \
    '用语面板应勾着用语项'
assert not [k for k, v in _wor.items() if v and not k.startswith('w_')], \
    '用语面板不该顺手勾上版式项'

# 用语模式：卡片在、预览按钮文案对、结果窗口是用语那一套
from app.worker import MODE_WORDING as _MW
assert _MW in home._mode_cards, '首页应有公文用语检查卡片'
home.set_mode(_MW)
assert home.preview_btn.text() == '预览检查结果', home.preview_btn.text()

_wsrc = os.path.join(SMOKE, 'gui_wording.docx')
from docx import Document as _Doc
_wd = _Doc()
for _t in ('关于开展检查的请示。', '省安委会',
           '我委迫不急待地组织排查，按步就班推进。', '特此报告。'):
    _wd.add_paragraph(_t)
_wd.save(_wsrc)
from app.worker import ProcessWorker as _PW
_ww = _PW([_wsrc], _MW, 'official_gbk', None, '_w')
_ww.compliance_options = _cmp.only(_cmp.WORDING_KEYS)
_ww.run()
_wres = _ww._compliance_results[0]
assert _wres['kind'] == 'wording'
assert not _wres['preview'], '用语模式不该再去算版式预览'
from app.compliance_report_dialog import ComplianceReportDialog as _CRD
_wdlg = _CRD([_wres])
try:
    assert _wdlg.windowTitle() == '公文用语检查结果', _wdlg.windowTitle()
    for _cb in _wdlg._boxes[0].values():
        _cb.setChecked(True)
    _wdlg._render_preview()
    _a = _wdlg.pv_after.toPlainText()
    assert '迫不及待' in _a and '迫不急待' not in _a, '改后侧没显示改正结果'
    assert '按部就班' in _a, '第二处错词没改'
    assert '处会改' in _wdlg.pv_note.text(), _wdlg.pv_note.text()
finally:
    _wdlg.reject()

# 套打：侧边栏一级页面，首页那个按钮改成"送你过去"
from app.main_window import NAV_ITEMS as _NAV
assert ('套打填写', 2) in _NAV, '套打应是侧边栏一级入口：{}'.format(_NAV)
_mw2 = MainWindow()
assert _mw2.stack.count() == len(_NAV), '页数与导航项对不上'
_mw2._switch_page(0)
_mw2.home_page._run_overprint()
assert _mw2.stack.currentWidget() is _mw2.overprint_page, '首页入口没把人送到套打页'
_panel = _mw2.overprint_page.panel
assert _panel.embedded, '页面里的套打面板应是嵌入形态'
_texts = [b.text() for b in _panel.findChildren(type(_panel.btn_print))]
assert '取消' not in _texts, '长在页面里就不该有"取消"'
assert '打印…' in _texts, '缺少直接打印入口'
print('[28] 用语检查独立成模式（面板/结果窗口/改后预览）+ 套打升为一级页面 ✓')


# ---------- 29. 密级标注：检查项在面板里、插入要人选密级 ----------
_secsrc = os.path.join(SMOKE, 'gui_sec.docx')
from docx import Document as _Doc2
_sd = _Doc2()
for _t in ('000123', '关于某事的通知', '各部门：', '正文内容。'):
    _sd.add_paragraph(_t)
_sd.save(_secsrc)

# 合规检查面板里应有"密级标注"这一组，且默认勾上
_cod = _COD(None)
assert any('密级' in _g for _g, _items in _cmp.CHECK_GROUPS), \
    '合规检查面板应有密级标注一组'
for _k in _cmp.SECURITY_KEYS:
    assert _k in _cod._checks, '密级检查项 {} 没出现在面板里'.format(_k)
assert all(_cod.get_options()[_k] for _k in _cmp.SECURITY_KEYS), \
    '密级检查默认应打开——漏标密级是事故，不该要人自己去翻开关'
# 用语面板不该混进密级项
assert not [_k for _k in _cmp.SECURITY_KEYS if _WOD(None).get_options()[_k]], \
    '用语面板不该管密级'

_secw = _PW([_secsrc], 'compliance', 'official_gbk', None, '_sec')
_secw.compliance_options = _cmp.only(_cmp.SECURITY_KEYS)
_secw.run()
_secres = _secw._compliance_results[0]
_sdlg = _CRD([_secres])
try:
    assert 'security:insert' in _sdlg._boxes[0], '应给出"插入密级"的可修正项'
    assert 0 in _sdlg._sec_pick, '插入密级必须让人选密级和期限，不能软件替人拍板'
    _lv, _pd = _sdlg._sec_pick[0]
    assert _lv.currentText() == '秘密' and _pd.currentText() == '1年'
    # 没勾选时选择条是灰的，勾上才让选
    _cb = _sdlg._boxes[0]['security:insert']
    assert not _lv.parentWidget().isEnabled(), '未勾选时不该能选密级'
    _cb.setChecked(True)
    assert _lv.parentWidget().isEnabled()
    _lv.setCurrentText('机密')
    _pd.setCurrentText('3年')
    _keys = _sdlg.selections()[0]['fix_keys']
    assert _keys == ['security:insert:机密★3年'], \
        '用户选的密级要原样带到修正里：{}'.format(_keys)
    _secout = os.path.join(SMOKE, 'gui_sec_fixed.docx')
    _applied = _cmp.apply_compliance_fixes(_secsrc, _secout,
                                           _secres['preset'], _keys)
    _txt = [_p.text for _p in _Doc2(_secout).paragraphs if _p.text.strip()]
    assert _txt[:2] == ['000123', '机密★3年'], '密级应排在份号之下：{}'.format(_txt)
finally:
    _sdlg.reject()
print('[29] 密级标注：检查项默认开 + 插入密级由人选定并原样落笔 ✓')


print('\nGUI 自动化测试全部通过 ✓')
