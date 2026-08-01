# -*- coding: utf-8 -*-
"""按实测尺寸重建「文件送审单」套打模板。

所有数字来自用户拿直尺量真实预印纸的结果（见 SPEC），不是反推的。
改尺寸只改 SPEC，重跑本脚本即可——模板不再是"改一次错一次"的手工活。

坐标约定：一律用"距纸张左边 / 距纸张上边"的厘米数；用户从下边量的值
在 SPEC 里按 H - d 换算好，注释里保留原始说法便于核对。

垂直定位靠"精确行距 + 段前距"，横向定位靠**制表位**（缇为单位、与字体
无关）。生成后用 tools/check_songshendan.py 渲染成 PDF 逐项实测复核。
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PT_PER_CM = 28.3465
W = 21.0                     # 纸宽（实测=标称）
# 纸高用标称 A4 的 29.7：用户尺子读数 29.6 是量纸误差，而"距纸下边多少"
# 这个关系必须相对**真实纸边**成立，所以按 29.7 换算（照 29.6 算会整体高 1mm）
H = 29.7
RED = 'CC2222'               # 仅生成"红头样张"时用；模板本身是白字

SPEC = {
    # ---- 文件头 ----
    'head1_top': 2.8,        # 单位名称行 上边线距纸上边
    'head1_left': 5.3,       # 距纸左侧
    'head1_right': 5.0,      # 距纸右侧
    'head1_gap_cm': 0.9,     # 两组单位名之间的空当
    'head2_top': 3.7,        # 「文件送审单」上边线
    'head2_bottom': 4.5,     # 下边线
    'head2_left': 7.7, 'head2_right': 7.5,

    # ---- 紧急程度 / 密级 ----
    'urgent_left': 2.5,      # 「紧急程度：」左边线
    'urgent_top': 5.2,       # 「紧急程度」上边线距纸上边（实测）
    'sec_right': 5.1,        # 「密级：」右边线距纸右侧
    'rule_after_urgent': 5.7,  # 该行下面的红线距纸上边

    # ---- 表格区（红线的位置就是表格行线）----
    'rule_after_title': 8.0,
    'title_text_top': 6.6, 'title_left': 2.4,
    # 标题栏那条竖线：实测 4.7。原先是拿「紧急程度：」冒号右边线（4.6）
    # 代替的——用户说"基本一致"，实际差 1mm，"基本"确实只是基本
    'title_vline': 4.7,
    'lead_text_top': 8.5, 'lead_left': 2.4, 'lead_right': 4.5,
    'rule_before_opinion': H - 9.5,   # 距纸下面 9.5
    'opinion_text_top': H - 9.0,      # 距纸下面 9.0
    # 拟办意见正文首行的左边线：实测 3.3，正好落在栏目名「意」字上。
    # "空两个字"说的是**栏目名**的两个字（2×0.42），不是正文小三的两个字
    # （2×0.529=1.06，那样会排到 3.46、越过「见」字）
    'opinion_body_left': 3.3,
    'rule_after_opinion': H - 4.9,    # 距纸下面 4.9
    'dept_text_top': H - 4.0,         # 承办部门 距纸下面 4.0
    'handler_text_top': H - 4.5,      # 经办人   距纸下面 4.5
    'rule_mid_right': H - 3.9,        # 经办人/文字校核之间 距纸下面 3.9
    'check_text_top': H - 3.5,        # 文字校核 距纸下面 3.5
    'rule_bottom': H - 2.9,           # 文字校核下面 距纸下面 2.9
    'vline_from_right': 12.0,         # 承办部门右侧竖线 距纸右侧
    'handler_right': 9.8,             # 「经办人：」右边线距纸右侧
    'handler_to_phone': 3.1,          # 经办人：与电话：之间
    'phone_right': 5.5,               # 「电话：」右边线距纸右侧

    # ---- 成文日期 / 落款 ----
    'ymd_year_left': 4.4, 'ymd_month_left': 5.6, 'ymd_day_left': 7.0,
    'sign_left_from_right': 8.1,      # 落款首字左边线距纸右侧
    'sign_right_from_right': 2.2,     # 落款尾字右边线距纸右侧
    'sign_bottom_from_bottom': 2.2,   # 该行下边线距纸下面
    'sign_text': '某地市某某单位的办公室制',

    # ---- 长红线左右 ----
    'rule_side': 2.1,
}

HEAD1_LEFT_TEXT = '中国某地市某单位'      # 8 字
HEAD1_RIGHT_TEXT = '某地市某单位'          # 6 字
HEAD2_TEXT = '文件送审单'

# ---------------------------------------------------------------- 字体
# 预印纸上的字体与字号由用户照实说明，公文号数换算：
#   小二 18pt   二号 22pt   三号 16pt   小三 15pt   四号 14pt
# 元组是 (中文字体, 西文字体, 字号pt, 是否加粗)。
#
# 白字栏目名的字体也得给准：它不显影，但它有多宽决定了紧跟其后的
# 黑字从哪儿开始印。字号猜小了，填的内容就整体左移。
F_HEAD1 = ('方正大标宋简体', 'Times New Roman', 18.0, True)   # 顶端单位名称
F_HEAD2 = ('方正大标宋简体', 'Times New Roman', 22.0, True)   # 文件送审单
F_LABEL = ('方正楷体_GBK', 'Times New Roman', 14.0, True)     # 各栏目名
F_SIGN = ('方正楷体_GBK', 'Times New Roman', 14.0, True)      # 落款
# ---- 以下是"要打印出来"的内容 ----
F_TITLE = ('方正小标宋_GBK', 'Times New Roman', 16.0, True)   # 标题
F_TEXT = ('方正楷体_GBK', 'Times New Roman', 14.0, True)      # 紧急程度/密级/承办部门/经办人/文字校核
F_OPINION = ('方正仿宋_GBK', 'Times New Roman', 15.0, True)   # 拟办意见
F_NUM = ('Times New Roman', 'Times New Roman', 14.0, True)    # 电话、年月日的数字

LABEL_PT = F_LABEL[2]
TITLE_PT = F_TITLE[2]
SIGN_PT = F_SIGN[2]


def _font(run, spec):
    """按 (中文, 西文, 磅, 粗) 设置字体。"""
    cn, en, pt, bold = spec
    run.font.size = Pt(pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:cs'), en)
    return run


def _track(total_cm, n_chars, pt):
    """由实测的整段宽度反推字距。

    文件头那两行在纸上是拉开排的，字距不是随便定的常数——量出来的左右
    边线之间要正好放下 n 个字，字距只能由它解出来：
        总宽 = n × 字宽 + (n-1) × 字距
    """
    if n_chars < 2:
        return 0.0
    return (total_cm - n_chars * pt / PT_PER_CM) / (n_chars - 1)


# 栏目名一个字占多宽，由实测反推：「领导批示：」左边线 2.4、冒号右边线 4.5，
# 五个字（含冒号）跨 2.10cm，一个字 0.42cm。四号足宽是 0.4939cm——纸上的
# 栏目名是**收着排**的，不是足宽。这个差别不能忽略：栏目名有多宽，决定了
# 紧跟其后的黑字从哪儿起印，按足宽算每栏都会右移两三毫米。
LABEL_UNIT_CM = (4.5 - 2.4) / 5
# 个别栏目另有实测，直接用实测值（键是栏目名）
LABEL_WIDTH_CM = {
    '领导批示：': 4.5 - 2.4,          # 实测
    # 「经办人：」和「文字校核：」在纸上宽度一致——「经办人」三个字之间
    # 排得更开，正是为了凑齐这个宽度。取五个字的宽。
    '经办人：': 5 * LABEL_UNIT_CM,
    '文字校核：': 5 * LABEL_UNIT_CM,
    # 「电话：」左边线 = 经办人冒号右边线 + 3.1，右边线距纸右 5.5 → 宽 1.2
    '电话：': (21.0 - 5.5) - ((21.0 - 9.8) + 3.1),
}


def _lw(label):
    """栏目名的宽度（cm）"""
    w = LABEL_WIDTH_CM.get(label)
    return w if w is not None else len(label) * LABEL_UNIT_CM


def _label(para, text, spec=None):
    """排一个预印栏目名：白字，且**宽度正好等于实测宽度**。

    字距由目标宽度反解——w:spacing 是每个字后面都加，所以
        总宽 = 字数 × (字宽 + 字距)
    解出字距即可。收着排是负值、撑开排是正值：「经办人：」四个字要凑成
    五个字的宽，解出来就是正的，正对应用户说的"经办人三个字中间更宽"。
    这样栏目名后面紧跟的填写位，自然就落在实测的冒号右边线上，不必再为
    它单设制表位（设了反而会因为笔位正好压在制表位上而被跳过）。
    """
    spec = spec or F_LABEL
    r = para.add_run(text)
    _font(r, spec)
    _white(r)
    n = len(text)
    if n:
        _set_spacing(r, _lw(text) / n - spec[2] / PT_PER_CM)
    return r


def _white(run):
    """预印在纸上的内容：白字占位，不显影但占准位置"""
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _set_spacing(run, cm):
    """字距（w:spacing，单位二十分之一磅）"""
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement('w:spacing')
    el.set(qn('w:val'), str(int(round(cm * 28.3465 * 20))))
    rPr.append(el)


def _exact_line(para, pt, before_pt=0.0):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt)
    # 段前距不能为负（OOXML 不允许），钳到 0；真要往上挪得靠上一个元素让位
    pf.space_before = Pt(max(0.0, before_pt))
    pf.space_after = Pt(0)
    # 关掉网格吸附，否则行高会被文档网格改写、纵向位置全乱。
    # 必须按 schema 顺序插入：w:pPr 的子元素次序是有约束的，直接 append
    # 会排到 w:spacing 之后，Word/LO 可能整段忽略后续设置（实测制表位失效）。
    pPr = para._p.get_or_add_pPr()
    sg = OxmlElement('w:snapToGrid')
    sg.set(qn('w:val'), '0')
    pPr.insert_element_before(
        sg, 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
        'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
        'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
        'w:rPr', 'w:sectPr', 'w:pPrChange')
    return para


def _tabs(para, stops):
    """stops: [(cm_from_left_margin, 'left'|'right')]"""
    ts = para.paragraph_format.tab_stops
    for pos, kind in stops:
        ts.add_tab_stop(Cm(pos), WD_TAB_ALIGNMENT.RIGHT if kind == 'right'
                        else WD_TAB_ALIGNMENT.LEFT)


def _cell_borders(cell, **sides):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        val = sides.get(side)
        if val:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '8')
            el.set(qn('w:color'), 'FFFFFF')   # 白线：占位不显影
        else:
            el.set(qn('w:val'), 'none')
        b.append(el)
    tcPr.append(b)


def _row_height(row, cm):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(int(round(cm * 566.93))))
    h.set(qn('w:hRule'), 'exact')
    trPr.append(h)


CALIB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                          'songshendan_calib.json')


def load_calib():
    """纵向修正量（cm）。字面高度与行盒顶端的差随字体、字号而变，
    与其猜公式，不如量出来：build → 渲染 → 实测 → 回填，迭代收敛。"""
    import json
    try:
        with open(CALIB_PATH, 'r', encoding='utf-8') as f:
            return {k: float(v) for k, v in json.load(f).items()}
    except (IOError, OSError, ValueError):
        return {}


def build(path, top_margin_cm=None, calib=None):
    S = SPEC
    C = load_calib() if calib is None else calib

    def cal(key):
        return C.get(key, 0.0)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(W), Cm(29.7)   # 打印用标称 A4
    sec.left_margin = Cm(S['rule_side'])
    sec.right_margin = Cm(S['rule_side'])
    sec.top_margin = Cm(top_margin_cm if top_margin_cm is not None
                        else S['head1_top'] + cal('head1'))
    sec.bottom_margin = Cm(0.6)
    # 去掉文档网格，纵向才由我们说了算
    sectPr = sec._sectPr
    for g in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(g)

    body_left = S['rule_side']          # 版心左边界（= 长红线左端）

    def rel(x_cm, origin=None):
        """纸面绝对横坐标 → 制表位坐标。

        制表位是相对**当前容器左沿**的：普通段落是左边距，表格单元格里
        则是该单元格的左沿。用错原点会整体偏出一大截（实测差 7cm）。
        """
        return x_cm - (body_left if origin is None else origin)

    # ---------- 文件头第一行：单位名称 ----------
    head1_pt = F_HEAD1[2]
    # 两组名称之间还夹着 0.9cm 空当，字距要从"净宽"里解
    head1_span = (W - S['head1_right']) - S['head1_left'] - S['head1_gap_cm']
    head1_track = _track(head1_span,
                         len(HEAD1_LEFT_TEXT) + len(HEAD1_RIGHT_TEXT), head1_pt)
    # 第二组的起点直接用制表位，不再拿"定宽空格"顶：空格有多宽随字体变，
    # 顶出来的空当也就跟着变（实测差 0.26cm）。制表位是绝对位置，不受影响。
    head1_g2 = (S['head1_left'] + len(HEAD1_LEFT_TEXT) * head1_pt / PT_PER_CM
                + (len(HEAD1_LEFT_TEXT) - 1) * head1_track + S['head1_gap_cm'])
    p = doc.add_paragraph()
    _exact_line(p, head1_pt * 1.0)
    _tabs(p, [(rel(S['head1_left']), 'left'), (rel(head1_g2), 'left')])
    p.add_run('\t')
    r = p.add_run(HEAD1_LEFT_TEXT)
    _font(r, F_HEAD1); _white(r); _set_spacing(r, head1_track)
    p.add_run('\t')
    r = p.add_run(HEAD1_RIGHT_TEXT)
    _font(r, F_HEAD1); _white(r); _set_spacing(r, head1_track)

    # ---------- 文件头第二行：文件送审单 ----------
    head2_pt = F_HEAD2[2]
    head2_track = _track((W - S['head2_right']) - S['head2_left'],
                         len(HEAD2_TEXT), head2_pt)
    p = doc.add_paragraph()
    _exact_line(p, head2_pt * 1.0,
                before_pt=(S['head2_top'] - S['head1_top']
                           - head1_pt / PT_PER_CM) * PT_PER_CM + cal('head2') * PT_PER_CM)
    _tabs(p, [(rel(S['head2_left']), 'left')])
    p.add_run('\t')
    r = p.add_run(HEAD2_TEXT)
    _font(r, F_HEAD2); _white(r); _set_spacing(r, head2_track)

    # ---------- 紧急程度 / 密级 ----------
    lp = LABEL_PT
    p = doc.add_paragraph()
    _exact_line(p, lp * 1.4,
                before_pt=(S['urgent_top'] - S['head2_bottom']) * PT_PER_CM
                + cal('urgent') * PT_PER_CM)
    # 「密级：」用户量的是**右边线**（冒号的右沿），填的字紧接其后。
    # 这里给栏目名和填写位各一个左制表位：填写位钉在实测的右边线上，
    # 栏目名往左退它自己的宽度。栏目名全是全角字，宽度 = 字数 × 字号，
    # 与具体字体无关，所以这样算是准的；右对齐制表位反而靠不住——
    # 一旦栏目名退到容器左沿以外，渲染器会把它顶回来，后面的制表位
    # 全跟着串位（实测「电话」错出 1.5cm）。
    sec_x = W - S['sec_right']
    _tabs(p, [(rel(S['urgent_left']), 'left'),
              (rel(sec_x - _lw('密级：')), 'left')])
    p.add_run('\t')
    _label(p, '紧急程度：')
    r = p.add_run('{{紧急程度}}'); _font(r, F_TEXT)
    p.add_run('\t')
    _label(p, '密级：')
    r = p.add_run('{{密级}}'); _font(r, F_TEXT)

    # ---------- 表格 ----------
    rows = [
        ('title', S['rule_after_title'] - S['rule_after_urgent']),
        ('lead', S['rule_before_opinion'] - S['rule_after_title']),
        ('opinion', S['rule_after_opinion'] - S['rule_before_opinion']),
        ('dept1', S['rule_mid_right'] - S['rule_after_opinion']),
        ('dept2', S['rule_bottom'] - S['rule_mid_right']),
    ]
    # 三列网格：标题栏的竖线和承办部门栏的竖线不在同一处，两条竖线各占
    # 一个网格线，用 gridSpan 合并出各行实际的分栏。
    #   ├ 2.10 ── 标题竖线 ── 承办部门竖线 ── 18.90 ┤
    title_vline = S['title_vline']
    dept_vline = W - S['vline_from_right']
    edges = [body_left, title_vline, dept_vline, W - S['rule_side']]
    widths = [edges[i + 1] - edges[i] for i in range(3)]
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    grid = table._tbl.find(qn('w:tblGrid'))
    for gc, wcm in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(round(wcm * 566.93))))
    # 表格整体不缩进。关键是把**表级**单元格内边距归零：Word/LO 会把
    # 表格左沿放在"左边距 − 单元格左内边距"处，默认 0.19cm，
    # 实测整张表左移到 1.91cm（应为 2.10），表内所有文字跟着差 0.19。
    tblPr = table._tbl.find(qn('w:tblPr'))
    # 表级白色框线：套打识别（cleaner.looks_like_overprint）和预览画线
    # 都读 tblBorders，只设单元格级的话它们看不到
    tb = OxmlElement('w:tblBorders')
    for side, on in (('top', True), ('left', False), ('bottom', True),
                     ('right', False), ('insideH', True), ('insideV', True)):
        e = OxmlElement('w:' + side)
        if on:
            e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '8')
            e.set(qn('w:color'), 'FFFFFF')
        else:
            e.set(qn('w:val'), 'none')
        tb.append(e)
    tblPr.append(tb)
    cm0 = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + side)
        e.set(qn('w:w'), '0'); e.set(qn('w:type'), 'dxa')
        cm0.append(e)
    tblPr.append(cm0)
    ind = OxmlElement('w:tblInd')
    ind.set(qn('w:w'), '0'); ind.set(qn('w:type'), 'dxa')
    tblPr.append(ind)

    for i, (kind, h) in enumerate(rows):
        _row_height(table.rows[i], h)
        # 显式写 tcW：只改 tblGrid 的话，读宽度的代码（自适应、预览）
        # 会退回"平均分"，与真实版面不符
        for c, wcm in zip(table.rows[i].cells, widths):
            c.width = Cm(wcm)
        for c in table.rows[i].cells:
            c.vertical_alignment = None
            tcPr = c._tc.get_or_add_tcPr()
            mar = OxmlElement('w:tcMar')
            for side in ('top', 'left', 'bottom', 'right'):
                e = OxmlElement('w:' + side)
                e.set(qn('w:w'), '0'); e.set(qn('w:type'), 'dxa')
                mar.append(e)
            tcPr.append(mar)

    def merge_row(i, a=0, b=2):
        c = table.rows[i].cells[a]
        for j in range(a + 1, b + 1):
            c = c.merge(table.rows[i].cells[j])
        return c

    # 标题行：栏目名单占一格，标题正文在竖线右边的格子里**居中**。
    # 竖线的位置来自实测——与「紧急程度：」冒号后第一个字的左边线对齐。
    c = table.rows[0].cells[0]
    _cell_borders(c, top=True, bottom=True, right=True)
    p = c.paragraphs[0]
    _exact_line(p, lp * 1.0,
                before_pt=(S['title_text_top'] - S['rule_after_urgent'])
                * PT_PER_CM + cal('title') * PT_PER_CM)
    _tabs(p, [(rel(S['title_left']), 'left')])
    p.add_run('\t')
    _label(p, '标  题')

    c = merge_row(0, 1, 2)
    _cell_borders(c, top=True, bottom=True)
    p = c.paragraphs[0]
    # 行距按**标题正文**字号设，不能按栏目名的字号：标题回成两行时
    # 小字号的行距装不下大字，两行会叠在一起还压出栏外
    _exact_line(p, TITLE_PT * 1.0,
                before_pt=(S['title_text_top'] - S['rule_after_urgent'])
                * PT_PER_CM + cal('title') * PT_PER_CM)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('{{标题}}'); _font(r, F_TITLE)

    # 领导批示行
    c = merge_row(1)
    _cell_borders(c, bottom=True)
    p = c.paragraphs[0]
    _exact_line(p, lp * 1.0,
                before_pt=(S['lead_text_top'] - S['rule_after_title']) * PT_PER_CM
                + cal('lead') * PT_PER_CM)
    _tabs(p, [(rel(S['lead_left']), 'left')])
    p.add_run('\t')
    _label(p, '领导批示：')

    # 拟办意见行
    c = merge_row(2)
    _cell_borders(c, bottom=True)
    p = c.paragraphs[0]
    _exact_line(p, lp * 1.0,
                before_pt=(S['opinion_text_top'] - S['rule_before_opinion'])
                * PT_PER_CM + cal('opinion') * PT_PER_CM)
    _tabs(p, [(rel(S['lead_left']), 'left')])
    p.add_run('\t')
    _label(p, '拟办意见：')
    p2 = c.add_paragraph()
    _exact_line(p2, F_OPINION[2] * 1.4)
    # 正文首行缩进两个字：公文行文惯例，拟办意见是成段的话。
    # 起算点是栏目名的左边线（不是格子左沿），这样"空两格"是相对
    # 「拟办意见：」那一列说的，看着才齐。
    p2.paragraph_format.left_indent = Cm(S['lead_left'] - body_left)
    p2.paragraph_format.first_line_indent = Cm(S['opinion_body_left']
                                               - S['lead_left'])
    r = p2.add_run('{{拟办意见}}'); _font(r, F_OPINION)

    # 承办部门（纵向合并两行）/ 经办人 / 文字校核
    left_cell = merge_row(3, 0, 1).merge(merge_row(4, 0, 1))
    _cell_borders(left_cell, bottom=True, right=True)
    p = left_cell.paragraphs[0]
    _exact_line(p, lp * 1.0,
                before_pt=(S['dept_text_top'] - S['rule_after_opinion'])
                * PT_PER_CM + cal('dept') * PT_PER_CM)
    _tabs(p, [(rel(S['lead_left']), 'left')])
    p.add_run('\t')
    _label(p, '承办部门：')
    r = p.add_run('{{承办部门}}'); _font(r, F_TEXT)

    # 用户量的是「经办人：」「电话：」的**右沿**（含冒号），填的字紧接其后。
    # 同上：填写位钉在实测右沿，栏目名往左退自己的宽度；退过了单元格
    # 左沿就贴着左沿放（「文字校核：」比「经办人：」长一个字，真会顶出去）。
    hr = W - S['handler_right']
    pr = W - S['phone_right']
    cell_left = dept_vline                     # 右列单元格左沿
    for ri, (label, key, top_cm, prev_rule, fld_font) in (
            (3, ('经办人：', '经办人', S['handler_text_top'],
                 S['rule_after_opinion'], F_TEXT)),
            (4, ('文字校核：', '文字校核', S['check_text_top'],
                 S['rule_mid_right'], F_TEXT))):
        c = table.rows[ri].cells[2]
        _cell_borders(c, bottom=True)
        p = c.paragraphs[0]
        _exact_line(p, lp * 1.0,
                    before_pt=(top_cm - prev_rule + cal('handler')) * PT_PER_CM)
        stops = [(max(0.02, rel(hr - _lw(label), cell_left)), 'left')]
        if ri == 3:
            stops.append((rel(pr - _lw('电话：'), cell_left), 'left'))
        _tabs(p, stops)
        p.add_run('\t')
        _label(p, label)
        r = p.add_run('{{%s}}' % key); _font(r, fld_font)
        if ri == 3:
            p.add_run('\t')
            _label(p, '电话：')
            r = p.add_run('{{电话}}'); _font(r, F_NUM)

    # ---------- 成文日期 + 落款 ----------
    p = doc.add_paragraph()
    _exact_line(p, SIGN_PT * 1.0,
                before_pt=((H - S['sign_bottom_from_bottom'] - SIGN_PT / PT_PER_CM)
                           - S['rule_bottom'] + cal('sign')) * PT_PER_CM)
    sign_left = W - S['sign_left_from_right']
    sign_track = _track((W - S['sign_right_from_right']) - sign_left,
                        len(S['sign_text']), SIGN_PT)
    # 数字用**右对齐**制表位顶到预印「年/月/日」的左沿，预印字紧随其后：
    # 这样一位数两位数都自动贴齐，不必按位数补空格
    ymd = [S['ymd_year_left'], S['ymd_month_left'], S['ymd_day_left']]
    stops = []
    for x in ymd:
        stops.append((rel(x), 'right'))        # 数字右沿顶到预印字左沿
        stops.append((rel(x + 0.01), 'left'))  # 预印字本身；错开 0.01cm 免得并成一个
    stops.append((rel(sign_left), 'left'))
    _tabs(p, stops)
    for key, ch in (('年', '年'), ('月', '月'), ('日', '日')):
        p.add_run('\t')
        r = p.add_run('{{%s}}' % key); _font(r, F_NUM)
        p.add_run('\t')
        r = p.add_run(ch); _font(r, F_LABEL); _white(r)
    p.add_run('\t')
    r = p.add_run(S['sign_text']); _font(r, F_SIGN); _white(r)
    _set_spacing(r, sign_track)

    doc.save(path)
    return path


if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 else \
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                     'templates', '套打', '文件送审单.docx')
    build(out)
    print('已生成', out)
